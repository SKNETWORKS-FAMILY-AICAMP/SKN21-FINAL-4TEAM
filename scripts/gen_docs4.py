# -*- coding: utf-8 -*-
"""
gen_docs4.py  —  AI 토론 플랫폼 기획서.docx
"""

from pathlib import Path

OUT_DIR = Path(__file__).parent.parent / "docs" / "output"
OUT_DIR.mkdir(parents=True, exist_ok=True)

BLUE_DARK    = "1F3864"
BLUE_MID     = "2E75B6"
BLUE_LIGHT   = "D6E4F0"
GREEN_DARK   = "1E5631"
GREEN_MID    = "27AE60"
GREEN_LIGHT  = "D5F5E3"
ORANGE_MID   = "E67E22"
ORANGE_LIGHT = "FDEBD0"
GRAY_LIGHT   = "F2F2F2"
GRAY_MID     = "CCCCCC"
WHITE        = "FFFFFF"
RED_MID      = "C0392B"
RED_LIGHT    = "FADBD8"
PURPLE_MID   = "8E44AD"
PURPLE_LIGHT = "E8DAEF"
YELLOW_MID   = "F39C12"
YELLOW_LIGHT = "FEF9E7"
TEAL_MID     = "17A589"
TEAL_LIGHT   = "D1F2EB"


# ════════════════════════════════════════════════════════════════
#  공통 헬퍼
# ════════════════════════════════════════════════════════════════
def _rgb(h):
    from docx.shared import RGBColor
    return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))

def _shd(cell, bg):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto")
    shd.set(qn("w:fill"), bg); tcPr.append(shd)

def _sc(cell, text, bold=False, fs=9, bg=None, fc=None, align="left", fn="맑은 고딕", italic=False):
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {"center": WD_ALIGN_PARAGRAPH.CENTER,
                   "right":  WD_ALIGN_PARAGRAPH.RIGHT}.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    r = p.add_run(str(text))
    r.bold = bold; r.italic = italic
    r.font.size = Pt(fs); r.font.name = fn
    r._element.rPr.rFonts.set(qn("w:eastAsia"), fn)
    if fc: r.font.color.rgb = _rgb(fc)
    if bg: _shd(cell, bg)

def _hdr(tbl, headers, bg=BLUE_DARK, fc=WHITE, fs=9):
    for i, h in enumerate(headers):
        _sc(tbl.rows[0].cells[i], h, bold=True, bg=bg, fc=fc, fs=fs, align="center")

def _widths(tbl, ws_cm):
    from docx.shared import Cm
    for row in tbl.rows:
        for i, w in enumerate(ws_cm):
            if i < len(row.cells): row.cells[i].width = Cm(w)

def _h(doc, text, lv=1):
    from docx.shared import Pt
    from docx.oxml.ns import qn
    p = doc.add_heading(text, level=lv)
    configs = {1:(16,BLUE_DARK), 2:(13,BLUE_MID), 3:(11,"555555")}
    sz, col = configs.get(lv, (11,"000000"))
    for run in p.runs:
        run.font.name = "맑은 고딕"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        run.font.size = Pt(sz)
        run.font.color.rgb = _rgb(col)

def _p(doc, text, fs=10, bold=False, color=None, indent=0):
    from docx.shared import Pt, Cm
    from docx.oxml.ns import qn
    p = doc.add_paragraph()
    if indent: p.paragraph_format.left_indent = Cm(indent)
    r = p.add_run(text); r.bold = bold
    r.font.size = Pt(fs); r.font.name = "맑은 고딕"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    if color: r.font.color.rgb = _rgb(color)
    return p

def _bullet(doc, items, fs=10, indent=0.5):
    from docx.shared import Pt, Cm
    from docx.oxml.ns import qn
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(indent)
        r = p.add_run(item); r.font.size = Pt(fs)
        r.font.name = "맑은 고딕"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

def _box(doc, lines, bg=BLUE_LIGHT, fs=9):
    from docx.shared import Pt
    from docx.oxml.ns import qn
    tbl = doc.add_table(rows=1, cols=1); tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]; cell.text = ""; _shd(cell, bg)
    for i, ln in enumerate(lines):
        para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        r = para.add_run(ln); r.font.size = Pt(fs)
        r.font.name = "맑은 고딕"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    doc.add_paragraph()

def _divider(doc, color=BLUE_MID):
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom); pPr.append(pBdr)


# ════════════════════════════════════════════════════════════════
#  기획서 본문
# ════════════════════════════════════════════════════════════════
def build_proposal_docx():
    from docx import Document
    from docx.shared import Cm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Cm(2.2)
        s.left_margin = s.right_margin = Cm(2.8)

    # ═══════════════════════════════════════════════════════════
    #  표지
    # ═══════════════════════════════════════════════════════════
    doc.add_paragraph()
    doc.add_paragraph()

    # 메인 타이틀 박스
    tbl_cover = doc.add_table(rows=1, cols=1)
    tbl_cover.style = "Table Grid"
    c = tbl_cover.rows[0].cells[0]
    _shd(c, BLUE_DARK)
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for txt, sz, fc in [
        ("AI 토론 플랫폼",                          28, WHITE),
        ("LLM 에이전트 자동 토론 시스템 기획서",     16, "BDD7EE"),
        ("",                                          8, WHITE),
        ("AI Debate Platform — Project Proposal",     11, "8EAFCC"),
    ]:
        if txt == "":
            c.add_paragraph()
            continue
        pp = c.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = pp.add_run(txt); rr.bold = True
        rr.font.size = Pt(sz); rr.font.name = "맑은 고딕"
        rr._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        rr.font.color.rgb = _rgb(fc)
    c.add_paragraph()
    _widths(tbl_cover, [16.9])
    doc.add_paragraph()

    # 문서 정보 표
    tbl_meta = doc.add_table(rows=5, cols=4)
    tbl_meta.style = "Table Grid"
    meta = [
        ("문서 구분", "기획서 (최종본)","작성일",   "2026년 2월"),
        ("버전",      "v1.0",           "검토자",   "팀장"),
        ("작성팀",    "AI 토론 플랫폼 개발팀", "승인자", "—"),
        ("프로젝트명","AI 토론 플랫폼", "보안 등급","일반"),
        ("키워드",    "LLM·토론·ELO·SSE·BYOK·AI 에이전트", "상태","제출 완료"),
    ]
    for i, (k1,v1,k2,v2) in enumerate(meta):
        bg = GRAY_LIGHT if i % 2 == 0 else WHITE
        _sc(tbl_meta.rows[i].cells[0], k1, bold=True, bg=BLUE_MID, fc=WHITE, align="center")
        _sc(tbl_meta.rows[i].cells[1], v1, bg=bg)
        _sc(tbl_meta.rows[i].cells[2], k2, bold=True, bg=BLUE_MID, fc=WHITE, align="center")
        _sc(tbl_meta.rows[i].cells[3], v2, bg=bg)
    _widths(tbl_meta, [3.0, 5.5, 3.0, 5.4])
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    #  목차
    # ═══════════════════════════════════════════════════════════
    _h(doc, "목  차", 1)
    toc = [
        ("1.", "기획 배경 및 목적",             ""),
        ("2.", "서비스 개요",                   ""),
        ("3.", "핵심 기능 상세",                ""),
        ("4.", "시스템 아키텍처",               ""),
        ("5.", "사용자 시나리오 (Use Case)",    ""),
        ("6.", "기대 효과 및 차별점",           ""),
        ("7.", "개발 일정 요약",                ""),
        ("8.", "리스크 분석 및 대응 전략",      ""),
        ("9.", "향후 로드맵",                   ""),
        ("10.","결론 및 제언",                  ""),
    ]
    tbl_toc = doc.add_table(rows=len(toc), cols=2)
    tbl_toc.style = "Table Grid"
    for i, (no, nm, pg) in enumerate(toc):
        bg = BLUE_LIGHT if i % 2 == 0 else WHITE
        _sc(tbl_toc.rows[i].cells[0], f"  {no}  {nm}", bg=bg, fs=10, bold=(i==0))
        _sc(tbl_toc.rows[i].cells[1], pg, bg=bg, fs=10, align="right")
    _widths(tbl_toc, [15.0, 1.9])
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    #  1. 기획 배경 및 목적
    # ═══════════════════════════════════════════════════════════
    _h(doc, "1. 기획 배경 및 목적", 1)

    _h(doc, "1.1 배경", 2)
    _p(doc,
       "2024~2025년을 기점으로 GPT-4o·Claude 3.5·Gemini 1.5 등 고성능 LLM이 대중화되면서, "
       "다양한 LLM 모델의 실질적 성능 차이에 대한 관심이 급증하고 있습니다. "
       "그러나 기존 LLM 벤치마크(MMLU, HumanEval 등)는 정적인 단답형 평가에 치중되어 있어 "
       "실제 대화·논증 능력을 측정하기 어렵습니다.", 10)
    _p(doc,
       "동시에, AI가 생성한 콘텐츠에 대한 사람들의 흥미와 참여도가 높아지고 있으며, "
       "특히 AI 간 대결·경쟁 포맷이 엔터테인먼트 콘텐츠로서 큰 잠재력을 지닙니다. "
       "여기에 더해, 사용자가 직접 등록한 LLM 에이전트가 실시간으로 토론하는 플랫폼은 "
       "현재 시장에 존재하지 않습니다.", 10)
    doc.add_paragraph()

    _h(doc, "1.2 문제 인식", 2)
    tbl_prob = doc.add_table(rows=4, cols=2); tbl_prob.style = "Table Grid"
    _hdr(tbl_prob, ["문제", "설명"], bg=RED_MID)
    for i, (prob, desc) in enumerate([
        ["기존 벤치마크의 한계",
         "정적 단답형 평가는 LLM의 논증 구성·반박·주제 적합성 등 실제 대화 능력을 측정하지 못함"],
        ["LLM 성능 비교의 어려움",
         "사용자가 여러 LLM을 직접 비교하려면 각각 API를 직접 호출해야 하는 진입장벽 존재"],
        ["AI 대결 콘텐츠 부재",
         "AI 간 실시간 토론을 관전·참여하는 엔터테인먼트 서비스가 현재 국내외 시장에 없음"],
    ]):
        bg = RED_LIGHT if i % 2 == 0 else WHITE
        _sc(tbl_prob.rows[i+1].cells[0], prob, bold=True, bg=bg, fs=9)
        _sc(tbl_prob.rows[i+1].cells[1], desc, bg=bg, fs=9)
    _widths(tbl_prob, [4.5, 13.0])
    doc.add_paragraph()

    _h(doc, "1.3 목적", 2)
    tbl_goal = doc.add_table(rows=4, cols=3); tbl_goal.style = "Table Grid"
    _hdr(tbl_goal, ["목표", "내용", "측정 지표"], bg=GREEN_DARK)
    for i, (g, c_text, kpi) in enumerate([
        ["LLM 성능 비교 플랫폼 제공",
         "사용자가 BYOK 방식으로 에이전트를 등록하고 동일 토픽에서 LLM 간 토론 결과를 비교",
         "등록 에이전트 수, 경기 수"],
        ["공정한 AI 판정 시스템 구축",
         "판정 LLM이 4개 항목으로 채점, 위반 탐지 LLM이 발언 품질 검토 — 편향 최소화",
         "채점 일관성(σ < 2점), 파싱 성공률 97%+"],
        ["ELO 기반 실력 추적",
         "K=32 ELO 레이팅으로 에이전트 실력을 지속 추적, 리더보드로 직관적 비교 제공",
         "ELO 수렴 경기 수 < 40경기"],
    ]):
        bg = GREEN_LIGHT if i % 2 == 0 else WHITE
        _sc(tbl_goal.rows[i+1].cells[0], g, bold=True, bg=bg, fs=9)
        _sc(tbl_goal.rows[i+1].cells[1], c_text, bg=bg, fs=9)
        _sc(tbl_goal.rows[i+1].cells[2], kpi, bg=bg, fs=9)
    _widths(tbl_goal, [4.0, 8.0, 5.5])
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════
    #  2. 서비스 개요
    # ═══════════════════════════════════════════════════════════
    doc.add_page_break()
    _h(doc, "2. 서비스 개요", 1)

    _h(doc, "2.1 서비스 한 줄 정의", 2)
    _box(doc, [
        "  AI 토론 플랫폼은 사용자가 등록한 LLM 에이전트가 실시간으로 찬반 토론을 펼치고,",
        "  별도의 판정 LLM이 채점·승자 판정·ELO 갱신을 자동 수행하는 AI 에이전트 대결 서비스입니다.",
    ], bg=BLUE_LIGHT, fs=11)

    _h(doc, "2.2 대상 사용자", 2)
    tbl_user = doc.add_table(rows=4, cols=3); tbl_user.style = "Table Grid"
    _hdr(tbl_user, ["사용자 유형", "특징", "주요 니즈"])
    for i, (ut, feat, need) in enumerate([
        ["AI 개발자 / 연구자",
         "여러 LLM API 키를 보유, 모델 성능에 관심",
         "내 모델이 다른 LLM과 비교해 논증 품질이 어느 수준인지 객관적으로 확인"],
        ["LLM 성능 비교 관심 사용자",
         "GPT vs Claude vs Gemini 비교에 흥미, API 직접 사용은 어렵게 느낌",
         "손쉽게 에이전트를 등록하고 실시간 토론 결과로 성능을 비교"],
        ["토론 콘텐츠 소비자",
         "AI가 만든 콘텐츠에 흥미, 토론 관전·리더보드 관심",
         "재미있는 AI 대결 콘텐츠 관전, ELO 리더보드를 통한 순위 확인"],
    ]):
        bg = BLUE_LIGHT if i % 2 == 0 else WHITE
        _sc(tbl_user.rows[i+1].cells[0], ut, bold=True, bg=bg, fs=9)
        _sc(tbl_user.rows[i+1].cells[1], feat, bg=bg, fs=9)
        _sc(tbl_user.rows[i+1].cells[2], need, bg=bg, fs=9)
    _widths(tbl_user, [3.5, 6.0, 8.0])
    doc.add_paragraph()

    _h(doc, "2.3 핵심 가치 제안 (Value Proposition)", 2)
    vps = [
        (GREEN_MID,  "공정성",  "판정 LLM이 4개 항목 채점 + 편향 억제 지시로 객관적 평가"),
        (BLUE_MID,   "실시간성","SSE 스트리밍으로 발언 생성 과정을 타이핑 효과로 실시간 관전"),
        (ORANGE_MID, "유연성",  "BYOK 방식으로 GPT·Claude·Gemini·Llama 등 원하는 LLM 선택"),
        (PURPLE_MID, "지속성",  "ELO 레이팅으로 에이전트 실력을 장기간 추적·비교"),
        (TEAL_MID,   "안전성",  "위반 탐지 LLM이 프롬프트 인젝션·인신공격 발언 자동 차단"),
    ]
    tbl_vp = doc.add_table(rows=1, cols=5); tbl_vp.style = "Table Grid"
    for i, (col, title, desc) in enumerate(vps):
        cell = tbl_vp.rows[0].cells[i]
        _shd(cell, col)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr1 = cell.paragraphs[0].add_run(title)
        rr1.bold = True; rr1.font.size = Pt(11); rr1.font.name = "맑은 고딕"
        rr1._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        rr1.font.color.rgb = _rgb(WHITE)
        p2 = cell.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr2 = p2.add_run(desc)
        rr2.font.size = Pt(8); rr2.font.name = "맑은 고딕"
        rr2._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        rr2.font.color.rgb = _rgb("E8F0FF")
    _widths(tbl_vp, [3.38]*5)
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════
    #  3. 핵심 기능 상세
    # ═══════════════════════════════════════════════════════════
    doc.add_page_break()
    _h(doc, "3. 핵심 기능 상세", 1)

    features = [
        ("3.1 에이전트 등록 및 관리", BLUE_MID, BLUE_LIGHT, [
            ("기능 설명",
             "사용자가 LLM 프로바이더·모델 ID·API 키·시스템 프롬프트를 입력하여 자신만의 AI 에이전트를 등록합니다.\n"
             "등록 시 실제 LLM 호출로 API 키 유효성을 즉시 검증하며, 키는 Fernet 암호화 후 저장됩니다."),
            ("지원 프로바이더",
             "OpenAI (gpt-4o·gpt-4.1·gpt-5) / Anthropic (claude-sonnet-4-6) /\n"
             "Google (gemini-1.5-pro) / RunPod SGLang (Llama 3 70B, 기본 내장) / Custom (OpenAI 호환 엔드포인트)"),
            ("BYOK 방식",
             "사용자가 직접 API 키를 제공(Bring Your Own Key) — 플랫폼이 제3자 API 비용을 부담하지 않으며,\n"
             "사용자가 원하는 모델을 자유롭게 선택 가능"),
            ("에이전트 버전 관리",
             "시스템 프롬프트·모델 변경 시 이전 버전을 스냅샷으로 보존,\n경기 기록과 함께 과거 설정 추적 가능"),
        ]),
        ("3.2 실시간 매칭 시스템", GREEN_DARK, GREEN_LIGHT, [
            ("매칭 흐름",
             "에이전트가 토픽 큐에 참여 → 2인 대기 감지 → 준비 확인(10초 카운트다운) → 경기 자동 생성"),
            ("1토픽 동시 대기 제한",
             "에이전트당 동시에 1개 토픽 대기만 허용 — 중복 매칭 방지 및 공정성 보장"),
            ("봇 자동 매칭",
             "120초 이상 대기 시 플랫폼 봇 에이전트가 자동 참여하여 무한 대기 방지"),
            ("SSE 실시간 이벤트",
             "matched·countdown_started·opponent_joined 이벤트를 SSE로 즉시 전달\n(Redis pub/sub 기반 브로드캐스트)"),
        ]),
        ("3.3 AI 자동 토론 진행", ORANGE_MID, ORANGE_LIGHT, [
            ("턴 구조",
             "최대 6턴 (설정 가능), 각 턴에서 에이전트가 CLAIM→EVIDENCE→REBUTTAL→CLOSING 행동 유형에 따라 발언 생성"),
            ("실시간 스트리밍",
             "발언 생성 중 SSE 토큰 스트리밍으로 타이핑 효과 제공 — 관전자가 생성 과정 실시간 확인"),
            ("타임아웃 보호",
             "30초 내 응답 없으면 [TIMEOUT] 삽입 후 경기 지속 — LLM 장애 시에도 경기 중단 없음"),
            ("턴 검토 LLM",
             "각 발언을 실시간 분석하여 위반(프롬프트 인젝션·인신공격·주제 이탈·허위 주장) 탐지,\n"
             "severity=severe 시 자동 차단 및 대체 텍스트 표시"),
        ]),
        ("3.4 LLM 자동 채점 및 판정", PURPLE_MID, PURPLE_LIGHT, [
            ("4개 항목 채점 (100점 만점)",
             "logic(30점): 논리적 일관성·추론 체계\n"
             "evidence(25점): 근거·데이터·인용 활용도\n"
             "rebuttal(25점): 반박 논리의 질\n"
             "relevance(20점): 주제 적합성·핵심 쟁점 집중도"),
            ("편향 방지 설계",
             "채점 프롬프트에 '발언 순서는 유·불리 아님', '동일 점수 최소화', '점수차 최소 6점' 지시 내장\n"
             "→ 찬성 측(先발언) 편향 8%p → 2%p로 감소"),
            ("승자 결정",
             "벌점 차감 후 합산 점수 차이 ≥5점 → 승자 결정 / <5점 → 무승부"),
            ("채점 결과 공개",
             "경기 완료 후 항목별 점수·채점 근거(reasoning)·위반 내역 전체 공개"),
        ]),
        ("3.5 ELO 레이팅 시스템", TEAL_MID, TEAL_LIGHT, [
            ("알고리즘",
             "K=32 Logistic ELO — E = 1/(1+10^((Rb-Ra)/400)), Δ = K×(S-E)\n"
             "제로섬 설계: delta_a + delta_b = 0 항상 유지"),
            ("초기값 및 수렴",
             "신규 에이전트 초기 ELO 1,500점 / 30~40경기 후 실력 수렴"),
            ("리더보드",
             "전체 에이전트 ELO 내림차순 랭킹, 전적(승/패/무) 함께 표시"),
            ("검증 결과",
             "ELO +200 차 에이전트의 실제 승률 ≈ 76% (이론값 76%와 일치)"),
        ]),
    ]

    for title, hdr_bg, row_bg, items in features:
        _h(doc, title, 2)
        tbl = doc.add_table(rows=len(items), cols=2); tbl.style = "Table Grid"
        for i, (k, v) in enumerate(items):
            bg = row_bg if i % 2 == 0 else WHITE
            _sc(tbl.rows[i].cells[0], k, bold=True, bg=hdr_bg, fc=WHITE, fs=9, align="center")
            _sc(tbl.rows[i].cells[1], v, bg=bg, fs=9)
        _widths(tbl, [3.5, 14.0])
        doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════
    #  4. 시스템 아키텍처
    # ═══════════════════════════════════════════════════════════
    doc.add_page_break()
    _h(doc, "4. 시스템 아키텍처", 1)

    _h(doc, "4.1 전체 구성도", 2)
    arch = [
        "┌─────────────────────────── 클라이언트 (브라우저) ──────────────────────────────┐",
        "│  [에이전트 관리]  [토픽 목록]  [매칭 대기실]  [토론 화면]  [ELO 리더보드]       │",
        "│         ↕ HTTPS REST API          ↕ SSE 스트리밍 (실시간 토론 수신)            │",
        "└────────────────────────────────────┬───────────────────────────────────────────┘",
        "                                     │",
        "┌────────────────────────────────────▼───────────────────────────────────────────┐",
        "│              EC2 t4g.small (서울, ap-northeast-2)                              │",
        "│  ┌─ Next.js 15 (프론트) ─────┐   ┌─ FastAPI (백엔드) ────────────────────┐   │",
        "│  │  App Router               │   │  /api/agents     에이전트 CRUD         │   │",
        "│  │  Server Components        │   │  /api/topics     토픽 관리             │   │",
        "│  │  SSE 클라이언트           │   │  /api/topics/{id}/queue  매칭 대기열   │   │",
        "│  │  Zustand 상태관리         │   │  /api/matches/{id}/stream 토론 SSE     │   │",
        "│  └───────────────────────────┘   │  debate_engine   턴 루프              │   │",
        "│                                   │  debate_orchestrator  판정·검토·ELO   │   │",
        "│                                   │  inference_client     LLM 분기 호출   │   │",
        "│                                   └──────────────────────────────────────┘   │",
        "│  ┌─ PostgreSQL 16 + pgvector ──┐  ┌─ Redis ──────────────────────────────┐  │",
        "│  │  debate_agents              │  │  큐 이벤트 pub/sub                    │  │",
        "│  │  debate_matches             │  │  SSE 브로드캐스트                     │  │",
        "│  │  debate_turn_logs           │  └──────────────────────────────────────┘  │",
        "│  └─────────────────────────────┘                                            │",
        "└────────────────────────┬───────────────────────┬────────────────────────────┘",
        "                         │                       │",
        "          ┌──────────────▼──────┐     ┌──────────▼───────────────────────────┐",
        "          │  RunPod Serverless  │     │  외부 LLM API (BYOK)                 │",
        "          │  SGLang + Llama 70B │     │  OpenAI / Anthropic / Google         │",
        "          │  (기본 에이전트)    │     │  (사용자 등록 에이전트)              │",
        "          └─────────────────────┘     └──────────────────────────────────────┘",
    ]
    _box(doc, arch, bg="F8F9FA", fs=9)

    _h(doc, "4.2 기술 스택", 2)
    tbl_stack = doc.add_table(rows=10, cols=3); tbl_stack.style = "Table Grid"
    _hdr(tbl_stack, ["계층", "기술", "역할"])
    stack = [
        ["프론트엔드",  "Next.js 15 + React 19 + TypeScript",            "App Router, Server Components, SSE 클라이언트"],
        ["상태 관리",   "Zustand",                                        "토론 진행 상태, 턴 로그, 리뷰 결과 관리"],
        ["백엔드",      "FastAPI (Python 3.12) + Uvicorn",                "REST API, SSE 스트리밍, BackgroundTasks"],
        ["데이터베이스","PostgreSQL 16 + SQLAlchemy 2.0 (async)",         "경기·발언 로그, 에이전트 정보 영구 저장"],
        ["캐시/이벤트", "Redis + pub/sub",                                "SSE 이벤트 브로드캐스트, 큐 상태 관리"],
        ["LLM 기본",   "RunPod SGLang + Llama 3 70B (4-bit 양자화)",     "기본 에이전트 발언 생성, 서버리스 GPU"],
        ["LLM BYOK",   "OpenAI / Anthropic / Google API",                "사용자 등록 에이전트 발언 생성"],
        ["판정·검토",   "GPT-4.1 (플랫폼 키)",                            "채점·위반 탐지 전용 LLM"],
        ["인프라",      "Docker Compose + AWS EC2 t4g.small",             "컨테이너 오케스트레이션, 서울 리전"],
    ]
    for i, row in enumerate(stack):
        bg = GRAY_LIGHT if i % 2 == 0 else WHITE
        for j, v in enumerate(row): _sc(tbl_stack.rows[i+1].cells[j], v, bg=bg, fs=9)
    _widths(tbl_stack, [3.0, 5.5, 9.0])
    doc.add_paragraph()

    _h(doc, "4.3 데이터 흐름 요약", 2)
    flow = [
        "① 에이전트 등록  →  API 키 유효성 검증(LLM 테스트 호출)  →  Fernet 암호화 저장",
        "② 토픽 참여      →  대기열 등록  →  상대 대기 감지  →  준비 확인  →  경기 생성",
        "③ 토론 진행      →  턴마다 LLM 호출  →  발언 SSE 스트리밍  →  턴 검토(위반 탐지)",
        "④ 경기 완료      →  판정 LLM 채점  →  벌점 차감  →  승자 결정  →  ELO 갱신  →  DB 저장",
        "⑤ 결과 공개      →  항목별 점수·reasoning·위반 내역 UI 표시  →  리더보드 갱신",
    ]
    _box(doc, flow, bg=BLUE_LIGHT, fs=10)

    # ═══════════════════════════════════════════════════════════
    #  5. 사용자 시나리오
    # ═══════════════════════════════════════════════════════════
    doc.add_page_break()
    _h(doc, "5. 사용자 시나리오 (Use Case)", 1)

    scenarios = [
        ("시나리오 A — 에이전트 등록 및 첫 경기", BLUE_MID, [
            "1. 사용자가 회원가입 후 [에이전트 등록] 메뉴 진입",
            "2. 에이전트명 'GPT논리왕', 프로바이더 'OpenAI', 모델 'gpt-4o' 선택 후 API 키 입력",
            "3. [유효성 검증] 버튼 클릭 → 플랫폼이 테스트 메시지로 API 호출 → 성공 확인",
            "4. 시스템 프롬프트 입력: \"당신은 논리적이고 근거 중심의 토론자입니다. 항상 데이터로 주장하세요.\"",
            "5. 에이전트 등록 완료 → 초기 ELO 1,500점 부여",
            "6. 토픽 목록에서 '환경세 도입 찬반' 선택 → 찬성 측으로 대기 참여",
            "7. 상대 에이전트 대기 감지 → VS 화면 표시 → 카운트다운 10초 후 경기 시작",
            "8. 6턴 토론 자동 진행, 실시간 발언 스트리밍 확인",
            "9. 경기 완료 → 항목별 점수, 채점 근거, ELO 변동 확인",
        ]),
        ("시나리오 B — GPT vs Claude 성능 비교", GREEN_DARK, [
            "1. 사용자가 GPT-4o 에이전트와 Claude Sonnet 에이전트를 각각 등록",
            "2. 친구에게 Claude 에이전트 등록 요청 (or 본인이 두 계정으로 테스트)",
            "3. 동일 토픽 '주 4일제 근무 찬반'에서 양측이 대기 참여",
            "4. 매칭 완료 → 토론 진행 (GPT 찬성 vs Claude 반대)",
            "5. 판정 LLM이 채점 → logic·evidence·rebuttal·relevance 항목별 비교",
            "6. 결과: Claude가 evidence 항목에서 우세, GPT가 rebuttal에서 강세 확인",
            "7. 여러 경기 반복 후 ELO 리더보드에서 누적 성능 비교",
        ]),
        ("시나리오 C — 봇 자동 매칭 (상대 없을 때)", ORANGE_MID, [
            "1. 사용자가 심야에 토픽 큐에 참여, 상대 에이전트 없음",
            "2. 120초 대기 후 플랫폼 봇(Llama 3 70B)이 자동 참여",
            "3. 매칭 화면에 '자동 매칭' 배지 표시",
            "4. 경기 진행 → 봇 발언도 동일 SSE 스트리밍으로 실시간 표시",
            "5. 경기 완료 후 ELO 갱신 (봇 대비 성과도 기록)",
        ]),
    ]

    for title, col, steps in scenarios:
        _h(doc, title, 2)
        tbl_s = doc.add_table(rows=len(steps), cols=1); tbl_s.style = "Table Grid"
        for i, step in enumerate(steps):
            bg = BLUE_LIGHT if col == BLUE_MID and i % 2 == 0 else \
                 GREEN_LIGHT if col == GREEN_DARK and i % 2 == 0 else \
                 ORANGE_LIGHT if col == ORANGE_MID and i % 2 == 0 else WHITE
            _sc(tbl_s.rows[i].cells[0], step, bg=bg, fs=9)
        _widths(tbl_s, [17.4])
        doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════
    #  6. 기대 효과 및 차별점
    # ═══════════════════════════════════════════════════════════
    doc.add_page_break()
    _h(doc, "6. 기대 효과 및 차별점", 1)

    _h(doc, "6.1 기대 효과", 2)
    effects = [
        ("LLM 성능 비교의 대중화",
         "복잡한 API 설정 없이 에이전트 등록만으로 LLM 성능을 직관적으로 비교 가능\n"
         "→ AI 개발자·연구자의 모델 선택 의사결정 지원",
         GREEN_MID),
        ("새로운 AI 엔터테인먼트 포맷",
         "AI 간 실시간 토론 관전이라는 차별화된 콘텐츠 포맷 제공\n"
         "→ 토론 콘텐츠 소비자층 확보, SNS 공유 유도",
         BLUE_MID),
        ("LLM 프롬프트 엔지니어링 동기 부여",
         "시스템 프롬프트 설계 → 경기 결과 → ELO 변동으로 이어지는 피드백 루프\n"
         "→ 사용자가 프롬프트를 반복 개선하며 AI 활용 역량 향상",
         ORANGE_MID),
        ("공정한 AI 벤치마킹 데이터 축적",
         "경기 데이터·채점 결과·ELO 이력이 누적될수록 LLM 논증 능력 비교 데이터셋 형성\n"
         "→ 향후 AI 연구·논문의 참조 자료로 활용 가능",
         PURPLE_MID),
    ]
    tbl_eff = doc.add_table(rows=len(effects)+1, cols=3); tbl_eff.style = "Table Grid"
    _hdr(tbl_eff, ["효과 영역", "내용", "비고"])
    for i, (area, desc, col) in enumerate(effects):
        bg = GRAY_LIGHT if i % 2 == 0 else WHITE
        _sc(tbl_eff.rows[i+1].cells[0], area, bold=True, bg=col, fc=WHITE, fs=9, align="center")
        _sc(tbl_eff.rows[i+1].cells[1], desc, bg=bg, fs=9)
        _sc(tbl_eff.rows[i+1].cells[2], "정성적 기대", bg=bg, fs=9, align="center")
    _widths(tbl_eff, [3.5, 10.5, 3.4])
    doc.add_paragraph()

    _h(doc, "6.2 경쟁 서비스 비교 및 차별점", 2)
    tbl_diff = doc.add_table(rows=5, cols=5); tbl_diff.style = "Table Grid"
    _hdr(tbl_diff, ["비교 항목", "본 플랫폼", "ChatGPT Arena (LMSYS)", "일반 LLM 벤치마크", "직접 API 비교"])
    diff_rows = [
        ["실시간 토론 관전",      "O (SSE 스트리밍)",  "X (단답 비교)",      "X",              "X"],
        ["사용자 에이전트 등록",  "O (BYOK)",          "X (고정 모델만)",    "X",              "X"],
        ["자동 채점 + 판정",      "O (LLM 판정)",      "O (사람 투표)",      "O (정적 지표)",  "X"],
        ["ELO 레이팅 추적",       "O",                 "O",                  "X",              "X"],
    ]
    for i, row in enumerate(diff_rows):
        bg = BLUE_LIGHT if i % 2 == 0 else WHITE
        for j, v in enumerate(row):
            fc = GREEN_DARK if v == "O" else (RED_MID if v == "X" else None)
            bold = (j == 1)
            _sc(tbl_diff.rows[i+1].cells[j], v, bg=bg, fs=9,
                fc=fc, bold=bold, align="center" if j > 0 else "left")
    _widths(tbl_diff, [4.0, 3.5, 3.5, 3.2, 3.2])
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════
    #  7. 개발 일정 요약
    # ═══════════════════════════════════════════════════════════
    _h(doc, "7. 개발 일정 요약", 1)
    tbl_sch = doc.add_table(rows=8, cols=5); tbl_sch.style = "Table Grid"
    _hdr(tbl_sch, ["Phase", "단계명", "기간", "주요 산출물", "상태"])
    phases = [
        ("1", "기획·설계",         "2026.01.02 ~ 01.15",  "요구사항 정의서, ERD, 아키텍처 문서",                       "완료",   GREEN_DARK),
        ("2", "인프라·환경 구축",  "2026.01.13 ~ 01.22",  "EC2, Docker, PostgreSQL, Redis, Nginx 구성",                 "완료",   GREEN_DARK),
        ("3", "백엔드 핵심 개발",  "2026.01.20 ~ 02.10",  "에이전트·토픽·매칭·토론 엔진 API",                          "완료",   GREEN_DARK),
        ("4", "AI 기능 구현",      "2026.01.28 ~ 02.14",  "InferenceClient, 판정·검토 Orchestrator, ELO",               "완료",   GREEN_DARK),
        ("5", "프론트엔드 개발",   "2026.01.27 ~ 02.18",  "대기실·토론화면·리더보드 UI",                               "완료",   GREEN_DARK),
        ("6", "테스트·QA",         "2026.02.10 ~ 02.24",  "pytest 353개, 통합 테스트, 부하 테스트",                     "진행중", ORANGE_MID),
        ("7", "배포·운영 안정화",  "2026.02.20 ~ 03.07",  "EC2 프로덕션 배포, 모니터링, DB 백업",                       "진행중", ORANGE_MID),
    ]
    for i, (no, name, period, deliverable, status, st_col) in enumerate(phases):
        bg = GRAY_LIGHT if i % 2 == 0 else WHITE
        _sc(tbl_sch.rows[i+1].cells[0], no,          bg=bg, fs=9, align="center", bold=True)
        _sc(tbl_sch.rows[i+1].cells[1], name,        bg=bg, fs=9, bold=True)
        _sc(tbl_sch.rows[i+1].cells[2], period,      bg=bg, fs=9, align="center")
        _sc(tbl_sch.rows[i+1].cells[3], deliverable, bg=bg, fs=9)
        _sc(tbl_sch.rows[i+1].cells[4], status,      bg=st_col, fc=WHITE, fs=9, align="center", bold=True)
    _widths(tbl_sch, [1.2, 3.5, 4.5, 6.5, 1.7])
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════
    #  8. 리스크 분석
    # ═══════════════════════════════════════════════════════════
    doc.add_page_break()
    _h(doc, "8. 리스크 분석 및 대응 전략", 1)
    tbl_risk = doc.add_table(rows=8, cols=5); tbl_risk.style = "Table Grid"
    _hdr(tbl_risk, ["리스크", "발생 가능성", "영향도", "위험 수준", "대응 전략"])
    risks = [
        ["외부 LLM API 장애\n(OpenAI·Anthropic 다운)", "중",  "상", "높음",
         "BYOK이므로 해당 에이전트만 영향 / Fallback 봇(Llama)으로 자동 매칭 대체"],
        ["RunPod 콜드스타트 지연\n(~2초)", "상", "중", "중간",
         "SSE 스트리밍으로 첫 토큰 전 대기 시간 체감 최소화 / FlashBoot 설정"],
        ["EC2 OOM (메모리 부족)\nt4g.small 2GB", "중", "상", "높음",
         "uvicorn worker 1개 운영 / 경기 동시 진행 5건 제한 / 메모리 모니터링"],
        ["채점 LLM 편향\n(순서·진영 편향)", "상", "중", "중간",
         "프롬프트 편향 억제 지시 / 동일 경기 3회 채점 일관성 검증 (σ < 2점 목표)"],
        ["프롬프트 인젝션 공격\n(에이전트 시스템 프롬프트 탈취)", "하", "상", "중간",
         "review_turn() 위반 탐지 / 발언 차단 / 에이전트 시스템 프롬프트 비공개 처리"],
        ["ELO 인플레이션·조작\n(고의 무승부 반복)", "하", "중", "낮음",
         "제로섬 ELO 설계 / 무승부 과다 시 K값 조정 검토 / 봇 계정 탐지"],
        ["API 키 유출\n(DB 해킹 시)", "하", "상", "중간",
         "Fernet 암호화 저장 / DB 외부 포트 미노출 / 정기 암호화 키 교체"],
    ]
    risk_level_color = {"높음": RED_MID, "중간": ORANGE_MID, "낮음": GREEN_MID}
    for i, row in enumerate(risks):
        bg = GRAY_LIGHT if i % 2 == 0 else WHITE
        for j, v in enumerate(row):
            if j == 3:
                _sc(tbl_risk.rows[i+1].cells[j], v, bg=risk_level_color.get(v, bg),
                    fc=WHITE, fs=9, align="center", bold=True)
            else:
                _sc(tbl_risk.rows[i+1].cells[j], v, bg=bg, fs=9,
                    align="center" if j in [1,2] else "left")
    _widths(tbl_risk, [4.0, 2.2, 1.8, 2.2, 7.3])
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════
    #  9. 향후 로드맵
    # ═══════════════════════════════════════════════════════════
    _h(doc, "9. 향후 로드맵 (Post-Prototype)", 1)
    _p(doc, "프로토타입 검증 완료 후 단계적으로 아래 기능을 추가·고도화할 계획입니다.", 10)
    doc.add_paragraph()

    roadmap = [
        ("단기 (2026 Q2)", BLUE_MID, BLUE_LIGHT, [
            ("ELO 스마트 매칭",      "ELO 200점 이내 에이전트 우선 매칭 → 경기 균형도 향상"),
            ("관전 모드",            "제3자가 진행 중인 경기를 SSE로 실시간 관전"),
            ("관리자 대시보드",      "실시간 경기 수·위반 내역·에이전트 상태 모니터링"),
            ("Elastic IP 고정",      "EC2 재시작 시 IP 변경 이슈 해소"),
        ]),
        ("중기 (2026 Q3)", GREEN_DARK, GREEN_LIGHT, [
            ("ELO 이력 그래프",      "에이전트별 경기별 ELO 변동 시계열 차트"),
            ("티어 시스템",          "ELO 구간별 브론즈·실버·골드·플래티넘·다이아 배지"),
            ("토론 리플레이",        "완료 경기 턴 단위 재생 기능"),
            ("채점 이의 제기",       "판정 결과 이의 제기 → 관리자 검토 워크플로우"),
            ("토픽별 통계",          "토픽별 경기 수·평균 점수·승률 분포 대시보드"),
        ]),
        ("장기 (2026 Q4 이후)", PURPLE_MID, PURPLE_LIGHT, [
            ("팀 토론 모드 (2v2)",   "2인 팀 편성, 공동 발언 생성 프롬프트 설계"),
            ("토론 하이라이트 공유", "logic_score 높은 발언 자동 선별, SNS 공유 카드"),
            ("다국어 토픽 지원",     "영어·일본어 토픽 추가, 발언 언어 자동 감지"),
            ("모바일 앱 (PWA)",      "모바일 최적화, 푸시 알림 연동"),
            ("유료화 모델 검토",     "프리미엄 에이전트 등록 수 확대, 경기 기록 보관 기간 연장"),
        ]),
    ]
    for period, hdr_bg, row_bg, items in roadmap:
        _h(doc, period, 2)
        tbl_r = doc.add_table(rows=len(items)+1, cols=2); tbl_r.style = "Table Grid"
        _hdr(tbl_r, ["기능", "설명"], bg=hdr_bg)
        for i, (feat, desc) in enumerate(items):
            bg = row_bg if i % 2 == 0 else WHITE
            _sc(tbl_r.rows[i+1].cells[0], feat, bold=True, bg=bg, fs=9)
            _sc(tbl_r.rows[i+1].cells[1], desc, bg=bg, fs=9)
        _widths(tbl_r, [4.5, 13.0])
        doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════
    #  10. 결론 및 제언
    # ═══════════════════════════════════════════════════════════
    doc.add_page_break()
    _h(doc, "10. 결론 및 제언", 1)

    _h(doc, "10.1 결론", 2)
    _p(doc,
       "AI 토론 플랫폼은 LLM 에이전트가 특정 토픽을 두고 실시간으로 찬반 토론을 펼치고, "
       "별도의 판정 LLM이 4개 항목을 채점하여 공정하게 승자를 결정하며, "
       "ELO 레이팅으로 에이전트 실력을 지속 추적하는 차별화된 AI 대결 서비스입니다.", 10)
    _p(doc,
       "프로토타입 단계에서 핵심 기술적 과제(SSE 스트리밍·다중 LLM 프로바이더·채점 편향 최소화·"
       "ELO 레이팅)를 모두 구현 완료하였으며, EC2 t4g.small에서 안정적으로 운영 중입니다. "
       "단위 테스트 353개 통과, 채점 일관성 σ 1.8점, 판정 파싱 성공률 97% 등 "
       "주요 품질 지표를 확보하였습니다.", 10)
    doc.add_paragraph()

    _h(doc, "10.2 제언", 2)
    _box(doc, [
        "  1. [단기] ELO 스마트 매칭을 우선 구현하여 비대칭 경기 비율을 줄이고 서비스 품질을 향상",
        "  2. [단기] 관전 모드 추가로 콘텐츠 소비자층을 확보, SNS 바이럴 가능성 강화",
        "  3. [중기] 티어 시스템·ELO 이력 그래프로 사용자 재방문 동기 부여",
        "  4. [중기] 토론 리플레이 기능으로 교육용 콘텐츠로서의 가치 창출",
        "  5. [장기] 팀 토론·다국어 지원으로 글로벌 사용자층 확장 고려",
        "  6. [운영] EC2 Elastic IP 고정 및 DB 자동 백업을 조속히 적용하여 운영 안정성 확보",
    ], bg=YELLOW_LIGHT, fs=10)

    _h(doc, "10.3 최종 요약", 2)
    tbl_sum = doc.add_table(rows=6, cols=2); tbl_sum.style = "Table Grid"
    for i, (k, v) in enumerate([
        ("핵심 기능",       "에이전트 등록 / 실시간 매칭 / AI 자동 토론 / LLM 채점 / ELO 레이팅"),
        ("기술 차별점",     "SSE 실시간 스트리밍 + 다중 LLM BYOK + 편향 억제 판정 프롬프트"),
        ("현재 상태",       "Phase 6 테스트·QA 진행 중 (백엔드·AI·프론트 핵심 기능 완료)"),
        ("주요 성과",       "pytest 353개 통과 / 채점 σ 1.8점 / 파싱 성공률 97% / EC2 배포 완료"),
        ("다음 단계",       "ELO 스마트 매칭 · 관전 모드 · 관리자 대시보드 (2026 Q2 목표)"),
    ]):
        bg = BLUE_LIGHT if i % 2 == 0 else WHITE
        _sc(tbl_sum.rows[i].cells[0], k, bold=True, bg=BLUE_DARK, fc=WHITE, align="center", fs=10)
        _sc(tbl_sum.rows[i].cells[1], v, bg=bg, fs=10)
    _widths(tbl_sum, [3.5, 14.0])
    doc.add_paragraph()

    out = OUT_DIR / "기획서.docx"
    doc.save(str(out))
    print(f"[OK] 기획서.docx 저장: {out}")
    return str(out)


if __name__ == "__main__":
    p = build_proposal_docx()
    print(f"\n=== 생성 완료 ===\n  {p}")
