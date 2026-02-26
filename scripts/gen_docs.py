# -*- coding: utf-8 -*-
"""
평가 제출용 문서 생성 스크립트  —  AI 토론 플랫폼 전용
  - 수집_데이터_명세서.docx
  - 데이터베이스_설계문서.xlsx
"""

from pathlib import Path

OUT_DIR = Path(__file__).parent.parent / "docs" / "output"
OUT_DIR.mkdir(parents=True, exist_ok=True)

# ── 색상 팔레트 ───────────────────────────────────────────────────
BLUE_DARK    = "1F3864"
BLUE_MID     = "2E75B6"
BLUE_LIGHT   = "D6E4F0"
GREEN_DARK   = "1E5631"
GREEN_MID    = "27AE60"
GREEN_LIGHT  = "D5F5E3"
ORANGE_MID   = "E67E22"
ORANGE_LIGHT = "FDEBD0"
GRAY_LIGHT   = "F2F2F2"
WHITE        = "FFFFFF"
RED_MID      = "C0392B"
RED_LIGHT    = "FADBD8"
PURPLE_MID   = "8E44AD"
PURPLE_LIGHT = "E8DAEF"
YELLOW_LIGHT = "FEF9E7"


# ════════════════════════════════════════════════════════════════
#  공통 헬퍼
# ════════════════════════════════════════════════════════════════
def _rgb(hex_str):
    from docx.shared import RGBColor
    r, g, b = int(hex_str[0:2], 16), int(hex_str[2:4], 16), int(hex_str[4:6], 16)
    return RGBColor(r, g, b)


def _set_cell(cell, text, bold=False, font_size=9, bg=None,
              font_color=None, align="left", font_name="맑은 고딕"):
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    cell.text = ""
    para = cell.paragraphs[0]
    if align == "center":
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = para.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if font_color:
        run.font.color.rgb = _rgb(font_color)
    if bg:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), bg)
        tcPr.append(shd)


def _header_row(table, headers, bg=BLUE_DARK, font_color=WHITE, font_size=9):
    row = table.rows[0]
    for i, h in enumerate(headers):
        _set_cell(row.cells[i], h, bold=True, bg=bg,
                  font_color=font_color, font_size=font_size, align="center")


def _set_col_widths(table, widths_cm):
    from docx.shared import Cm
    for row in table.rows:
        for i, w in enumerate(widths_cm):
            if i < len(row.cells):
                row.cells[i].width = Cm(w)


def _add_heading(doc, text, level=1):
    from docx.shared import Pt
    from docx.oxml.ns import qn
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "맑은 고딕"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        if level == 1:
            run.font.size = Pt(14)
            run.font.color.rgb = _rgb(BLUE_DARK)
        elif level == 2:
            run.font.size = Pt(12)
            run.font.color.rgb = _rgb(BLUE_MID)
        else:
            run.font.size = Pt(11)
            run.font.color.rgb = _rgb("555555")
    return p


def _add_para(doc, text, font_size=10, bold=False, color=None):
    from docx.shared import Pt
    from docx.oxml.ns import qn
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    if color:
        run.font.color.rgb = _rgb(color)
    return p


def _add_box(doc, lines, bg=BLUE_LIGHT):
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    cell.text = ""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), bg)
    tcPr.append(shd)
    for i, line in enumerate(lines):
        para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        run = para.add_run(line)
        run.font.size = Pt(9)
        run.font.name = "맑은 고딕"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    doc.add_paragraph()


def _cover(doc, title, subtitle, date="2026년 2월"):
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(title)
    r.bold = True; r.font.size = Pt(24); r.font.name = "맑은 고딕"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    r.font.color.rgb = _rgb(BLUE_DARK)
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = s.add_run(subtitle)
    rs.font.size = Pt(13); rs.font.name = "맑은 고딕"
    rs._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    rs.font.color.rgb = _rgb("555555")
    d = doc.add_paragraph()
    d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rd = d.add_run(f"작성일: {date}   |   버전: 1.0   |   분류: 평가 제출용")
    rd.font.size = Pt(9); rd.font.name = "맑은 고딕"
    rd._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    rd.font.color.rgb = _rgb("888888")
    doc.add_page_break()


# ════════════════════════════════════════════════════════════════
#  1. 수집 데이터 명세서 (.docx)  —  AI 토론 플랫폼 전용
# ════════════════════════════════════════════════════════════════
def build_docx():
    from docx import Document
    from docx.shared import Cm
    doc = Document()
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Cm(2.0)
        section.left_margin = section.right_margin = Cm(2.5)

    _cover(doc, "수집 데이터 명세서", "AI 토론 플랫폼 — LLM 에이전트 자동 토론 시스템")

    # ── 1. 수집 데이터 개요 ──────────────────────────────────────
    _add_heading(doc, "1. 수집 데이터 개요", 1)
    _add_para(doc,
        "AI 토론 플랫폼은 사용자가 등록한 LLM 에이전트(찬성/반대 측)가 주어진 토픽에 대해 자동으로 "
        "토론을 진행하고, 별도의 판정 LLM이 채점·승자 판정·ELO 갱신을 수행하는 시스템이다. "
        "수집·생성되는 데이터는 3가지 유형으로 분류된다.", 10)
    doc.add_paragraph()

    _add_heading(doc, "1.1 데이터 수집 유형 분류", 2)
    tbl = doc.add_table(rows=4, cols=4)
    tbl.style = "Table Grid"
    _header_row(tbl, ["유형", "대상", "수집 방식", "저장 테이블"])
    type_rows = [
        ["유형 1\n외부·관리자 입력",
         "토론 토픽(주제·설명·규칙)",
         "관리자 또는 인가 사용자가 API를 통해 직접 입력",
         "debate_topics"],
        ["유형 2\n사용자 생성",
         "사용자 계정, 에이전트 등록\n(LLM 종류·API 키·시스템 프롬프트), 매칭 대기",
         "플랫폼 API 실시간 저장\n(에이전트 등록·큐 참여 즉시)",
         "users\ndebate_agents\ndebate_match_queue"],
        ["유형 3\n플랫폼 자동 생성",
         "토론 매치 기록, 턴별 발언 로그\n(LLM 검토 결과·벌점 포함), ELO 변동",
         "debate_engine / orchestrator가\n경기 진행 중 자동 기록",
         "debate_matches\ndebate_turn_logs\ndebate_agents (ELO)"],
    ]
    for i, row in enumerate(type_rows):
        bg = BLUE_LIGHT if i % 2 == 0 else WHITE
        for j, val in enumerate(row):
            _set_cell(tbl.rows[i+1].cells[j], val, bg=bg, font_size=8)
    _set_col_widths(tbl, [2.5, 5.0, 5.5, 4.5])
    doc.add_paragraph()

    _add_heading(doc, "1.2 시스템 인프라 개요", 2)
    tbl2 = doc.add_table(rows=7, cols=2)
    tbl2.style = "Table Grid"
    _header_row(tbl2, ["항목", "내용"])
    infra = [
        ["서비스 단계", "프로토타입 (동시 접속 10명 이하, 초기 사용자 ~100명)"],
        ["서버", "AWS EC2 t4g.small (서울 리전, ap-northeast-2)"],
        ["데이터베이스", "PostgreSQL 16 + pgvector 확장 (Docker, EC2 내부)"],
        ["캐시 / 이벤트", "Redis — 큐 이벤트 pub/sub (SSE 브로드캐스트용)"],
        ["발언 생성 LLM", "RunPod SGLang Serverless (Llama 3 70B, 기본) + 사용자 BYOK\n(OpenAI / Anthropic / Google API)"],
        ["판정·검토 LLM", "GPT-4.1 (플랫폼 API 키, OpenAI) — 채점·위반 탐지 전용"],
    ]
    for i, (k, v) in enumerate(infra):
        bg = GRAY_LIGHT if i % 2 == 0 else WHITE
        _set_cell(tbl2.rows[i+1].cells[0], k, bold=True, bg=bg, font_size=9)
        _set_cell(tbl2.rows[i+1].cells[1], v, bg=bg, font_size=9)
    _set_col_widths(tbl2, [4.0, 13.5])
    doc.add_paragraph()

    # ── 2. 수집 방법 및 자동화 절차 ─────────────────────────────
    doc.add_page_break()
    _add_heading(doc, "2. 수집 방법 및 자동화 절차", 1)

    _add_heading(doc, "2.1 실시간 파이프라인 — 에이전트 등록", 2)
    tbl3 = doc.add_table(rows=6, cols=3)
    tbl3.style = "Table Grid"
    _header_row(tbl3, ["단계", "처리 내용", "관련 컴포넌트"])
    reg_steps = [
        ["1", "사용자가 에이전트명·LLM 프로바이더·모델 ID·API 키·시스템 프롬프트 입력",
         "POST /api/agents"],
        ["2", "API 키 유효성 테스트 — 실제 LLM 호출(테스트 메시지) 후 응답 확인",
         "inference_client.py"],
        ["3", "API 키 Fernet 대칭 암호화 후 DB 저장",
         "debate_agents.encrypted_api_key"],
        ["4", "에이전트 버전 스냅샷 생성 (시스템 프롬프트·모델 변경 이력 보존)",
         "debate_agent_versions"],
        ["5", "초기 ELO 1,500점 설정",
         "debate_agents.elo_rating"],
    ]
    for i, row in enumerate(reg_steps):
        bg = BLUE_LIGHT if i % 2 == 0 else WHITE
        for j, val in enumerate(row):
            _set_cell(tbl3.rows[i+1].cells[j], val, bg=bg, font_size=8,
                      align="center" if j == 0 else "left")
    _set_col_widths(tbl3, [1.2, 10.5, 5.8])
    doc.add_paragraph()

    _add_heading(doc, "2.2 실시간 파이프라인 — 매칭 및 토론 진행", 2)
    tbl4 = doc.add_table(rows=9, cols=3)
    tbl4.style = "Table Grid"
    _header_row(tbl4, ["단계", "처리 내용", "관련 테이블/컴포넌트"])
    match_steps = [
        ["1", "에이전트가 특정 토픽 큐에 참여 (에이전트당 1 토픽 동시 대기 제한)",
         "debate_match_queue\nPOST /topics/{id}/queue/join"],
        ["2", "2명 대기 감지 → 준비 확인(ready_up) → 10초 카운트다운 후 자동 매치 생성",
         "debate_matching_service.py\n_auto_match_safe()"],
        ["3", "SSE 이벤트 브로드캐스트 (matched / countdown_started / opponent_joined)",
         "Redis pub/sub\n→ SSE /topics/{id}/queue/stream"],
        ["4", "debate_matches 레코드 생성 (status=pending → in_progress)",
         "debate_matches"],
        ["5", "debate_engine이 턴 루프 시작: 에이전트별 LLM 호출 → 발언 생성 (SSE 스트리밍)",
         "debate_engine.py\n_generate_turn()"],
        ["6", "각 턴 발언을 turn_review LLM이 실시간 검토\n(logic_score 1~10, 위반 유형·심각도, block 여부)",
         "debate_orchestrator.review_turn()\ndebate_turn_logs.review_result"],
        ["7", "모든 턴 완료 후 판정 LLM이 양측 발언 번들 채점\n(logic/evidence/rebuttal/relevance + reasoning)",
         "debate_orchestrator.judge()\ndebate_matches.scorecard"],
        ["8", "벌점 차감 → 승자 판정 → ELO 갱신 → DB 저장",
         "debate_agents.elo_rating\ndebate_matches.winner"],
    ]
    for i, row in enumerate(match_steps):
        bg = GREEN_LIGHT if i % 2 == 0 else WHITE
        for j, val in enumerate(row):
            _set_cell(tbl4.rows[i+1].cells[j], val, bg=bg, font_size=8,
                      align="center" if j == 0 else "left")
    _set_col_widths(tbl4, [1.2, 10.5, 5.8])
    doc.add_paragraph()

    _add_heading(doc, "2.3 배치 및 자동화 절차", 2)
    tbl5 = doc.add_table(rows=5, cols=4)
    tbl5.style = "Table Grid"
    _header_row(tbl5, ["자동화 작업", "실행 주기", "처리 내용", "출력"])
    batch = [
        ["타임아웃 에이전트 처리",
         "토론 진행 중\n(30s 타임아웃)",
         "LLM 응답 초과 시 [TIMEOUT] 텍스트 삽입,\n해당 턴 0~5점 자동 부여, 경기 지속",
         "debate_turn_logs\n(content='[TIMEOUT]')"],
        ["자동 매칭\n(봇 대기)",
         "큐 대기 120초 초과",
         "플랫폼 봇 에이전트가 자동 매칭 참여,\n사용자 대기 무한 방지",
         "debate_matches\n(auto_matched=true)"],
        ["큐 정리\n(abandoned)",
         "매칭 완료 또는\n취소 즉시",
         "debate_match_queue 레코드 삭제,\nSSE cancelled 이벤트 브로드캐스트",
         "debate_match_queue"],
        ["ELO 일관성 검증",
         "매칭 완료 직후",
         "제로섬 확인: delta_a + delta_b = 0\n위반 시 로그 기록 후 수동 검토",
         "debate_agents.elo_rating"],
    ]
    for i, row in enumerate(batch):
        bg = ORANGE_LIGHT if i % 2 == 0 else WHITE
        for j, val in enumerate(row):
            _set_cell(tbl5.rows[i+1].cells[j], val, bg=bg, font_size=8)
    _set_col_widths(tbl5, [3.5, 2.8, 7.0, 4.2])
    doc.add_paragraph()

    # ── 3. 데이터 설명 및 구성 ──────────────────────────────────
    doc.add_page_break()
    _add_heading(doc, "3. 데이터 설명 및 구성", 1)
    _add_para(doc, "AI 토론 플랫폼의 핵심 테이블 7종과 주요 필드를 정의한다.", 10)
    doc.add_paragraph()

    _add_heading(doc, "3.1 도메인별 테이블 구성", 2)
    tbl6 = doc.add_table(rows=4, cols=3)
    tbl6.style = "Table Grid"
    _header_row(tbl6, ["도메인", "테이블 목록", "설명"])
    domains = [
        ["사용자·에이전트",
         "users\ndebate_agents\ndebate_agent_versions",
         "계정 관리, 에이전트 등록 (LLM 설정·암호화 키), 버전 이력"],
        ["토픽·매칭",
         "debate_topics\ndebate_match_queue",
         "토론 주제 정의, 실시간 매칭 대기열"],
        ["경기·기록",
         "debate_matches\ndebate_turn_logs",
         "매치 결과·채점 JSONB, 턴별 발언·검토 결과·벌점"],
    ]
    for i, (k, v, d) in enumerate(domains):
        bg = BLUE_LIGHT if i % 2 == 0 else WHITE
        _set_cell(tbl6.rows[i+1].cells[0], k, bold=True, bg=bg, font_size=9)
        _set_cell(tbl6.rows[i+1].cells[1], v, bg=bg, font_size=9)
        _set_cell(tbl6.rows[i+1].cells[2], d, bg=bg, font_size=9)
    _set_col_widths(tbl6, [3.5, 5.5, 8.5])
    doc.add_paragraph()

    # 핵심 테이블 필드 정의
    def field_table(doc, title, headers, rows, widths, bg_hdr=BLUE_DARK):
        _add_heading(doc, title, 3)
        tbl = doc.add_table(rows=len(rows)+1, cols=len(headers))
        tbl.style = "Table Grid"
        _header_row(tbl, headers, bg=bg_hdr)
        for i, row in enumerate(rows):
            bg = GRAY_LIGHT if i % 2 == 0 else WHITE
            for j, val in enumerate(row):
                _set_cell(tbl.rows[i+1].cells[j], val, bg=bg, font_size=8)
        _set_col_widths(tbl, widths)
        doc.add_paragraph()

    field_table(doc, "① debate_agents — AI 에이전트 등록 (유형 2: 사용자 생성)",
        ["컬럼명", "타입", "설명", "비고"],
        [
            ["id",                  "UUID",         "에이전트 고유 ID",                         "PK"],
            ["owner_id",            "UUID",         "소유 사용자 참조",                         "FK → users"],
            ["name",                "VARCHAR(80)",  "에이전트 이름",                            "NOT NULL"],
            ["provider",            "VARCHAR(30)",  "LLM 프로바이더 (openai/anthropic/google/runpod/custom)", "NOT NULL"],
            ["model_id",            "VARCHAR(100)", "모델 식별자 (gpt-4o, claude-sonnet-4-6 등)", "NOT NULL"],
            ["encrypted_api_key",   "TEXT",         "Fernet 암호화된 사용자 API 키",            "NULL=플랫폼 기본 모델"],
            ["system_prompt",       "TEXT",         "에이전트 고유 토론 지침·역할 프롬프트",    ""],
            ["elo_rating",          "REAL",         "현재 ELO 점수 (초기 1,500)",               "DEFAULT 1500"],
            ["wins",                "INTEGER",      "승리 수",                                  "DEFAULT 0"],
            ["losses",              "INTEGER",      "패배 수",                                  "DEFAULT 0"],
            ["draws",               "INTEGER",      "무승부 수",                                "DEFAULT 0"],
            ["is_active",           "BOOLEAN",      "활성 여부 (비활성 시 매칭 제외)",           "DEFAULT true"],
            ["image_url",           "TEXT",         "에이전트 프로필 이미지 URL",               "NULLABLE"],
        ],
        [3.5, 3.5, 6.5, 4.0])

    field_table(doc, "② debate_topics — 토론 토픽 (유형 1: 관리자 입력)",
        ["컬럼명", "타입", "설명", "비고"],
        [
            ["id",          "UUID",         "토픽 고유 ID",                 "PK"],
            ["title",       "VARCHAR(200)", "토론 주제 제목",               "NOT NULL"],
            ["description", "TEXT",         "토픽 상세 설명",               ""],
            ["category",    "VARCHAR(50)",  "카테고리 (사회/경제/기술/환경 등)", ""],
            ["difficulty",  "VARCHAR(20)",  "난이도 (beginner/intermediate/advanced)", ""],
            ["max_turns",   "INTEGER",      "최대 턴 수 (기본 6턴)",        "DEFAULT 6"],
            ["time_limit",  "INTEGER",      "턴당 발언 생성 제한 시간(초)", "DEFAULT 30"],
            ["is_active",   "BOOLEAN",      "활성 상태 (매칭 허용 여부)",   "DEFAULT true"],
            ["created_at",  "TIMESTAMPTZ",  "생성 시각",                    "NOT NULL"],
        ],
        [3.5, 3.5, 6.5, 4.0])

    field_table(doc, "③ debate_matches — 경기 기록 (유형 3: 플랫폼 자동 생성)",
        ["컬럼명", "타입", "설명", "비고"],
        [
            ["id",              "UUID",         "경기 고유 ID",                         "PK"],
            ["topic_id",        "UUID",         "토픽 참조",                            "FK → debate_topics"],
            ["agent_a_id",      "UUID",         "찬성 측 에이전트 참조",                "FK → debate_agents"],
            ["agent_b_id",      "UUID",         "반대 측 에이전트 참조",                "FK → debate_agents"],
            ["status",          "VARCHAR(20)",  "경기 상태 (pending/in_progress/completed/cancelled)", ""],
            ["winner",          "VARCHAR(10)",  "승자 (agent_a/agent_b/draw)",         "NULLABLE"],
            ["scorecard",       "JSONB",        '{"agent_a":{logic,evidence,rebuttal,relevance},'
                                                '"agent_b":{...},"reasoning":"채점근거"}', "판정 LLM 출력"],
            ["elo_delta_a",     "REAL",         "에이전트A ELO 변동량 (+/-)",           ""],
            ["elo_delta_b",     "REAL",         "에이전트B ELO 변동량 (+/-)",           ""],
            ["auto_matched",    "BOOLEAN",      "봇 자동 매칭 여부",                    "DEFAULT false"],
            ["started_at",      "TIMESTAMPTZ",  "경기 시작 시각",                       ""],
            ["completed_at",    "TIMESTAMPTZ",  "경기 완료 시각",                       "NULLABLE"],
        ],
        [3.5, 3.5, 7.0, 3.5])

    field_table(doc, "④ debate_turn_logs — 턴별 발언 로그 (유형 3: 플랫폼 자동 생성)",
        ["컬럼명", "타입", "설명", "비고"],
        [
            ["id",              "BIGINT",       "턴 로그 ID",                           "PK IDENTITY"],
            ["match_id",        "UUID",         "경기 참조",                            "FK → debate_matches CASCADE"],
            ["agent_id",        "UUID",         "발언 에이전트 참조",                   "FK → debate_agents"],
            ["turn_number",     "INTEGER",      "턴 순번 (1부터)",                      "NOT NULL"],
            ["side",            "VARCHAR(10)",  "진영 (pro/con)",                       "NOT NULL"],
            ["action",          "VARCHAR(20)",  "행동 유형 (CLAIM/EVIDENCE/REBUTTAL/CLOSING)", ""],
            ["content",         "TEXT",         "발언 내용 (LLM 생성 원문)",            "NOT NULL"],
            ["is_blocked",      "BOOLEAN",      "LLM 검토 결과 차단 여부\n(severe 위반 시 true)", "DEFAULT false"],
            ["review_result",   "JSONB",        '{"logic_score":7,"violations":[],"severity":"none",'
                                                '"feedback":"한줄평","block":false}',   "턴 검토 LLM 출력"],
            ["penalty_points",  "INTEGER",      "위반 유형별 벌점 합산",                "DEFAULT 0"],
            ["input_tokens",    "INTEGER",      "발언 생성 입력 토큰 수",               ""],
            ["output_tokens",   "INTEGER",      "발언 생성 출력 토큰 수",               ""],
            ["created_at",      "TIMESTAMPTZ",  "발언 생성 시각",                       "NOT NULL"],
        ],
        [3.5, 3.2, 7.3, 3.5])

    field_table(doc, "⑤ debate_match_queue — 매칭 대기열 (유형 2: 사용자 생성)",
        ["컬럼명", "타입", "설명", "비고"],
        [
            ["id",              "BIGINT",       "대기열 항목 ID",                   "PK IDENTITY"],
            ["topic_id",        "UUID",         "대기 중인 토픽 참조",              "FK → debate_topics"],
            ["agent_id",        "UUID",         "대기 중인 에이전트 참조",          "FK → debate_agents"],
            ["user_id",         "UUID",         "소유 사용자 참조",                 "FK → users"],
            ["is_ready",        "BOOLEAN",      "준비 완료 여부 (ready_up 호출 후)", "DEFAULT false"],
            ["joined_at",       "TIMESTAMPTZ",  "대기 시작 시각",                   "NOT NULL"],
        ],
        [3.5, 3.5, 6.5, 4.0])

    # ── 4. 데이터 양 ──────────────────────────────────────────────
    doc.add_page_break()
    _add_heading(doc, "4. 데이터 양", 1)
    _add_para(doc, "프로토타입 단계(동시 접속 10명 이하) 기준 예상 데이터 규모와 6개월 후 예측값이다.", 10)
    doc.add_paragraph()

    _add_heading(doc, "4.1 테이블별 예상 행 수", 2)
    tbl7 = doc.add_table(rows=8, cols=5)
    tbl7.style = "Table Grid"
    _header_row(tbl7, ["테이블", "초기(1개월)", "3개월", "6개월", "증가 패턴"])
    vol_rows = [
        ["users",               "~50명",     "~150명",    "~300명",    "사용자 가입"],
        ["debate_agents",       "~80개",     "~250개",    "~500개",    "사용자당 평균 1~2개"],
        ["debate_topics",       "~20개",     "~50개",     "~100개",    "관리자 추가"],
        ["debate_matches",      "~500건",    "~2,000건",  "~5,000건",  "일 평균 5~15경기"],
        ["debate_turn_logs",    "~3,000건",  "~12,000건", "~30,000건", "경기당 평균 6~8턴"],
        ["debate_match_queue",  "일시적",    "일시적",    "일시적",    "매칭 완료 후 즉시 삭제"],
        ["debate_agent_versions","~80건",    "~350건",    "~700건",    "에이전트 수정 시"],
    ]
    for i, row in enumerate(vol_rows):
        bg = BLUE_LIGHT if i % 2 == 0 else WHITE
        for j, val in enumerate(row):
            _set_cell(tbl7.rows[i+1].cells[j], val, bg=bg, font_size=8,
                      align="center" if j > 0 else "left")
    _set_col_widths(tbl7, [4.5, 2.8, 2.8, 2.8, 4.6])
    doc.add_paragraph()

    _add_heading(doc, "4.2 데이터 저장 용량 추정", 2)
    tbl8 = doc.add_table(rows=5, cols=4)
    tbl8.style = "Table Grid"
    _header_row(tbl8, ["데이터 유형", "단위 크기(추정)", "6개월 예상 건수", "총 용량(추정)"])
    size_rows = [
        ["발언 로그 (debate_turn_logs)",    "~2KB/건\n(content + JSONB)",  "~30,000건", "~60MB"],
        ["채점 결과 (scorecard JSONB)",     "~1KB/건",                     "~5,000건",  "~5MB"],
        ["암호화 API 키",                   "~200B/건",                    "~500건",    "<1MB"],
        ["합계 (DB 전체 추정)",             "—",                           "—",         "~100MB\n(인덱스 포함)"],
    ]
    for i, row in enumerate(size_rows):
        bg = GRAY_LIGHT if i % 2 == 0 else WHITE
        is_total = i == 3
        for j, val in enumerate(row):
            _set_cell(tbl8.rows[i+1].cells[j], val, bg=BLUE_DARK if is_total else bg,
                      font_color=WHITE if is_total else None,
                      bold=is_total, font_size=8, align="center" if j > 1 else "left")
    _set_col_widths(tbl8, [5.0, 3.5, 3.5, 3.5])
    doc.add_paragraph()

    # ── 5. 품질 및 정합성 관리 ──────────────────────────────────
    doc.add_page_break()
    _add_heading(doc, "5. 품질 및 정합성 관리", 1)

    _add_heading(doc, "5.1 데이터 정합성 제약조건", 2)
    tbl9 = doc.add_table(rows=7, cols=3)
    tbl9.style = "Table Grid"
    _header_row(tbl9, ["제약 유형", "적용 대상", "내용"])
    constraints = [
        ["FK 참조 무결성",
         "debate_matches\ndebate_turn_logs",
         "agent_a_id·agent_b_id → debate_agents CASCADE DELETE\nmatch_id → debate_matches CASCADE DELETE"],
        ["에이전트당 1토픽 대기",
         "debate_match_queue",
         "agent_id에 UNIQUE 제약 — 동일 에이전트의 복수 토픽 동시 대기 불허"],
        ["점수 범위 CHECK",
         "채점 로직 (코드)",
         "logic 0~30, evidence 0~25, rebuttal 0~25, relevance 0~20\n코드 clamp로 강제"],
        ["ELO 제로섬 검증",
         "경기 완료 직후",
         "delta_a + delta_b = 0 검증, 불일치 시 경고 로그 기록"],
        ["API 키 암호화",
         "debate_agents.encrypted_api_key",
         "Fernet 대칭 암호화 필수 — 평문 저장 절대 금지"],
        ["winner 유효값",
         "debate_matches.winner",
         "CHECK (winner IN ('agent_a', 'agent_b', 'draw') OR winner IS NULL)"],
    ]
    for i, row in enumerate(constraints):
        bg = GREEN_LIGHT if i % 2 == 0 else WHITE
        for j, val in enumerate(row):
            _set_cell(tbl9.rows[i+1].cells[j], val, bg=bg, font_size=8)
    _set_col_widths(tbl9, [3.5, 4.5, 9.5])
    doc.add_paragraph()

    _add_heading(doc, "5.2 오류 대응 및 Fallback 정책", 2)
    tbl10 = doc.add_table(rows=5, cols=3)
    tbl10.style = "Table Grid"
    _header_row(tbl10, ["오류 유형", "탐지 방법", "Fallback 처리"])
    fallback = [
        ["LLM 타임아웃\n(발언 생성)",
         "asyncio.wait_for(timeout=30s)",
         "[TIMEOUT] 텍스트 삽입, 해당 턴 채점 시 0~5점 자동 부여, 경기 지속"],
        ["채점 JSON 파싱 실패",
         "json.loads() 예외 → 정규식 재추출 시도",
         "최종 실패 시 fallback 점수 사용\n{logic:15, evidence:12, rebuttal:12, relevance:10}"],
        ["턴 검토 타임아웃\n(review_turn)",
         "asyncio.wait_for(timeout=10s)",
         "fallback: {logic_score:5, violations:[], block:false}\n경기 중단 없이 계속"],
        ["SSE 연결 끊김",
         "es.onerror 이벤트",
         "최대 10회 재연결 시도 → 실패 시 폴링(4초 간격)으로 자동 전환"],
    ]
    for i, row in enumerate(fallback):
        bg = ORANGE_LIGHT if i % 2 == 0 else WHITE
        for j, val in enumerate(row):
            _set_cell(tbl10.rows[i+1].cells[j], val, bg=bg, font_size=8)
    _set_col_widths(tbl10, [3.5, 4.5, 9.5])
    doc.add_paragraph()

    # ── 6. 법적·윤리적 고려사항 ─────────────────────────────────
    doc.add_page_break()
    _add_heading(doc, "6. 법적·윤리적 고려사항", 1)
    tbl11 = doc.add_table(rows=5, cols=3)
    tbl11.style = "Table Grid"
    _header_row(tbl11, ["항목", "적용 규정 / 원칙", "플랫폼 대응"])
    legal = [
        ["개인정보 최소 수집",
         "개인정보보호법 제16조 (필요 최소한 수집)",
         "이메일·닉네임만 수집, 회원 탈퇴 시 즉시 파기 정책"],
        ["API 키 보호",
         "개인정보보호법 (인증정보 안전 조치 의무)",
         "Fernet 대칭 암호화 저장, 복호화는 LLM 호출 시점에만 수행"],
        ["AI 생성 콘텐츠 책임",
         "AI 기본법 (2026 시행 예정) / 정보통신망법",
         "LLM 위반 탐지(prompt_injection·ad_hominem 등) 및 자동 차단,\n이용약관에 AI 생성 발언 고지"],
        ["제3자 LLM API 이용",
         "OpenAI / Anthropic / Google 이용약관",
         "BYOK 방식으로 사용자가 직접 API 키 사용,\n플랫폼이 제3자 API 호출 책임 분리"],
    ]
    for i, row in enumerate(legal):
        bg = YELLOW_LIGHT if i % 2 == 0 else WHITE
        for j, val in enumerate(row):
            _set_cell(tbl11.rows[i+1].cells[j], val, bg=bg, font_size=8)
    _set_col_widths(tbl11, [3.5, 5.0, 9.0])
    doc.add_paragraph()

    out_path = OUT_DIR / "수집_데이터_명세서.docx"
    doc.save(str(out_path))
    print(f"[OK] 수집_데이터_명세서.docx 저장: {out_path}")
    return str(out_path)


# ════════════════════════════════════════════════════════════════
#  2. 데이터베이스 설계문서 (.xlsx)  —  AI 토론 플랫폼 전용
# ════════════════════════════════════════════════════════════════
def build_xlsx():
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    wb = Workbook()
    wb.remove(wb.active)

    def tfill(c): return PatternFill("solid", fgColor=c)
    def tb():
        s = Side(style="thin", color="AAAAAA")
        return Border(left=s, right=s, top=s, bottom=s)

    def sc(ws, row, col, val, bold=False, bg=None, fc="000000",
           sz=10, ah="left", av="center", wrap=False):
        c = ws.cell(row=row, column=col, value=val)
        c.font = Font(bold=bold, color=fc, size=sz, name="맑은 고딕")
        if bg:
            c.fill = tfill(bg)
        c.alignment = Alignment(horizontal=ah, vertical=av, wrap_text=wrap)
        c.border = tb()
        return c

    def mtitle(ws, row, text, cols=7):
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=cols)
        c = ws.cell(row=row, column=1, value=text)
        c.font = Font(bold=True, size=14, color=WHITE, name="맑은 고딕")
        c.fill = tfill(BLUE_DARK)
        c.alignment = Alignment(horizontal="center", vertical="center")
        ws.row_dimensions[row].height = 32

    def stitle(ws, row, text, bg=BLUE_MID, cols=7):
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=cols)
        c = ws.cell(row=row, column=1, value=text)
        c.font = Font(bold=True, size=12, color=WHITE, name="맑은 고딕")
        c.fill = tfill(bg)
        c.alignment = Alignment(horizontal="left", vertical="center")
        ws.row_dimensions[row].height = 22

    # ── 시트 1: ERD 개요 ─────────────────────────────────────────
    ws1 = wb.create_sheet("1. ERD 개요")
    ws1.sheet_view.showGridLines = False
    ws1.column_dimensions["A"].width = 22
    ws1.column_dimensions["B"].width = 20
    ws1.column_dimensions["C"].width = 30
    ws1.column_dimensions["D"].width = 25
    ws1.column_dimensions["E"].width = 20

    mtitle(ws1, 1, "AI 토론 플랫폼 — 데이터베이스 ERD 개요", cols=5)
    ws1.row_dimensions[2].height = 5

    # ERD 텍스트
    erd_lines = [
        "users ─────────────────────────────────────────────────────────",
        "  │ (owner_id FK)                                              ",
        "  ├── debate_agents ─────────────────────────────────────────  ",
        "  │      │ (agent_a_id / agent_b_id FK)                       ",
        "  │      ├── debate_matches ──────────────────────────────────",
        "  │      │       │ (match_id FK)                              ",
        "  │      │       └── debate_turn_logs                         ",
        "  │      │                                                    ",
        "  │      └── debate_match_queue ─ (topic_id FK) ─ debate_topics",
        "  │                                                           ",
        "  └── debate_agent_versions (agent_id FK)                    ",
        "                                                              ",
        "* debate_matches.topic_id ──── debate_topics                 ",
        "* debate_matches.scorecard ─── JSONB (판정 결과)             ",
        "* debate_turn_logs.review_result ─ JSONB (검토 결과)         ",
    ]
    ws1.merge_cells("A3:E17")
    cell = ws1["A3"]
    cell.value = "\n".join(erd_lines)
    cell.font = Font(name="Courier New", size=10, color="000000")
    cell.fill = tfill("F8F9FA")
    cell.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
    ws1.row_dimensions[3].height = 200

    ws1.row_dimensions[18].height = 8
    for j, h in enumerate(["테이블명", "행 수(6개월)", "주요 PK 타입", "주요 관계", "비고"], 1):
        sc(ws1, 19, j, h, bold=True, bg=BLUE_MID, fc=WHITE, ah="center")
    ws1.row_dimensions[19].height = 18

    erd_meta = [
        ["users",                   "~300",     "UUID",           "1:N → debate_agents",              "계정"],
        ["debate_agents",           "~500",     "UUID",           "N:M → debate_matches (양측)",       "Fernet 암호화 API 키"],
        ["debate_agent_versions",   "~700",     "BIGINT IDENTITY","N:1 → debate_agents",              "변경 이력"],
        ["debate_topics",           "~100",     "UUID",           "1:N → debate_matches",             "관리자 입력"],
        ["debate_match_queue",      "일시적",   "BIGINT IDENTITY","N:1 → debate_topics, debate_agents","매칭 완료 후 삭제"],
        ["debate_matches",          "~5,000",   "UUID",           "1:N → debate_turn_logs",           "scorecard JSONB"],
        ["debate_turn_logs",        "~30,000",  "BIGINT IDENTITY","N:1 → debate_matches",             "review_result JSONB"],
    ]
    for i, row in enumerate(erd_meta):
        r = i + 20
        bg = BLUE_LIGHT if i % 2 == 0 else WHITE
        ws1.row_dimensions[r].height = 20
        for j, val in enumerate(row, 1):
            sc(ws1, r, j, val, bg=bg, ah="center" if j == 2 else "left")

    # ── 시트 2: 테이블 정의서 ─────────────────────────────────────
    ws2 = wb.create_sheet("2. 테이블 정의서")
    ws2.sheet_view.showGridLines = False
    for col, w in zip("ABCDEFG", [22, 22, 14, 8, 8, 14, 22]):
        ws2.column_dimensions[col].width = w

    mtitle(ws2, 1, "2. 테이블 정의서 — AI 토론 플랫폼", cols=7)
    ws2.row_dimensions[2].height = 5

    tables_def = {
        "debate_agents (AI 에이전트)": [
            ("id",               "UUID",          "PK",           "Y", "N", "에이전트 고유 ID"),
            ("owner_id",         "UUID",          "FK→users",     "Y", "N", "소유 사용자"),
            ("name",             "VARCHAR(80)",   "",             "Y", "N", "에이전트 이름"),
            ("provider",         "VARCHAR(30)",   "CHECK",        "Y", "N", "openai/anthropic/google/runpod/custom"),
            ("model_id",         "VARCHAR(100)",  "",             "Y", "N", "모델 식별자"),
            ("encrypted_api_key","TEXT",          "",             "N", "Y", "Fernet 암호화 API 키"),
            ("system_prompt",    "TEXT",          "",             "N", "Y", "에이전트 토론 지침"),
            ("elo_rating",       "REAL",          "DEFAULT 1500", "Y", "N", "현재 ELO 점수"),
            ("wins",             "INTEGER",       "DEFAULT 0",    "Y", "N", "승리 수"),
            ("losses",           "INTEGER",       "DEFAULT 0",    "Y", "N", "패배 수"),
            ("draws",            "INTEGER",       "DEFAULT 0",    "Y", "N", "무승부 수"),
            ("is_active",        "BOOLEAN",       "DEFAULT true", "Y", "N", "활성 여부"),
            ("created_at",       "TIMESTAMPTZ",   "NOT NULL",     "Y", "N", "생성 시각"),
        ],
        "debate_topics (토론 토픽)": [
            ("id",          "UUID",         "PK",            "Y", "N", "토픽 고유 ID"),
            ("title",       "VARCHAR(200)", "NOT NULL",      "Y", "N", "토론 주제 제목"),
            ("description", "TEXT",         "",              "N", "Y", "토픽 상세 설명"),
            ("category",    "VARCHAR(50)",  "",              "N", "Y", "카테고리"),
            ("difficulty",  "VARCHAR(20)",  "CHECK",         "N", "Y", "beginner/intermediate/advanced"),
            ("max_turns",   "INTEGER",      "DEFAULT 6",     "Y", "N", "최대 턴 수"),
            ("time_limit",  "INTEGER",      "DEFAULT 30",    "Y", "N", "턴당 제한시간(초)"),
            ("is_active",   "BOOLEAN",      "DEFAULT true",  "Y", "N", "활성 여부"),
            ("created_at",  "TIMESTAMPTZ",  "NOT NULL",      "Y", "N", "생성 시각"),
        ],
        "debate_matches (경기 기록)": [
            ("id",            "UUID",        "PK",           "Y", "N", "경기 고유 ID"),
            ("topic_id",      "UUID",        "FK→topics",    "Y", "N", "토픽 참조"),
            ("agent_a_id",    "UUID",        "FK→agents",    "Y", "N", "찬성 측 에이전트"),
            ("agent_b_id",    "UUID",        "FK→agents",    "Y", "N", "반대 측 에이전트"),
            ("status",        "VARCHAR(20)", "CHECK",        "Y", "N", "pending/in_progress/completed/cancelled"),
            ("winner",        "VARCHAR(10)", "CHECK",        "N", "Y", "agent_a/agent_b/draw"),
            ("scorecard",     "JSONB",       "",             "N", "Y", "판정 LLM 채점 결과"),
            ("elo_delta_a",   "REAL",        "",             "N", "Y", "A 에이전트 ELO 변동"),
            ("elo_delta_b",   "REAL",        "",             "N", "Y", "B 에이전트 ELO 변동"),
            ("auto_matched",  "BOOLEAN",     "DEFAULT false","Y", "N", "봇 자동 매칭 여부"),
            ("started_at",    "TIMESTAMPTZ", "",             "N", "Y", "경기 시작 시각"),
            ("completed_at",  "TIMESTAMPTZ", "",             "N", "Y", "경기 완료 시각"),
        ],
        "debate_turn_logs (턴 발언 로그)": [
            ("id",             "BIGINT",      "PK IDENTITY",  "Y", "N", "턴 로그 ID"),
            ("match_id",       "UUID",        "FK→matches",   "Y", "N", "경기 참조 (CASCADE)"),
            ("agent_id",       "UUID",        "FK→agents",    "Y", "N", "발언 에이전트"),
            ("turn_number",    "INTEGER",     "NOT NULL",     "Y", "N", "턴 순번"),
            ("side",           "VARCHAR(10)", "CHECK",        "Y", "N", "pro/con"),
            ("action",         "VARCHAR(20)", "CHECK",        "Y", "N", "CLAIM/EVIDENCE/REBUTTAL/CLOSING"),
            ("content",        "TEXT",        "NOT NULL",     "Y", "N", "발언 내용"),
            ("is_blocked",     "BOOLEAN",     "DEFAULT false","Y", "N", "발언 차단 여부"),
            ("review_result",  "JSONB",       "",             "N", "Y", "턴 검토 결과 JSONB"),
            ("penalty_points", "INTEGER",     "DEFAULT 0",    "Y", "N", "위반 벌점 합산"),
            ("input_tokens",   "INTEGER",     "",             "N", "Y", "입력 토큰 수"),
            ("output_tokens",  "INTEGER",     "",             "N", "Y", "출력 토큰 수"),
            ("created_at",     "TIMESTAMPTZ", "NOT NULL",     "Y", "N", "발언 시각"),
        ],
        "debate_match_queue (매칭 대기열)": [
            ("id",        "BIGINT",      "PK IDENTITY",        "Y", "N", "대기 항목 ID"),
            ("topic_id",  "UUID",        "FK→topics",          "Y", "N", "대기 토픽"),
            ("agent_id",  "UUID",        "FK→agents, UNIQUE",  "Y", "N", "대기 에이전트 (1토픽 제한)"),
            ("user_id",   "UUID",        "FK→users",           "Y", "N", "소유 사용자"),
            ("is_ready",  "BOOLEAN",     "DEFAULT false",      "Y", "N", "준비 완료 여부"),
            ("joined_at", "TIMESTAMPTZ", "NOT NULL",           "Y", "N", "대기 시작 시각"),
        ],
    }

    cur_row = 3
    for tname, fields in tables_def.items():
        stitle(ws2, cur_row, f"[테이블] {tname}", bg=BLUE_MID, cols=7)
        cur_row += 1
        for j, h in enumerate(["컬럼명", "데이터 타입", "제약조건", "NOT NULL", "NULLABLE", "설명"], 1):
            sc(ws2, cur_row, j, h, bold=True, bg=BLUE_DARK, fc=WHITE, ah="center")
        ws2.row_dimensions[cur_row].height = 18
        cur_row += 1
        for i, field in enumerate(fields):
            bg = GRAY_LIGHT if i % 2 == 0 else WHITE
            ws2.row_dimensions[cur_row].height = 18
            for j, val in enumerate(field, 1):
                sc(ws2, cur_row, j, val, bg=bg,
                   ah="center" if j in [4, 5] else "left")
            cur_row += 1
        ws2.row_dimensions[cur_row].height = 6
        cur_row += 1

    # ── 시트 3: 제약조건 명세 ────────────────────────────────────
    ws3 = wb.create_sheet("3. 제약조건 명세")
    ws3.sheet_view.showGridLines = False
    for col, w in zip("ABCDE", [18, 22, 16, 28, 22]):
        ws3.column_dimensions[col].width = w

    mtitle(ws3, 1, "3. 제약조건 명세", cols=5)
    ws3.row_dimensions[2].height = 5
    for j, h in enumerate(["제약 유형", "테이블", "컬럼/대상", "내용", "목적"], 1):
        sc(ws3, 3, j, h, bold=True, bg=BLUE_MID, fc=WHITE, ah="center")
    ws3.row_dimensions[3].height = 18

    constraints = [
        ("PK",              "debate_agents",       "id",                    "UUID PRIMARY KEY",                        "에이전트 고유 식별"),
        ("PK",              "debate_matches",       "id",                    "UUID PRIMARY KEY",                        "경기 고유 식별"),
        ("PK",              "debate_turn_logs",     "id",                    "BIGINT GENERATED ALWAYS AS IDENTITY",     "순번 자동 생성"),
        ("FK CASCADE",      "debate_turn_logs",     "match_id",              "REFERENCES debate_matches(id) ON DELETE CASCADE", "경기 삭제 시 턴 자동 삭제"),
        ("FK",              "debate_agents",        "owner_id",              "REFERENCES users(id)",                    "에이전트 소유자"),
        ("FK",              "debate_matches",       "agent_a_id, agent_b_id","REFERENCES debate_agents(id)",            "양측 에이전트 참조"),
        ("UNIQUE",          "debate_match_queue",   "agent_id",              "UNIQUE — 에이전트당 1토픽 동시 대기만 허용","복수 대기 방지"),
        ("CHECK",           "debate_matches",       "winner",                "IN ('agent_a','agent_b','draw') OR NULL", "유효 승자값만 허용"),
        ("CHECK",           "debate_agents",        "provider",              "IN ('openai','anthropic','google','runpod','custom')", "지원 프로바이더만"),
        ("CHECK",           "debate_turn_logs",     "side",                  "IN ('pro','con')",                        "진영 구분"),
        ("CHECK",           "debate_turn_logs",     "action",                "IN ('CLAIM','EVIDENCE','REBUTTAL','CLOSING')", "행동 유형"),
        ("DEFAULT",         "debate_agents",        "elo_rating",            "DEFAULT 1500",                            "신규 에이전트 초기 ELO"),
        ("DEFAULT",         "debate_turn_logs",     "is_blocked",            "DEFAULT false",                           "기본 비차단"),
        ("INDEX",           "debate_turn_logs",     "match_id",              "idx_turn_logs_match_id",                  "경기별 조회 성능"),
        ("INDEX",           "debate_matches",       "agent_a_id, agent_b_id","idx_matches_agents",                      "에이전트 전적 조회"),
        ("INDEX",           "debate_agents",        "elo_rating DESC",       "idx_agents_elo",                          "ELO 리더보드 정렬"),
    ]
    for i, row in enumerate(constraints):
        r = i + 4
        bg = BLUE_LIGHT if row[0] == "PK" else \
             GREEN_LIGHT if row[0] == "FK CASCADE" else \
             ORANGE_LIGHT if row[0] == "FK" else \
             PURPLE_LIGHT if row[0] == "UNIQUE" else \
             RED_LIGHT if row[0] == "CHECK" else \
             YELLOW_LIGHT if row[0] == "INDEX" else \
             (GRAY_LIGHT if i % 2 == 0 else WHITE)
        ws3.row_dimensions[r].height = 22
        for j, val in enumerate(row, 1):
            sc(ws3, r, j, val, bg=bg, wrap=True)

    # ── 시트 4: 샘플 데이터 ──────────────────────────────────────
    ws4 = wb.create_sheet("4. 샘플 데이터")
    ws4.sheet_view.showGridLines = False
    for col, w in zip("ABCDEFGH", [20, 14, 12, 12, 10, 10, 10, 30]):
        ws4.column_dimensions[col].width = w

    mtitle(ws4, 1, "4. 샘플 데이터 (INSERT 예시)", cols=8)
    ws4.row_dimensions[2].height = 5

    # debate_agents 샘플
    stitle(ws4, 3, "[debate_agents] 에이전트 샘플 5건", GREEN_DARK, cols=8)
    for j, h in enumerate(["id(일부)", "owner_id(일부)", "name", "provider", "model_id", "elo_rating", "wins/losses/draws", "비고"], 1):
        sc(ws4, 4, j, h, bold=True, bg=GREEN_MID, fc=WHITE, ah="center")
    ws4.row_dimensions[4].height = 18

    agent_samples = [
        ["a1b2-...", "u001-...", "논리왕-GPT4",      "openai",    "gpt-4o",             "1623", "8/3/1", "BYOK"],
        ["a2c3-...", "u001-...", "반박의달인",         "anthropic", "claude-sonnet-4-6",  "1554", "5/4/2", "BYOK"],
        ["a3d4-...", "u002-...", "증거수집기",         "google",    "gemini-1.5-pro",     "1487", "4/6/1", "BYOK"],
        ["a4e5-...", "u003-...", "Llama-토론봇",       "runpod",    "meta-llama/Llama-3-70b", "1501", "2/2/0", "기본 내장"],
        ["a5f6-...", "u004-...", "논증왕2025",         "openai",    "gpt-4.1",            "1612", "7/2/1", "BYOK"],
    ]
    for i, row in enumerate(agent_samples):
        r = i + 5
        bg = GREEN_LIGHT if i % 2 == 0 else WHITE
        ws4.row_dimensions[r].height = 20
        for j, val in enumerate(row, 1):
            sc(ws4, r, j, val, bg=bg, ah="center" if j in [4, 5, 6, 7] else "left")

    ws4.row_dimensions[10].height = 8

    # debate_matches 샘플
    stitle(ws4, 11, "[debate_matches] 경기 기록 샘플 3건", BLUE_MID, cols=8)
    for j, h in enumerate(["match_id(일부)", "topic", "agent_a", "agent_b", "winner", "score_a", "score_b", "completed_at"], 1):
        sc(ws4, 12, j, h, bold=True, bg=BLUE_MID, fc=WHITE, ah="center")
    ws4.row_dimensions[12].height = 18

    match_samples = [
        ["m001-...", "환경세 도입 찬반", "논리왕-GPT4(찬)", "반박의달인(반)", "agent_a", "82점", "74점", "2026-02-20 14:32"],
        ["m002-...", "주 4일제 근무",    "논증왕2025(찬)",  "Llama-토론봇(반)","draw",   "65점", "63점", "2026-02-21 10:15"],
        ["m003-...", "AI 규제 찬반",     "증거수집기(찬)",  "논리왕-GPT4(반)", "agent_b","70점", "79점", "2026-02-22 16:44"],
    ]
    for i, row in enumerate(match_samples):
        r = i + 13
        bg = BLUE_LIGHT if i % 2 == 0 else WHITE
        ws4.row_dimensions[r].height = 20
        for j, val in enumerate(row, 1):
            sc(ws4, r, j, val, bg=bg)

    ws4.row_dimensions[16].height = 8

    # debate_turn_logs 샘플
    stitle(ws4, 17, "[debate_turn_logs] 발언 로그 샘플 5건", PURPLE_MID, cols=8)
    for j, h in enumerate(["match_id(일부)", "agent", "turn", "side", "action", "is_blocked", "logic_score", "발언 내용(일부)"], 1):
        sc(ws4, 18, j, h, bold=True, bg=PURPLE_MID, fc=WHITE, ah="center")
    ws4.row_dimensions[18].height = 18

    turn_samples = [
        ["m001-...", "논리왕-GPT4",   "1", "pro",  "CLAIM",    "false", "8", "환경세 도입은 탄소 배출 감소에 효과적입니다..."],
        ["m001-...", "반박의달인",     "1", "con",  "CLAIM",    "false", "7", "환경세는 중소기업 원가 부담을 가중시킵니다..."],
        ["m001-...", "논리왕-GPT4",   "2", "pro",  "EVIDENCE", "false", "9", "OECD 12개국 탄소세 도입 후 배출량 18% 감소..."],
        ["m001-...", "반박의달인",     "2", "con",  "REBUTTAL", "false", "8", "해당 통계는 고소득 국가 평균으로..."],
        ["m002-...", "논증왕2025",    "3", "pro",  "CLOSING",  "false", "7", "따라서 주 4일제는 생산성과 복지를 동시에..."],
    ]
    for i, row in enumerate(turn_samples):
        r = i + 19
        bg = PURPLE_LIGHT if i % 2 == 0 else WHITE
        ws4.row_dimensions[r].height = 20
        for j, val in enumerate(row, 1):
            sc(ws4, r, j, val, bg=bg, ah="center" if j in [3, 4, 5, 6, 7] else "left",
               wrap=(j == 8))

    # ── 시트 5: 쿼리 성능 ────────────────────────────────────────
    ws5 = wb.create_sheet("5. 쿼리 성능 고려사항")
    ws5.sheet_view.showGridLines = False
    for col, w in zip("ABCDE", [26, 34, 18, 16, 14]):
        ws5.column_dimensions[col].width = w

    mtitle(ws5, 1, "5. 쿼리 성능 고려사항", cols=5)
    ws5.row_dimensions[2].height = 5
    for j, h in enumerate(["쿼리 유형", "SQL 핵심 패턴", "인덱스 전략", "예상 응답시간", "비고"], 1):
        sc(ws5, 3, j, h, bold=True, bg=BLUE_MID, fc=WHITE, ah="center")
    ws5.row_dimensions[3].height = 18

    queries = [
        ["ELO 리더보드 조회\n(상위 50명)",
         "SELECT * FROM debate_agents\nORDER BY elo_rating DESC LIMIT 50",
         "idx_agents_elo\n(elo_rating DESC)",
         "< 10ms", "인덱스 스캔"],
        ["경기별 턴 로그 조회",
         "SELECT * FROM debate_turn_logs\nWHERE match_id = $1\nORDER BY turn_number",
         "idx_turn_logs_match_id",
         "< 5ms", "에이전트당 최대 수십 턴"],
        ["에이전트 전적 집계",
         "SELECT wins, losses, draws, elo_rating\nFROM debate_agents WHERE id = $1",
         "PK 조회",
         "< 2ms", "단순 PK 조회"],
        ["진행 중 경기 조회",
         "SELECT * FROM debate_matches\nWHERE status = 'in_progress'",
         "idx_matches_status\n(status)",
         "< 10ms", "소수 행 예상"],
        ["에이전트 히스토리\n(최근 10경기)",
         "SELECT * FROM debate_matches\nWHERE agent_a_id=$1 OR agent_b_id=$1\nORDER BY started_at DESC LIMIT 10",
         "idx_matches_agents",
         "< 20ms", "복합 OR 조건"],
        ["매칭 대기열 상태 확인",
         "SELECT * FROM debate_match_queue\nWHERE topic_id=$1 ORDER BY joined_at",
         "idx_queue_topic\n(topic_id)",
         "< 5ms", "항상 소수 행"],
        ["ELO 제로섬 검증\n(완료 경기)",
         "SELECT SUM(elo_delta_a + elo_delta_b)\nFROM debate_matches\nWHERE status='completed'",
         "idx_matches_status",
         "< 50ms", "집계 함수"],
    ]
    for i, row in enumerate(queries):
        r = i + 4
        bg = BLUE_LIGHT if i % 2 == 0 else WHITE
        ws5.row_dimensions[r].height = 40
        for j, val in enumerate(row, 1):
            sc(ws5, r, j, val, bg=bg, wrap=True, av="top",
               ah="center" if j in [3, 4, 5] else "left")

    # ── 시트 6: 데이터 보안 ──────────────────────────────────────
    ws6 = wb.create_sheet("6. 데이터 보안")
    ws6.sheet_view.showGridLines = False
    for col, w in zip("ABCDE", [22, 22, 30, 26, 10]):
        ws6.column_dimensions[col].width = w

    mtitle(ws6, 1, "6. 데이터 보안 및 접근 제어", cols=5)
    ws6.row_dimensions[2].height = 5
    for j, h in enumerate(["보안 항목", "적용 대상", "구현 방법", "목적·효과", "준거"], 1):
        sc(ws6, 3, j, h, bold=True, bg=RED_MID, fc=WHITE, ah="center")
    ws6.row_dimensions[3].height = 18

    security = [
        ["API 키 암호화",
         "debate_agents.encrypted_api_key",
         "Fernet 대칭 암호화\n(settings.fernet_key)",
         "사용자 LLM API 키 평문 저장 금지\n복호화는 LLM 호출 시점에만",
         "개인정보보호법"],
        ["RBAC 접근 제어",
         "관리자 토픽 생성/수정\n에이전트 소유자만 수정 가능",
         "FastAPI Depends(require_user)\n소유자 검증 로직",
         "타인의 에이전트 무단 수정 방지",
         "서비스 정책"],
        ["발언 위반 탐지",
         "debate_turn_logs",
         "review_turn() LLM 검토\nprompt_injection·ad_hominem·\noff_topic·false_claim",
         "악의적 프롬프트 인젝션 차단\n심각 위반 시 발언 자동 차단",
         "AI 안전"],
        ["이벤트 스트림 인증",
         "SSE /topics/{id}/queue/stream\nSSE /matches/{id}/stream",
         "withCredentials=true\nCookie 기반 세션 인증",
         "인가된 사용자만 실시간 이벤트 수신",
         "인증 정책"],
        ["DB 접근 격리",
         "PostgreSQL",
         "Docker 내부 네트워크 전용\n외부 포트 미노출",
         "외부 직접 DB 접근 차단",
         "인프라 보안"],
        ["개인정보 파기",
         "users (회원탈퇴)",
         "CASCADE DELETE 연쇄 삭제\ndebate_agents → 연관 matches 익명화",
         "잔여 개인정보 제거",
         "개인정보보호법"],
    ]
    for i, row in enumerate(security):
        r = i + 4
        bg = RED_LIGHT if i % 2 == 0 else WHITE
        ws6.row_dimensions[r].height = 45
        for j, val in enumerate(row, 1):
            sc(ws6, r, j, val, bg=bg, wrap=True, av="top")

    out_path = OUT_DIR / "데이터베이스_설계문서.xlsx"
    wb.save(str(out_path))
    print(f"[OK] 데이터베이스_설계문서.xlsx 저장: {out_path}")
    return str(out_path)


# ════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    p1 = build_docx()
    p2 = build_xlsx()
    print("\n=== 생성 완료 ===")
    print(f"  {p1}")
    print(f"  {p2}")
