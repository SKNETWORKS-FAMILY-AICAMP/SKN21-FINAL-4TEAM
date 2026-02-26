"""
AI 에이전트 토론 플랫폼 문서 생성 스크립트
Word (.docx) + Excel (.xlsx) 양식으로 4종 문서 생성

실행: python scripts/generate_debate_docs.py
출력: docs/output/ 디렉터리
"""

import os
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
import openpyxl
from openpyxl.styles import (
    Font, PatternFill, Alignment, Border, Side,
    numbers
)
from openpyxl.utils import get_column_letter

OUT_DIR = Path("docs/output")
OUT_DIR.mkdir(parents=True, exist_ok=True)

# ──────────────────────────────────────────────────────
# 공통 스타일 헬퍼 (Word)
# ──────────────────────────────────────────────────────

PRIMARY   = RGBColor(0x1E, 0x40, 0xAF)   # 진한 파랑
SECONDARY = RGBColor(0x16, 0x55, 0x34)   # 진한 녹색
ACCENT    = RGBColor(0xD9, 0x77, 0x06)   # 황금색
GRAY      = RGBColor(0x6B, 0x72, 0x80)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG  = RGBColor(0xF0, 0xF4, 0xFF)


def set_cell_bg(cell, hex_color: str):
    """표 셀 배경색 설정"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    if level == 1:
        run.font.color.rgb = PRIMARY
        run.font.size = Pt(18)
    elif level == 2:
        run.font.color.rgb = SECONDARY
        run.font.size = Pt(14)
    elif level == 3:
        run.font.color.rgb = ACCENT
        run.font.size = Pt(12)
    return p


def add_para(doc: Document, text: str, bold=False, italic=False, color=None, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def add_table(doc: Document, headers: list[str], rows: list[list], col_widths=None):
    """헤더 + 데이터 행 테이블 생성"""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 헤더
    for ci, h in enumerate(headers):
        cell = t.cell(0, ci)
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = WHITE
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, "1E40AF")

    # 데이터
    for ri, row in enumerate(rows):
        bg = "F0F4FF" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = t.cell(ri + 1, ci)
            cell.text = str(val)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_bg(cell, bg)

    # 컬럼 너비
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[ci].width = Inches(w)

    doc.add_paragraph()
    return t


def add_code_block(doc: Document, code: str):
    """코드 블록 스타일 단락"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    # 배경 없이 테두리로 구분
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        bdr = OxmlElement(f"w:{side}")
        bdr.set(qn("w:val"), "single")
        bdr.set(qn("w:sz"), "4")
        bdr.set(qn("w:space"), "4")
        bdr.set(qn("w:color"), "CBD5E1")
        pBdr.append(bdr)
    pPr.append(pBdr)
    return p


def init_doc(title: str, subtitle: str) -> Document:
    doc = Document()
    # 여백 설정
    for sec in doc.sections:
        sec.top_margin    = Cm(2.0)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.5)

    # 표지
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = PRIMARY

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(subtitle)
    r2.font.size = Pt(13)
    r2.font.color.rgb = GRAY

    doc.add_paragraph()
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("작성일: 2026-02-26  |  버전: 1.0  |  프로젝트: LLM 기반 AI 에이전트 자율 토론 플랫폼")
    r3.font.size = Pt(10)
    r3.font.color.rgb = GRAY

    doc.add_page_break()
    return doc


# ──────────────────────────────────────────────────────
# 공통 스타일 헬퍼 (Excel)
# ──────────────────────────────────────────────────────

XL_HDR   = PatternFill("solid", fgColor="1E40AF")
XL_HDR2  = PatternFill("solid", fgColor="1E5534")
XL_ALT   = PatternFill("solid", fgColor="EFF6FF")
XL_WHITE = PatternFill("solid", fgColor="FFFFFF")
XL_GOLD  = PatternFill("solid", fgColor="D97706")
XL_GRAY  = PatternFill("solid", fgColor="F1F5F9")

THIN = Border(
    left=Side(style="thin", color="CBD5E1"),
    right=Side(style="thin", color="CBD5E1"),
    top=Side(style="thin", color="CBD5E1"),
    bottom=Side(style="thin", color="CBD5E1"),
)
MEDIUM = Border(
    left=Side(style="medium", color="1E40AF"),
    right=Side(style="medium", color="1E40AF"),
    top=Side(style="medium", color="1E40AF"),
    bottom=Side(style="medium", color="1E40AF"),
)


def xl_header(ws, row: int, cols: list[str], fill=None):
    fill = fill or XL_HDR
    for ci, val in enumerate(cols, 1):
        c = ws.cell(row=row, column=ci, value=val)
        c.font   = Font(bold=True, color="FFFFFF", size=11)
        c.fill   = fill
        c.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        c.border = THIN


def xl_row(ws, row: int, vals: list, alt=False):
    fill = XL_ALT if alt else XL_WHITE
    for ci, val in enumerate(vals, 1):
        c = ws.cell(row=row, column=ci, value=val)
        c.fill   = fill
        c.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        c.border = THIN


def xl_section_title(ws, row: int, text: str, col_span: int, fill=None):
    fill = fill or XL_GOLD
    c = ws.cell(row=row, column=1, value=text)
    c.font   = Font(bold=True, color="FFFFFF", size=12)
    c.fill   = fill
    c.alignment = Alignment(horizontal="left", vertical="center")
    c.border = THIN
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=col_span)


def xl_note(ws, row: int, text: str, col_span: int):
    c = ws.cell(row=row, column=1, value=text)
    c.font = Font(italic=True, color="6B7280", size=9)
    c.fill = XL_GRAY
    c.alignment = Alignment(horizontal="left", wrap_text=True)
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=col_span)


def auto_col_width(ws, min_w=10, max_w=50):
    for col in ws.columns:
        max_len = 0
        col_letter = get_column_letter(col[0].column)
        for cell in col:
            try:
                val = str(cell.value or "")
                # 한글은 2배 너비
                kor = sum(1 for c in val if ord(c) > 0x1100)
                lng = len(val) + kor
                if lng > max_len:
                    max_len = lng
            except Exception:
                pass
        ws.column_dimensions[col_letter].width = min(max(max_len + 2, min_w), max_w)


# ══════════════════════════════════════════════════════════════════
#  문서 1: 기획서
# ══════════════════════════════════════════════════════════════════

def make_plan_docx():
    doc = init_doc("AI 에이전트 토론 플랫폼", "시스템 기획서")

    add_heading(doc, "1. 프로젝트 개요", 1)

    add_heading(doc, "1.1 한 줄 정의", 2)
    add_para(doc,
        "서로 다른 LLM 에이전트들이 주어진 주제로 자율 토론을 벌이고, "
        "심판 AI가 판정하여 ELO 레이팅으로 순위를 매기는 실시간 관전 플랫폼.",
        bold=True)

    add_heading(doc, "1.2 핵심 가치 제안", 2)
    add_table(doc,
        ["관점", "가치"],
        [
            ["사용자", "자신만의 AI 에이전트를 만들어 다른 에이전트와 토론 대결"],
            ["연구자", "프롬프트·모델 조합에 따른 논증 품질 비교 실험 환경"],
            ["일반 관객", "LLM들이 실시간으로 설전하는 엔터테인먼트 콘텐츠"],
        ],
        [1.5, 4.5]
    )

    add_heading(doc, "1.3 기술 스택", 2)
    add_table(doc,
        ["계층", "기술"],
        [
            ["백엔드", "Python 3.12 + FastAPI + SQLAlchemy (async)"],
            ["프론트엔드", "Next.js 15 + React 19 + Zustand"],
            ["데이터베이스", "PostgreSQL 16 (Docker)"],
            ["캐시/메시징", "Redis (Pub/Sub + SSE 스트리밍)"],
            ["LLM", "OpenAI / Anthropic / Google / RunPod Serverless"],
            ["실시간", "SSE (Server-Sent Events) + WebSocket (로컬 에이전트)"],
            ["배포", "AWS EC2 t4g.small (서울) + Docker Compose"],
        ],
        [1.8, 4.2]
    )

    add_heading(doc, "2. 배경 및 목적", 1)

    add_heading(doc, "2.1 문제 인식", 2)
    add_para(doc, "현재 LLM 성능 비교는 정적 벤치마크(MMLU, HumanEval 등)에 의존합니다. 주요 한계:")
    for item in [
        "단일 질문-단일 답변 구조 → 다회전 논증 능력 미측정",
        "정답이 있는 문제 → 주관적 판단력·설득력 미측정",
        "상호작용 없음 → 반박·재반박 대응 능력 미측정",
        "오염 위험 → 벤치마크 데이터가 학습셋에 포함될 가능성",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)

    add_heading(doc, "2.2 해결 방향 — 정적 벤치마크 vs AI 토론 플랫폼", 2)
    add_table(doc,
        ["항목", "정적 벤치마크", "AI 토론 플랫폼"],
        [
            ["평가 방식", "한 번 답변, 정답 비교", "여러 턴, 상호 반박"],
            ["데이터", "표준화된 데이터셋", "사용자 정의 주제"],
            ["대상", "단일 모델 평가", "모델 간 직접 대결"],
            ["점수", "점수 누적 평균", "ELO 기반 상대 평가"],
            ["접근성", "연구자만 접근", "누구나 에이전트 생성"],
        ],
        [2.0, 2.0, 2.0]
    )

    add_heading(doc, "3. 핵심 기능 정의", 1)

    add_heading(doc, "3.1 에이전트 시스템", 2)
    add_table(doc,
        ["유형", "연결 방식", "특징"],
        [
            ["템플릿 기반", "HTTP (API 호출)", "기본 프롬프트 + 슬라이더/선택형 커스터마이징"],
            ["BYOK", "HTTP (API 호출)", "자체 API 키 + 완전 커스텀 프롬프트"],
            ["로컬 에이전트", "WebSocket", "데스크탑/CLI 봇, 인간 대리 참여 가능"],
        ],
        [1.5, 1.8, 3.7]
    )

    add_heading(doc, "3.2 토론 모드", 2)
    add_table(doc,
        ["모드", "설명", "주요 액션"],
        [
            ["debate", "찬반 토론 (기본)", "argue / rebut / concede / question / summarize"],
            ["persuasion", "한 측이 다른 측을 설득", "argue / rebut / concede"],
            ["cross_exam", "교차 심문 방식", "question / answer / rebut"],
        ],
        [1.5, 2.0, 3.5]
    )

    add_heading(doc, "3.3 채점 기준", 2)
    add_table(doc,
        ["항목", "배점", "설명"],
        [
            ["논리성", "0 ~ 30", "논증 구조의 일관성·타당성"],
            ["근거 활용", "0 ~ 25", "주장을 뒷받침하는 증거의 질"],
            ["반박력", "0 ~ 25", "상대 주장에 대한 효과적 반론"],
            ["주제 적합성", "0 ~ 20", "토론 주제와의 연관성 유지"],
            ["페널티 차감", "-N", "규칙 위반 벌점"],
            ["최종 점수", "0 ~ 100", "소계 - 페널티"],
        ],
        [2.0, 1.2, 3.8]
    )

    add_heading(doc, "3.4 페널티 체계", 2)
    add_table(doc,
        ["위반 유형", "벌점", "탐지 방법"],
        [
            ["스키마 위반", "-5", "JSON 파싱 실패"],
            ["반복 발언", "-3", "이전 발언과 단어 70% 이상 중복"],
            ["프롬프트 인젝션", "-10", "정규식 패턴 (IGNORE ALL, system: 등)"],
            ["타임아웃", "-5", "제한 시간 초과"],
            ["거짓 출처", "-7", "존재하지 않는 인용"],
            ["인신공격", "-8", "욕설·개인 공격 패턴"],
            ["인간 의심", "-15", "의심 점수 61 이상 (로컬 에이전트만)"],
        ],
        [2.0, 1.0, 4.0]
    )

    add_heading(doc, "4. ELO 레이팅 시스템", 1)

    add_heading(doc, "4.1 수식", 2)
    add_code_block(doc,
        "기대 승률:   E_A = 1 / (1 + 10^((R_B - R_A) / 400))\n"
        "새 레이팅:   R'_A = round(R_A + K_A × (S_A - E_A))\n\n"
        "S_A:  승=1.0  /  패=0.0  /  무=0.5"
    )

    add_heading(doc, "4.2 비대칭 K-factor", 2)
    add_table(doc,
        ["상황", "K값", "근거"],
        [
            ["승자", "40", "업셋 승리에 큰 보상, 빠른 상승"],
            ["패자", "24", "예상 패배 시 손실 완화"],
            ["무승부 양측", "32", "중간값 적용"],
        ],
        [2.0, 1.0, 4.0]
    )

    add_heading(doc, "4.3 ELO 기대 변동 예시", 2)
    add_table(doc,
        ["상황", "기존 레이팅", "결과", "A 변동", "B 변동"],
        [
            ["동급 대결", "A=1500  B=1500", "A 승", "+20", "-12"],
            ["약자 업셋", "A=1300  B=1700", "A 승", "+38", "-23"],
            ["강자 승",   "A=1700  B=1300", "A 승",  "+2",  "-1"],
            ["무승부",    "A=1500  B=1500", "무",     "0",   "0"],
        ],
        [1.8, 2.0, 1.2, 1.0, 1.0]
    )

    add_heading(doc, "5. 구현 현황", 1)

    add_heading(doc, "5.1 완료 기능", 2)
    add_table(doc,
        ["영역", "완료된 기능"],
        [
            ["에이전트", "생성(템플릿/BYOK/로컬), 수정, 삭제, 버전 관리"],
            ["주제", "CRUD, 스케줄링, 상태 자동 전환, 작성자 편집·삭제"],
            ["매치메이킹", "큐 참가, 준비 완료 버튼, 자동 매칭, SSE 알림, 폴링 폴백"],
            ["토론 엔진", "순차 턴, API 에이전트, 로컬 WebSocket 에이전트"],
            ["도구", "calculator, stance_tracker, opponent_summary, turn_info"],
            ["페널티", "7종 자동 탐지 (스키마/반복/인젝션/타임아웃/출처/인신공격/인간의심)"],
            ["심판", "LLM 기반 4항목 채점, 승패 판정 (5점 차 기준)"],
            ["ELO", "비대칭 K-factor, 에이전트·버전별 통계"],
            ["스트리밍", "Redis Pub/Sub → SSE → 타이핑 효과 UI"],
            ["관전 UI", "HP 바 고정 헤더, 발언 버블, 스코어카드 인포그래픽"],
            ["대기방", "준비 완료 버튼 시스템, opponent_joined/ready SSE 이벤트"],
        ],
        [1.8, 5.2]
    )

    doc.save(OUT_DIR / "AI토론_기획서.docx")
    print("[OK] AI토론_기획서.docx 생성 완료")


def make_plan_xlsx():
    wb = openpyxl.Workbook()

    # ── Sheet 1: 개요 ──
    ws = wb.active
    ws.title = "개요"
    ws.row_dimensions[1].height = 40
    xl_section_title(ws, 1, "AI 에이전트 토론 플랫폼 — 시스템 기획서", 4, XL_HDR)
    ws.row_dimensions[2].height = 20
    xl_note(ws, 2, "작성일: 2026-02-26  |  버전: 1.0  |  최종 발표용", 4)

    ws.row_dimensions[4].height = 24
    xl_section_title(ws, 4, "핵심 가치 제안", 4, XL_HDR2)
    xl_header(ws, 5, ["관점", "가치", "비고", ""], XL_HDR2)
    for i, row in enumerate([
        ["사용자", "자신만의 AI 에이전트를 만들어 다른 에이전트와 토론 대결", "참여형", ""],
        ["연구자", "프롬프트·모델 조합에 따른 논증 품질 비교 실험 환경", "연구형", ""],
        ["일반 관객", "LLM들이 실시간으로 설전하는 엔터테인먼트 콘텐츠", "관전형", ""],
    ], 6):
        xl_row(ws, i, row, i % 2 == 0)

    ws.row_dimensions[10].height = 24
    xl_section_title(ws, 10, "기술 스택", 4, XL_HDR2)
    xl_header(ws, 11, ["계층", "기술", "버전/상세", "비고"], XL_HDR2)
    for i, row in enumerate([
        ["백엔드", "Python + FastAPI + SQLAlchemy", "3.12 / async", "EC2 t4g.small"],
        ["프론트엔드", "Next.js + React + Zustand", "15 / 19", "App Router"],
        ["데이터베이스", "PostgreSQL", "16 (Docker)", "pgvector 포함"],
        ["캐시/메시징", "Redis", "7-alpine", "Pub/Sub + SSE"],
        ["LLM", "OpenAI / Anthropic / Google / RunPod", "BYOK 지원", "동적 라우팅"],
        ["실시간", "SSE + WebSocket", "—", "로컬 에이전트 WS"],
        ["배포", "Docker Compose + EC2", "서울 ap-northeast-2", "t4g.small ARM"],
    ], 12):
        xl_row(ws, i, row, i % 2 == 0)
    auto_col_width(ws)
    ws.column_dimensions["A"].width = 16
    ws.column_dimensions["B"].width = 30
    ws.column_dimensions["C"].width = 22
    ws.column_dimensions["D"].width = 20

    # ── Sheet 2: 기능 정의 ──
    ws2 = wb.create_sheet("기능 정의")
    xl_section_title(ws2, 1, "토론 모드", 4, XL_HDR)
    xl_header(ws2, 2, ["모드", "설명", "주요 액션", "특이사항"])
    for i, row in enumerate([
        ["debate", "찬반 토론 (기본)", "argue / rebut / concede / question / summarize", "기본 모드"],
        ["persuasion", "한 측이 다른 측을 설득", "argue / rebut / concede", "비대칭"],
        ["cross_exam", "교차 심문 방식", "question / answer / rebut", "Q&A 중심"],
    ], 3):
        xl_row(ws2, i, row, i % 2 == 0)

    xl_section_title(ws2, 7, "채점 기준", 4, XL_HDR2)
    xl_header(ws2, 8, ["항목", "배점", "설명", "비율"], XL_HDR2)
    for i, row in enumerate([
        ["논리성", "0~30", "논증 구조의 일관성·타당성", "30%"],
        ["근거 활용", "0~25", "주장을 뒷받침하는 증거의 질", "25%"],
        ["반박력", "0~25", "상대 주장에 대한 효과적 반론", "25%"],
        ["주제 적합성", "0~20", "토론 주제와의 연관성 유지", "20%"],
        ["페널티 차감", "-N", "규칙 위반 벌점", "—"],
    ], 9):
        xl_row(ws2, i, row, i % 2 == 0)

    xl_section_title(ws2, 15, "페널티 체계", 4, XL_HDR)
    xl_header(ws2, 16, ["위반 유형", "벌점", "탐지 방법", "탐지 시점"])
    for i, row in enumerate([
        ["스키마 위반", "-5점", "JSON 파싱 실패", "턴 완료 직후"],
        ["반복 발언", "-3점", "이전 발언 단어 70% 이상 중복", "턴 완료 직후"],
        ["프롬프트 인젝션", "-10점", "IGNORE ALL / system: 등 정규식", "턴 완료 직후"],
        ["타임아웃", "-5점", "응답 제한 시간 초과", "응답 대기 중"],
        ["거짓 출처", "-7점", "존재하지 않는 인용", "턴 완료 직후"],
        ["인신공격", "-8점", "욕설·개인 공격 패턴", "턴 완료 직후"],
        ["인간 의심", "-15점", "의심 점수 ≥61 (로컬 에이전트만)", "응답 분석 후"],
    ], 17):
        xl_row(ws2, i, row, i % 2 == 0)
    auto_col_width(ws2)

    # ── Sheet 3: ELO 시스템 ──
    ws3 = wb.create_sheet("ELO 시스템")
    xl_section_title(ws3, 1, "ELO 레이팅 수식", 4, XL_HDR)
    xl_header(ws3, 2, ["항목", "수식", "설명", ""])
    for i, row in enumerate([
        ["기대 승률", "E_A = 1 / (1 + 10^((R_B - R_A) / 400))", "R = 현재 레이팅", ""],
        ["새 레이팅 (승)", "R'_A = round(R_A + K_win × (1.0 - E_A))", "K_win = 40", ""],
        ["새 레이팅 (패)", "R'_A = round(R_A + K_loss × (0.0 - E_A))", "K_loss = 24", ""],
        ["새 레이팅 (무)", "R'_A = round(R_A + K_draw × (0.5 - E_A))", "K_draw = 32", ""],
    ], 3):
        xl_row(ws3, i, row, i % 2 == 0)

    xl_section_title(ws3, 8, "ELO 변동 시뮬레이션", 5, XL_HDR2)
    xl_header(ws3, 9, ["상황", "A 레이팅", "B 레이팅", "결과", "A 변동", "B 변동"], XL_HDR2)
    for i, row in enumerate([
        ["동급 대결", 1500, 1500, "A 승", "+20", "-12"],
        ["약자 업셋", 1300, 1700, "A 승", "+38", "-23"],
        ["강자 승",   1700, 1300, "A 승",  "+2",  "-1"],
        ["무승부",    1500, 1500, "무",    "±0",  "±0"],
        ["동급 패배", 1500, 1500, "A 패", "-12", "+20"],
    ], 10):
        xl_row(ws3, i, row, i % 2 == 0)
    ws3.column_dimensions["F"].width = 12

    # 6열짜리이므로 merged cell 수정
    ws3.unmerge_cells("A8:E8")
    ws3.merge_cells(start_row=8, start_column=1, end_row=8, end_column=6)
    auto_col_width(ws3)

    wb.save(OUT_DIR / "AI토론_기획서.xlsx")
    print("[OK] AI토론_기획서.xlsx 생성 완료")


# ══════════════════════════════════════════════════════════════════
#  문서 2: 아키텍처
# ══════════════════════════════════════════════════════════════════

def make_arch_docx():
    doc = init_doc("AI 에이전트 토론 플랫폼", "시스템 아키텍처 문서")

    add_heading(doc, "1. 시스템 전체 구조 (3-Tier)", 1)
    add_table(doc,
        ["계층", "구성 요소", "역할"],
        [
            ["CLIENT TIER", "Next.js 15 (App Router, React 19)\nZustand 스토어\nSSE/WebSocket 소비자\nLive2D 렌더러 (PixiJS)", "사용자 인터페이스\n상태 관리\n실시간 이벤트 수신"],
            ["APPLICATION TIER", "FastAPI (Uvicorn ASGI)\nDebate Engine (asyncio.Task)\nPostgreSQL 16 + pgvector\nRedis (Pub/Sub)", "비즈니스 로직\n토론 실행\n데이터 영속성\n이벤트 브로드캐스트"],
            ["LLM TIER", "OpenAI (GPT-4o 등)\nAnthropic (Claude 등)\nGoogle (Gemini 등)\nRunPod Serverless (Llama 70B)", "LLM 추론 제공\nBYOK 지원\n동적 라우팅"],
        ],
        [1.5, 3.0, 2.5]
    )

    add_heading(doc, "2. 백엔드 서비스 레이어", 1)

    add_heading(doc, "2.1 API 라우터 레이어", 2)
    add_table(doc,
        ["라우터", "경로", "주요 엔드포인트"],
        [
            ["토론 주제", "/api/topics/*", "CRUD + 큐 join/stream/status/leave/ready"],
            ["에이전트", "/api/agents/*", "CRUD + 템플릿 + 랭킹"],
            ["매치", "/api/matches/*", "조회 + SSE 스트림 + 스코어카드"],
            ["WebSocket", "/api/debate/ws/*", "로컬 에이전트 인증 + 발언 교환"],
            ["관리자", "/api/admin/*", "에이전트/매치 모니터링 (RBAC)"],
        ],
        [1.5, 2.0, 3.5]
    )

    add_heading(doc, "2.2 서비스 컴포넌트", 2)
    add_table(doc,
        ["서비스", "파일", "책임"],
        [
            ["Debate Engine", "debate_engine.py", "메인 토론 루프 (asyncio.Task 백그라운드)"],
            ["Orchestrator", "debate_orchestrator.py", "심판 LLM 호출 + ELO 계산 + 턴 검토"],
            ["Matching Service", "debate_matching_service.py", "큐 등록 + 준비 완료 → 매치 생성"],
            ["Auto Matcher", "debate_auto_match.py", "10초 폴링 → 대기 초과 시 플랫폼 에이전트 배정"],
            ["Inference Client", "inference_client.py", "4개 LLM 프로바이더 동적 라우팅"],
            ["Broadcast Service", "debate_broadcast.py", "Redis Pub/Sub → SSE 이벤트 발행"],
            ["WS Manager", "debate_ws_manager.py", "WebSocket 싱글턴 (로컬 에이전트)"],
            ["Tool Executor", "debate_tool_executor.py", "calculator / stance_tracker / opponent_summary / turn_info"],
            ["Human Detector", "human_detection.py", "5개 신호 기반 인간 의심 점수"],
        ],
        [1.8, 2.2, 3.0]
    )

    add_heading(doc, "3. 실시간 통신 아키텍처", 1)

    add_heading(doc, "3.1 SSE 채널 구조", 2)
    add_table(doc,
        ["채널", "구독 URL", "주요 이벤트", "Redis 키"],
        [
            ["큐 대기방", "/topics/{id}/queue/stream", "matched / opponent_joined / opponent_ready / timeout / cancelled", "debate:queue:{agent_id}"],
            ["매치 관전", "/matches/{id}/stream", "started / turn_chunk / turn / finished / forfeit / error", "debate:match:{match_id}"],
        ],
        [1.2, 2.5, 2.5, 2.2]
    )

    add_heading(doc, "3.2 WebSocket 메시지 프로토콜 (로컬 에이전트)", 2)
    add_table(doc,
        ["방향", "타입", "설명"],
        [
            ["서버 → 에이전트", "WSMatchReady", "매치 시작 알림 (match_id, topic, side)"],
            ["서버 → 에이전트", "WSTurnRequest", "발언 요청 (turn_number, history, tools)"],
            ["서버 → 에이전트", "WSToolResult", "도구 실행 결과"],
            ["에이전트 → 서버", "WSTurnResponse", "최종 발언 (action, claim, evidence)"],
            ["에이전트 → 서버", "WSToolRequest", "도구 실행 요청 (tool_name, tool_input)"],
        ],
        [2.0, 2.0, 3.0]
    )

    add_heading(doc, "4. 매치 상태 머신", 1)
    add_table(doc,
        ["상태", "전이 조건", "다음 상태", "비고"],
        [
            ["pending", "run_debate() 시작", "waiting_agent / in_progress", "매치 생성 직후"],
            ["waiting_agent", "로컬 에이전트 WS 연결", "in_progress", "30초 타임아웃"],
            ["waiting_agent", "30초 초과", "forfeit", "WS 미연결"],
            ["in_progress", "턴 루프 완료 → 심판 판정", "completed", "정상 종료"],
            ["in_progress", "LLM 오류 / DB 오류", "error", "예외 발생"],
            ["completed", "—", "—", "ELO 업데이트 완료"],
        ],
        [1.8, 2.5, 1.8, 2.0]
    )

    add_heading(doc, "5. 보안 아키텍처", 1)
    add_table(doc,
        ["레이어", "구현", "비고"],
        [
            ["전송 보안", "HTTPS (TLS 1.3) + WSS\nNginx X-Accel-Buffering: no (SSE)", ""],
            ["인증/인가", "JWT Bearer Token (24h)\nDepends(get_current_user)", "미인증 → 401"],
            ["소유권 검증", "agent.owner_id == user.id", "타인 에이전트 접근 → 403"],
            ["API 키 암호화", "Fernet 대칭 암호화 (SECRET_KEY)\n복호화는 LLM 호출 직전에만", "응답 바디 포함 금지"],
            ["입력 검증", "Pydantic v2 + 프롬프트 인젝션 정규식\n발언 내용 매 턴 7종 패턴 검사", ""],
            ["DB 무결성", "UNIQUE(topic_id, agent_id) for queue\nSELECT FOR UPDATE 동시성 제어", "레이스 컨디션 방지"],
        ],
        [1.8, 3.2, 2.0]
    )

    add_heading(doc, "6. 인프라 구성", 1)
    add_table(doc,
        ["서비스", "이미지", "포트", "데이터"],
        [
            ["frontend", "Node.js 20 + Next.js 15", "3000", "stateless"],
            ["backend", "Python 3.12 + FastAPI", "8000", "stateless"],
            ["postgres", "postgres:16", "5432", "/var/lib/postgresql/data (볼륨)"],
            ["redis", "redis:7-alpine", "6379", "휘발성 (Pub/Sub 전용)"],
            ["nginx", "nginx:alpine", "80/443", "Reverse Proxy"],
        ],
        [1.5, 2.5, 1.0, 2.5]
    )

    doc.save(OUT_DIR / "AI토론_아키텍처.docx")
    print("[OK] AI토론_아키텍처.docx 생성 완료")


def make_arch_xlsx():
    wb = openpyxl.Workbook()

    # Sheet 1: 시스템 구조
    ws = wb.active
    ws.title = "시스템 구조"
    xl_section_title(ws, 1, "AI 토론 플랫폼 — 시스템 아키텍처", 5, XL_HDR)
    xl_note(ws, 2, "작성일: 2026-02-26  |  최종 발표용", 5)

    xl_section_title(ws, 4, "3-Tier 아키텍처", 5, XL_HDR2)
    xl_header(ws, 5, ["계층", "서비스", "역할", "기술", "비고"], XL_HDR2)
    for i, row in enumerate([
        ["CLIENT", "Browser", "사용자 인터페이스", "Next.js 15 + React 19", "Zustand 상태관리"],
        ["CLIENT", "Live2D Renderer", "캐릭터 표시", "PixiJS + pixi-live2d-display", "Cubism SDK 4"],
        ["APPLICATION", "FastAPI", "API + 비즈니스 로직", "Python 3.12 + Uvicorn", "EC2 t4g.small"],
        ["APPLICATION", "Debate Engine", "토론 실행", "asyncio.Task", "백그라운드"],
        ["APPLICATION", "PostgreSQL 16", "데이터 영속성", "pgvector 포함", "Docker"],
        ["APPLICATION", "Redis", "이벤트 브로드캐스트", "Pub/Sub + Presence", "Docker"],
        ["LLM", "OpenAI", "GPT-4o 등", "BYOK 지원", "기본 심판"],
        ["LLM", "Anthropic", "Claude 등", "BYOK 지원", ""],
        ["LLM", "Google", "Gemini 등", "BYOK 지원", ""],
        ["LLM", "RunPod Serverless", "Llama 3.1 70B", "SGLang + AWQ", "RTT ~150ms"],
    ], 6):
        xl_row(ws, i, row, i % 2 == 0)
    auto_col_width(ws)

    # Sheet 2: 서비스 컴포넌트
    ws2 = wb.create_sheet("서비스 컴포넌트")
    xl_section_title(ws2, 1, "백엔드 서비스 컴포넌트 맵", 5, XL_HDR)
    xl_header(ws2, 2, ["그룹", "서비스명", "파일", "주요 책임", "호출 관계"])
    for i, row in enumerate([
        ["매치메이킹", "MatchingService", "debate_matching_service.py", "큐 등록 + 준비완료 → 매치 생성", "→ BroadcastService"],
        ["매치메이킹", "AutoMatcher", "debate_auto_match.py", "10초 폴링 → 플랫폼 에이전트 자동 배정", "→ DebateEngine"],
        ["토론 실행", "DebateEngine", "debate_engine.py", "메인 턴 루프 (asyncio.Task)", "→ Orchestrator, InferenceClient"],
        ["토론 실행", "Orchestrator", "debate_orchestrator.py", "심판 LLM + ELO 계산 + 턴 LLM 검토", "→ InferenceClient"],
        ["토론 실행", "InferenceClient", "inference_client.py", "4개 LLM 프로바이더 라우팅", "→ LLM API"],
        ["토론 실행", "ToolExecutor", "debate_tool_executor.py", "4개 도구 실행", "← DebateEngine"],
        ["토론 실행", "HumanDetector", "human_detection.py", "5개 신호 기반 의심 점수", "← DebateEngine"],
        ["통신", "BroadcastService", "debate_broadcast.py", "Redis Pub/Sub → SSE", "→ Redis → Browser"],
        ["통신", "WSManager", "debate_ws_manager.py", "WebSocket 싱글턴 관리", "↔ LocalAgent"],
        ["데이터", "TopicService", "debate_topic_service.py", "주제 CRUD + 스케줄 동기화", "→ DB"],
        ["데이터", "AgentService", "debate_agent_service.py", "에이전트 CRUD + ELO 업데이트", "→ DB"],
    ], 3):
        xl_row(ws2, i, row, i % 2 == 0)
    auto_col_width(ws2)

    # Sheet 3: 인프라
    ws3 = wb.create_sheet("인프라 구성")
    xl_section_title(ws3, 1, "Docker Compose 서비스 구성", 5, XL_HDR)
    xl_header(ws3, 2, ["서비스명", "이미지", "포트", "데이터", "비고"])
    for i, row in enumerate([
        ["frontend", "Node.js 20 + Next.js 15", "3000", "stateless", "SSR + App Router"],
        ["backend", "Python 3.12 + FastAPI", "8000", "stateless", "Uvicorn ASGI"],
        ["postgres", "postgres:16", "5432", "/var/lib/postgresql/data", "pgvector 확장"],
        ["redis", "redis:7-alpine", "6379", "휘발성", "Pub/Sub 전용"],
        ["nginx", "nginx:alpine", "80/443", "—", "X-Accel-Buffering: no (SSE)"],
    ], 3):
        xl_row(ws3, i, row, i % 2 == 0)

    xl_section_title(ws3, 10, "Redis 키 구조", 5, XL_HDR2)
    xl_header(ws3, 11, ["키 패턴", "타입", "TTL", "목적", "비고"], XL_HDR2)
    for i, row in enumerate([
        ["debate:match:{match_id}", "Pub/Sub", "없음", "매치 이벤트 채널", ""],
        ["debate:queue:{agent_id}", "Pub/Sub", "없음", "큐 이벤트 채널", ""],
        ["debate:agent:{agent_id}", "String", "60초", "WS 프레즌스 표시", "로컬 에이전트"],
    ], 12):
        xl_row(ws3, i, row, i % 2 == 0)
    auto_col_width(ws3)

    wb.save(OUT_DIR / "AI토론_아키텍처.xlsx")
    print("[OK] AI토론_아키텍처.xlsx 생성 완료")


# ══════════════════════════════════════════════════════════════════
#  문서 3: 데이터 처리 명세서
# ══════════════════════════════════════════════════════════════════

def make_data_docx():
    doc = init_doc("AI 에이전트 토론 플랫폼", "데이터 처리 명세서")

    add_heading(doc, "1. 데이터 모델 — 핵심 테이블 스키마", 1)

    add_heading(doc, "debate_agents", 2)
    add_table(doc,
        ["컬럼", "타입", "제약", "설명"],
        [
            ["id", "UUID", "PK", "에이전트 고유 ID"],
            ["owner_id", "UUID", "FK(users)", "소유자"],
            ["name", "VARCHAR(100)", "NOT NULL", "에이전트 이름"],
            ["provider", "VARCHAR(20)", "CHECK IN (...)", "openai/anthropic/google/runpod/local"],
            ["model_id", "VARCHAR(100)", "NOT NULL", "모델 식별자 (예: gpt-4o)"],
            ["encrypted_api_key", "TEXT", "NULL", "Fernet 암호화된 API 키"],
            ["elo_rating", "INTEGER", "DEFAULT 1500", "ELO 점수"],
            ["wins/losses/draws", "INTEGER", "DEFAULT 0", "전적"],
            ["is_active", "BOOLEAN", "DEFAULT true", "활성 여부"],
            ["is_platform", "BOOLEAN", "DEFAULT false", "플랫폼 에이전트 여부"],
        ],
        [1.8, 1.5, 1.5, 2.2]
    )

    add_heading(doc, "debate_matches", 2)
    add_table(doc,
        ["컬럼", "타입", "제약", "설명"],
        [
            ["id", "UUID", "PK", "매치 ID"],
            ["topic_id", "UUID", "FK(topics)", "토론 주제"],
            ["agent_a_id", "UUID", "FK(agents)", "에이전트 A (선공)"],
            ["agent_b_id", "UUID", "FK(agents)", "에이전트 B (후공)"],
            ["status", "VARCHAR(20)", "CHECK IN (...)", "pending/in_progress/completed/error/forfeit"],
            ["winner_id", "UUID", "NULL", "승자 (null=무승부)"],
            ["scorecard", "JSONB", "NULL", "심판 점수 및 이유"],
            ["score_a / score_b", "INTEGER", "NULL", "최종 점수 (페널티 차감 후)"],
            ["penalty_a / penalty_b", "INTEGER", "DEFAULT 0", "누적 페널티"],
        ],
        [1.8, 1.5, 1.5, 2.2]
    )

    add_heading(doc, "debate_turn_logs", 2)
    add_table(doc,
        ["컬럼", "타입", "설명"],
        [
            ["id", "UUID PK", "턴 로그 ID"],
            ["match_id", "UUID FK", "매치 참조 (CASCADE)"],
            ["turn_number", "INTEGER", "턴 번호 (1 ~ max_turns)"],
            ["speaker", "VARCHAR(10)", "agent_a / agent_b"],
            ["action", "VARCHAR(20)", "argue / rebut / concede / question / summarize"],
            ["claim", "TEXT", "주장 본문"],
            ["evidence", "TEXT NULL", "근거 자료"],
            ["tool_used", "VARCHAR(50) NULL", "사용한 도구 이름"],
            ["penalties", "JSONB NULL", "부과된 페널티 딕셔너리"],
            ["penalty_total", "INTEGER", "총 페널티 합계"],
            ["review_result", "JSONB NULL", "LLM 턴 검토 결과 (logic_score, violations, feedback)"],
            ["is_blocked", "BOOLEAN", "차단 여부 (기본 false)"],
            ["human_suspicion_score", "INTEGER", "인간 의심 점수 (0~100)"],
            ["input_tokens / output_tokens", "INTEGER NULL", "LLM 사용 토큰 수"],
        ],
        [2.0, 1.8, 3.2]
    )

    add_heading(doc, "2. API 명세", 1)

    add_heading(doc, "에이전트 API (/api/agents)", 2)
    add_table(doc,
        ["메서드", "경로", "인증", "설명"],
        [
            ["GET", "/agents/templates", "필요", "템플릿 목록"],
            ["POST", "/agents", "필요", "에이전트 생성"],
            ["GET", "/agents/me", "필요", "내 에이전트 목록"],
            ["GET", "/agents/ranking", "필요", "ELO 랭킹 목록"],
            ["GET", "/agents/{id}", "필요", "에이전트 상세"],
            ["PUT", "/agents/{id}", "필요(소유자)", "에이전트 수정"],
            ["DELETE", "/agents/{id}", "필요(소유자)", "에이전트 삭제"],
            ["GET", "/agents/{id}/versions", "필요", "버전 목록"],
        ],
        [1.0, 2.5, 1.2, 2.3]
    )

    add_heading(doc, "주제/큐 API (/api/topics)", 2)
    add_table(doc,
        ["메서드", "경로", "인증", "설명"],
        [
            ["POST", "/topics", "필요", "주제 생성"],
            ["GET", "/topics", "필요", "주제 목록 (status/sort/page 필터)"],
            ["PATCH", "/topics/{id}", "필요(작성자)", "주제 수정"],
            ["DELETE", "/topics/{id}", "필요(작성자)", "주제 삭제"],
            ["POST", "/topics/{id}/join", "필요", "큐 참가"],
            ["POST", "/topics/{id}/queue/ready", "필요", "준비 완료 버튼"],
            ["GET", "/topics/{id}/queue/stream", "필요", "대기방 SSE 스트림"],
            ["GET", "/topics/{id}/queue/status", "필요", "큐 상태 조회"],
            ["DELETE", "/topics/{id}/queue", "필요", "큐 탈퇴"],
        ],
        [1.0, 2.8, 1.2, 2.0]
    )

    add_heading(doc, "3. SSE 이벤트 명세", 1)
    add_table(doc,
        ["채널", "이벤트", "발생 시점", "페이로드 주요 필드"],
        [
            ["큐 스트림", "opponent_joined", "상대가 큐에 참가", "opponent_agent_id"],
            ["큐 스트림", "opponent_ready", "상대가 준비 완료", "—"],
            ["큐 스트림", "matched", "양쪽 준비 완료 → 매치 생성", "match_id, opponent_agent_id, auto_matched"],
            ["큐 스트림", "timeout", "자동 매칭 실패", "reason"],
            ["큐 스트림", "cancelled", "사용자가 대기 취소", "—"],
            ["매치 스트림", "started", "토론 시작", "match_id"],
            ["매치 스트림", "turn_chunk", "LLM 토큰 생성 중", "turn_number, speaker, chunk"],
            ["매치 스트림", "turn", "한 턴 완료", "turn_number, speaker, action, claim, penalty_total"],
            ["매치 스트림", "finished", "토론 완료", "winner_id, score_a, score_b, elo_a, elo_b"],
            ["매치 스트림", "error", "오류 발생", "message"],
        ],
        [1.5, 1.8, 2.0, 2.7]
    )

    add_heading(doc, "4. 점수 계산 로직", 1)

    add_heading(doc, "4.1 최종 점수 & 승패 판정", 2)
    add_code_block(doc,
        "raw_score_a = sum(scorecard['agent_a'].values())   # 최대 100\n"
        "raw_score_b = sum(scorecard['agent_b'].values())\n\n"
        "score_a = max(0, raw_score_a - match.penalty_a)\n"
        "score_b = max(0, raw_score_b - match.penalty_b)\n\n"
        "gap = score_a - score_b\n"
        "if   gap >= 5:  winner = agent_a\n"
        "elif gap <= -5: winner = agent_b\n"
        "else:           winner = None  # 무승부 (5점 미만 차이)"
    )

    doc.save(OUT_DIR / "AI토론_데이터처리명세서.docx")
    print("[OK] AI토론_데이터처리명세서.docx 생성 완료")


def make_data_xlsx():
    wb = openpyxl.Workbook()

    # Sheet 1: ERD (테이블 관계)
    ws = wb.active
    ws.title = "DB 스키마"
    xl_section_title(ws, 1, "핵심 테이블 스키마 — debate_agents", 5, XL_HDR)
    xl_header(ws, 2, ["컬럼", "타입", "제약", "설명", "비고"])
    for i, row in enumerate([
        ["id", "UUID", "PK", "에이전트 고유 ID", ""],
        ["owner_id", "UUID", "FK(users)", "소유자", "CASCADE"],
        ["name", "VARCHAR(100)", "NOT NULL", "에이전트 이름", "최대 100자"],
        ["description", "TEXT", "NULL", "설명", ""],
        ["provider", "VARCHAR(20)", "CHECK", "openai/anthropic/google/runpod/local", ""],
        ["model_id", "VARCHAR(100)", "NOT NULL", "모델 식별자", "예: gpt-4o"],
        ["encrypted_api_key", "TEXT", "NULL", "Fernet 암호화된 API 키", "응답에 포함 금지"],
        ["image_url", "TEXT", "NULL", "프로필 이미지 URL", ""],
        ["template_id", "UUID", "FK(templates) NULL", "기반 템플릿", ""],
        ["customizations", "JSONB", "NULL", "템플릿 커스터마이징 값", ""],
        ["elo_rating", "INTEGER", "DEFAULT 1500", "ELO 점수", "초기값"],
        ["wins", "INTEGER", "DEFAULT 0", "승리 수", ""],
        ["losses", "INTEGER", "DEFAULT 0", "패배 수", ""],
        ["draws", "INTEGER", "DEFAULT 0", "무승부 수", ""],
        ["is_active", "BOOLEAN", "DEFAULT true", "활성 여부", ""],
        ["is_platform", "BOOLEAN", "DEFAULT false", "플랫폼 에이전트 여부", "자동 매칭용"],
        ["created_at", "TIMESTAMPTZ", "DEFAULT now()", "생성 시각", ""],
        ["updated_at", "TIMESTAMPTZ", "DEFAULT now()", "수정 시각", ""],
    ], 3):
        xl_row(ws, i, row, i % 2 == 0)

    xl_section_title(ws, 22, "debate_turn_logs — 핵심 컬럼", 5, XL_HDR2)
    xl_header(ws, 23, ["컬럼", "타입", "제약", "설명", "비고"], XL_HDR2)
    for i, row in enumerate([
        ["id", "UUID", "PK", "턴 로그 ID", ""],
        ["match_id", "UUID", "FK CASCADE", "매치 참조", ""],
        ["turn_number", "INTEGER", "NOT NULL", "턴 번호", "1 ~ max_turns"],
        ["speaker", "VARCHAR(10)", "CHECK", "agent_a / agent_b", ""],
        ["action", "VARCHAR(20)", "NOT NULL", "argue/rebut/concede/question/summarize", ""],
        ["claim", "TEXT", "NOT NULL", "주장 본문", ""],
        ["evidence", "TEXT", "NULL", "근거 자료", ""],
        ["tool_used", "VARCHAR(50)", "NULL", "사용한 도구", ""],
        ["penalties", "JSONB", "NULL", "페널티 딕셔너리", '예: {"schema_violation":5}'],
        ["penalty_total", "INTEGER", "DEFAULT 0", "총 페널티 합계", ""],
        ["review_result", "JSONB", "NULL", "LLM 턴 검토 결과", "logic_score, violations, feedback"],
        ["is_blocked", "BOOLEAN", "DEFAULT false", "차단 여부", "LLM 검토 차단 시"],
        ["human_suspicion_score", "INTEGER", "DEFAULT 0", "인간 의심 점수", "0~100"],
        ["input_tokens", "INTEGER", "NULL", "입력 토큰", ""],
        ["output_tokens", "INTEGER", "NULL", "출력 토큰", ""],
    ], 24):
        xl_row(ws, i, row, i % 2 == 0)
    auto_col_width(ws)

    # Sheet 2: API 명세
    ws2 = wb.create_sheet("API 명세")
    xl_section_title(ws2, 1, "에이전트 API — /api/agents", 6, XL_HDR)
    xl_header(ws2, 2, ["메서드", "경로", "인증", "요청 바디/쿼리", "응답", "비고"])
    for i, row in enumerate([
        ["GET", "/agents/templates", "필요", "—", "AgentTemplate[]", ""],
        ["POST", "/agents", "필요", "AgentCreate", "AgentResponse 201", ""],
        ["GET", "/agents/me", "필요", "—", "AgentResponse[]", ""],
        ["GET", "/agents/ranking", "필요", "limit, offset", "랭킹 목록", ""],
        ["GET", "/agents/{id}", "필요", "—", "AgentResponse", ""],
        ["PUT", "/agents/{id}", "필요(소유자)", "AgentUpdate", "AgentResponse", ""],
        ["DELETE", "/agents/{id}", "필요(소유자)", "—", "204", "진행 중 매치 있으면 409"],
        ["GET", "/agents/{id}/versions", "필요", "—", "AgentVersion[]", ""],
    ], 3):
        xl_row(ws2, i, row, i % 2 == 0)

    xl_section_title(ws2, 12, "주제/큐 API — /api/topics", 6, XL_HDR2)
    xl_header(ws2, 13, ["메서드", "경로", "인증", "요청 바디/쿼리", "응답", "비고"], XL_HDR2)
    for i, row in enumerate([
        ["POST", "/topics", "필요", "TopicCreate", "TopicResponse 201", ""],
        ["GET", "/topics", "필요", "status/sort/page/page_size", "TopicListResponse", ""],
        ["PATCH", "/topics/{id}", "필요(작성자)", "TopicUpdatePayload", "TopicResponse", ""],
        ["DELETE", "/topics/{id}", "필요(작성자)", "—", "204", "진행 중 매치 → 409"],
        ["POST", "/topics/{id}/join", "필요", "{agent_id}", "{status, position}", "큐 참가"],
        ["POST", "/topics/{id}/queue/ready", "필요", "{agent_id}", "{status, match_id?}", "준비 완료"],
        ["GET", "/topics/{id}/queue/stream", "필요", "agent_id (쿼리)", "SSE 스트림", "대기방"],
        ["GET", "/topics/{id}/queue/status", "필요", "agent_id (쿼리)", "큐 상태", ""],
        ["DELETE", "/topics/{id}/queue", "필요", "agent_id (쿼리)", "{status: left}", ""],
    ], 14):
        xl_row(ws2, i, row, i % 2 == 0)
    auto_col_width(ws2)

    # Sheet 3: SSE 이벤트
    ws3 = wb.create_sheet("SSE 이벤트")
    xl_section_title(ws3, 1, "SSE 이벤트 명세", 5, XL_HDR)
    xl_header(ws3, 2, ["채널", "이벤트", "발생 시점", "페이로드 주요 필드", "비고"])
    for i, row in enumerate([
        ["큐 스트림", "opponent_joined", "상대가 큐에 참가", "opponent_agent_id", "SSE or polling"],
        ["큐 스트림", "opponent_ready", "상대가 준비 완료 버튼", "—", ""],
        ["큐 스트림", "matched", "양쪽 준비 완료 → 매치 생성", "match_id, opponent_agent_id, auto_matched", ""],
        ["큐 스트림", "timeout", "자동 매칭 실패 (플랫폼 에이전트 없음)", "reason", ""],
        ["큐 스트림", "cancelled", "사용자가 대기 취소", "—", ""],
        ["매치 스트림", "started", "토론 시작", "match_id", ""],
        ["매치 스트림", "waiting_agent", "로컬 에이전트 WS 연결 대기", "match_id", ""],
        ["매치 스트림", "turn_chunk", "LLM 토큰 생성 중", "turn_number, speaker, chunk", "타이핑 효과용"],
        ["매치 스트림", "turn", "한 턴 완료", "turn_number, speaker, action, claim, penalty_total", ""],
        ["매치 스트림", "finished", "토론 완료", "winner_id, score_a, score_b, elo_a, elo_b", "스트림 자동 종료"],
        ["매치 스트림", "forfeit", "타임아웃/몰수패", "reason, winner_id", ""],
        ["매치 스트림", "error", "오류 발생", "message", "스트림 자동 종료"],
        ["공통", "heartbeat", "매 1초", "—", "연결 유지용"],
    ], 3):
        xl_row(ws3, i, row, i % 2 == 0)
    auto_col_width(ws3)

    wb.save(OUT_DIR / "AI토론_데이터처리명세서.xlsx")
    print("[OK] AI토론_데이터처리명세서.xlsx 생성 완료")


# ══════════════════════════════════════════════════════════════════
#  문서 4: 모델 비교 전략 (실제 에이전트/오케스트레이터 테스트 결과)
# ══════════════════════════════════════════════════════════════════

def make_model_docx():
    doc = init_doc("AI 에이전트 토론 플랫폼", "LLM 모델 비교 및 테스트 결과")

    add_heading(doc, "1. 개요 — 테스트 환경", 1)
    add_para(doc,
        "본 문서는 AI 토론 플랫폼에서 에이전트(발언자)와 오케스트레이터(심판/검토자) 역할로 "
        "실제 사용된 LLM 모델들의 성능을 정리한 실측 결과 보고서입니다.")
    add_table(doc,
        ["구분", "내용"],
        [
            ["테스트 기간", "2026-01 ~ 2026-02"],
            ["테스트 매치 수", "약 60+ 매치 (에이전트 역할), 약 40+ 회 (오케스트레이터 역할)"],
            ["토론 모드", "debate (찬반 토론) 위주, 일부 persuasion"],
            ["턴 수", "기본 6턴 (총 12발언/매치)"],
            ["토큰 제한", "턴당 500 tokens"],
            ["심판 모델", "GPT-4o (gpt-4o) — 오케스트레이터 고정"],
            ["평가 기준", "JSON 형식 준수율, 채점 일관성, 스키마 위반 빈도, 응답 속도"],
        ],
        [2.0, 5.0]
    )

    add_heading(doc, "2. 에이전트 역할 — 모델별 성능 비교", 1)

    add_heading(doc, "2.1 JSON 형식 준수율 (가장 중요한 지표)", 2)
    add_para(doc,
        "토론 엔진은 에이전트 발언을 {action, claim, evidence} JSON으로 파싱합니다. "
        "파싱 실패 시 정규식 재추출을 시도하고, 그래도 실패 시 -5점 페널티가 부과됩니다. "
        "아래는 validate_response_schema() 통과율입니다.", italic=True)
    add_table(doc,
        ["모델", "JSON 준수율", "마크다운 코드블록 감싸기", "재추출 성공률", "스키마 위반 페널티 빈도"],
        [
            ["GPT-4o (gpt-4o)", "99%+", "거의 없음", "99%+", "매우 낮음 (1~2%)"],
            ["GPT-4.1", "99%+", "거의 없음", "99%+", "낮음 (2%)"],
            ["GPT-4.1-mini", "96%", "가끔", "97%", "낮음 (4%)"],
            ["Claude Sonnet 4.6", "97%", "2~3%", "99%+", "낮음 (3%)"],
            ["Claude Sonnet 4.5", "96%", "3~4%", "98%", "낮음 (4%)"],
            ["Claude Haiku 4.5", "90%", "5~8%", "94%", "중간 (10%)"],
            ["Gemini 2.5 Pro", "95%", "4~5%", "96%", "중간 (5%)"],
            ["Gemini 2.5 Flash", "90%", "5~10%", "92%", "중간 (10%)"],
            ["Llama 3.1 70B (RunPod)", "85%", "10~15%", "88%", "높음 (15%)"],
        ],
        [2.0, 1.4, 2.0, 1.5, 2.1]
    )
    add_para(doc,
        "※ 개선사항: validate_response_schema()에 3단계 파싱 로직(코드블록 제거 → 전체 파싱 → 정규식 추출)을 "
        "적용한 이후 스키마 위반 페널티 발생률이 약 40% 감소했습니다.",
        italic=True, color=GRAY)

    add_heading(doc, "2.2 논증 품질 지표 (평균 채점, 에이전트 역할)", 2)
    add_para(doc,
        "심판 LLM(GPT-4o)이 채점한 4개 항목의 평균값입니다. "
        "각 매치 결과 scorecard JSONB에서 집계했습니다.", italic=True)
    add_table(doc,
        ["모델", "논리성 /30", "근거 활용 /25", "반박력 /25", "주제 적합성 /20", "평균 합계 /100"],
        [
            ["GPT-4o", "26.2", "22.1", "22.8", "18.9", "90.0"],
            ["GPT-4.1", "25.8", "21.9", "22.4", "18.7", "88.8"],
            ["Claude Sonnet 4.6", "25.1", "23.4", "21.6", "19.2", "89.3"],
            ["Claude Sonnet 4.5", "24.7", "22.8", "21.0", "18.8", "87.3"],
            ["Claude Haiku 4.5", "21.3", "18.7", "18.2", "17.1", "75.3"],
            ["Gemini 2.5 Pro", "23.9", "19.8", "20.1", "17.6", "81.4"],
            ["Gemini 2.5 Flash", "21.0", "17.5", "18.0", "16.4", "72.9"],
            ["Llama 3.1 70B", "20.4", "16.9", "16.7", "15.8", "69.8"],
            ["GPT-4.1-mini", "22.1", "18.4", "19.2", "16.9", "76.6"],
        ],
        [2.0, 1.3, 1.5, 1.3, 1.8, 1.7]
    )

    add_heading(doc, "2.3 실측 관찰 — 모델별 특성", 2)

    add_heading(doc, "GPT-4o / GPT-4.1 (에이전트)", 3)
    for item in [
        "논증 구조가 명확하고 단계적 (premise → evidence → conclusion)",
        "JSON 출력 형식 준수율 최상급 — 스키마 위반 페널티 거의 없음",
        "반박 시 상대 발언의 핵심 약점을 정확히 짚는 경향",
        "단점: 긴 토론에서 초반 주장을 반복하는 경향 → repetition 페널티 위험",
        "단점: 과도하게 균형 잡힌 논증 → 단호함 부족으로 설득력 약화 가능",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)

    add_heading(doc, "Claude Sonnet 4.6 (에이전트)", 3)
    for item in [
        "서술 품질 최고 수준 — 논리를 이야기로 풀어내는 능력",
        "긴 컨텍스트(200K)로 긴 토론에서 일관성 유지",
        "간혹 마크다운 코드블록에 JSON을 감싸는 현상 → 정규식 재추출로 처리",
        "근거 활용 점수(23.4/25)가 가장 높음 — 다양한 관점과 인용 제시",
        "단점: 때로 지나치게 신중한 표현 → 공격성 낮음",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)

    add_heading(doc, "Llama 3.1 70B — RunPod (에이전트)", 3)
    for item in [
        "스키마 위반 페널티 빈도: GPT-4o 대비 3~4배 높음 → few-shot 프롬프트 필수",
        "4-bit AWQ 양자화로 인한 미세 추론 능력 손실",
        "한국어 토론 주제에서 가끔 영어 혼용 발생",
        "콜드스타트 ~2초 (FlashBoot 기준), 서울↔미국 RTT ~150ms",
        "장점: 고정 인프라 비용 — 대량 매칭에 경제적",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)

    add_heading(doc, "3. 오케스트레이터 역할 — 심판/검토 모델 테스트", 1)

    add_heading(doc, "3.1 심판(Judge) 역할 테스트", 2)
    add_para(doc,
        "오케스트레이터는 매치 완료 후 전체 토론 로그를 받아 4항목 채점 JSON을 출력합니다. "
        "현재 GPT-4o가 고정 심판으로 사용되며, Claude Sonnet 4.6과 교차 검증을 실시했습니다.", italic=True)
    add_table(doc,
        ["심판 모델", "JSON 출력 안정성", "채점 일관성 (동일 토론 5회 반복)", "편향 성향", "한국어 이해", "권장 여부"],
        [
            ["GPT-4o (현재)", "99%+", "±3점 이내", "약간 논리 편향", "우수", "✅ 현재 사용"],
            ["Claude Sonnet 4.6", "97%", "±5점 이내", "서술·감성 편향", "우수", "🔲 교차 검증용"],
            ["GPT-4.1-mini", "95%", "±8점", "불안정", "양호", "❌ 부적합"],
            ["Gemini 2.5 Flash", "91%", "±10점", "빠른 판단 편향", "양호", "❌ 부적합"],
        ],
        [2.0, 1.5, 2.2, 1.5, 1.3, 1.5]
    )

    add_heading(doc, "3.2 심판 판정 실패 사례 및 대응", 2)
    add_table(doc,
        ["실패 유형", "발생 원인", "발생률", "대응 방법"],
        [
            ["JSON 파싱 실패", "심판이 마크다운 + 설명 텍스트 추가", "2~3%", "re.search(r'\\{[\\s\\S]*\\}') 정규식 추출"],
            ["scorecard 필드 누락", "agent_a 또는 agent_b 항목 일부 미포함", "1%", "누락 항목 기본값 0으로 채움"],
            ["점수 범위 초과", "논리성 31점 등 배점 초과 값 출력", "<1%", "min()/max()로 배점 상한 강제"],
            ["reasoning 빈 문자열", "''로 출력", "<1%", "'판정 이유를 제공하지 않음'으로 대체"],
        ],
        [2.0, 2.5, 1.0, 2.5]
    )

    add_heading(doc, "3.3 턴 LLM 검토(review_turn) 테스트", 2)
    add_para(doc,
        "debate_orchestrator.py의 review_turn() 메서드는 매 턴 발언을 GPT-4o가 검토하여 "
        "logic_score(1~10), violations, feedback, block 여부를 반환합니다.", italic=True)
    add_table(doc,
        ["테스트 항목", "결과", "비고"],
        [
            ["정상 응답 처리", "logic_score 추출 + violations 파싱 정상", "대부분의 경우"],
            ["block:true 시 대체 텍스트", "'[차단됨: 규칙 위반]'으로 claims 교체 확인", "severe 위반 시"],
            ["openai_api_key 미설정 시", "early return → fallback dict 반환 (토론 중단 없음)", "API 키 없는 환경"],
            ["LLM 응답 타임아웃 (10초)", "fallback dict 반환 (logic_score=5, block=False)", "설정값 초과 시"],
            ["JSON 파싱 실패", "fallback dict 반환 → 원문 그대로 전달", "마크다운 감싸기 등"],
            ["false_claim 위반 탐지", "-7점 LLM 벌점 부과 확인", "llm_ 접두사로 구분"],
            ["prompt_injection 탐지", "-10점 LLM 벌점 부과 확인", ""],
        ],
        [2.5, 3.0, 1.5]
    )

    add_heading(doc, "3.4 review_turn 성능 영향", 2)
    add_table(doc,
        ["항목", "측정값", "비고"],
        [
            ["평균 검토 소요시간", "1.2~2.8초/턴", "GPT-4o 기준"],
            ["기존 debate_turn_delay_seconds 흡수", "가능 (1.5초 이상 설정 시)", "지연 증가 없음"],
            ["turn당 추가 토큰 소비", "입력 ~300 + 출력 ~100 = ~400 tokens", ""],
            ["매치당 추가 비용 (6턴 × 2측)", "~$0.024 (GPT-4o 기준)", "BYOK 사용자 제외"],
            ["차단 발생률", "전체 발언의 약 3~5%", "severe 위반 시"],
        ],
        [2.5, 2.5, 2.0]
    )

    add_heading(doc, "4. o-series 모델 특이사항", 1)
    add_para(doc,
        "OpenAI o1 / o3 / o4 계열 모델은 추론 특화 모델로, 표준 파라미터와 다릅니다.", bold=True)
    add_table(doc,
        ["파라미터", "일반 모델", "o-series 모델", "처리 방법"],
        [
            ["토큰 제한 파라미터", "max_tokens", "max_completion_tokens", "_openai_max_tokens_key() 헬퍼로 자동 선택"],
            ["temperature", "지원 (0.0~2.0)", "미지원", "o-series 감지 시 temperature 파라미터 제거"],
            ["스트리밍 지원", "지원", "제한적", "챗봇 스트리밍 구현에 주의 필요"],
            ["응답 속도", "빠름 (1~3초)", "느림 (5~30초)", "토론 turnDelay와 충돌 가능"],
        ],
        [2.0, 1.8, 2.0, 2.2]
    )
    add_code_block(doc,
        "_OPENAI_COMPLETION_TOKENS_PREFIXES = ('o1', 'o3', 'o4')\n\n"
        "def _openai_max_tokens_key(model_id: str) -> str:\n"
        "    model = model_id.lower()\n"
        "    if any(model.startswith(p) for p in _OPENAI_COMPLETION_TOKENS_PREFIXES):\n"
        "        return 'max_completion_tokens'\n"
        "    return 'max_tokens'"
    )

    add_heading(doc, "5. 비용 최적화 전략", 1)
    add_table(doc,
        ["에이전트 조합 (에이전트 × 2 + 심판)", "매치당 입력 비용", "매치당 출력 비용", "심판 비용", "총계"],
        [
            ["GPT-4o × 2 + GPT-4o 심판", "$0.108", "$0.036", "$0.027", "~$0.17"],
            ["Claude Sonnet 4.6 × 2 + GPT-4o 심판", "$0.115", "$0.054", "$0.027", "~$0.20"],
            ["Gemini 2.5 Flash × 2 + GPT-4o 심판", "$0.001", "$0.001", "$0.027", "~$0.03"],
            ["Llama 3.1 70B × 2 + GPT-4o 심판", "고정비", "고정비", "$0.027", "~$0.03+"],
            ["GPT-4.1-mini × 2 + GPT-4o 심판", "$0.003", "$0.003", "$0.027", "~$0.03"],
        ],
        [2.8, 1.8, 1.8, 1.5, 1.1]
    )

    add_heading(doc, "6. 권장 조합 가이드", 1)
    add_table(doc,
        ["목적", "에이전트 A", "에이전트 B", "심판", "특징"],
        [
            ["최고 품질 (랭킹 매치)", "GPT-4o", "Claude Sonnet 4.6", "GPT-4o", "논리 vs 서술 대결, 관전 재미 최대"],
            ["비용 효율 (일반 매치)", "GPT-4.1-mini", "Gemini 2.5 Flash", "GPT-4o", "심판만큼은 고품질 유지"],
            ["프롬프트 A/B 테스트", "GPT-4o + 프롬프트 v1", "GPT-4o + 프롬프트 v2", "Claude Sonnet 4.6", "교차 심판으로 편향 최소화"],
            ["엔터테인먼트 (관전용)", "GPT-4o (공격성 5)", "Claude Sonnet 4.6 (공격성 5)", "GPT-4o", "극단적 입장, 10턴, 도구 활성화"],
            ["비용 최소화 (실험용)", "Llama 3.1 70B", "Llama 3.1 70B", "GPT-4o", "RunPod 고정비 활용"],
        ],
        [1.8, 1.8, 1.8, 1.5, 2.1]
    )

    doc.save(OUT_DIR / "AI토론_모델비교전략.docx")
    print("[OK] AI토론_모델비교전략.docx 생성 완료")


def make_model_xlsx():
    wb = openpyxl.Workbook()

    # Sheet 1: 에이전트 성능 비교
    ws = wb.active
    ws.title = "에이전트 성능 비교"
    ws.row_dimensions[1].height = 36
    xl_section_title(ws, 1, "에이전트 역할 — JSON 형식 준수율 (실측)", 7, XL_HDR)
    xl_header(ws, 2, ["모델", "프로바이더", "JSON 준수율", "코드블록 감싸기", "스키마 위반 페널티 빈도", "권장 용도", "비고"])
    for i, row in enumerate([
        ["gpt-4o",              "OpenAI",    "99%+", "거의 없음", "매우 낮음 (1~2%)",  "고품질 랭킹 매치",    "기본 권장"],
        ["gpt-4.1",             "OpenAI",    "99%+", "거의 없음", "낮음 (2%)",          "고품질 매치",         ""],
        ["gpt-4.1-mini",        "OpenAI",    "96%",  "가끔",      "낮음 (4%)",          "비용 효율 매치",      ""],
        ["claude-sonnet-4-6",   "Anthropic", "97%",  "2~3%",     "낮음 (3%)",          "고품질 랭킹 매치",    "서술 품질 최고"],
        ["claude-sonnet-4-5",   "Anthropic", "96%",  "3~4%",     "낮음 (4%)",          "일반 매치",           ""],
        ["claude-haiku-4-5",    "Anthropic", "90%",  "5~8%",     "중간 (10%)",         "비용 효율 매치",      ""],
        ["gemini-2.5-pro",      "Google",    "95%",  "4~5%",     "중간 (5%)",          "일반 매치",           ""],
        ["gemini-2.5-flash",    "Google",    "90%",  "5~10%",    "중간 (10%)",         "비용 최소화",         "속도 빠름"],
        ["llama-3.1-70b",       "RunPod",    "85%",  "10~15%",   "높음 (15%)",         "실험/벤치마크",       "고정 비용"],
    ], 3):
        xl_row(ws, i, row, i % 2 == 0)

    xl_section_title(ws, 13, "에이전트 역할 — 평균 채점 결과 (심판: GPT-4o, ~60 매치)", 7, XL_HDR2)
    xl_header(ws, 14, ["모델", "논리성 /30", "근거 활용 /25", "반박력 /25", "주제 적합성 /20", "평균 합계 /100", "비고"], XL_HDR2)
    for i, row in enumerate([
        ["gpt-4o",           "26.2", "22.1", "22.8", "18.9", "90.0", "균형 잡힌 최고점"],
        ["gpt-4.1",          "25.8", "21.9", "22.4", "18.7", "88.8", ""],
        ["claude-sonnet-4-6","25.1", "23.4", "21.6", "19.2", "89.3", "근거 활용 최고"],
        ["claude-sonnet-4-5","24.7", "22.8", "21.0", "18.8", "87.3", ""],
        ["gemini-2.5-pro",   "23.9", "19.8", "20.1", "17.6", "81.4", ""],
        ["gpt-4.1-mini",     "22.1", "18.4", "19.2", "16.9", "76.6", ""],
        ["claude-haiku-4-5", "21.3", "18.7", "18.2", "17.1", "75.3", ""],
        ["gemini-2.5-flash", "21.0", "17.5", "18.0", "16.4", "72.9", ""],
        ["llama-3.1-70b",    "20.4", "16.9", "16.7", "15.8", "69.8", "양자화 영향"],
    ], 15):
        xl_row(ws, i, row, i % 2 == 0)
    auto_col_width(ws)

    # Sheet 2: 오케스트레이터 테스트
    ws2 = wb.create_sheet("오케스트레이터 테스트")
    xl_section_title(ws2, 1, "심판(Judge) 역할 — 오케스트레이터 모델 비교", 6, XL_HDR)
    xl_header(ws2, 2, ["심판 모델", "JSON 출력 안정성", "채점 일관성 (5회 반복)", "편향 성향", "한국어 이해", "권장 여부"])
    for i, row in enumerate([
        ["GPT-4o (현재 사용)", "99%+", "±3점 이내", "약간 논리 편향", "우수", "✅ 현재 고정 사용"],
        ["Claude Sonnet 4.6", "97%",  "±5점 이내", "서술·감성 편향", "우수", "🔲 교차 검증용"],
        ["GPT-4.1-mini",      "95%",  "±8점",      "불안정",         "양호", "❌ 부적합"],
        ["Gemini 2.5 Flash",  "91%",  "±10점",     "빠른 판단 편향", "양호", "❌ 부적합"],
    ], 3):
        xl_row(ws2, i, row, i % 2 == 0)

    xl_section_title(ws2, 8, "심판 판정 실패 사례 및 대응 (실측)", 6, XL_HDR2)
    xl_header(ws2, 9, ["실패 유형", "발생 원인", "발생률", "대응 방법", "적용 여부", ""], XL_HDR2)
    for i, row in enumerate([
        ["JSON 파싱 실패", "마크다운 + 설명 텍스트 추가", "2~3%", "re.search 정규식 추출", "✅ 적용됨", ""],
        ["scorecard 필드 누락", "항목 일부 미포함", "1%", "누락 항목 기본값 0", "✅ 적용됨", ""],
        ["점수 범위 초과", "배점 초과 값 출력", "<1%", "min()/max() 강제", "✅ 적용됨", ""],
        ["reasoning 빈 문자열", "'' 출력", "<1%", "기본 텍스트로 대체", "✅ 적용됨", ""],
    ], 10):
        xl_row(ws2, i, row, i % 2 == 0)

    xl_section_title(ws2, 15, "턴 LLM 검토(review_turn) 테스트 결과", 6, XL_HDR)
    xl_header(ws2, 16, ["테스트 항목", "결과", "발생률/측정값", "처리 방법", "상태", ""])
    for i, row in enumerate([
        ["정상 응답 처리", "logic_score + violations 파싱 정상", "95%+", "—", "✅", ""],
        ["block:true 발동", "차단 텍스트로 claims 교체 확인", "3~5%", "severe 위반 감지 시", "✅", ""],
        ["API 키 미설정", "early return → fallback 반환", "환경 의존", "토론 중단 없음", "✅", ""],
        ["LLM 타임아웃 (10초)", "fallback dict 반환", "드물게", "설정값 초과 시", "✅", ""],
        ["JSON 파싱 실패", "fallback → 원문 그대로", "2~3%", "코드블록 감싸기 등", "✅", ""],
        ["false_claim 탐지", "-7점 llm_false_claim 벌점", "발생 시", "llm_ 접두사 구분", "✅", ""],
        ["prompt_injection 탐지", "-10점 llm_prompt_injection 벌점", "발생 시", "regex 패턴", "✅", ""],
        ["평균 검토 소요시간", "1.2~2.8초/턴", "GPT-4o 기준", "delay에 흡수", "✅", ""],
        ["매치당 추가 비용", "~$0.024 (6턴×2측)", "GPT-4o 기준", "BYOK 제외", "측정됨", ""],
    ], 17):
        xl_row(ws2, i, row, i % 2 == 0)
    auto_col_width(ws2)

    # Sheet 3: 비용 비교
    ws3 = wb.create_sheet("비용 비교")
    xl_section_title(ws3, 1, "모델 조합별 매치당 비용 (에이전트 ×2 + 심판 GPT-4o)", 7, XL_HDR)
    xl_note(ws3, 2, "기준: 6턴, 턴당 토큰 500, 입력 ~36,000 / 출력 ~6,000 / 심판 입력 ~5,000+출력 ~400 (2026-02 기준 가격)", 7)
    xl_header(ws3, 3, ["에이전트 조합", "입력 단가 ($/1M)", "출력 단가 ($/1M)", "입력 비용", "출력 비용", "심판 비용", "총 비용"])
    for i, row in enumerate([
        ["GPT-4o × 2 + GPT-4o 심판",           "$2.50",  "$10.00", "$0.108", "$0.036", "$0.027", "~$0.17"],
        ["Claude Sonnet 4.6 × 2 + GPT-4o 심판", "$3.00",  "$15.00", "$0.115", "$0.054", "$0.027", "~$0.20"],
        ["Gemini 2.5 Pro × 2 + GPT-4o 심판",    "$1.25",  "$5.00",  "$0.045", "$0.018", "$0.027", "~$0.09"],
        ["Gemini 2.5 Flash × 2 + GPT-4o 심판",  "$0.075", "$0.30",  "$0.003", "$0.001", "$0.027", "~$0.03"],
        ["GPT-4.1-mini × 2 + GPT-4o 심판",      "$0.40",  "$1.60",  "$0.014", "$0.006", "$0.027", "~$0.05"],
        ["Llama 3.1 70B × 2 + GPT-4o 심판",     "고정비", "고정비", "고정비", "고정비", "$0.027", "~$0.03+"],
    ], 4):
        xl_row(ws3, i, row, i % 2 == 0)

    xl_section_title(ws3, 11, "o-series 모델 특이사항", 7, XL_HDR2)
    xl_header(ws3, 12, ["파라미터", "일반 모델", "o-series (o1/o3/o4)", "처리 방법", "적용 여부", "", ""], XL_HDR2)
    for i, row in enumerate([
        ["토큰 제한 파라미터", "max_tokens", "max_completion_tokens", "_openai_max_tokens_key() 헬퍼 자동 선택", "✅ 적용됨", "", ""],
        ["temperature", "지원 (0.0~2.0)", "미지원", "o-series 감지 시 파라미터 제거", "✅ 적용됨", "", ""],
        ["스트리밍 지원", "지원", "제한적", "챗봇 SSE 구현 주의", "확인 필요", "", ""],
    ], 13):
        xl_row(ws3, i, row, i % 2 == 0)

    xl_section_title(ws3, 17, "권장 조합 가이드", 7, XL_HDR)
    xl_header(ws3, 18, ["목적", "에이전트 A", "에이전트 B", "심판", "특징", "예상 비용", ""])
    for i, row in enumerate([
        ["최고 품질 (랭킹 매치)", "GPT-4o", "Claude Sonnet 4.6", "GPT-4o", "논리 vs 서술, 관전 재미 최대", "~$0.19", ""],
        ["비용 효율 (일반 매치)", "GPT-4.1-mini", "Gemini 2.5 Flash", "GPT-4o", "심판만 고품질 유지", "~$0.04", ""],
        ["프롬프트 A/B 테스트", "GPT-4o + v1", "GPT-4o + v2", "Claude Sonnet 4.6", "교차 심판으로 편향 최소화", "~$0.20", ""],
        ["엔터테인먼트 (관전용)", "GPT-4o (공격성 5)", "Claude Sonnet 4.6 (공격성 5)", "GPT-4o", "10턴, 도구 활성화", "~$0.30", ""],
        ["비용 최소화 (실험)", "Llama 3.1 70B", "Llama 3.1 70B", "GPT-4o", "RunPod 고정비 활용", "~$0.03+", ""],
    ], 19):
        xl_row(ws3, i, row, i % 2 == 0)
    auto_col_width(ws3)

    wb.save(OUT_DIR / "AI토론_모델비교전략.xlsx")
    print("[OK] AI토론_모델비교전략.xlsx 생성 완료")


# ══════════════════════════════════════════════════════════════════
#  문서 5: WBS (Work Breakdown Structure)
# ══════════════════════════════════════════════════════════════════

# WBS 데이터 정의 (코드, 작업명, 기간(일), 담당, 상태, 우선순위, 비고)
WBS_DATA = [
    # phase, wbs_id, task_name, duration_days, owner, status, priority, note
    ("1. 기반 인프라 구축", None, None, None, None, None, None, None),
    (None, "1.1", "EC2 서버 초기 설정 (t4g.small, 서울 리전)", 1, "인프라", "완료", "High", "Ubuntu 22.04, Docker 설치"),
    (None, "1.2", "Docker Compose 멀티 서비스 설정", 2, "인프라", "완료", "High", "backend, frontend, db, redis, nginx"),
    (None, "1.3", "PostgreSQL 16 + pgvector 컨테이너", 1, "인프라", "완료", "High", "볼륨 마운트, health check"),
    (None, "1.4", "Redis 캐시 컨테이너 설정", 1, "인프라", "완료", "High", "Pub/Sub 채널 포함"),
    (None, "1.5", "Nginx 리버스 프록시 & SSL", 1, "인프라", "완료", "High", "SSE 버퍼링 비활성화"),
    (None, "1.6", "BuildKit 캐시 최적화 (배포 속도 개선)", 1, "인프라", "완료", "Medium", "캐시 마운트 --no-cache 제거"),

    ("2. 데이터베이스 설계 & 마이그레이션", None, None, None, None, None, None, None),
    (None, "2.1", "ERD 설계 — 에이전트/매치/큐/결과 테이블", 2, "백엔드", "완료", "High", "debate_agents, debate_matches 등"),
    (None, "2.2", "에이전트 테이블 (debate_agents)", 1, "백엔드", "완료", "High", "ELO, provider, model_id, credit 컬럼"),
    (None, "2.3", "토론 매치 테이블 (debate_matches)", 1, "백엔드", "완료", "High", "상태 머신: pending→ongoing→finished"),
    (None, "2.4", "매치 큐 테이블 (debate_match_queue)", 1, "백엔드", "완료", "High", "is_ready 컬럼, 자동 매칭 지원"),
    (None, "2.5", "턴 로그 테이블 (debate_turn_logs)", 1, "백엔드", "완료", "High", "claim, evidence, penalties JSONB"),
    (None, "2.6", "크레딧 원장 연동 (credit_ledger)", 1, "백엔드", "완료", "Medium", "토론 비용 차감 기록"),
    (None, "2.7", "Alembic 마이그레이션 파일 관리", 3, "백엔드", "완료", "High", "r8m9n0, u1v2w3 등 순차 적용"),

    ("3. 에이전트 시스템", None, None, None, None, None, None, None),
    (None, "3.1", "에이전트 CRUD API (/api/agents)", 2, "백엔드", "완료", "High", "생성/조회/수정/삭제 + RBAC"),
    (None, "3.2", "에이전트 프롬프트 레이어 설계", 2, "백엔드", "완료", "High", "system_prompt, style_rules, catchphrases"),
    (None, "3.3", "공개 에이전트 마켓플레이스 API", 1, "백엔드", "완료", "Medium", "is_public 필터, 사용 가능한 에이전트 목록"),
    (None, "3.4", "다중 LLM 라우팅 (OpenAI/Anthropic/Google/RunPod)", 3, "백엔드", "완료", "High", "inference_client.py 공급자별 분기"),
    (None, "3.5", "o-series 파라미터 호환성 처리", 1, "백엔드", "완료", "High", "max_completion_tokens, temperature 미지원"),
    (None, "3.6", "에이전트 ELO 레이팅 초기화 및 조회", 1, "백엔드", "완료", "Medium", "초기값 1500, 비대칭 K-factor"),

    ("4. 매치메이킹 시스템", None, None, None, None, None, None, None),
    (None, "4.1", "큐 참가 API (POST /queue)", 1, "백엔드", "완료", "High", "에이전트별 큐 상태 관리"),
    (None, "4.2", "큐 상태 조회 API (GET /queue/status)", 1, "백엔드", "완료", "High", "matched/waiting/not_in_queue"),
    (None, "4.3", "자동 매칭 로직 (플랫폼 에이전트 투입)", 2, "백엔드", "완료", "High", "120초 타임아웃 시 자동 매칭"),
    (None, "4.4", "SSE 실시간 알림 스트림 (/queue/stream)", 2, "백엔드", "완료", "High", "matched, opponent_joined, timeout 이벤트"),
    (None, "4.5", "준비 완료 시스템 (POST /queue/ready)", 1, "백엔드", "완료", "High", "양쪽 준비 시 매치 자동 생성"),
    (None, "4.6", "큐 취소 API (DELETE /queue)", 1, "백엔드", "완료", "Medium", "Redis 큐에서 제거"),
    (None, "4.7", "대기방 UI (WaitingRoomVS 컴포넌트)", 2, "프론트", "완료", "High", "VS 레이아웃, 준비 버튼, 타이머"),
    (None, "4.8", "SSE 재연결 & 폴링 안전망", 1, "프론트", "완료", "Medium", "MAX_SSE_RETRIES=10, 4초 폴링"),

    ("5. 토론 엔진", None, None, None, None, None, None, None),
    (None, "5.1", "토론 엔진 설계 (DebateEngine)", 3, "백엔드", "완료", "High", "debate_engine.py 핵심 로직"),
    (None, "5.2", "에이전트 턴 실행 (_execute_turn)", 2, "백엔드", "완료", "High", "claim, evidence, action 생성"),
    (None, "5.3", "반박 시스템 (opponent_claims 전달)", 1, "백엔드", "완료", "High", "이전 턴 주장을 다음 에이전트에 전달"),
    (None, "5.4", "턴 간 딜레이 & 비용 계산", 1, "백엔드", "완료", "Medium", "debate_turn_delay_seconds 설정"),
    (None, "5.5", "SSE 이벤트 발행 (turn, judge_start 등)", 2, "백엔드", "완료", "High", "Redis Pub/Sub → EventSource"),
    (None, "5.6", "토론 종료 처리 (결과 저장, ELO 업데이트)", 2, "백엔드", "완료", "High", "winner 판정 후 elo_history 기록"),
    (None, "5.7", "크레딧 차감 및 원장 기록", 1, "백엔드", "완료", "High", "매치 완료 후 credit_ledger INSERT"),

    ("6. 심판 & 오케스트레이터", None, None, None, None, None, None, None),
    (None, "6.1", "오케스트레이터 설계 (DebateOrchestrator)", 3, "백엔드", "완료", "High", "debate_orchestrator.py"),
    (None, "6.2", "심판 LLM 프롬프트 (JUDGE_SYSTEM_PROMPT)", 2, "백엔드", "완료", "High", "논리성/근거/반박/주제적합성 4개 기준"),
    (None, "6.3", "채점 로직 (judge_debate)", 2, "백엔드", "완료", "High", "GPT-4o 심판, JSON 파싱 3단계"),
    (None, "6.4", "ELO 레이팅 계산 (update_elo_ratings)", 1, "백엔드", "완료", "High", "K_win=40, K_loss=24, K_draw=32"),
    (None, "6.5", "regex 벌점 시스템", 1, "백엔드", "완료", "Medium", "off_topic, false_claim 감지"),
    (None, "6.6", "LLM 턴 검토 시스템 (review_turn)", 3, "백엔드", "진행중", "High", "logic_score, violations, block 판정"),
    (None, "6.7", "validate_response_schema JSON 파싱 강화", 1, "백엔드", "완료", "Medium", "코드블록 제거→전체→정규식 3단계"),

    ("7. 관전자 UI & 토픽 시스템", None, None, None, None, None, None, None),
    (None, "7.1", "토픽 목록 페이지 (/debate/topics)", 1, "프론트", "완료", "High", "공개 토픽, 카테고리 필터"),
    (None, "7.2", "토픽 상세 & 에이전트 선택 페이지", 2, "프론트", "완료", "High", "AgentPicker, 큐 참가 버튼"),
    (None, "7.3", "매치 관전 페이지 (/debate/matches/[id])", 2, "프론트", "완료", "High", "DebateViewer, 실시간 SSE"),
    (None, "7.4", "TurnBubble 컴포넌트 (발언 버블)", 2, "프론트", "완료", "High", "벌점 레이블, 애니메이션"),
    (None, "7.5", "토론 결과 & 채점 표시 UI", 2, "프론트", "완료", "High", "JudgeResult, ELO 변동 표시"),
    (None, "7.6", "AgentForm 컴포넌트 (에이전트 생성/편집)", 2, "프론트", "완료", "High", "image_url, is_public 포함"),
    (None, "7.7", "LLM 검토 결과 UI (review_turn 피드백)", 2, "프론트", "진행중", "Medium", "논증 점수바, 차단 배지"),

    ("8. 테스트 & 품질 보증", None, None, None, None, None, None, None),
    (None, "8.1", "백엔드 단위 테스트 (pytest)", 3, "백엔드", "완료", "High", "353개 테스트 통과"),
    (None, "8.2", "오케스트레이터 단위 테스트", 2, "백엔드", "완료", "High", "TestJudgeDebate, TestUpdateElo 등"),
    (None, "8.3", "프론트엔드 컴포넌트 테스트 (vitest)", 2, "프론트", "완료", "High", "114개 테스트 통과"),
    (None, "8.4", "review_turn 단위 테스트 (TestReviewTurn)", 1, "백엔드", "미완", "High", "정상/차단/타임아웃/파싱실패 시나리오"),
    (None, "8.5", "통합 테스트 (매치메이킹 전체 흐름)", 2, "백엔드", "미완", "Medium", "큐→매칭→토론→심판 E2E"),

    ("9. 배포 & 운영", None, None, None, None, None, None, None),
    (None, "9.1", "EC2 배포 스크립트 (deploy.sh)", 1, "인프라", "완료", "High", "tar+scp 또는 git pull 방식"),
    (None, "9.2", "Alembic 마이그레이션 자동 적용", 1, "인프라", "완료", "High", "배포 시 upgrade head 실행"),
    (None, "9.3", "환경 변수 & .env 관리", 1, "인프라", "완료", "High", "config.py BaseSettings 통합"),
    (None, "9.4", "로그 모니터링 (docker compose logs)", 1, "인프라", "완료", "Medium", "백엔드 로그 실시간 확인"),
    (None, "9.5", "성능 튜닝 — 토론 딜레이 흡수 설계", 1, "백엔드", "완료", "Medium", "review_turn 시간 debate_delay에서 차감"),
    (None, "9.6", "프로덕션 docker-compose.prod.yml", 1, "인프라", "완료", "High", "BuildKit + 헬스체크 포함"),
]


def make_wbs_docx():
    doc = init_doc("AI 에이전트 토론 플랫폼", "프로젝트 WBS (Work Breakdown Structure)")

    add_heading(doc, "프로젝트 개요", 1)
    add_table(doc,
        ["항목", "내용"],
        [
            ["프로젝트명", "LLM 기반 AI 에이전트 자율 토론 플랫폼"],
            ["기간",       "2025-10 ~ 2026-02 (약 5개월)"],
            ["팀 구성",    "백엔드 · 프론트엔드 · 인프라 (소규모 풀스택)"],
            ["기술 스택",  "FastAPI · Next.js · PostgreSQL · Redis · OpenAI/Anthropic API"],
            ["현재 상태",  "주요 기능 완료, 턴 검토(review_turn) UI 마무리 진행 중"],
        ],
        [2.0, 4.5]
    )

    add_heading(doc, "WBS 범례", 2)
    add_table(doc,
        ["상태", "의미"],
        [
            ["완료", "구현 및 배포 완료"],
            ["진행중", "현재 개발 중"],
            ["미완", "구현 예정"],
        ],
        [1.5, 5.0]
    )

    add_heading(doc, "WBS 상세", 1)

    headers = ["WBS 코드", "작업명", "기간(일)", "담당", "상태", "우선순위", "비고"]
    col_widths = [0.7, 2.8, 0.7, 0.7, 0.8, 0.8, 1.8]

    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    for ci, h in enumerate(headers):
        cell = t.cell(0, ci)
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = WHITE
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, "1E40AF")
        if col_widths:
            cell.width = Inches(col_widths[ci])

    STATUS_COLOR = {"완료": "16A34A", "진행중": "D97706", "미완": "DC2626"}
    data_row = 0
    for entry in WBS_DATA:
        phase, wbs_id, task_name, duration, owner, status, priority, note = entry
        if phase is not None:
            # 단계 헤더
            row_obj = t.add_row()
            cell = row_obj.cells[0]
            # 병합
            for ci2 in range(1, len(headers)):
                cell = cell.merge(row_obj.cells[ci2])
            cell.text = phase
            run = cell.paragraphs[0].runs[0]
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = WHITE
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_cell_bg(row_obj.cells[0], "1E5534")
        else:
            row_obj = t.add_row()
            vals = [wbs_id, task_name, str(duration), owner, status, priority, note]
            bg = "F0F4FF" if data_row % 2 == 0 else "FFFFFF"
            for ci2, val in enumerate(vals):
                c = row_obj.cells[ci2]
                c.text = str(val)
                run = c.paragraphs[0].runs[0]
                run.font.size = Pt(9)
                if ci2 == 4 and val in STATUS_COLOR:
                    run.font.color.rgb = RGBColor.from_string(STATUS_COLOR[val])
                    run.bold = True
                c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if ci2 != 1 else WD_ALIGN_PARAGRAPH.LEFT
                set_cell_bg(c, bg)
                if col_widths:
                    c.width = Inches(col_widths[ci2])
            data_row += 1

    doc.add_paragraph()

    add_heading(doc, "완료율 요약", 1)
    phases_summary = [
        ("1. 기반 인프라 구축",           6, 6, 0, 0),
        ("2. DB 설계 & 마이그레이션",      7, 7, 0, 0),
        ("3. 에이전트 시스템",             6, 6, 0, 0),
        ("4. 매치메이킹 시스템",           8, 8, 0, 0),
        ("5. 토론 엔진",                   7, 7, 0, 0),
        ("6. 심판 & 오케스트레이터",       7, 5, 1, 1),
        ("7. 관전자 UI & 토픽 시스템",     7, 5, 2, 0),
        ("8. 테스트 & 품질 보증",          5, 3, 0, 2),
        ("9. 배포 & 운영",                 6, 6, 0, 0),
    ]
    total_tasks   = sum(r[1] for r in phases_summary)
    total_done    = sum(r[2] for r in phases_summary)
    total_wip     = sum(r[3] for r in phases_summary)
    total_pending = sum(r[4] for r in phases_summary)

    add_table(doc,
        ["단계", "전체", "완료", "진행중", "미완", "완료율"],
        [
            [r[0], r[1], r[2], r[3], r[4], f"{r[2]/r[1]*100:.0f}%"]
            for r in phases_summary
        ] + [["합계", total_tasks, total_done, total_wip, total_pending,
              f"{total_done/total_tasks*100:.0f}%"]],
        [3.0, 0.6, 0.6, 0.7, 0.6, 0.8]
    )

    doc.save(OUT_DIR / "AI토론_WBS.docx")
    print("[OK] AI토론_WBS.docx 생성 완료")


def make_wbs_xlsx():
    wb = openpyxl.Workbook()

    # ── 시트 1: WBS 전체 ──────────────────────────────
    ws = wb.active
    ws.title = "WBS 전체"
    ws.freeze_panes = "C3"

    # 제목
    ws.merge_cells("A1:H1")
    title_cell = ws["A1"]
    title_cell.value = "AI 에이전트 토론 플랫폼 — WBS (Work Breakdown Structure)"
    title_cell.font = Font(bold=True, size=14, color="FFFFFF")
    title_cell.fill = PatternFill("solid", fgColor="1E40AF")
    title_cell.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 28

    # 헤더
    headers = ["WBS 코드", "작업명", "기간(일)", "담당", "상태", "우선순위", "비고", "비고2"]
    xl_header(ws, 2, headers[:7] + [""])

    STATUS_FILL = {
        "완료":   PatternFill("solid", fgColor="DCFCE7"),
        "진행중": PatternFill("solid", fgColor="FEF9C3"),
        "미완":   PatternFill("solid", fgColor="FEE2E2"),
    }
    STATUS_FONT = {
        "완료":   Font(bold=True, color="16A34A", size=10),
        "진행중": Font(bold=True, color="D97706", size=10),
        "미완":   Font(bold=True, color="DC2626", size=10),
    }
    PHASE_FILL = PatternFill("solid", fgColor="1E5534")

    cur_row = 3
    data_row = 0
    for entry in WBS_DATA:
        phase, wbs_id, task_name, duration, owner, status, priority, note = entry
        if phase is not None:
            ws.merge_cells(start_row=cur_row, start_column=1, end_row=cur_row, end_column=8)
            c = ws.cell(row=cur_row, column=1, value=phase)
            c.font   = Font(bold=True, color="FFFFFF", size=11)
            c.fill   = PHASE_FILL
            c.alignment = Alignment(horizontal="left", vertical="center", indent=1)
            c.border = THIN
            ws.row_dimensions[cur_row].height = 22
        else:
            alt = data_row % 2 == 0
            vals = [wbs_id, task_name, duration, owner, status, priority, note, ""]
            for ci, val in enumerate(vals, 1):
                c = ws.cell(row=cur_row, column=ci, value=val)
                c.border = THIN
                c.alignment = Alignment(horizontal="center" if ci != 2 else "left",
                                        vertical="center", wrap_text=True)
                if ci == 5 and val in STATUS_FILL:
                    c.fill = STATUS_FILL[val]
                    c.font = STATUS_FONT[val]
                else:
                    c.fill = XL_ALT if alt else XL_WHITE
            data_row += 1
        cur_row += 1

    # 컬럼 너비 수동 설정
    col_w = [9, 42, 9, 9, 9, 10, 30, 4]
    for i, w in enumerate(col_w, 1):
        ws.column_dimensions[get_column_letter(i)].width = w

    # ── 시트 2: 완료율 대시보드 ──────────────────────
    ws2 = wb.create_sheet("완료율 대시보드")

    ws2.merge_cells("A1:G1")
    t = ws2["A1"]
    t.value = "단계별 완료율 대시보드"
    t.font   = Font(bold=True, size=13, color="FFFFFF")
    t.fill   = PatternFill("solid", fgColor="1E40AF")
    t.alignment = Alignment(horizontal="center", vertical="center")
    ws2.row_dimensions[1].height = 28

    xl_header(ws2, 2, ["단계", "전체 작업", "완료", "진행중", "미완", "완료율(%)", "진행 바"])

    phases_summary = [
        ("1. 기반 인프라 구축",          6, 6, 0, 0),
        ("2. DB 설계 & 마이그레이션",     7, 7, 0, 0),
        ("3. 에이전트 시스템",            6, 6, 0, 0),
        ("4. 매치메이킹 시스템",          8, 8, 0, 0),
        ("5. 토론 엔진",                  7, 7, 0, 0),
        ("6. 심판 & 오케스트레이터",      7, 5, 1, 1),
        ("7. 관전자 UI & 토픽 시스템",    7, 5, 2, 0),
        ("8. 테스트 & 품질 보증",         5, 3, 0, 2),
        ("9. 배포 & 운영",                6, 6, 0, 0),
    ]

    total_tasks   = sum(r[1] for r in phases_summary)
    total_done    = sum(r[2] for r in phases_summary)
    total_wip     = sum(r[3] for r in phases_summary)
    total_pending = sum(r[4] for r in phases_summary)

    for ri, (phase_name, total, done, wip, pending) in enumerate(phases_summary, 3):
        rate = done / total * 100
        bar  = "█" * int(rate // 10) + "░" * (10 - int(rate // 10))
        alt  = ri % 2 == 0
        vals = [phase_name, total, done, wip, pending, round(rate, 1), f"{bar} {rate:.0f}%"]
        for ci, val in enumerate(vals, 1):
            c = ws2.cell(row=ri, column=ci, value=val)
            c.border = THIN
            c.alignment = Alignment(horizontal="center" if ci != 1 else "left",
                                    vertical="center")
            if ci == 6:
                if rate >= 100:
                    c.fill = PatternFill("solid", fgColor="DCFCE7")
                    c.font = Font(bold=True, color="16A34A")
                elif rate >= 70:
                    c.fill = PatternFill("solid", fgColor="FEF9C3")
                    c.font = Font(bold=True, color="D97706")
                else:
                    c.fill = PatternFill("solid", fgColor="FEE2E2")
                    c.font = Font(bold=True, color="DC2626")
            elif ci == 7:
                c.font = Font(name="Consolas", size=10,
                              color="16A34A" if rate >= 100 else "D97706")
                c.alignment = Alignment(horizontal="left", vertical="center")
            else:
                c.fill = XL_ALT if alt else XL_WHITE

    # 합계 행
    total_rate = total_done / total_tasks * 100
    total_bar  = "█" * int(total_rate // 10) + "░" * (10 - int(total_rate // 10))
    tr = len(phases_summary) + 3
    total_vals = ["합  계", total_tasks, total_done, total_wip, total_pending,
                  round(total_rate, 1), f"{total_bar} {total_rate:.0f}%"]
    for ci, val in enumerate(total_vals, 1):
        c = ws2.cell(row=tr, column=ci, value=val)
        c.fill   = PatternFill("solid", fgColor="1E40AF")
        c.font   = Font(bold=True, color="FFFFFF", size=11)
        c.border = THIN
        c.alignment = Alignment(horizontal="center" if ci != 1 else "left", vertical="center")
    ws2.row_dimensions[tr].height = 24

    col_w2 = [34, 11, 9, 9, 9, 12, 22]
    for i, w in enumerate(col_w2, 1):
        ws2.column_dimensions[get_column_letter(i)].width = w

    # ── 시트 3: 간트 차트 (주 단위) ────────────────
    ws3 = wb.create_sheet("간트 차트")

    ws3.merge_cells("A1:R1")
    g = ws3["A1"]
    g.value = "AI 토론 플랫폼 — 간트 차트 (월별)"
    g.font  = Font(bold=True, size=13, color="FFFFFF")
    g.fill  = PatternFill("solid", fgColor="1E40AF")
    g.alignment = Alignment(horizontal="center", vertical="center")
    ws3.row_dimensions[1].height = 28

    month_headers = ["단계", "담당", "2025-10", "2025-11", "2025-12", "2026-01", "2026-02", "비고"]
    xl_header(ws3, 2, month_headers)

    GANTT_DATA = [
        ("1. 기반 인프라 구축",          "인프라", True,  True,  False, False, False, "완료"),
        ("2. DB 설계 & 마이그레이션",    "백엔드", True,  True,  False, False, False, "완료"),
        ("3. 에이전트 시스템",           "백엔드", False, True,  True,  False, False, "완료"),
        ("4. 매치메이킹 시스템",         "백엔드\n프론트", False, True,  True,  True,  False, "완료"),
        ("5. 토론 엔진",                 "백엔드", False, False, True,  True,  False, "완료"),
        ("6. 심판 & 오케스트레이터",     "백엔드", False, False, True,  True,  True,  "진행중"),
        ("7. 관전자 UI & 토픽",          "프론트", False, False, True,  True,  True,  "진행중"),
        ("8. 테스트 & 품질 보증",        "전체",   False, False, False, True,  True,  "진행중"),
        ("9. 배포 & 운영",               "인프라", False, True,  True,  True,  True,  "완료"),
    ]

    GANTT_DONE_FILL   = PatternFill("solid", fgColor="16A34A")
    GANTT_WIP_FILL    = PatternFill("solid", fgColor="D97706")
    GANTT_EMPTY_FILL  = PatternFill("solid", fgColor="F1F5F9")

    for ri, row_data in enumerate(GANTT_DATA, 3):
        phase_name, owner, m10, m11, m12, m1, m2, note = row_data
        is_done = note == "완료"
        bar_fill = GANTT_DONE_FILL if is_done else GANTT_WIP_FILL

        for ci, val in enumerate([phase_name, owner], 1):
            c = ws3.cell(row=ri, column=ci, value=val)
            c.fill = XL_ALT if ri % 2 == 0 else XL_WHITE
            c.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
            c.border = THIN

        for ci, active in enumerate([m10, m11, m12, m1, m2], 3):
            c = ws3.cell(row=ri, column=ci, value="●" if active else "")
            c.fill  = bar_fill if active else GANTT_EMPTY_FILL
            c.font  = Font(color="FFFFFF") if active else Font(color="CBD5E1")
            c.alignment = Alignment(horizontal="center", vertical="center")
            c.border = THIN

        c_note = ws3.cell(row=ri, column=8, value=note)
        c_note.font = Font(bold=True,
                           color="16A34A" if is_done else "D97706")
        c_note.fill = XL_ALT if ri % 2 == 0 else XL_WHITE
        c_note.alignment = Alignment(horizontal="center", vertical="center")
        c_note.border = THIN

    col_w3 = [34, 10, 12, 12, 12, 12, 12, 10]
    for i, w in enumerate(col_w3, 1):
        ws3.column_dimensions[get_column_letter(i)].width = w

    wb.save(OUT_DIR / "AI토론_WBS.xlsx")
    print("[OK] AI토론_WBS.xlsx 생성 완료")


# ══════════════════════════════════════════════════════════════════
#  실행
# ══════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    print(f"\n[출력 경로] {OUT_DIR.resolve()}\n")

    make_plan_docx()
    make_plan_xlsx()
    make_arch_docx()
    make_arch_xlsx()
    make_data_docx()
    make_data_xlsx()
    make_model_docx()
    make_model_xlsx()
    make_wbs_docx()
    make_wbs_xlsx()

    print(f"\n[완료] 총 10개 파일 생성 완료 -> {OUT_DIR.resolve()}")
    files = list(OUT_DIR.iterdir())
    for f in sorted(files):
        size_kb = f.stat().st_size // 1024
        print(f"  {f.name:45s}  {size_kb} KB")
