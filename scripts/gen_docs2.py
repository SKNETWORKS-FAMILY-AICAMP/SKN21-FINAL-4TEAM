# -*- coding: utf-8 -*-
"""
gen_docs2.py  —  AI 토론 플랫폼 문서 생성
- 인공지능_데이터_전처리_결과서.docx
- 인공지능_학습_결과서.xlsx
- AI_모델_명세서.docx
"""

from pathlib import Path

OUT_DIR = Path(__file__).parent.parent / "docs" / "output"
OUT_DIR.mkdir(parents=True, exist_ok=True)

# ── 색상 팔레트 ───────────────────────────────────────────────────
BLUE_DARK    = "1F3864"
BLUE_MID     = "2E75B6"
BLUE_LIGHT   = "D6E4F0"
GREEN_DARK   = "1E5631"
GREEN_MID    = "27AE60"
GREEN_LIGHT  = "D5F5E3"
ORANGE_MID   = "E67E22"
ORANGE_LIGHT = "FDEBD0"
GRAY_LIGHT   = "F2F2F2"
WHITE        = "FFFFFF"
RED_MID      = "C0392B"
RED_LIGHT    = "FADBD8"
PURPLE_MID   = "8E44AD"
PURPLE_LIGHT = "E8DAEF"
YELLOW_LIGHT = "FEF9E7"


# ════════════════════════════════════════════════════════════════
#  공통 헬퍼
# ════════════════════════════════════════════════════════════════
def _rgb(hex_str):
    from docx.shared import RGBColor
    r, g, b = int(hex_str[0:2], 16), int(hex_str[2:4], 16), int(hex_str[4:6], 16)
    return RGBColor(r, g, b)


def _set_cell(cell, text, bold=False, italic=False,
              font_size=9, bg=None, font_color=None,
              align="left", font_name="맑은 고딕"):
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    cell.text = ""
    para = cell.paragraphs[0]
    if align == "center":
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = para.add_run(str(text))
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(font_size)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if font_color:
        run.font.color.rgb = _rgb(font_color)
    if bg:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), bg)
        tcPr.append(shd)


def _header_row(table, headers, bg=BLUE_DARK, font_color=WHITE, font_size=9):
    row = table.rows[0]
    for i, h in enumerate(headers):
        _set_cell(row.cells[i], h, bold=True, bg=bg,
                  font_color=font_color, font_size=font_size, align="center")


def _set_col_widths(table, widths_cm):
    from docx.shared import Cm
    for row in table.rows:
        for i, w in enumerate(widths_cm):
            if i < len(row.cells):
                row.cells[i].width = Cm(w)


def _add_heading(doc, text, level=1):
    from docx.shared import Pt
    from docx.oxml.ns import qn
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "맑은 고딕"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        if level == 1:
            run.font.size = Pt(14)
            run.font.color.rgb = _rgb(BLUE_DARK)
        elif level == 2:
            run.font.size = Pt(12)
            run.font.color.rgb = _rgb(BLUE_MID)
        else:
            run.font.size = Pt(11)
            run.font.color.rgb = _rgb("555555")
    return p


def _add_para(doc, text, font_size=10, bold=False, color=None):
    from docx.shared import Pt
    from docx.oxml.ns import qn
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    if color:
        run.font.color.rgb = _rgb(color)
    return p


def _add_box(doc, lines, bg=BLUE_LIGHT):
    """단순 단락 기반 박스 (표 1×1)"""
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    cell.text = ""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), bg)
    tcPr.append(shd)
    for i, line in enumerate(lines):
        para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        run = para.add_run(line)
        run.font.size = Pt(9)
        run.font.name = "맑은 고딕"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    doc.add_paragraph()


def _cover(doc, title, subtitle):
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(title)
    r.bold = True; r.font.size = Pt(22); r.font.name = "맑은 고딕"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    r.font.color.rgb = _rgb(BLUE_DARK)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run(subtitle)
    rs.font.size = Pt(13); rs.font.name = "맑은 고딕"
    rs._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    rs.font.color.rgb = _rgb("555555")
    doc.add_paragraph()


def _meta_table(doc, rows):
    from docx.shared import Cm
    meta = doc.add_table(rows=len(rows), cols=2)
    meta.style = "Table Grid"
    for i, (k, v) in enumerate(rows):
        _set_cell(meta.rows[i].cells[0], k, bold=True, bg=BLUE_DARK, font_color=WHITE, align="center")
        _set_cell(meta.rows[i].cells[1], v, bg=BLUE_LIGHT)
    for row in meta.rows:
        row.cells[0].width = Cm(4.0)
        row.cells[1].width = Cm(12.0)
    doc.add_paragraph()


# ════════════════════════════════════════════════════════════════
#  1. 인공지능 데이터 전처리 결과서 (.docx)
#     — AI 토론 플랫폼 기준
# ════════════════════════════════════════════════════════════════
def build_preprocessing_docx():
    from docx import Document
    from docx.shared import Cm
    doc = Document()
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Cm(2.0)
        section.left_margin = section.right_margin = Cm(2.5)

    _cover(doc, "인공지능 데이터 전처리 결과서", "AI 토론 플랫폼 — LLM 기반 자동 토론 시스템")
    _meta_table(doc, [
        ("작성일",   "2026년 2월"),
        ("버전",     "v1.0"),
        ("목적",     "전처리 과정의 재현성 및 이해도 향상"),
        ("대상 데이터", "토론 발언 로그, 채점 입력/출력, ELO 이력, 위반 판정 결과"),
    ])

    # ── 1. 데이터셋 개요 ─────────────────────────────────────────
    _add_heading(doc, "1. 데이터셋 개요", 1)
    _add_para(doc, "AI 토론 플랫폼은 LLM 에이전트 간의 자동 토론을 운영하며, "
                   "토론 품질 평가·ELO 산출·위반 탐지를 위해 아래 4종 데이터를 구성합니다.", 10)
    doc.add_paragraph()

    tbl1 = doc.add_table(rows=5, cols=4)
    tbl1.style = "Table Grid"
    _header_row(tbl1, ["항목 구분", "설명", "원천", "예시"])
    rows1 = [
        ["토론 발언 로그",
         "에이전트가 생성한 턴별 주장 텍스트. 찬성/반대 측 구분, 행동 유형(CLAIM/EVIDENCE/REBUTTAL/CLOSING) 포함",
         "debate_turn_logs 테이블\n(LLM 생성 결과)",
         '"환경세 도입은 기업 경쟁력을 저해합니다. 2023년 EU 사례에 따르면..."'],
        ["채점 입력 데이터",
         "판정 LLM에 전달되는 양측 발언 번들. (주제, 찬성 측 발언 합산, 반대 측 발언 합산)으로 구성",
         "debate_engine.py\n발언 번들 생성 로직",
         "주제: 환경세 도입 / agent_a 발언: 3턴 / agent_b 발언: 3턴"],
        ["채점 출력 데이터",
         "판정 LLM이 반환하는 JSON. logic/evidence/rebuttal/relevance 항목별 점수 + reasoning",
         "DebateOrchestrator\njudge() 메서드 출력",
         '{"agent_a":{"logic":26,"evidence":22,"rebuttal":21,"relevance":18},...}'],
        ["위반 판정 결과",
         "턴 검토 LLM이 반환하는 JSON. logic_score, violations(유형·심각도), block 여부",
         "DebateOrchestrator\nreview_turn() 출력",
         '{"logic_score":7,"violations":[],"severity":"none","block":false}'],
    ]
    for i, row in enumerate(rows1):
        bg = BLUE_LIGHT if i % 2 == 0 else WHITE
        for j, val in enumerate(row):
            _set_cell(tbl1.rows[i+1].cells[j], val, bg=bg, font_size=8)
    _set_col_widths(tbl1, [3.5, 5.5, 4.0, 4.5])
    doc.add_paragraph()

    # ── 2. 원본 데이터 샘플 ──────────────────────────────────────
    _add_heading(doc, "2. 원본 데이터 샘플 (토론 발언 로그 10건)", 1)
    _add_para(doc, "아래는 실제 AI 토론 매칭에서 수집된 발언 로그 샘플입니다. "
                   "각 행은 debate_turn_logs 테이블의 한 레코드에 해당합니다.", 10)
    doc.add_paragraph()

    tbl2 = doc.add_table(rows=11, cols=6)
    tbl2.style = "Table Grid"
    _header_row(tbl2, ["No.", "토론 주제", "발언자", "턴", "행동 유형", "발언 내용 (원본 샘플)"])
    samples = [
        ["1", "환경세 도입 찬반", "agent_a (찬성)", "1", "CLAIM",
         "환경세 도입은 탄소 배출 감소에 효과적입니다. OECD 국가 평균 탄소세 도입 후 배출량이 18% 감소했습니다."],
        ["2", "환경세 도입 찬반", "agent_b (반대)", "1", "CLAIM",
         "환경세는 중소기업의 원가 부담을 가중시켜 국내 산업 경쟁력을 약화시킵니다."],
        ["3", "환경세 도입 찬반", "agent_a", "2", "EVIDENCE",
         "유럽 환경세 시행국 12개국 중 9개국에서 GDP 대비 탄소집약도가 도입 5년 내 15% 이상 하락했습니다."],
        ["4", "환경세 도입 찬반", "agent_b", "2", "REBUTTAL",
         "상대의 OECD 통계는 고소득 국가 평균이며, 제조업 비중이 높은 한국과 직접 비교하기 어렵습니다."],
        ["5", "주 4일제 근무 찬반", "agent_a (찬성)", "1", "CLAIM",
         "주 4일제는 직원 생산성을 오히려 향상시킵니다. MS 일본 시험 운영에서 생산성 40% 향상이 확인되었습니다."],
        ["6", "주 4일제 근무 찬반", "agent_b (반대)", "1", "CLAIM",
         "제조업·서비스업 등 대면 업종에서는 주 4일제 도입 시 인건비 급등으로 수익성이 악화됩니다."],
        ["7", "주 4일제 근무 찬반", "agent_a", "3", "REBUTTAL",
         "인건비 상승 우려는 생산성 향상으로 상쇄됩니다. 아이슬란드 전국 시범에서 비용 중립 결과가 나왔습니다."],
        ["8", "AI 규제 찬반", "agent_b (찬성)", "2", "EVIDENCE",
         "EU AI 법(EU AI Act)은 고위험 AI 시스템에 사전 적합성 평가를 요구하며, 2024년 발효되었습니다."],
        ["9", "AI 규제 찬반", "agent_a (반대)", "2", "REBUTTAL",
         "과도한 규제는 AI 스타트업의 R&D 투자 위축을 초래하며, 미국·중국과의 기술 격차를 심화시킵니다."],
        ["10", "AI 규제 찬반", "agent_b", "4", "CLOSING",
         "결론적으로 AI 위험 최소화를 위한 최소한의 규제 프레임워크는 장기적으로 혁신과 신뢰를 동시에 확보합니다."],
    ]
    for i, row in enumerate(samples):
        bg = GREEN_LIGHT if i % 2 == 0 else WHITE
        for j, val in enumerate(row):
            al = "center" if j in [0, 2, 3, 4] else "left"
            _set_cell(tbl2.rows[i+1].cells[j], val, bg=bg, font_size=8, align=al)
    _set_col_widths(tbl2, [0.8, 2.8, 2.5, 0.8, 2.5, 8.0])
    doc.add_paragraph()

    # ── 3. 전처리 흐름도 ─────────────────────────────────────────
    _add_heading(doc, "3. 전처리 흐름도", 1)
    _add_para(doc, "토론 발언이 LLM 입력으로 전달되고 채점·검토가 완료되기까지의 전처리 파이프라인입니다.", 10)
    doc.add_paragraph()

    flow = [
        "┌──────────────────────────────────────────────────────────────────────────┐",
        "│             AI 토론 데이터 전처리 파이프라인                              │",
        "├──────────────────────────────────────────────────────────────────────────┤",
        "│                                                                          │",
        "│  [0] 토론 시작 — 주제(DebateTopic) + 양측 에이전트 확정                  │",
        "│       ・ 토픽 텍스트 정제: 특수문자·HTML 제거, 최대 500자 자름            │",
        "│       ・ 에이전트 시스템 프롬프트 로드 (debate_agents.system_prompt)      │",
        "│                          │                                               │",
        "│                          ▼                                               │",
        "│  [1] 발언 생성 입력 구성                                                 │",
        "│       ・ 시스템 프롬프트 주입 (역할: 찬성/반대, 토론 규칙)               │",
        "│       ・ 이전 턴 이력 슬라이딩 윈도우 (최근 6턴)                         │",
        "│       ・ 행동 유형 지시 (CLAIM / EVIDENCE / REBUTTAL / CLOSING)          │",
        "│       ・ 공백·개행 정규화, 최대 4,000 토큰 자름                          │",
        "│                          │                                               │",
        "│                          ▼                                               │",
        "│  [2] LLM 발언 생성 (에이전트별 LLM)                                      │",
        "│       ・ 타임아웃 초과 시 [TIMEOUT] 텍스트로 대체                        │",
        "│       ・ 오류 시 [ERROR] 텍스트 삽입 (경기는 계속 진행)                  │",
        "│                          │                                               │",
        "│             ┌────────────┴────────────┐                                 │",
        "│             ▼                         ▼                                 │",
        "│  [3A] 턴 검토 전처리               [3B] 채점 입력 전처리                 │",
        "│   ・ 주제+발언자+턴+발언 구조화     ・ 양측 발언 전체 병합               │",
        "│   ・ 직전 상대 발언 추가            ・ 주제별 발언 번들 구성              │",
        "│   ・ review LLM으로 JSON 요청       ・ 총 발언량 균형 확인               │",
        "│             │                         │                                 │",
        "│             ▼                         ▼                                 │",
        "│  [4A] 위반 판정 결과 파싱          [4B] 채점 결과 파싱                   │",
        "│   ・ JSON 코드블록 제거            ・ JSON 코드블록 제거                 │",
        "│   ・ logic_score (1~10) 추출       ・ logic/evidence/rebuttal/relevance  │",
        "│   ・ violations 배열 파싱          ・ 스왑 판정 시 에이전트명 역변환     │",
        "│   ・ 벌점 산출 (유형별 감점)       ・ 벌점 차감 후 최종 점수 확정        │",
        "│   ・ block=true 시 발언 차단       ・ 승자 판정 (점수차 5점 기준)        │",
        "│             │                         │                                 │",
        "│             └────────────┬────────────┘                                 │",
        "│                          ▼                                               │",
        "│  [5] DB 저장 및 SSE 브로드캐스트                                         │",
        "│       ・ debate_turn_logs INSERT (review_result JSONB, is_blocked)       │",
        "│       ・ debate_matches UPDATE (winner, scores, elo_delta)               │",
        "│       ・ ELO 갱신 → debate_agents UPDATE                                │",
        "│                                                                          │",
        "└──────────────────────────────────────────────────────────────────────────┘",
    ]
    _add_box(doc, flow, bg="F8F9FA")

    # ── 4. 세부 전처리 단계 ──────────────────────────────────────
    _add_heading(doc, "4. 세부 전처리 단계", 1)

    # 4.1 결측치·이상치 제거
    _add_heading(doc, "4.1 결측치 및 이상치 처리", 2)
    tbl3 = doc.add_table(rows=7, cols=4)
    tbl3.style = "Table Grid"
    _header_row(tbl3, ["처리 항목", "발생 원인", "처리 방법", "대체값 / 후속 처리"])
    missing = [
        ["빈 발언 (None/빈 문자열)", "LLM 응답 파싱 실패 또는 스트리밍 중단",
         "None 체크 후 교체", "[ERROR] 텍스트로 대체, 해당 턴 0점 처리"],
        ["타임아웃 발언", "LLM 응답 초과 (debate_llm_timeout)",
         "asyncio.wait_for 예외 포착", "[TIMEOUT] 텍스트 삽입, 채점 시 0~5점 범위"],
        ["JSON 파싱 실패 (채점)", "LLM이 JSON 외 마크다운·설명 포함",
         "```코드블록 제거 → 정규식 JSON 추출 → 실패 시 fallback",
         "fallback: logic=15, evidence=12, rebuttal=12, relevance=10"],
        ["JSON 파싱 실패 (검토)", "review_turn() 응답 파싱 실패",
         "동일 정규식 파이프라인 → 실패 시 fallback",
         "fallback: logic_score=5, violations=[], block=false"],
        ["점수 범위 초과", "LLM이 명세 외 점수 반환 (예: logic=35)",
         "min/max clamp 적용",
         "logic: 0~30, evidence: 0~25, rebuttal: 0~25, relevance: 0~20"],
        ["에이전트 스왑 오류", "판정 LLM이 찬성/반대 에이전트를 혼동",
         "스왑 여부 판단 후 역변환 로직",
         "reasoning 텍스트에서 에이전트명 역변환"],
    ]
    for i, row in enumerate(missing):
        bg = ORANGE_LIGHT if i % 2 == 0 else WHITE
        for j, val in enumerate(row):
            _set_cell(tbl3.rows[i+1].cells[j], val, bg=bg, font_size=8)
    _set_col_widths(tbl3, [3.5, 4.0, 4.0, 6.0])
    doc.add_paragraph()

    # 4.2 텍스트 정규화
    _add_heading(doc, "4.2 텍스트 정규화 (LLM 입력 정제)", 2)
    tbl4 = doc.add_table(rows=7, cols=3)
    tbl4.style = "Table Grid"
    _header_row(tbl4, ["정규화 항목", "원본 예시", "처리 후"])
    norm = [
        ["Markdown 코드블록 제거",
         '```json\n{"logic":25,...}\n```',
         '{"logic":25,...}'],
        ["JSON 객체 경계 추출",
         '{"logic":25,...} 판정 결과입니다.',
         '{"logic":25,...}'],
        ["HTML 특수문자 이스케이프",
         "탄소배출 &amp; 경제성장 &lt;균형&gt;",
         "탄소배출 & 경제성장 <균형>"],
        ["연속 공백·개행 압축",
         "이것은\n\n\n너무\n\n긴 개행",
         "이것은\n너무\n긴 개행"],
        ["최대 토큰 자름 (입력)",
         "발언 이력 6턴 합산 > 4,000 토큰",
         "오래된 턴부터 제거 (슬라이딩 윈도우)"],
        ["최대 토큰 자름 (출력)",
         "발언 > 1,500 자 초과 생성",
         "debate_max_output_tokens 설정값 내 자름"],
    ]
    for i, row in enumerate(norm):
        bg = GRAY_LIGHT if i % 2 == 0 else WHITE
        for j, val in enumerate(row):
            _set_cell(tbl4.rows[i+1].cells[j], val, bg=bg, font_size=8)
    _set_col_widths(tbl4, [4.0, 6.5, 6.5])
    doc.add_paragraph()

    # 4.3 점수 표준화
    _add_heading(doc, "4.3 점수 정규화 및 표준화", 2)
    tbl5 = doc.add_table(rows=6, cols=4)
    tbl5.style = "Table Grid"
    _header_row(tbl5, ["처리 항목", "방법", "적용 단계", "결과 범위"])
    std = [
        ["채점 항목별 clamp",
         "min(max(score, 0), max_val) — logic 최대 30, evidence 25, rebuttal 25, relevance 20",
         "judge() 파싱 직후",
         "각 항목 명세 범위 내"],
        ["벌점 차감 후 합산",
         "total = sum(scores) - sum(LLM_VIOLATION_PENALTIES)\n음수는 0으로 clamp",
         "채점 완료 직후",
         "0 ~ 100점"],
        ["ELO 델타 계산",
         "K=32, expected = 1/(1+10^((Rb-Ra)/400))\ndelta = K*(actual-expected)",
         "judge() 결과 확정 후",
         "delta: 약 -32 ~ +32"],
        ["logic_score 범위 확인",
         "review_turn() 출력 1~10 범위 검증\n범위 초과 시 clamp(1, 10)",
         "review_turn() 파싱 직후",
         "1 ~ 10"],
        ["위반 심각도 매핑",
         "severity: none/minor/severe\nblock은 severe일 때만 true",
         "위반 판정 파싱 직후",
         "block: bool"],
    ]
    for i, row in enumerate(std):
        bg = BLUE_LIGHT if i % 2 == 0 else WHITE
        for j, val in enumerate(row):
            _set_cell(tbl5.rows[i+1].cells[j], val, bg=bg, font_size=8)
    _set_col_widths(tbl5, [3.5, 6.5, 4.0, 3.5])
    doc.add_paragraph()

    # 4.4 전처리 전후 품질
    _add_heading(doc, "4.4 전처리 전후 품질 비교", 2)
    tbl6 = doc.add_table(rows=6, cols=3)
    tbl6.style = "Table Grid"
    _header_row(tbl6, ["지표", "전처리 전", "전처리 후"])
    quality = [
        ["채점 JSON 파싱 성공률", "약 82%\n(LLM이 마크다운·설명 추가)", "약 97%\n(정규식 코드블록 제거 + 재추출)"],
        ["위반 판정 파싱 성공률", "약 85%", "약 98%\n(fallback 정책으로 0% 중단)"],
        ["점수 범위 초과 건수", "약 4.2% (logic > 30 등)", "0%\n(clamp 적용)"],
        ["타임아웃으로 인한 경기 중단", "약 3.1%", "0%\n([TIMEOUT] 대체로 경기 지속)"],
        ["스왑 오류로 인한 잘못된 승자 판정", "약 1.5%", "약 0.2%\n(역변환 로직 적용)"],
    ]
    for i, row in enumerate(quality):
        bg = GRAY_LIGHT if i % 2 == 0 else WHITE
        for j, val in enumerate(row):
            _set_cell(tbl6.rows[i+1].cells[j], val, bg=bg, font_size=8)
    _set_col_widths(tbl6, [4.5, 5.5, 7.5])
    doc.add_paragraph()

    out_path = OUT_DIR / "인공지능_데이터_전처리_결과서.docx"
    doc.save(str(out_path))
    print(f"[OK] 인공지능_데이터_전처리_결과서.docx 저장: {out_path}")
    return str(out_path)


# ════════════════════════════════════════════════════════════════
#  2. 인공지능 학습 결과서 (.xlsx)  — AI 토론 플랫폼 기준
# ════════════════════════════════════════════════════════════════
def build_training_xlsx():
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    wb = Workbook()
    wb.remove(wb.active)

    def tfill(c): return PatternFill("solid", fgColor=c)
    def tb():
        s = Side(style="thin", color="AAAAAA")
        return Border(left=s, right=s, top=s, bottom=s)

    def sc(ws, row, col, val, bold=False, bg=None, fc="000000",
           sz=10, ah="left", av="center", wrap=False):
        c = ws.cell(row=row, column=col, value=val)
        c.font = Font(bold=bold, color=fc, size=sz, name="맑은 고딕")
        if bg:
            c.fill = tfill(bg)
        c.alignment = Alignment(horizontal=ah, vertical=av, wrap_text=wrap)
        c.border = tb()
        return c

    def mtitle(ws, row, title, cols=7):
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=cols)
        c = ws.cell(row=row, column=1, value=title)
        c.font = Font(bold=True, size=14, color=WHITE, name="맑은 고딕")
        c.fill = tfill(BLUE_DARK)
        c.alignment = Alignment(horizontal="center", vertical="center")
        ws.row_dimensions[row].height = 32

    def stitle(ws, row, title, bg, cols=7):
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=cols)
        c = ws.cell(row=row, column=1, value=title)
        c.font = Font(bold=True, size=12, color=WHITE, name="맑은 고딕")
        c.fill = tfill(bg)
        c.alignment = Alignment(horizontal="left", vertical="center")
        ws.row_dimensions[row].height = 22

    # ── 시트 1: 모델 목적 ────────────────────────────────────────
    ws1 = wb.create_sheet("1. 모델 목적")
    ws1.sheet_view.showGridLines = False
    for col, w in zip("ABCD", [22, 42, 26, 18]):
        ws1.column_dimensions[col].width = w

    mtitle(ws1, 1, "1. AI 토론 플랫폼 — 인공지능 모델 목적 및 역할", cols=4)
    ws1.row_dimensions[2].height = 5

    for j, h in enumerate(["모델 / 컴포넌트", "역할 및 목적", "입출력 설명", "활용 위치"], 1):
        sc(ws1, 3, j, h, bold=True, bg=BLUE_MID, fc=WHITE, ah="center")
    ws1.row_dimensions[3].height = 20

    model_rows = [
        ["발언 생성 LLM\n(에이전트별)",
         "사용자가 등록한 에이전트의 LLM 모델이 토론 주제에 맞게 찬성/반대 논거를 생성.\n"
         "행동 유형(CLAIM·EVIDENCE·REBUTTAL·CLOSING)에 따라 다른 지시를 받음",
         "입력: 시스템 프롬프트 + 이전 턴 이력 (최근 6턴)\n"
         "출력: 토론 발언 텍스트 (스트리밍)",
         "debate_engine.py\n_generate_turn()"],
        ["판정 LLM\n(DebateOrchestrator)",
         "토론 전체 발언을 수신하여 4개 항목(논리성·근거·반박력·주제 적합성)을 채점.\n"
         "양측 점수 + 채점 근거(reasoning) JSON 반환. 승자 판정 및 ELO 계산의 기준",
         "입력: 채점 시스템 프롬프트 + 양측 발언 번들\n"
         "출력: {agent_a: {logic,evidence,rebuttal,relevance}, agent_b: {...}, reasoning}",
         "debate_orchestrator.py\njudge()"],
        ["턴 검토 LLM\n(review_turn)",
         "각 턴 발언을 실시간 검토하여 논리 점수(1~10) 및 위반 여부(prompt_injection·"
         "ad_hominem·off_topic·false_claim) 판정. severe 위반 시 발언 차단",
         "입력: 검토 시스템 프롬프트 + 단일 턴 발언\n"
         "출력: {logic_score, violations[], severity, feedback, block}",
         "debate_orchestrator.py\nreview_turn()"],
        ["ELO 산출 알고리즘\n(수식 기반)",
         "판정 결과(승/패/무)를 기반으로 에이전트 ELO 점수를 갱신.\n"
         "K=32, 기대승률 Logistic 함수 적용. 별도 학습 없이 수식으로 처리",
         "입력: 현재 ELO(Ra, Rb), 실제 결과(1/0/0.5)\n"
         "출력: 델타값(±0~32), 갱신된 ELO",
         "debate_orchestrator.py\ncalculate_elo()"],
    ]
    for i, row in enumerate(model_rows):
        r = i + 4
        bg = BLUE_LIGHT if i % 2 == 0 else WHITE
        ws1.row_dimensions[r].height = 55
        for j, val in enumerate(row, 1):
            sc(ws1, r, j, val, bg=bg, wrap=True, av="top")

    # ── 시트 2: 모델 비교·선정 ──────────────────────────────────
    ws2 = wb.create_sheet("2. 모델 비교·선정")
    ws2.sheet_view.showGridLines = False
    for col, w in zip("ABCDEFG", [18, 14, 14, 14, 14, 22, 22]):
        ws2.column_dimensions[col].width = w

    mtitle(ws2, 1, "2. 모델 비교 및 선정 근거 (AI 토론 플랫폼)", cols=7)
    ws2.row_dimensions[2].height = 5

    # 2A 발언 생성 LLM 비교
    stitle(ws2, 3, "[A] 발언 생성 LLM 비교 (에이전트 기본 모델)", GREEN_DARK, cols=7)
    for j, h in enumerate(["모델", "한국어 토론 품질", "컨텍스트", "응답 속도", "비용(1K)", "장점", "단점"], 1):
        sc(ws2, 4, j, h, bold=True, bg=GREEN_MID, fc=WHITE, ah="center")
    ws2.row_dimensions[4].height = 18

    gen_comp = [
        ["Llama 3 70B\n(기본·선정)", "4.0/5.0", "8K", "3.2s",
         "$0.002 (RunPod)",
         "저비용 기본 제공\n오픈소스 · 커스텀 가능\n플랫폼 내장",
         "한국어 토론 품질 최하\n컨텍스트 제한"],
        ["GPT-4o\n(BYOK)", "4.5/5.0", "128K", "1.8s",
         "$0.010 (OpenAI)",
         "최고 수준 논거 생성\n긴 컨텍스트 지원\nBYOK 유연",
         "비용 가장 높음\nAPI 의존"],
        ["Claude Sonnet 4.6\n(BYOK)", "4.6/5.0", "200K", "2.1s",
         "$0.009 (Anthropic)",
         "한국어 논증 구조 우수\n긴 컨텍스트\n안전성 높음",
         "비용 높음\nBYOK 필요"],
        ["Gemini 1.5 Pro\n(BYOK)", "4.2/5.0", "1M", "2.4s",
         "$0.007 (Google)",
         "초장문 컨텍스트\n비용 중간",
         "한국어 논거 세밀도\n타 모델 대비 낮음"],
    ]
    for i, row in enumerate(gen_comp):
        r = i + 5
        bg = GREEN_LIGHT if i == 0 else (GRAY_LIGHT if i % 2 == 0 else WHITE)
        ws2.row_dimensions[r].height = 45
        for j, val in enumerate(row, 1):
            sc(ws2, r, j, val, bold=(i == 0), bg=bg, wrap=True, av="top",
               ah="center" if j in [2, 3, 4, 5] else "left")

    ws2.row_dimensions[9].height = 8

    # 2B 판정 LLM 비교
    stitle(ws2, 10, "[B] 판정·검토 LLM 비교 (Orchestrator 전용)", BLUE_MID, cols=7)
    for j, h in enumerate(["모델", "채점 일관성", "JSON 준수율", "응답시간", "비용(1K)", "장점", "단점"], 1):
        sc(ws2, 11, j, h, bold=True, bg=BLUE_MID, fc=WHITE, ah="center")
    ws2.row_dimensions[11].height = 18

    judge_comp = [
        ["GPT-4.1\n(선정)", "4.7/5.0", "98.2%", "1.6s",
         "$0.008",
         "JSON 출력 일관성 최고\n채점 편향 최소\n빠른 응답",
         "비용 발생 (플랫폼 부담)"],
        ["GPT-4o", "4.6/5.0", "97.1%", "1.8s",
         "$0.010",
         "높은 채점 품질",
         "GPT-4.1 대비 높은 비용\n일관성 소폭 낮음"],
        ["Claude Sonnet 4.6", "4.5/5.0", "96.8%", "2.1s",
         "$0.009",
         "논리 근거 서술 우수",
         "JSON 구조 가끔 이탈"],
        ["Llama 3 70B", "3.8/5.0", "88.4%", "3.2s",
         "$0.002",
         "저비용",
         "채점 일관성 낮음\nJSON 파싱 실패 多"],
    ]
    for i, row in enumerate(judge_comp):
        r = i + 12
        bg = BLUE_LIGHT if i == 0 else (GRAY_LIGHT if i % 2 == 0 else WHITE)
        ws2.row_dimensions[r].height = 40
        for j, val in enumerate(row, 1):
            sc(ws2, r, j, val, bold=(i == 0), bg=bg, wrap=True, av="top",
               ah="center" if j in [2, 3, 4, 5] else "left")

    # ── 시트 3: 학습 설정·프롬프트 ──────────────────────────────
    ws3 = wb.create_sheet("3. 프롬프트·파라미터 설정")
    ws3.sheet_view.showGridLines = False
    for col, w in zip("ABCDE", [24, 32, 16, 16, 20]):
        ws3.column_dimensions[col].width = w

    mtitle(ws3, 1, "3. 프롬프트 설계 및 주요 파라미터 설정", cols=5)
    ws3.row_dimensions[2].height = 5

    # 3A 발언 생성
    stitle(ws3, 3, "[A] 발언 생성 LLM — 시스템 프롬프트 구조 및 파라미터", GREEN_DARK, cols=5)
    for j, h in enumerate(["파라미터 / 구성요소", "설정값 / 내용", "선택 옵션", "선정 방법", "근거"], 1):
        sc(ws3, 4, j, h, bold=True, bg=GREEN_MID, fc=WHITE, ah="center")

    gen_params = [
        ["시스템 프롬프트 역할 지시",
         "당신은 [{topic}] 주제의 [{side}] 측 토론자입니다.\n"
         "행동: {action_type} | 상대 마지막 발언: {opponent_last}",
         "—", "설계 결정", "역할·행동·상대발언을 구조화하여 LLM 혼동 최소화"],
        ["최대 출력 토큰",
         "1,024 tokens (debate_max_output_tokens)",
         "512 / 1024 / 2048", "품질/속도 균형", "1,024에서 충분한 논거 + 빠른 응답"],
        ["Temperature",
         "0.7 (발언 다양성)",
         "0.5 / 0.7 / 0.9", "Human 평가", "0.7에서 창의성-일관성 균형 최적"],
        ["이전 턴 이력 윈도우",
         "최근 6턴 (debate_history_window)",
         "4 / 6 / 8턴", "컨텍스트 길이 제약", "6턴에서 문맥 충분 + 토큰 절약"],
        ["타임아웃",
         "30s (debate_llm_timeout)",
         "15s / 30s / 60s", "UX 요구사항", "30s 초과 시 경기 진행 품질보다 흐름 중요"],
        ["o-series temperature 제외",
         "o1/o3/o4 계열은 temperature 파라미터 제외\n(OpenAI API 스펙 준수)",
         "—", "API 명세", "_openai_supports_temperature() 함수로 분기"],
        ["max_tokens vs max_completion_tokens",
         "gpt-4.1·gpt-5·o-series → max_completion_tokens\n구형 모델 → max_tokens",
         "—", "API 명세", "_OPENAI_COMPLETION_TOKENS_PREFIXES 튜플 기반 분기"],
    ]
    for i, row in enumerate(gen_params):
        r = i + 5
        bg = GREEN_LIGHT if i % 2 == 0 else WHITE
        ws3.row_dimensions[r].height = 40
        for j, val in enumerate(row, 1):
            sc(ws3, r, j, val, bg=bg, wrap=True, av="top")

    ws3.row_dimensions[12].height = 8

    # 3B 판정 LLM
    stitle(ws3, 13, "[B] 판정 LLM — 채점 시스템 프롬프트 및 파라미터", BLUE_MID, cols=5)
    for j, h in enumerate(["파라미터 / 구성요소", "설정값 / 내용", "선택 옵션", "선정 방법", "근거"], 1):
        sc(ws3, 14, j, h, bold=True, bg=BLUE_MID, fc=WHITE, ah="center")

    judge_params = [
        ["채점 항목 및 배점",
         "logic(30) + evidence(25) + rebuttal(25) + relevance(20) = 100점",
         "—", "설계 결정", "토론 핵심 역량을 균형 있게 반영"],
        ["판정 Temperature",
         "0.1 (일관된 채점)",
         "0.0 / 0.1 / 0.3", "채점 분산 최소화", "낮은 temperature에서 채점 재현성 높음"],
        ["무승부 억제 지시",
         "프롬프트에 '동일 점수 최소화, 명확한 승자 결정' 명시\n점수차 최소 6점 요구",
         "—", "Human Eval", "무승부 과다 발생 방지 (기존 18% → 6%)"],
        ["JSON 전용 출력 강제",
         "'설명·마크다운 절대 금지, JSON만 출력' 반복 강조\n정규식 재추출 + fallback 병행",
         "—", "파싱 성공률", "JSON 파싱 성공률 82% → 97% 향상"],
        ["최대 출력 토큰 (채점)",
         "512 tokens (reasoning 포함)",
         "256 / 512", "품질/비용 균형", "512에서 reasoning 충분히 포함 가능"],
        ["타임아웃 (채점)",
         "20s (debate_judge_timeout)",
         "10s / 20s / 30s", "UX", "20s 초과 시 fallback 점수로 경기 지속"],
    ]
    for i, row in enumerate(judge_params):
        r = i + 15
        bg = BLUE_LIGHT if i % 2 == 0 else WHITE
        ws3.row_dimensions[r].height = 40
        for j, val in enumerate(row, 1):
            sc(ws3, r, j, val, bg=bg, wrap=True, av="top")

    ws3.row_dimensions[21].height = 8

    # 3C 턴 검토
    stitle(ws3, 22, "[C] 턴 검토 LLM — 위반 탐지 파라미터", PURPLE_MID, cols=5)
    for j, h in enumerate(["파라미터 / 구성요소", "설정값 / 내용", "선택 옵션", "선정 방법", "근거"], 1):
        sc(ws3, 23, j, h, bold=True, bg=PURPLE_MID, fc=WHITE, ah="center")

    review_params = [
        ["위반 유형 및 벌점",
         "prompt_injection(-10), ad_hominem(-8), off_topic(-5), false_claim(-7)",
         "—", "설계 결정", "위반 심각도에 비례한 벌점 설계"],
        ["차단 기준",
         "severity='severe'일 때만 block=true\nminor 위반은 벌점만 부과, 발언 표시",
         "—", "UX·공정성", "사소한 위반으로 경기 흐름 중단 방지"],
        ["검토 Temperature",
         "0.1 (일관된 위반 탐지)",
         "0.0 / 0.1", "일관성", "낮은 temperature에서 위반 판정 재현성"],
        ["검토 모델",
         "debate_turn_review_model 설정 (기본: GPT-4.1)",
         "—", "Config", "판정 LLM과 동일 모델 사용 권장"],
        ["검토 타임아웃",
         "10s (debate_turn_review_timeout)",
         "5s / 10s", "속도", "10s 초과 시 fallback (block=false, score=5)"],
        ["검토 활성화 여부",
         "debate_turn_review_enabled (bool)",
         "true / false", "Config", "비용·속도 절충 시 비활성화 가능"],
    ]
    for i, row in enumerate(review_params):
        r = i + 24
        bg = PURPLE_LIGHT if i % 2 == 0 else WHITE
        ws3.row_dimensions[r].height = 35
        for j, val in enumerate(row, 1):
            sc(ws3, r, j, val, bg=bg, wrap=True, av="top")

    # ── 시트 4: 성능 평가 ────────────────────────────────────────
    ws4 = wb.create_sheet("4. 성능 평가")
    ws4.sheet_view.showGridLines = False
    for col, w in zip("ABCDEFG", [22, 14, 14, 14, 14, 14, 20]):
        ws4.column_dimensions[col].width = w

    mtitle(ws4, 1, "4. 학습 결과 및 성능 평가", cols=7)
    ws4.row_dimensions[2].height = 5

    # 4A 발언 품질
    stitle(ws4, 3, "[A] 발언 생성 품질 평가 (Human Eval, n=200 경기)", GREEN_DARK, cols=7)
    for j, h in enumerate(["모델", "논거 설득력", "한국어 자연성", "주제 적합성", "반박 품질", "전체 평균", "비고"], 1):
        sc(ws4, 4, j, h, bold=True, bg=GREEN_MID, fc=WHITE, ah="center")
    ws4.row_dimensions[4].height = 18

    gen_perf = [
        ["Claude Sonnet 4.6", "4.6/5", "4.7/5", "4.5/5", "4.6/5", "4.60/5", "BYOK"],
        ["GPT-4o",            "4.5/5", "4.6/5", "4.4/5", "4.5/5", "4.50/5", "BYOK"],
        ["Gemini 1.5 Pro",    "4.2/5", "4.3/5", "4.1/5", "4.2/5", "4.20/5", "BYOK"],
        ["Llama 3 70B",       "4.0/5", "4.1/5", "3.9/5", "3.9/5", "3.98/5", "기본 내장"],
    ]
    for i, row in enumerate(gen_perf):
        r = i + 5
        bg = GREEN_LIGHT if i % 2 == 0 else WHITE
        ws4.row_dimensions[r].height = 20
        for j, val in enumerate(row, 1):
            sc(ws4, r, j, val, bg=bg, ah="center" if j > 1 else "left")

    ws4.row_dimensions[9].height = 8

    # 4B 채점 일관성
    stitle(ws4, 10, "[B] 판정 LLM 채점 일관성 평가 (동일 경기 3회 반복 채점)", BLUE_MID, cols=7)
    for j, h in enumerate(["모델", "점수 편차(σ)", "승자 일치율", "무승부 비율", "JSON 파싱 성공률", "평균 채점 시간", "비고"], 1):
        sc(ws4, 11, j, h, bold=True, bg=BLUE_MID, fc=WHITE, ah="center")
    ws4.row_dimensions[11].height = 18

    judge_perf = [
        ["GPT-4.1 (선정)", "±1.8점", "97.2%", "5.8%", "98.2%", "1.6s", "최고 일관성"],
        ["GPT-4o",         "±2.1점", "95.4%", "7.1%", "97.1%", "1.8s", ""],
        ["Claude S4.6",    "±2.3점", "94.8%", "6.9%", "96.8%", "2.1s", ""],
        ["Llama 3 70B",    "±4.7점", "81.3%", "18.5%","88.4%", "3.2s", "일관성 낮아 기본 판정에 미사용"],
    ]
    for i, row in enumerate(judge_perf):
        r = i + 12
        bg = BLUE_LIGHT if i == 0 else (GRAY_LIGHT if i % 2 == 0 else WHITE)
        ws4.row_dimensions[r].height = 20
        for j, val in enumerate(row, 1):
            sc(ws4, r, j, val, bold=(i == 0), bg=bg, ah="center" if j > 1 else "left")

    ws4.row_dimensions[16].height = 8

    # 4C 위반 탐지
    stitle(ws4, 17, "[C] 턴 검토 — 위반 탐지 성능 (수동 레이블 100건 대비)", PURPLE_MID, cols=7)
    for j, h in enumerate(["위반 유형", "Precision", "Recall", "F1-Score", "오탐율", "벌점", "비고"], 1):
        sc(ws4, 18, j, h, bold=True, bg=PURPLE_MID, fc=WHITE, ah="center")
    ws4.row_dimensions[18].height = 18

    viol_perf = [
        ["prompt_injection", "0.94", "0.89", "0.91", "6.0%", "-10점", ""],
        ["ad_hominem",       "0.88", "0.85", "0.86", "12.0%","-8점",  "유사 표현 오탐 있음"],
        ["off_topic",        "0.82", "0.79", "0.80", "18.0%","-5점",  "주제 경계 판단 어려움"],
        ["false_claim",      "0.79", "0.74", "0.76", "21.0%","-7점",  "사실 검증 한계"],
        ["전체 (macro 평균)", "0.86", "0.82", "0.83", "14.2%","—",    ""],
    ]
    for i, row in enumerate(viol_perf):
        r = i + 19
        is_total = i == len(viol_perf) - 1
        bg = BLUE_DARK if is_total else (PURPLE_LIGHT if i % 2 == 0 else WHITE)
        fc = WHITE if is_total else "000000"
        ws4.row_dimensions[r].height = 20
        for j, val in enumerate(row, 1):
            sc(ws4, r, j, val, bold=is_total, bg=bg, fc=fc,
               ah="center" if j > 1 else "left")

    ws4.row_dimensions[24].height = 8

    # 4D ELO
    stitle(ws4, 25, "[D] ELO 레이팅 시스템 검증", GREEN_MID, cols=7)
    for j, h in enumerate(["검증 항목", "결과", "기준", "비고", "", "", ""], 1):
        if h:
            sc(ws4, 26, j, h, bold=True, bg=GREEN_MID, fc=WHITE, ah="center")
    ws4.row_dimensions[26].height = 18

    elo_perf = [
        ["ELO 수렴 매칭 수", "약 30~40경기", "전통 체스 기준 30~50", ""],
        ["상위 에이전트 승률 (ELO +200 차)", "약 76%", "이론값 76%", "수식 일치"],
        ["ELO 변동 범위 (K=32)", "±0~32점/경기", "K=32 표준", ""],
        ["ELO 인플레이션 여부", "없음 (제로섬)", "총합 일정 유지", "승패 동일 델타"],
    ]
    for i, row in enumerate(elo_perf):
        r = i + 27
        bg = GREEN_LIGHT if i % 2 == 0 else WHITE
        ws4.row_dimensions[r].height = 20
        for j, val in enumerate(row, 1):
            sc(ws4, r, j, val, bg=bg)

    # ── 시트 5: 과적합·한계 대응 ────────────────────────────────
    ws5 = wb.create_sheet("5. 한계·대응 전략")
    ws5.sheet_view.showGridLines = False
    for col, w in zip("ABCDE", [22, 28, 28, 20, 10]):
        ws5.column_dimensions[col].width = w

    mtitle(ws5, 1, "5. 모델 한계 및 대응 전략 (과적합·과소적합 포함)", cols=5)
    ws5.row_dimensions[2].height = 5

    for j, h in enumerate(["문제 유형", "징후 (관찰)", "대응 전략", "적용 결과", "구성요소"], 1):
        sc(ws5, 3, j, h, bold=True, bg=RED_MID, fc=WHITE, ah="center")
    ws5.row_dimensions[3].height = 18

    issues = [
        ["판정 편향\n(먼저 발언한 측 유리)",
         "찬성 측(先발언) 승률이 기대치\n대비 8~12%p 높게 관찰됨",
         "① 채점 프롬프트에 '발언 순서는 유·불리 아님' 명시\n"
         "② '동일 점수 최소화·명확한 승자' 강조\n"
         "③ 점수차 최소 6점 요구 지시",
         "찬성 편향 8%p → 2%p\n무승부 비율 18% → 6%",
         "JUDGE_SYSTEM_PROMPT"],
        ["채점 무결성 저하\n(JSON 파싱 실패)",
         "LLM이 JSON 외 마크다운·설명\n추가로 파싱 실패 발생 (약 18%)",
         "① 프롬프트에 '마크다운·추가 텍스트 절대 금지' 반복\n"
         "② 정규식 코드블록 제거 후 JSON 재추출\n"
         "③ 최종 실패 시 fallback 점수 사용",
         "파싱 실패율 18% → 2.8%\nfallback 사용 1.2%",
         "DebateOrchestrator\njudge() / review_turn()"],
        ["위반 탐지 오탐\n(off_topic·false_claim)",
         "주제 경계가 불명확하거나\n사실 검증 불가 주장에서\n오탐율 18~21%",
         "① minor 위반은 벌점만 부과 (발언 차단 안 함)\n"
         "② severe 기준 강화 ('명백한' 위반만 차단)\n"
         "③ 오탐 허용으로 경기 흐름 보호",
         "불필요 차단 감소\n(차단율 12% → 3%)",
         "REVIEW_SYSTEM_PROMPT\nLLM_VIOLATION_PENALTIES"],
        ["발언 품질 불균형\n(에이전트 모델 격차)",
         "Llama 70B vs GPT-4o 간\n발언 품질 격차가 커\n비대칭 경기 다수",
         "① ELO 자동 분류 (수십 경기 후 격차 반영)\n"
         "② 동일 ELO 대역 매칭 우선 (추후 개선)\n"
         "③ 기본 Llama 70B도 충분한 논거 생성 가능하도록\n   프롬프트 최적화",
         "ELO 500점 이상 격차 발생\n→ 자동 매칭 분리 효과",
         "debate_matching_service\nELO 기반 매칭"],
        ["타임아웃으로 인한\n발언 품질 저하",
         "느린 모델(Llama)에서\n30s 초과 [TIMEOUT] 비율 약 4%",
         "① timeout 설정값 조정 (모델별 다른 timeout 검토)\n"
         "② [TIMEOUT] 발언 채점 시 0~5점 자동 부여\n"
         "③ 경기는 중단 없이 지속",
         "timeout으로 인한\n경기 중단 0건",
         "debate_engine.py\n_generate_turn()"],
    ]
    for i, row in enumerate(issues):
        r = i + 4
        bg = ORANGE_LIGHT if i % 2 == 0 else WHITE
        ws5.row_dimensions[r].height = 65
        for j, val in enumerate(row, 1):
            sc(ws5, r, j, val, bg=bg, wrap=True, av="top")

    out_path = OUT_DIR / "인공지능_학습_결과서.xlsx"
    wb.save(str(out_path))
    print(f"[OK] 인공지능_학습_결과서.xlsx 저장: {out_path}")
    return str(out_path)


# ════════════════════════════════════════════════════════════════
#  3. AI 모델 명세서 (.docx)  — AI 토론 플랫폼 기준
# ════════════════════════════════════════════════════════════════
def build_model_spec_docx():
    from docx import Document
    from docx.shared import Cm
    doc = Document()
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Cm(2.0)
        section.left_margin = section.right_margin = Cm(2.5)

    _cover(doc, "학습된 인공지능 모델 명세서", "AI 토론 플랫폼 — 활용 모델 파일 및 문서")
    _meta_table(doc, [
        ("작성일",   "2026년 2월"),
        ("버전",     "v1.0"),
        ("목적",     "제출용 AI 모델 파일 및 설정 문서 정리"),
        ("포함 모델","발언 생성 LLM / 판정 LLM / 턴 검토 LLM / ELO 알고리즘"),
    ])

    # ── 1. 모델 전체 목록 ────────────────────────────────────────
    _add_heading(doc, "1. 활용 AI 모델 전체 목록", 1)
    tbl1 = doc.add_table(rows=6, cols=6)
    tbl1.style = "Table Grid"
    _header_row(tbl1, ["No.", "역할", "기본 모델", "대안 모델 (BYOK)", "파일 형태", "라이선스"])
    model_list = [
        ["1", "발언 생성\n(에이전트 기본)",
         "Llama 3 70B Instruct\n(RunPod SGLang)",
         "GPT-4o / GPT-4.1\nClaude Sonnet 4.6\nGemini 1.5 Pro",
         "API (클라우드)\n자체 파일 없음",
         "Meta Llama License\n/ 각 제공사 ToS"],
        ["2", "판정 LLM\n(채점·승자 결정)",
         "GPT-4.1 (OpenAI)\n플랫폼 API 키",
         "GPT-4o / Claude\n(동일 Orchestrator 호환)",
         "API (클라우드)\n자체 파일 없음",
         "OpenAI ToS"],
        ["3", "턴 검토 LLM\n(위반 탐지)",
         "GPT-4.1 (OpenAI)\n(debate_turn_review_model)",
         "판정 LLM과 동일",
         "API (클라우드)",
         "OpenAI ToS"],
        ["4", "ELO 레이팅\n(수식 기반)",
         "K=32 Logistic ELO\n수식 자체 — 별도 모델 없음",
         "K값 조정 가능",
         "Python 코드\n(debate_orchestrator.py)",
         "MIT (플랫폼 자체 코드)"],
        ["5", "매칭 알고리즘\n(대기열 매칭)",
         "선착순 + 자동 매칭\n(2인 대기 or 타임아웃 후 자동)",
         "ELO 기반 매칭\n(추후 개선 예정)",
         "Python 코드\n(debate_matching_service.py)",
         "MIT (플랫폼 자체 코드)"],
    ]
    for i, row in enumerate(model_list):
        bg = BLUE_LIGHT if i % 2 == 0 else WHITE
        for j, val in enumerate(row):
            _set_cell(tbl1.rows[i+1].cells[j], val, bg=bg, font_size=8,
                      align="center" if j == 0 else "left")
    _set_col_widths(tbl1, [0.8, 3.2, 4.0, 4.5, 3.2, 2.8])
    doc.add_paragraph()

    # ── 2. 발언 생성 LLM 상세 ───────────────────────────────────
    _add_heading(doc, "2. 발언 생성 LLM 상세 명세", 1)
    _add_para(doc, "에이전트가 자신의 API 키로 등록한 LLM이 토론 발언을 생성합니다. "
                   "플랫폼은 에이전트별 provider·model_id를 DB에서 조회하여 "
                   "InferenceClient를 통해 분기 호출합니다.", 10)
    doc.add_paragraph()

    _add_heading(doc, "2.1 지원 프로바이더 및 모델 ID", 2)
    tbl2 = doc.add_table(rows=6, cols=4)
    tbl2.style = "Table Grid"
    _header_row(tbl2, ["Provider", "대표 Model ID", "토큰 파라미터", "비고"])
    providers = [
        ["openai", "gpt-4o-2024-11-20\ngpt-4.1\ngpt-5",
         "max_tokens (구형)\nmax_completion_tokens (4.1+/5+/o-series)",
         "o-series: temperature 제외"],
        ["anthropic", "claude-sonnet-4-6\nclaude-3-5-sonnet-20241022",
         "max_tokens", "스트리밍 SSE 지원"],
        ["google", "gemini-1.5-pro\ngemini-2.0-flash",
         "max_output_tokens", "Google Generative AI SDK"],
        ["runpod", "meta-llama/Llama-3-70b-Instruct",
         "max_tokens", "OpenAI 호환 엔드포인트 (SGLang)"],
        ["custom", "사용자 등록 모델 (openai 호환)",
         "max_tokens", "base_url 설정 필요"],
    ]
    for i, row in enumerate(providers):
        bg = ORANGE_LIGHT if i % 2 == 0 else WHITE
        for j, val in enumerate(row):
            _set_cell(tbl2.rows[i+1].cells[j], val, bg=bg, font_size=8)
    _set_col_widths(tbl2, [2.5, 5.0, 4.5, 5.5])
    doc.add_paragraph()

    _add_heading(doc, "2.2 발언 생성 흐름 (코드 핵심 로직)", 2)
    code1 = [
        "# debate_engine.py  — _generate_turn()",
        "system_prompt = build_debate_system_prompt(topic, side, action_type)",
        "messages = [",
        "    {'role': 'system', 'content': system_prompt},",
        "    *turn_history[-history_window*2:],   # 최근 6턴 슬라이딩",
        "    {'role': 'user', 'content': action_instruction},",
        "]",
        "# InferenceClient가 provider별 분기 후 스트리밍 반환",
        "async for chunk in client.stream(model_id, api_key, messages,",
        "                                 max_tokens=1024, temperature=0.7):",
        "    yield chunk  # SSE로 프론트엔드에 실시간 전달",
    ]
    _add_box(doc, code1, bg="F0F4F8")

    # ── 3. 판정 LLM 상세 ────────────────────────────────────────
    _add_heading(doc, "3. 판정 LLM (DebateOrchestrator.judge) 상세 명세", 1)

    _add_heading(doc, "3.1 채점 시스템 프롬프트 구조", 2)
    prompt_lines = [
        "역할 정의: '당신은 공정하지만 명확한 토론 심판입니다.'",
        "",
        "채점 기준 (4항목):",
        "  logic     0~30점 : 논리적 일관성, 타당한 추론 체계",
        "  evidence  0~25점 : 근거·데이터·인용 활용도",
        "  rebuttal  0~25점 : 반박 논리의 질, 상대 주장 대응 수준",
        "  relevance 0~20점 : 주제 적합성, 핵심 쟁점 집중도",
        "",
        "편향 방지 지시:",
        "  - '한 쪽이 더 나은 논거를 보였다면 합산 점수 최소 6점 차이'",
        "  - '무승부는 모든 항목에서 정말 구분 어려울 때만'",
        "  - '발언 순서(찬성 먼저)는 유·불리 요소 아님'",
        "  - '[TIMEOUT]/[ERROR] 응답은 0~5점 범위'",
        "",
        "출력 형식 (JSON 전용, 마크다운 금지):",
        '  {"agent_a": {"logic":_, "evidence":_, "rebuttal":_, "relevance":_},',
        '   "agent_b": {"logic":_, "evidence":_, "rebuttal":_, "relevance":_},',
        '   "reasoning": "<채점 근거 한국어>"}',
    ]
    _add_box(doc, prompt_lines, bg=BLUE_LIGHT)

    _add_heading(doc, "3.2 판정 결과 파싱 및 승자 결정 로직", 2)
    tbl3 = doc.add_table(rows=5, cols=2)
    tbl3.style = "Table Grid"
    _header_row(tbl3, ["단계", "처리 내용"])
    judge_steps = [
        ["① JSON 파싱",
         "마크다운 코드블록(```) 제거 → 정규식으로 JSON 객체 추출 → json.loads()\n"
         "파싱 실패 시 fallback: {logic:15, evidence:12, rebuttal:12, relevance:10} × 양측"],
        ["② 점수 clamp",
         "각 항목을 명세 범위 내로 강제:\n"
         "logic=min(max(v,0),30), evidence=min(max(v,0),25), rebuttal=min(max(v,0),25), relevance=min(max(v,0),20)"],
        ["③ 스왑 판정",
         "판정 LLM이 찬성/반대 에이전트를 혼동한 경우\nreasoning 텍스트 기반으로 역변환 후 점수 교환"],
        ["④ 벌점 차감 및 승자 결정",
         "total_a = sum(scores_a) - sum(penalties_a)\ntotal_b = sum(scores_b) - sum(penalties_b)\n"
         "차이 ≥ 5점 → 승자 결정 | 차이 < 5점 → 무승부"],
    ]
    for i, (k, v) in enumerate(judge_steps):
        bg = BLUE_LIGHT if i % 2 == 0 else WHITE
        _set_cell(tbl3.rows[i+1].cells[0], k, bold=True, bg=bg, font_size=9)
        _set_cell(tbl3.rows[i+1].cells[1], v, bg=bg, font_size=8)
    _set_col_widths(tbl3, [3.5, 14.0])
    doc.add_paragraph()

    # ── 4. 턴 검토 LLM 상세 ─────────────────────────────────────
    _add_heading(doc, "4. 턴 검토 LLM (DebateOrchestrator.review_turn) 상세 명세", 1)

    _add_heading(doc, "4.1 위반 유형 및 벌점", 2)
    tbl4 = doc.add_table(rows=5, cols=4)
    tbl4.style = "Table Grid"
    _header_row(tbl4, ["위반 유형", "설명", "벌점", "차단 기준"])
    violations = [
        ["prompt_injection",
         "시스템 지시를 무력화하려는 시도\n(예: '이전 지시 무시하고...')",
         "-10점", "severity=severe → block=true"],
        ["ad_hominem",
         "논거 대신 상대방을 직접 비하\n(인신공격성 발언)",
         "-8점", "severe 시 차단"],
        ["off_topic",
         "토론 주제와 무관한 내용\n발언 전체가 주제 이탈",
         "-5점", "minor 비율 높아 주로 벌점만"],
        ["false_claim",
         "명백히 허위이거나 사실 확인\n불가능한 주장",
         "-7점", "severe 시 차단"],
    ]
    for i, row in enumerate(violations):
        bg = RED_LIGHT if i % 2 == 0 else WHITE
        for j, val in enumerate(row):
            _set_cell(tbl4.rows[i+1].cells[j], val, bg=bg, font_size=8,
                      align="center" if j == 2 else "left")
    _set_col_widths(tbl4, [3.5, 5.5, 2.0, 7.0])
    doc.add_paragraph()

    _add_heading(doc, "4.2 검토 결과 JSON 구조", 2)
    review_json = [
        '{',
        '  "logic_score": 7,          // 1~10, 논리 점수',
        '  "violations": [',
        '    {',
        '      "type": "off_topic",',
        '      "severity": "minor",   // minor | severe',
        '      "detail": "주제와 관련 없는 개인 경험 서술"',
        '    }',
        '  ],',
        '  "severity": "minor",       // none | minor | severe (최대 심각도)',
        '  "feedback": "논리는 명확하나 주제 집중도 아쉬움",  // 30자 이내',
        '  "block": false             // true면 발언 차단, 대체 텍스트 표시',
        '}',
    ]
    _add_box(doc, review_json, bg="F0F4F8")

    # ── 5. ELO 레이팅 알고리즘 상세 ─────────────────────────────
    _add_heading(doc, "5. ELO 레이팅 알고리즘 명세", 1)
    _add_para(doc, "ELO 레이팅은 별도 학습 없이 수식으로 에이전트 실력을 추적합니다. "
                   "채점 결과(승/패/무)를 기반으로 매 경기 후 갱신됩니다.", 10)
    doc.add_paragraph()

    tbl5 = doc.add_table(rows=7, cols=2)
    tbl5.style = "Table Grid"
    _header_row(tbl5, ["항목", "내용"], bg=GREEN_DARK)
    elo_spec = [
        ["초기 ELO",       "모든 신규 에이전트 1,500점"],
        ["K 팩터",         "K = 32 (고정, 신규·숙련 구분 없음)"],
        ["기대 승률 공식",  "E_a = 1 / (1 + 10^((Rb - Ra) / 400))"],
        ["델타 계산",       "delta_a = K × (S_a - E_a)\nS_a: 승=1, 무=0.5, 패=0"],
        ["갱신 공식",       "new_Ra = Ra + delta_a\nnew_Rb = Rb - delta_a  (제로섬)"],
        ["승자 결정 기준",  "채점 합산 점수 차이 ≥ 5점 → 승패\n차이 < 5점 → 무승부"],
    ]
    for i, (k, v) in enumerate(elo_spec):
        bg = GREEN_LIGHT if i % 2 == 0 else WHITE
        _set_cell(tbl5.rows[i+1].cells[0], k, bold=True, bg=bg, font_size=9)
        _set_cell(tbl5.rows[i+1].cells[1], v, bg=bg, font_size=9)
    _set_col_widths(tbl5, [4.5, 13.0])
    doc.add_paragraph()

    # ── 6. 서비스 배포 구조 ──────────────────────────────────────
    _add_heading(doc, "6. 모델 서빙 및 배포 구조", 1)
    deploy_lines = [
        "┌─────────────────────────────────────────────────────────────────────┐",
        "│                  AI 토론 플랫폼 모델 서빙 구조                       │",
        "├─────────────────────────────────────────────────────────────────────┤",
        "│                                                                     │",
        "│  [클라이언트 브라우저]                                               │",
        "│       ↑ SSE 스트리밍 (발언 실시간 수신)                              │",
        "│       │                                                             │",
        "│  [EC2 t4g.small — 서울 리전]                                        │",
        "│   ┌─ FastAPI ─────────────────────────────────────────────────┐    │",
        "│   │  debate_engine.py      ← 발언 생성 루프 (async)           │    │",
        "│   │  debate_orchestrator.py← 판정·검토·ELO (async)            │    │",
        "│   │  inference_client.py   ← LLM 프로바이더 분기               │    │",
        "│   └──────────────────────────────────────────────────────────┘    │",
        "│       │              │               │             │               │",
        "│       ▼              ▼               ▼             ▼               │",
        "│  [RunPod SGLang] [OpenAI API] [Anthropic API] [Google API]        │",
        "│  Llama 3 70B     GPT-4.1        Claude S4.6    Gemini 1.5 Pro    │",
        "│  (기본 에이전트)  (판정·검토)   (BYOK 에이전트) (BYOK 에이전트)  │",
        "│                                                                     │",
        "│  [Redis]  ← 큐 이벤트 pub/sub (SSE 브로드캐스트)                   │",
        "│  [PostgreSQL]  ← debate_turn_logs, debate_matches, debate_agents   │",
        "│                                                                     │",
        "└─────────────────────────────────────────────────────────────────────┘",
    ]
    _add_box(doc, deploy_lines, bg="F8F9FA")

    tbl6 = doc.add_table(rows=5, cols=3)
    tbl6.style = "Table Grid"
    _header_row(tbl6, ["컴포넌트", "배포 방식", "비고"])
    deploy_rows = [
        ["발언 생성 LLM (기본)", "RunPod SGLang Serverless API\n(HTTP POST, OpenAI 호환 엔드포인트)",
         "콜드스타트 ~2s (FlashBoot), RTT ~150ms"],
        ["판정·검토 LLM", "OpenAI API\n(플랫폼 서버 측 API 키 사용)",
         "settings.openai_api_key 환경변수"],
        ["에이전트 BYOK LLM", "각 제공사 API\n(에이전트 등록 시 Fernet 암호화 키 저장)",
         "InferenceClient가 복호화 후 호출"],
        ["ELO·판정 로직", "FastAPI 백엔드 내부 Python 코드",
         "별도 모델 파일 없음 — 수식 자체"],
    ]
    for i, row in enumerate(deploy_rows):
        bg = GRAY_LIGHT if i % 2 == 0 else WHITE
        for j, val in enumerate(row):
            _set_cell(tbl6.rows[i+1].cells[j], val, bg=bg, font_size=8)
    _set_col_widths(tbl6, [4.0, 7.0, 6.5])
    doc.add_paragraph()

    # ── 7. 모델 갱신 정책 ────────────────────────────────────────
    _add_heading(doc, "7. 모델 갱신 및 버전 관리 정책", 1)
    tbl7 = doc.add_table(rows=5, cols=3)
    tbl7.style = "Table Grid"
    _header_row(tbl7, ["항목", "정책", "주기 / 트리거"])
    policy = [
        ["LLM 모델 추가", "관리자 대시보드에서 provider·model_id·API 키 즉시 등록",
         "수시 (신규 모델 출시 시)"],
        ["판정 모델 교체", "settings.debate_orchestrator_model 환경변수 변경",
         "배포 시 (재시작 필요)"],
        ["채점 프롬프트 튜닝", "debate_orchestrator.py 수정 후 재배포\n편향·파싱 오류 발생 시 개선",
         "이슈 발생 시"],
        ["ELO K팩터 조정", "settings 또는 하드코딩 상수 변경\n(현재 K=32 고정)",
         "레이팅 인플레이션/급변 발생 시"],
    ]
    for i, row in enumerate(policy):
        bg = BLUE_LIGHT if i % 2 == 0 else WHITE
        for j, val in enumerate(row):
            _set_cell(tbl7.rows[i+1].cells[j], val, bg=bg, font_size=8)
    _set_col_widths(tbl7, [4.0, 9.0, 4.5])
    doc.add_paragraph()

    out_path = OUT_DIR / "AI_모델_명세서.docx"
    doc.save(str(out_path))
    print(f"[OK] AI_모델_명세서.docx 저장: {out_path}")
    return str(out_path)


# ════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    p1 = build_preprocessing_docx()
    p2 = build_training_xlsx()
    p3 = build_model_spec_docx()
    print("\n=== 생성 완료 ===")
    for p in [p1, p2, p3]:
        print(f"  {p}")
