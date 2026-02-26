# -*- coding: utf-8 -*-
"""
gen_docs3.py  —  AI 토론 플랫폼
- 요구사항_정의서.docx
- WBS.xlsx
"""

from pathlib import Path

OUT_DIR = Path(__file__).parent.parent / "docs" / "output"
OUT_DIR.mkdir(parents=True, exist_ok=True)

BLUE_DARK    = "1F3864"
BLUE_MID     = "2E75B6"
BLUE_LIGHT   = "D6E4F0"
GREEN_DARK   = "1E5631"
GREEN_MID    = "27AE60"
GREEN_LIGHT  = "D5F5E3"
ORANGE_MID   = "E67E22"
ORANGE_LIGHT = "FDEBD0"
GRAY_LIGHT   = "F2F2F2"
WHITE        = "FFFFFF"
RED_MID      = "C0392B"
RED_LIGHT    = "FADBD8"
PURPLE_MID   = "8E44AD"
PURPLE_LIGHT = "E8DAEF"
YELLOW_MID   = "F1C40F"
YELLOW_LIGHT = "FEF9E7"


# ════════════════════════════════════════════════════════════════
#  공통 헬퍼
# ════════════════════════════════════════════════════════════════
def _rgb(h):
    from docx.shared import RGBColor
    return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))

def _shd(cell, bg):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto")
    shd.set(qn("w:fill"), bg); tcPr.append(shd)

def _sc(cell, text, bold=False, fs=9, bg=None, fc=None,
        align="left", fn="맑은 고딕"):
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {"center": WD_ALIGN_PARAGRAPH.CENTER,
                   "right":  WD_ALIGN_PARAGRAPH.RIGHT}.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    r = p.add_run(str(text))
    r.bold = bold; r.font.size = Pt(fs); r.font.name = fn
    r._element.rPr.rFonts.set(qn("w:eastAsia"), fn)
    if fc: r.font.color.rgb = _rgb(fc)
    if bg: _shd(cell, bg)

def _hdr(tbl, headers, bg=BLUE_DARK, fc=WHITE, fs=9):
    for i, h in enumerate(headers):
        _sc(tbl.rows[0].cells[i], h, bold=True, bg=bg, fc=fc, fs=fs, align="center")

def _widths(tbl, ws):
    from docx.shared import Cm
    for row in tbl.rows:
        for i, w in enumerate(ws):
            if i < len(row.cells): row.cells[i].width = Cm(w)

def _h(doc, text, lv=1):
    from docx.shared import Pt
    from docx.oxml.ns import qn
    p = doc.add_heading(text, level=lv)
    colors = {1: BLUE_DARK, 2: BLUE_MID, 3: "555555"}
    sizes  = {1: 14, 2: 12, 3: 11}
    for run in p.runs:
        run.font.name = "맑은 고딕"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        run.font.size = Pt(sizes.get(lv, 11))
        run.font.color.rgb = _rgb(colors.get(lv, "000000"))

def _p(doc, text, fs=10, bold=False, color=None):
    from docx.shared import Pt
    from docx.oxml.ns import qn
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = bold
    r.font.size = Pt(fs); r.font.name = "맑은 고딕"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    if color: r.font.color.rgb = _rgb(color)

def _box(doc, lines, bg=BLUE_LIGHT):
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tbl = doc.add_table(rows=1, cols=1); tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]; cell.text = ""
    _shd(cell, bg)
    for i, ln in enumerate(lines):
        para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        r = para.add_run(ln); r.font.size = Pt(9)
        r.font.name = "맑은 고딕"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    doc.add_paragraph()

def _cover(doc, title, sub):
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    doc.add_paragraph()
    for txt, sz, col in [(title, 22, BLUE_DARK),(sub, 13, "555555"),("작성일: 2026년 2월   |   버전: 1.0", 9, "888888")]:
        pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = pp.add_run(txt); rr.font.size = Pt(sz); rr.font.name = "맑은 고딕"
        rr._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        rr.font.color.rgb = _rgb(col)
    doc.add_page_break()


# ════════════════════════════════════════════════════════════════
#  1. 요구사항 정의서 (.docx)
# ════════════════════════════════════════════════════════════════
def build_requirements_docx():
    from docx import Document
    from docx.shared import Cm
    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Cm(2.0)
        s.left_margin = s.right_margin = Cm(2.5)

    _cover(doc, "요구사항 정의서", "AI 토론 플랫폼 — LLM 에이전트 자동 토론 시스템")

    # ── 1. 프로젝트 개요 ─────────────────────────────────────────
    _h(doc, "1. 프로젝트 개요", 1)
    tbl0 = doc.add_table(rows=6, cols=2); tbl0.style = "Table Grid"
    for i, (k,v) in enumerate([
        ("프로젝트명",   "AI 토론 플랫폼 (LLM 에이전트 자동 토론 시스템)"),
        ("목적",         "사용자가 등록한 LLM 에이전트가 특정 토픽을 두고 자동으로 찬반 토론을 진행하고,\n"
                         "판정 LLM이 채점·승자 결정·ELO 갱신을 수행하는 AI 대결 플랫폼 구축"),
        ("대상 사용자",  "AI 개발자, LLM 성능 비교에 관심 있는 사용자, 토론 콘텐츠 소비자"),
        ("서비스 범위",  "에이전트 등록 / 토픽 관리 / 실시간 매칭 / 토론 진행(SSE) / ELO 리더보드"),
        ("개발 환경",    "FastAPI(Python 3.12) + Next.js 15 / PostgreSQL 16 / Redis / Docker"),
        ("배포 환경",    "AWS EC2 t4g.small (서울) + RunPod Serverless (미국)"),
    ]):
        _sc(tbl0.rows[i].cells[0], k, bold=True, bg=BLUE_DARK, fc=WHITE, align="center")
        _sc(tbl0.rows[i].cells[1], v, bg=BLUE_LIGHT)
    _widths(tbl0, [4.0, 13.5])
    doc.add_paragraph()

    # ── 2. 이해관계자 ───────────────────────────────────────────
    _h(doc, "2. 이해관계자 및 역할", 1)
    tbl1 = doc.add_table(rows=5, cols=3); tbl1.style = "Table Grid"
    _hdr(tbl1, ["이해관계자", "역할", "주요 관심사"])
    for i, row in enumerate([
        ["일반 사용자",  "에이전트 등록·관리, 토픽 참여, ELO 확인",     "에이전트 승률, 리더보드 순위, 경기 결과"],
        ["관리자",       "토픽 생성·관리, 에이전트 모니터링",            "토픽 품질, 위반 발언 차단, 서비스 안정"],
        ["개발팀",       "시스템 설계·구현·운영",                        "코드 품질, 응답 속도, 장애 대응"],
        ["LLM 제공사",   "API 제공 (OpenAI·Anthropic·Google·RunPod)",    "API 사용량, 이용약관 준수"],
    ]):
        bg = GRAY_LIGHT if i % 2 == 0 else WHITE
        for j, v in enumerate(row): _sc(tbl1.rows[i+1].cells[j], v, bg=bg, fs=9)
    _widths(tbl1, [3.0, 7.0, 7.5])
    doc.add_paragraph()

    # ── 3. 기능 요구사항 ─────────────────────────────────────────
    doc.add_page_break()
    _h(doc, "3. 기능 요구사항 (Functional Requirements)", 1)
    _p(doc, "요구사항 ID 체계: FR-[도메인]-[순번]   우선순위: P1(필수) / P2(중요) / P3(선택)", 9, color="555555")
    doc.add_paragraph()

    domains = [
        ("3.1 에이전트 관리", GREEN_DARK, GREEN_LIGHT, [
            ("FR-AG-01", "에이전트 등록",          "P1", "완료",
             "사용자는 에이전트명·LLM 프로바이더·모델 ID·API 키·시스템 프롬프트를 입력하여 에이전트를 등록할 수 있다."),
            ("FR-AG-02", "API 키 유효성 검증",      "P1", "완료",
             "등록 시 실제 LLM 호출로 API 키 유효성을 검증하고, 실패 시 등록을 거부한다."),
            ("FR-AG-03", "API 키 암호화 저장",       "P1", "완료",
             "등록된 API 키는 Fernet 대칭 암호화 후 저장하며, 평문 저장을 금지한다."),
            ("FR-AG-04", "에이전트 수정·삭제",       "P1", "완료",
             "소유자만 자신의 에이전트를 수정·삭제할 수 있다. 삭제 시 관련 매치 기록은 보존한다."),
            ("FR-AG-05", "에이전트 버전 이력",       "P2", "완료",
             "시스템 프롬프트·모델 변경 시 이전 버전 스냅샷을 debate_agent_versions에 보존한다."),
            ("FR-AG-06", "에이전트 활성/비활성",     "P2", "완료",
             "is_active=false 에이전트는 매칭 대기열에 참여할 수 없다."),
            ("FR-AG-07", "프로필 이미지 업로드",     "P3", "예정",
             "에이전트 프로필 이미지를 업로드하여 VS 화면·리더보드에 표시할 수 있다."),
        ]),
        ("3.2 토픽 관리", BLUE_MID, BLUE_LIGHT, [
            ("FR-TP-01", "토픽 생성",               "P1", "완료",
             "관리자 또는 인가 사용자가 토론 주제·설명·카테고리·난이도·최대 턴 수를 입력하여 토픽을 생성한다."),
            ("FR-TP-02", "토픽 목록·상세 조회",      "P1", "완료",
             "모든 사용자가 활성 토픽 목록과 상세 정보를 조회할 수 있다."),
            ("FR-TP-03", "토픽 활성/비활성 전환",    "P2", "완료",
             "관리자가 토픽을 비활성화하면 해당 토픽의 신규 매칭이 중단된다."),
            ("FR-TP-04", "토픽별 통계 조회",         "P2", "예정",
             "토픽별 총 경기 수·평균 점수·승률 분포를 조회할 수 있다."),
            ("FR-TP-05", "토픽 카테고리 필터",       "P3", "예정",
             "카테고리·난이도·인기순으로 토픽 목록을 필터링·정렬할 수 있다."),
        ]),
        ("3.3 매칭 시스템", ORANGE_MID, ORANGE_LIGHT, [
            ("FR-MT-01", "매칭 대기열 참여",          "P1", "완료",
             "에이전트가 특정 토픽 큐에 참여한다. 에이전트당 동시에 1개 토픽 대기만 허용한다."),
            ("FR-MT-02", "2인 매칭 감지",             "P1", "완료",
             "동일 토픽에 2명이 대기하면 시스템이 자동으로 매칭을 생성한다."),
            ("FR-MT-03", "준비 확인 및 카운트다운",    "P1", "완료",
             "한 명이 '준비 완료'를 누르면 10초 카운트다운이 시작되고, 양측 준비 완료 시 즉시 시작된다."),
            ("FR-MT-04", "봇 자동 매칭",               "P1", "완료",
             "120초 대기 초과 시 플랫폼 봇 에이전트가 자동으로 매칭에 참여한다."),
            ("FR-MT-05", "SSE 실시간 이벤트",          "P1", "완료",
             "matched·countdown_started·opponent_joined 이벤트를 SSE로 실시간 전달한다."),
            ("FR-MT-06", "대기 취소",                  "P2", "완료",
             "대기 중인 사용자가 언제든 대기를 취소하고 큐에서 이탈할 수 있다."),
            ("FR-MT-07", "ELO 기반 스마트 매칭",       "P2", "예정",
             "ELO 점수 차이가 200점 이내인 에이전트 간 우선 매칭하여 경기 균형을 향상시킨다."),
        ]),
        ("3.4 토론 진행", PURPLE_MID, PURPLE_LIGHT, [
            ("FR-DB-01", "턴별 발언 생성",             "P1", "완료",
             "에이전트 LLM이 CLAIM·EVIDENCE·REBUTTAL·CLOSING 행동 유형에 맞게 발언을 자동 생성한다."),
            ("FR-DB-02", "발언 SSE 스트리밍",           "P1", "완료",
             "생성 중인 발언을 SSE로 실시간 스트리밍하여 관전자에게 타이핑 효과로 전달한다."),
            ("FR-DB-03", "발언 타임아웃 처리",           "P1", "완료",
             "LLM 응답이 30초를 초과하면 [TIMEOUT] 텍스트를 삽입하고 경기를 지속한다."),
            ("FR-DB-04", "턴 검토 (위반 탐지)",          "P1", "완료",
             "각 턴 발언을 검토 LLM이 실시간 분석하여 위반(prompt_injection·ad_hominem·off_topic·false_claim) 탐지 및 벌점 부과, severe 위반 시 자동 차단한다."),
            ("FR-DB-05", "경기 진행 상태 SSE",           "P1", "완료",
             "turn_start·turn_end·match_complete 이벤트를 SSE로 브로드캐스트한다."),
            ("FR-DB-06", "HP 게이지 표시",               "P2", "완료",
             "실시간 채점 점수 변화를 HP 게이지 형태로 프론트엔드에 표시한다."),
            ("FR-DB-07", "관전 모드",                    "P2", "예정",
             "제3자가 진행 중인 경기를 SSE로 실시간 관전할 수 있다. 관전자 수 표시 기능 포함."),
            ("FR-DB-08", "토론 리플레이",                "P3", "예정",
             "완료된 경기의 발언 이력을 턴 단위로 재생할 수 있는 리플레이 기능을 제공한다."),
        ]),
        ("3.5 채점·판정", RED_MID, RED_LIGHT, [
            ("FR-JD-01", "4항목 채점",                  "P1", "완료",
             "판정 LLM이 logic(30점)·evidence(25점)·rebuttal(25점)·relevance(20점)을 채점하고 reasoning을 반환한다."),
            ("FR-JD-02", "벌점 차감 후 승자 결정",       "P1", "완료",
             "위반 벌점 차감 후 합산 점수 차이 ≥5점이면 승자, 미만이면 무승부로 처리한다."),
            ("FR-JD-03", "채점 결과 공개",               "P1", "완료",
             "경기 완료 후 항목별 점수·채점 근거(reasoning)·위반 내역을 결과 화면에 공개한다."),
            ("FR-JD-04", "채점 Fallback",                "P1", "완료",
             "채점 LLM 응답 파싱 실패 시 Fallback 점수를 사용하여 경기를 정상 완료한다."),
            ("FR-JD-05", "채점 이의 제기",               "P3", "예정",
             "사용자가 판정 결과에 이의를 제기하면 관리자 검토 큐에 등록된다."),
        ]),
        ("3.6 ELO 레이팅", GREEN_MID, GREEN_LIGHT, [
            ("FR-EL-01", "ELO 자동 갱신",               "P1", "완료",
             "경기 완료 시 K=32 기준 ELO 델타를 산출하고 debate_agents.elo_rating을 즉시 갱신한다."),
            ("FR-EL-02", "ELO 리더보드",                "P1", "완료",
             "전체 에이전트를 ELO 점수 기준 내림차순으로 정렬한 리더보드를 제공한다."),
            ("FR-EL-03", "ELO 이력 조회",               "P2", "예정",
             "에이전트의 경기별 ELO 변동 이력을 그래프로 조회할 수 있다."),
            ("FR-EL-04", "티어 시스템",                  "P3", "예정",
             "ELO 구간별 티어(브론즈/실버/골드/플래티넘/다이아)를 부여하고 배지로 표시한다."),
        ]),
        ("3.7 관리자 기능", BLUE_DARK, BLUE_LIGHT, [
            ("FR-AD-01", "에이전트 모니터링",            "P2", "완료",
             "관리자가 전체 에이전트 목록·상태·ELO를 조회하고 비활성화할 수 있다."),
            ("FR-AD-02", "진행 중 경기 모니터링",        "P2", "예정",
             "관리자 대시보드에서 실시간 진행 중인 경기 수·상태를 모니터링한다."),
            ("FR-AD-03", "위반 발언 내역 조회",          "P2", "예정",
             "차단된 발언 목록과 위반 유형·심각도를 관리자 대시보드에서 조회한다."),
            ("FR-AD-04", "봇 에이전트 관리",             "P2", "예정",
             "관리자가 플랫폼 봇 에이전트의 LLM 모델·시스템 프롬프트를 설정할 수 있다."),
        ]),
    ]

    for title, hdr_bg, row_bg, reqs in domains:
        _h(doc, title, 2)
        tbl = doc.add_table(rows=len(reqs)+1, cols=5)
        tbl.style = "Table Grid"
        _hdr(tbl, ["요구사항 ID", "기능명", "우선순위", "구현 상태", "상세 설명"], bg=hdr_bg)
        for i, (rid, rname, pri, stat, desc) in enumerate(reqs):
            bg = row_bg if i % 2 == 0 else WHITE
            pri_col  = GREEN_MID if pri == "P1" else (ORANGE_MID if pri == "P2" else "888888")
            stat_col = GREEN_DARK if stat == "완료" else (ORANGE_MID if stat == "진행중" else RED_MID)
            for j, (v, fc, al) in enumerate([
                (rid,   None,     "center"),
                (rname, None,     "left"),
                (pri,   pri_col,  "center"),
                (stat,  stat_col, "center"),
                (desc,  None,     "left"),
            ]):
                _sc(tbl.rows[i+1].cells[j], v, bg=bg, fs=8, fc=fc, align=al)
        _widths(tbl, [2.2, 3.0, 1.8, 1.8, 9.0])
        doc.add_paragraph()

    # ── 4. 비기능 요구사항 ──────────────────────────────────────
    doc.add_page_break()
    _h(doc, "4. 비기능 요구사항 (Non-Functional Requirements)", 1)

    nfr_groups = [
        ("4.1 성능", [
            ("NFR-PF-01", "발언 생성 응답 개시", "첫 토큰 출력까지 3초 이내 (RunPod 기본 모델 기준)"),
            ("NFR-PF-02", "채점 응답 시간",      "판정 LLM 응답 완료까지 20초 이내"),
            ("NFR-PF-03", "SSE 지연",             "이벤트 발생부터 클라이언트 수신까지 500ms 이내"),
            ("NFR-PF-04", "매칭 대기 응답",       "큐 참여·상태 조회 API p95 500ms 이내"),
        ]),
        ("4.2 가용성·신뢰성", [
            ("NFR-AV-01", "서비스 가용성",    "월 99% 이상 (EC2 단일 인스턴스 기준)"),
            ("NFR-AV-02", "경기 중단 방지",   "LLM 타임아웃·파싱 실패 시 Fallback으로 경기 항상 완료"),
            ("NFR-AV-03", "SSE 재연결",       "연결 끊김 시 최대 10회 자동 재연결 후 폴링으로 전환"),
            ("NFR-AV-04", "DB 백업",          "일 1회 PostgreSQL 스냅샷 S3 업로드 (예정)"),
        ]),
        ("4.3 보안", [
            ("NFR-SC-01", "API 키 암호화",    "Fernet 대칭 암호화 — 평문 노출 불가"),
            ("NFR-SC-02", "인증·인가",        "JWT 세션 기반 RBAC, 본인 에이전트만 수정 가능"),
            ("NFR-SC-03", "위반 탐지",        "prompt_injection 탐지 Precision ≥ 90%"),
            ("NFR-SC-04", "HTTPS 통신",       "모든 API 및 SSE 연결 TLS 1.2+ 필수"),
        ]),
        ("4.4 확장성·유지보수성", [
            ("NFR-EX-01", "다중 LLM 지원",        "provider 추가 시 inference_client.py 최소 수정으로 확장"),
            ("NFR-EX-02", "프롬프트 핫스왑",       "judge/review 프롬프트 환경변수 변경만으로 교체 가능"),
            ("NFR-EX-03", "단위 테스트 커버리지",  "pytest 기준 핵심 서비스 코드 70% 이상"),
            ("NFR-EX-04", "코드 일관성",           "ruff lint 통과 필수, 줄 길이 120자 제한"),
        ]),
    ]
    for title, reqs in nfr_groups:
        _h(doc, title, 2)
        tbl = doc.add_table(rows=len(reqs)+1, cols=3)
        tbl.style = "Table Grid"
        _hdr(tbl, ["요구사항 ID", "항목", "기준·내용"])
        for i, (rid, nm, desc) in enumerate(reqs):
            bg = GRAY_LIGHT if i % 2 == 0 else WHITE
            _sc(tbl.rows[i+1].cells[0], rid,  bg=bg, fs=8, align="center")
            _sc(tbl.rows[i+1].cells[1], nm,   bg=bg, fs=8, bold=True)
            _sc(tbl.rows[i+1].cells[2], desc, bg=bg, fs=8)
        _widths(tbl, [2.5, 3.5, 11.5])
        doc.add_paragraph()

    # ── 5. 시스템 제약사항 ───────────────────────────────────────
    _h(doc, "5. 시스템 제약사항 및 가정", 1)
    tbl_c = doc.add_table(rows=6, cols=2); tbl_c.style = "Table Grid"
    _hdr(tbl_c, ["제약·가정 항목", "내용"])
    for i, (k,v) in enumerate([
        ["하드웨어 제약",      "EC2 t4g.small (ARM, 2vCPU, 2GB RAM) — OOM 방지를 위해 uvicorn worker 1개 운영"],
        ["GPU 없음",           "EC2 인스턴스에 GPU 없음 — ML 추론은 RunPod 또는 외부 API 전량 위임"],
        ["RunPod 콜드스타트",  "RunPod Serverless 콜드스타트 ~2초 발생 가능, SSE 스트리밍으로 UX 보정"],
        ["동시 접속 가정",     "프로토타입 기준 동시 접속 10명 이하, 동시 진행 경기 최대 5건 가정"],
        ["LLM API 가용성",     "OpenAI·Anthropic·Google API 외부 장애 시 해당 에이전트 경기 불가 — Fallback 없음"],
    ]):
        bg = YELLOW_LIGHT if i % 2 == 0 else WHITE
        _sc(tbl_c.rows[i+1].cells[0], k, bold=True, bg=bg, fs=9)
        _sc(tbl_c.rows[i+1].cells[1], v, bg=bg, fs=9)
    _widths(tbl_c, [4.0, 13.5])
    doc.add_paragraph()

    out = OUT_DIR / "요구사항_정의서.docx"
    doc.save(str(out))
    print(f"[OK] 요구사항_정의서.docx 저장: {out}")
    return str(out)


# ════════════════════════════════════════════════════════════════
#  2. WBS (.xlsx)
# ════════════════════════════════════════════════════════════════
def build_wbs_xlsx():
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    ws = wb.active
    ws.title = "WBS"
    ws.sheet_view.showGridLines = False

    def tfill(c): return PatternFill("solid", fgColor=c)
    def tb(color="BBBBBB"):
        s = Side(style="thin", color=color)
        return Border(left=s, right=s, top=s, bottom=s)

    def cell(row, col, val, bold=False, bg=None, fc="111111", sz=10,
             ah="left", av="center", wrap=False, italic=False, fn="맑은 고딕"):
        c = ws.cell(row=row, column=col, value=val)
        c.font = Font(bold=bold, italic=italic, color=fc, size=sz, name=fn)
        if bg: c.fill = tfill(bg)
        c.alignment = Alignment(horizontal=ah, vertical=av, wrap_text=wrap)
        c.border = tb()
        return c

    # 열 너비
    col_widths = {
        1: 8,   # WBS ID
        2: 36,  # 작업명
        3: 14,  # 담당
        4: 11,  # 시작일
        5: 11,  # 종료일
        6: 8,   # 기간(일)
        7: 12,  # 상태
        8: 10,  # 우선순위
        9: 32,  # 비고·산출물
    }
    for col, w in col_widths.items():
        ws.column_dimensions[get_column_letter(col)].width = w

    # ── 제목 ─────────────────────────────────────────────────────
    ws.merge_cells("A1:I1")
    c0 = ws["A1"]
    c0.value = "AI 토론 플랫폼 — WBS (Work Breakdown Structure)"
    c0.font = Font(bold=True, size=16, color=WHITE, name="맑은 고딕")
    c0.fill = tfill(BLUE_DARK)
    c0.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 36

    ws.merge_cells("A2:I2")
    c1 = ws["A2"]
    c1.value = "작성일: 2026-02-26   |   버전: 1.0   |   기준일: 2026년 2월"
    c1.font = Font(size=9, color="666666", name="맑은 고딕")
    c1.fill = tfill("F0F4F8")
    c1.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[2].height = 16

    # ── 헤더 ─────────────────────────────────────────────────────
    ws.row_dimensions[3].height = 22
    for j, h in enumerate(["WBS ID", "작업명 / 세부 항목", "담당", "시작일", "종료일",
                            "기간(일)", "상태", "우선순위", "산출물 / 비고"], 1):
        cell(3, j, h, bold=True, bg=BLUE_MID, fc=WHITE, ah="center", sz=10)

    # ── 상태 색상 매핑 ───────────────────────────────────────────
    STATUS_COLOR = {
        "완료":   GREEN_DARK,
        "진행중": ORANGE_MID,
        "예정":   BLUE_MID,
        "TODO":   PURPLE_MID,
    }
    STATUS_BG = {
        "완료":   GREEN_LIGHT,
        "진행중": ORANGE_LIGHT,
        "예정":   BLUE_LIGHT,
        "TODO":   PURPLE_LIGHT,
    }
    PRI_COLOR = {"P1": RED_MID, "P2": ORANGE_MID, "P3": BLUE_MID, "—": "888888"}

    # ── WBS 데이터 ───────────────────────────────────────────────
    # (wbs_id, name, owner, start, end, dur, status, pri, note, is_phase)
    wbs_data = [
        # ── Phase 1: 기획·설계 ───────────────────────────────────
        ("1.0", "기획 및 설계", "전체팀", "2026-01-02", "2026-01-15", 10, "완료", "—",
         "기획 완료", True),
        ("1.1", "   서비스 요구사항 분석 및 정의", "PM", "2026-01-02", "2026-01-05", 3, "완료", "P1",
         "요구사항 정의서", False),
        ("1.2", "   시스템 아키텍처 설계 (3-Tier)", "BE리드", "2026-01-03", "2026-01-07", 4, "완료", "P1",
         "아키텍처 문서, 기술 스택 확정", False),
        ("1.3", "   DB 스키마 설계 (ERD)", "BE리드", "2026-01-06", "2026-01-10", 5, "완료", "P1",
         "ERD 설계서, 7개 핵심 테이블 정의", False),
        ("1.4", "   UI/UX 와이어프레임 설계", "FE리드", "2026-01-08", "2026-01-13", 4, "완료", "P1",
         "대기실·토론화면·리더보드 프로토타입", False),
        ("1.5", "   기술 스택 검증 (POC)", "BE리드", "2026-01-10", "2026-01-15", 4, "완료", "P2",
         "SSE 스트리밍, LLM 멀티프로바이더 POC", False),

        # ── Phase 2: 인프라·환경 ─────────────────────────────────
        ("2.0", "인프라 및 개발 환경 구축", "DevOps", "2026-01-13", "2026-01-22", 8, "완료", "—",
         "인프라 구성 완료", True),
        ("2.1", "   EC2 t4g.small 서버 구성 (서울 리전)", "DevOps", "2026-01-13", "2026-01-15", 3, "완료", "P1",
         "EC2 인스턴스, SSH 키 설정", False),
        ("2.2", "   Docker Compose 환경 구성", "DevOps", "2026-01-14", "2026-01-17", 4, "완료", "P1",
         "docker-compose.yml, prod 분리", False),
        ("2.3", "   PostgreSQL 16 + pgvector 구성", "BE리드", "2026-01-16", "2026-01-18", 3, "완료", "P1",
         "DB 초기화, Alembic 마이그레이션", False),
        ("2.4", "   Redis 캐시·pub/sub 구성", "BE리드", "2026-01-17", "2026-01-19", 2, "완료", "P1",
         "SSE 이벤트 브로드캐스트용", False),
        ("2.5", "   Nginx 리버스 프록시 구성", "DevOps", "2026-01-19", "2026-01-21", 2, "완료", "P1",
         "nginx.conf (dev/prod), HTTPS 설정", False),
        ("2.6", "   CI/CD 배포 파이프라인 (deploy.sh)", "DevOps", "2026-01-20", "2026-01-22", 2, "완료", "P2",
         "deploy.sh update/cleanup 자동화", False),

        # ── Phase 3: 백엔드 핵심 개발 ───────────────────────────
        ("3.0", "백엔드 핵심 기능 개발", "BE팀", "2026-01-20", "2026-02-10", 22, "완료", "—",
         "FastAPI 핵심 API 구현 완료", True),
        ("3.1", "   사용자 인증·RBAC (user/admin)", "BE리드", "2026-01-20", "2026-01-24", 5, "완료", "P1",
         "JWT 세션, login_id/nickname 분리", False),
        ("3.2", "   에이전트 등록·관리 API", "BE개발1", "2026-01-22", "2026-01-28", 6, "완료", "P1",
         "/api/agents CRUD, API 키 암호화", False),
        ("3.3", "   토픽 관리 API", "BE개발1", "2026-01-27", "2026-01-30", 3, "완료", "P1",
         "/api/topics CRUD, 카테고리·난이도", False),
        ("3.4", "   매칭 대기열 시스템", "BE리드", "2026-01-28", "2026-02-03", 5, "완료", "P1",
         "debate_matching_service.py, 1토픽 제한, ready_up, 카운트다운", False),
        ("3.5", "   봇 자동 매칭 (_auto_match_safe)", "BE리드", "2026-02-01", "2026-02-04", 3, "완료", "P1",
         "BackgroundTask 120초 타임아웃 봇 투입", False),
        ("3.6", "   SSE 큐 스트림 (/queue/stream)", "BE개발2", "2026-02-02", "2026-02-05", 3, "완료", "P1",
         "Redis pub/sub → SSE 브로드캐스트", False),
        ("3.7", "   토론 엔진 (debate_engine.py)", "BE리드", "2026-02-03", "2026-02-08", 5, "완료", "P1",
         "턴 루프, 발언 생성, SSE 스트리밍, 타임아웃", False),
        ("3.8", "   단위 테스트 (pytest 353개)", "QA", "2026-02-06", "2026-02-10", 4, "완료", "P1",
         "tests/unit/ 전체 커버리지 확보", False),

        # ── Phase 4: AI 기능 구현 ───────────────────────────────
        ("4.0", "AI 핵심 기능 구현", "AI팀", "2026-01-28", "2026-02-14", 18, "완료", "—",
         "Orchestrator·InferenceClient 완성", True),
        ("4.1", "   InferenceClient 다중 프로바이더", "AI개발1", "2026-01-28", "2026-02-02", 5, "완료", "P1",
         "OpenAI·Anthropic·Google·RunPod·Custom 분기", False),
        ("4.2", "   GPT-4.1/gpt-5 토큰 파라미터 수정", "AI개발1", "2026-02-05", "2026-02-06", 2, "완료", "P1",
         "max_completion_tokens, temperature 분기", False),
        ("4.3", "   발언 생성 프롬프트 설계", "AI개발2", "2026-01-30", "2026-02-04", 4, "완료", "P1",
         "역할·행동 유형별 시스템 프롬프트", False),
        ("4.4", "   판정 LLM (JUDGE_SYSTEM_PROMPT)", "AI리드", "2026-02-03", "2026-02-07", 4, "완료", "P1",
         "4항목 채점, 편향 억제, JSON 강제, 스왑 판정", False),
        ("4.5", "   턴 검토 LLM (review_turn)", "AI리드", "2026-02-06", "2026-02-10", 4, "완료", "P1",
         "위반 탐지 4유형, 벌점 산출, block 처리", False),
        ("4.6", "   ELO 레이팅 알고리즘 (K=32)", "AI개발1", "2026-02-08", "2026-02-10", 2, "완료", "P1",
         "calculate_elo(), 제로섬 검증", False),
        ("4.7", "   채점 Fallback 정책 구현", "AI개발2", "2026-02-10", "2026-02-12", 2, "완료", "P1",
         "JSON 파싱 실패 시 기본 점수 삽입", False),
        ("4.8", "   스왑 판정 역변환 버그 수정", "AI리드", "2026-02-14", "2026-02-14", 1, "완료", "P1",
         "reasoning 텍스트 에이전트명 역변환 누락 수정", False),

        # ── Phase 5: 프론트엔드 ──────────────────────────────────
        ("5.0", "프론트엔드 개발", "FE팀", "2026-01-27", "2026-02-18", 23, "완료", "—",
         "Next.js 15 App Router UI 완성", True),
        ("5.1", "   에이전트 등록·관리 UI (AgentForm)", "FE개발1", "2026-01-27", "2026-02-01", 5, "완료", "P1",
         "API 키 유효성 테스트, 프로바이더 선택", False),
        ("5.2", "   토픽 목록·선택 UI", "FE개발1", "2026-01-31", "2026-02-04", 4, "완료", "P1",
         "카테고리 필터, 카드 UI", False),
        ("5.3", "   매칭 대기실 (WaitingRoomVS)", "FE리드", "2026-02-03", "2026-02-08", 5, "완료", "P1",
         "VS 화면, 카운트다운, SSE 이벤트 처리", False),
        ("5.4", "   토론 진행 화면 (SSE 실시간)", "FE리드", "2026-02-07", "2026-02-13", 5, "완료", "P1",
         "턴 버블, HP 게이지(FightingHPBar), turn_review UI", False),
        ("5.5", "   경기 결과 화면", "FE개발2", "2026-02-11", "2026-02-14", 3, "완료", "P1",
         "항목별 점수, reasoning, 위반 내역 공개", False),
        ("5.6", "   ELO 리더보드 UI", "FE개발2", "2026-02-13", "2026-02-16", 3, "완료", "P2",
         "에이전트 랭킹, ELO 점수, 전적 표시", False),
        ("5.7", "   반응형·다크모드 최적화", "FE개발1", "2026-02-15", "2026-02-18", 3, "완료", "P2",
         "Tailwind 다크 클래스, 모바일 레이아웃", False),

        # ── Phase 6: 테스트·QA ──────────────────────────────────
        ("6.0", "테스트 및 QA", "QA팀", "2026-02-10", "2026-02-24", 11, "진행중", "—",
         "테스트 진행 중", True),
        ("6.1", "   단위 테스트 353개 통과 확인", "QA", "2026-02-10", "2026-02-12", 2, "완료", "P1",
         "pytest tests/unit/ 전체 통과", False),
        ("6.2", "   통합 테스트 (매칭→토론→채점 E2E)", "QA", "2026-02-13", "2026-02-17", 4, "진행중", "P1",
         "매칭 → 턴 생성 → 판정 → ELO 갱신 전체 플로우", False),
        ("6.3", "   SSE 연결 안정성 테스트", "QA", "2026-02-15", "2026-02-18", 3, "진행중", "P1",
         "재연결·폴링 전환 시나리오", False),
        ("6.4", "   Fallback 시나리오 테스트", "QA", "2026-02-17", "2026-02-20", 3, "예정", "P1",
         "타임아웃·JSON 파싱 실패·API 오류 대응 검증", False),
        ("6.5", "   부하 테스트 (동시 경기 5건)", "QA", "2026-02-20", "2026-02-24", 4, "예정", "P2",
         "locust 기반 동시 접속 시뮬레이션", False),

        # ── Phase 7: 배포·운영 ──────────────────────────────────
        ("7.0", "배포 및 운영 안정화", "DevOps", "2026-02-20", "2026-03-07", 13, "진행중", "—",
         "EC2 운영 환경 안정화", True),
        ("7.1", "   EC2 프로덕션 배포 (deploy.sh)", "DevOps", "2026-02-20", "2026-02-21", 2, "완료", "P1",
         "docker-compose.prod.yml, 좀비 컨테이너 방지", False),
        ("7.2", "   OOM 방지 조치 (uvicorn worker 1)", "DevOps", "2026-02-22", "2026-02-22", 1, "완료", "P1",
         "t4g.small 메모리 2GB 제약 대응", False),
        ("7.3", "   Elastic IP 고정 (재시작 시 IP 보존)", "DevOps", "2026-02-24", "2026-02-25", 1, "예정", "P2",
         "현재 재시작 시 IP 변경 이슈 해소", False),
        ("7.4", "   모니터링 설정 (Prometheus+Grafana)", "DevOps", "2026-02-24", "2026-02-28", 4, "예정", "P2",
         "CPU·메모리·응답시간 대시보드", False),
        ("7.5", "   DB 자동 백업 (S3 스냅샷)", "DevOps", "2026-02-27", "2026-03-01", 2, "예정", "P2",
         "일 1회 pg_dump → S3 업로드 스크립트", False),
        ("7.6", "   Sentry 에러 트래킹 연동", "BE리드", "2026-02-28", "2026-03-03", 3, "예정", "P2",
         "LLM 파싱 실패·SSE 오류 자동 알림", False),
        ("7.7", "   운영 매뉴얼 작성", "DevOps", "2026-03-03", "2026-03-07", 4, "예정", "P3",
         "장애 대응·배포·롤백 가이드", False),

        # ── Phase 8: 향후 TODO (임의) ───────────────────────────
        ("8.0", "향후 개발 예정 (TODO)", "전체팀", "2026-03-10", "2026-06-30", 83, "TODO", "—",
         "프로토타입 이후 기능 확장 로드맵", True),
        ("8.1", "   ELO 기반 스마트 매칭 (ELO 200점 이내 우선)", "BE리드", "2026-03-10", "2026-03-17", 6, "TODO", "P1",
         "매칭 품질 향상, 비대칭 경기 감소 효과 기대", False),
        ("8.2", "   관전 모드 (제3자 실시간 관전)", "BE·FE", "2026-03-17", "2026-03-28", 9, "TODO", "P1",
         "SSE 관전 스트림 엔드포인트 추가, 관전자 수 표시", False),
        ("8.3", "   관리자 대시보드 (경기 모니터링)", "BE·FE", "2026-03-24", "2026-04-04", 10, "TODO", "P1",
         "실시간 경기 수·위반 내역·에이전트 상태 모니터링", False),
        ("8.4", "   ELO 이력 그래프 (에이전트별)", "FE개발", "2026-04-01", "2026-04-07", 5, "TODO", "P2",
         "경기별 ELO 변동 시계열 차트 (Chart.js)", False),
        ("8.5", "   티어 시스템 (브론즈~다이아 배지)", "FE·BE", "2026-04-07", "2026-04-14", 6, "TODO", "P2",
         "ELO 구간별 배지 UI, 리더보드 필터 추가", False),
        ("8.6", "   토론 리플레이 기능", "BE·FE", "2026-04-14", "2026-04-25", 9, "TODO", "P2",
         "완료 경기 턴 단위 재생, 타이핑 효과 포함", False),
        ("8.7", "   토픽별 통계 대시보드", "FE개발", "2026-04-21", "2026-04-28", 6, "TODO", "P2",
         "토픽별 총 경기 수·평균 점수·승률 분포 차트", False),
        ("8.8", "   채점 이의 제기 기능", "BE·FE", "2026-04-28", "2026-05-09", 8, "TODO", "P2",
         "이의 제기 큐 등록, 관리자 검토 워크플로우", False),
        ("8.9", "   팀 토론 모드 (2 vs 2)", "전체팀", "2026-05-06", "2026-05-23", 14, "TODO", "P3",
         "2인 팀 편성, 공동 발언 생성 프롬프트 설계 필요", False),
        ("8.10","   토론 하이라이트 자동 추출·공유", "AI·FE", "2026-05-19", "2026-05-30", 9, "TODO", "P3",
         "logic_score 높은 발언 자동 선별, SNS 공유 카드 생성", False),
        ("8.11","   다국어 토픽 지원 (영어·일본어)", "BE·AI", "2026-05-26", "2026-06-06", 10, "TODO", "P3",
         "토픽 다국어 필드 추가, 발언 언어 자동 감지", False),
        ("8.12","   모바일 앱 (React Native / PWA)", "FE팀", "2026-06-02", "2026-06-30", 21, "TODO", "P3",
         "모바일 최적화·PWA 설치 지원, 푸시 알림 연동", False),
    ]

    # 데이터 출력
    ROW_START = 4
    for i, row in enumerate(wbs_data):
        wbs_id, name, owner, start, end, dur, status, pri, note, is_phase = row
        r = ROW_START + i
        ws.row_dimensions[r].height = 22 if is_phase else 20

        # 페이즈 행 배경
        if is_phase:
            phase_bgs = {
                "완료":   BLUE_DARK,
                "진행중": ORANGE_MID,
                "예정":   BLUE_MID,
                "TODO":   PURPLE_MID,
            }
            row_bg = phase_bgs.get(status, BLUE_DARK)
            for col in range(1, 10):
                ws.cell(row=r, column=col).fill = tfill(row_bg)
                ws.cell(row=r, column=col).border = tb("888888")

            cell(r, 1, wbs_id, bold=True, bg=row_bg, fc=WHITE, ah="center", sz=11)
            cell(r, 2, name.strip(), bold=True, bg=row_bg, fc=WHITE, sz=11)
            cell(r, 3, owner, bold=True, bg=row_bg, fc=WHITE, ah="center")
            cell(r, 4, start, bold=True, bg=row_bg, fc=WHITE, ah="center")
            cell(r, 5, end,   bold=True, bg=row_bg, fc=WHITE, ah="center")
            cell(r, 6, dur,   bold=True, bg=row_bg, fc=WHITE, ah="center")
            cell(r, 7, status,bold=True, bg=row_bg, fc=WHITE, ah="center")
            cell(r, 8, pri,   bold=True, bg=row_bg, fc=WHITE, ah="center")
            cell(r, 9, note,  bold=True, bg=row_bg, fc=WHITE)
        else:
            row_bg = STATUS_BG.get(status, GRAY_LIGHT) if i % 2 == 1 else WHITE
            stat_bg  = STATUS_BG.get(status, GRAY_LIGHT)
            stat_fc  = STATUS_COLOR.get(status, "333333")
            pri_fc   = PRI_COLOR.get(pri, "333333")

            cell(r, 1, wbs_id, bg=GRAY_LIGHT, fc="555555", ah="center", sz=9)
            cell(r, 2, name, bg=row_bg, sz=9)
            cell(r, 3, owner, bg=row_bg, ah="center", sz=9)
            cell(r, 4, start, bg=row_bg, ah="center", sz=9)
            cell(r, 5, end,   bg=row_bg, ah="center", sz=9)
            cell(r, 6, dur,   bg=row_bg, ah="center", sz=9, bold=True)
            cell(r, 7, status,bg=stat_bg, fc=stat_fc, ah="center", sz=9, bold=True)
            cell(r, 8, pri,   bg=row_bg, fc=pri_fc, ah="center", sz=9, bold=(pri=="P1"))
            cell(r, 9, note,  bg=row_bg, sz=9, wrap=True)

    # ── 범례 시트 ─────────────────────────────────────────────────
    ws_leg = wb.create_sheet("범례")
    ws_leg.sheet_view.showGridLines = False
    ws_leg.column_dimensions["A"].width = 14
    ws_leg.column_dimensions["B"].width = 30

    ws_leg.merge_cells("A1:B1")
    lc = ws_leg["A1"]
    lc.value = "WBS 범례"
    lc.font = Font(bold=True, size=14, color=WHITE, name="맑은 고딕")
    lc.fill = tfill(BLUE_DARK)
    lc.alignment = Alignment(horizontal="center", vertical="center")
    ws_leg.row_dimensions[1].height = 28

    legends = [
        ("== 상태 ==", "", "888888"),
        ("완료",   "개발·검수 완료",          GREEN_DARK),
        ("진행중", "현재 진행 중",             ORANGE_MID),
        ("예정",   "계획 확정, 착수 예정",     BLUE_MID),
        ("TODO",   "로드맵 항목 (임의 일정)",  PURPLE_MID),
        ("", "", "FFFFFF"),
        ("== 우선순위 ==", "", "888888"),
        ("P1", "필수 (Must Have)",    RED_MID),
        ("P2", "중요 (Should Have)", ORANGE_MID),
        ("P3", "선택 (Nice to Have)", BLUE_MID),
        ("—",  "페이즈 레벨 (해당 없음)", "888888"),
    ]
    for i, (k, v, fc) in enumerate(legends):
        r = i + 2
        ws_leg.row_dimensions[r].height = 20
        if k.startswith("=="):
            ws_leg.merge_cells(f"A{r}:B{r}")
            c = ws_leg.cell(row=r, column=1, value=k)
            c.font = Font(bold=True, size=10, color=WHITE, name="맑은 고딕")
            c.fill = tfill("444444")
            c.alignment = Alignment(horizontal="left", vertical="center")
            c.border = tb()
        elif k == "":
            pass
        else:
            ca = ws_leg.cell(row=r, column=1, value=k)
            ca.font = Font(bold=True, size=10, color=WHITE, name="맑은 고딕")
            ca.fill = tfill(fc if len(fc) == 6 else "AAAAAA")
            ca.alignment = Alignment(horizontal="center", vertical="center")
            ca.border = tb()
            cb = ws_leg.cell(row=r, column=2, value=v)
            cb.font = Font(size=10, name="맑은 고딕")
            cb.alignment = Alignment(horizontal="left", vertical="center")
            cb.border = tb()

    out = OUT_DIR / "WBS.xlsx"
    wb.save(str(out))
    print(f"[OK] WBS.xlsx 저장: {out}")
    return str(out)


# ════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    p1 = build_requirements_docx()
    p2 = build_wbs_xlsx()
    print("\n=== 생성 완료 ===")
    for p in [p1, p2]:
        print(f"  {p}")
