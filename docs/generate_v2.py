"""SKN21기 4Team (우공이산) 제출 문서 v2 생성 스크립트"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent))

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch
import numpy as np

from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from generate_docs import (
    new_doc, h1, h2, h3, para, code_para, add_table,
    add_image, fig_to_buf,
    C_BLUE, C_NAVY, C_GREEN, C_ORANGE, C_GRAY, C_RED, PALETTE,
    OUTPUT_DIR,
)

matplotlib.rcParams["font.family"] = "Malgun Gothic"
matplotlib.rcParams["axes.unicode_minus"] = False
plt.rcParams["figure.dpi"] = 130

TEAM_MEMBERS = [
    ("박수빈", "PM, 백엔드 로직 구현"),
    ("이성진", "데이터 수집 및 처리, 유저 시나리오 테스트"),
    ("이의정", "Frontend, 문서 작성, 유저 시나리오 작성"),
    ("정덕규", "오케스트레이션, 에이전트 설계 및 튜닝"),
]


# ── 공통: 표지 팀원 표 삽입 ────────────────────────────────────────────────────
def add_cover_info(doc):
    """표지 아래 팀 정보 + 팀원 역할 표를 삽입한다."""
    add_table(
        doc,
        ["이름", "역할"],
        TEAM_MEMBERS,
    )


# ── 문서 1: 기획서 v2 ─────────────────────────────────────────────────────────
def build_기획서_v2():
    try:
        doc = new_doc("SKN21기 4팀 기획서 v2")
        add_cover_info(doc)

        # 1. 프로젝트 개요
        h1(doc, "1. 프로젝트 개요")
        para(doc, (
            "서로 다른 LLM 에이전트들이 주어진 주제로 자율 토론을 벌이고, "
            "심판 AI가 판정하여 ELO 레이팅으로 순위를 매기는 실시간 관전 플랫폼."
        ))

        # 2. 시장 현황 및 문제점
        h1(doc, "2. 시장 현황 및 문제점")

        h2(doc, "2.1 기존 LLM 평가의 한계")
        para(doc, (
            "현재 주류 벤치마크(MMLU, HumanEval 등)는 단일 질문-정답 비교 방식으로 설계되어 있다. "
            "이 방식은 다음 세 가지 근본적 한계를 갖는다."
        ))
        add_table(
            doc,
            ["한계", "설명"],
            [
                ["단일 질문/정답 비교", "고정된 정답셋과 비교 → 창의적·논증적 능력 미측정"],
                ["다회전 논증 미측정", "여러 턴에 걸친 반박·근거 보강 능력 평가 불가"],
                ["상호작용 없음", "에이전트 간 동적 토론·전략적 반응 평가 불가"],
            ],
        )

        h2(doc, "2.2 시장 현황")
        para(doc, (
            "AI 대화형 서비스의 일평균 사용 시간은 2시간 40분으로 급증하고 있으며, "
            "LLM 서비스 경쟁이 심화되면서 모델 성능의 '공정한 비교 수단'에 대한 수요가 높아지고 있다. "
            "그러나 현존 플랫폼은 모두 수동 투표 또는 고정 모델 비교 방식에 머물러 있다."
        ))

        h2(doc, "2.3 경쟁 플랫폼 분석")
        add_table(
            doc,
            ["플랫폼", "특징", "한계"],
            [
                ["AIDebateArena (iOS)", "2 AI 배틀 + 관전자 평가", "프롬프트 커스터마이징 없음"],
                ["Gemicha (Android)", "Gemini vs ChatGPT + 20+ 페르소나", "크레딧 기반, 고정 모델"],
                ["LLM Debate Arena (오픈소스)", "ELO + 다중 심판 + SSE", "UI 미비, 에이전트 등록 불가"],
                ["Chatbot Arena", "쌍대 비교 투표", "수동 입력, 실시간 관전 불가"],
            ],
        )

        # 3. 차별점
        h1(doc, "3. 차별점")
        add_table(
            doc,
            ["차별점", "설명"],
            [
                [
                    "UGC 제작자 중심",
                    "Persona Pack(프롬프트+말투+금칙)이 경쟁 단위 — 모델이 아닌 프롬프트가 실력을 결정",
                ],
                [
                    "BYOK(Bring Your Own Key)",
                    "사용자 자신의 LLM API 키로 직접 에이전트를 구동, 특정 공급자에 종속되지 않음",
                ],
                [
                    "실시간 관전 + SSE 스트리밍",
                    "토큰 단위 타이핑 효과로 현장감 있는 토론 관전 경험 제공",
                ],
                [
                    "신뢰 설계",
                    "LLM 심판 + 페널티 시스템(regex 7종 + LLM 검토 4종)으로 공정성 보장",
                ],
            ],
        )

        # 4. ELO 랭킹 시스템 선택 근거
        h1(doc, "4. ELO 랭킹 시스템 선택 근거")

        h2(doc, "4.1 왜 랭킹 시스템이 필요한가")
        para(doc, (
            "모델·프롬프트 간 실력 비교를 객관화하고, 반복 대결을 통해 신뢰도를 누적하기 위해 "
            "레이팅 시스템이 필요하다. 단순 승률은 대결 상대 강도를 반영하지 못하므로 "
            "상대 강도를 가중하는 ELO 계열 방식이 적합하다."
        ))

        h2(doc, "4.2 왜 ELO인가")
        para(doc, (
            "본 시스템은 표준 ELO(승/패 이진)를 확장하여 LLM 심판 점수(0-100)를 "
            "이전량(transfer) 결정에 직접 활용한다. 이를 통해 '얼마나 압도적으로 이겼는가'가 "
            "레이팅 변동 폭에 반영된다."
        ))
        para(doc, "ELO 선택의 구체적 이유:")
        add_table(
            doc,
            ["이유", "설명"],
            [
                [
                    "1. 점수 직접 활용",
                    "LLM 심판 점수(0-100) → 이전량 결정. 표준 ELO(이진 결과)와의 핵심 차별점",
                ],
                [
                    "2. 단순성/투명성 우선",
                    "프로토타입 단계: 알고리즘 투명성이 Glicko-2·TrueSkill 대비 높음",
                ],
                [
                    "3. time decay 불필요",
                    "AI 에이전트는 비활동 기간에도 실력 불변 → Glicko-2 time decay 불필요",
                ],
                [
                    "4. 제로섬 설계",
                    "인플레이션 없이 장기 운영 가능 (이긴 쪽 증가 = 진 쪽 감소)",
                ],
            ],
        )

        h2(doc, "4.3 레이팅 시스템 비교")
        add_table(
            doc,
            ["시스템", "핵심 특징", "적합성", "채택 여부"],
            [
                ["ELO (확장)", "비대칭 K-factor + 점수 이전량", "AI 에이전트 특성에 최적", "채택"],
                ["Glicko-2", "불확실성(RD) + time decay", "장기 체스 리그에 최적", "미채택"],
                ["TrueSkill", "베이즈 추론, 팀 게임 최적", "복잡도 높음, 투명성 낮음", "미채택"],
                ["Bradley-Terry", "쌍대 비교 확률 모델", "투표 기반 시스템에 적합", "미채택"],
                ["단순 승률", "W/(W+L+D)", "상대 강도 미반영", "미채택"],
            ],
        )

        # 5. 핵심 기능
        h1(doc, "5. 핵심 기능")
        add_table(
            doc,
            ["기능", "설명"],
            [
                ["에이전트 생성", "이름, LLM 모델, API 키, 시스템 프롬프트 기반 에이전트 등록 및 버전 관리"],
                ["주제 관리", "토론 제목/배경/턴 수/시간 제한 설정, 상태 자동 전환(scheduled→open→closed)"],
                ["매치메이킹", "큐 기반 자동 매칭, 셀프 매칭 방지, 대기 초과 시 플랫폼 에이전트 자동 투입"],
                ["토론 엔진", "비동기 턴 실행, SSE 스트리밍, regex+LLM 이중 페널티 시스템"],
                ["ELO/티어 시스템", "비대칭 K-factor ELO 갱신, 7단계 티어(Iron~Master), 강등 보호"],
            ],
        )

        # 6. 시스템 구조
        h1(doc, "6. 시스템 구조 (3계층 아키텍처)")
        add_table(
            doc,
            ["계층", "역할", "주요 컴포넌트"],
            [
                ["정책/제어 계층", "인증, RBAC, 매치 상태 머신 관리", "FastAPI, JWT, debate_matching_service"],
                ["비즈니스/오케스트레이션 계층", "토론 실행, 검토, 판정, ELO 갱신", "debate_engine, debate_orchestrator"],
                ["데이터/인프라 계층", "상태 저장, 이벤트 전파, LLM 호출", "PostgreSQL, Redis Pub/Sub, InferenceClient"],
            ],
        )

        # 7. 기술 스택
        h1(doc, "7. 기술 스택")
        add_table(
            doc,
            ["역할", "기술"],
            [
                ["Backend", "Python 3.12 + FastAPI + SQLAlchemy async"],
                ["Frontend", "Next.js 15 + React 19 + Zustand"],
                ["Database", "PostgreSQL 16 (Docker)"],
                ["Cache / Pub·Sub", "Redis (Docker)"],
                ["LLM 공급자", "OpenAI / Anthropic / Google / RunPod Serverless"],
                ["Judge 모델", "gpt-4.1 (채점 품질·비용 균형 최우수)"],
                ["Review 모델", "gpt-5-nano (속도·비용·정확도 균형)"],
                ["Streaming", "SSE (Server-Sent Events)"],
                ["Infra", "AWS EC2 t4g.small (서울) + RunPod Serverless (미국)"],
            ],
        )

        # 8. 역할 분담
        h1(doc, "8. 역할 분담")
        add_table(
            doc,
            ["이름", "역할", "주요 기여"],
            [
                ["박수빈", "PM, 백엔드 로직 구현", "전체 일정 관리, FastAPI 엔드포인트, 서비스 레이어"],
                ["이성진", "데이터 수집 및 처리, 유저 시나리오 테스트", "벤치마크 데이터셋 구축, E2E 시나리오 검증"],
                ["이의정", "Frontend, 문서 작성, 유저 시나리오 작성", "Next.js UI, SSE 연동, 문서 전반"],
                ["정덕규", "오케스트레이션, 에이전트 설계 및 튜닝", "OptimizedOrchestrator, 페널티 시스템, ELO"],
            ],
        )

        out = OUTPUT_DIR / "SKN21기_4Team_기획서_v2.docx"
        doc.save(str(out))
        print(f"[OK] {out}")
    except Exception as e:
        print(f"[ERROR] 기획서_v2: {e}")


# ── 문서 2: 요구사항 정의서 v2 ───────────────────────────────────────────────
def build_요구사항_정의서_v2():
    try:
        doc = new_doc("SKN21기 4팀 요구사항 정의서 v2")
        add_cover_info(doc)

        # ID 체계 설명
        h1(doc, "1. 요구사항 ID 체계")
        para(doc, "요구사항 ID 형식: REQ-{카테고리}-{번호:03d}")
        add_table(
            doc,
            ["카테고리 코드", "설명"],
            [
                ["AGENT", "에이전트 관련 기능 요구사항"],
                ["MATCH", "매치/토론 실행 관련"],
                ["TOPIC", "토론 주제 관련"],
                ["UI", "프론트엔드/화면 관련"],
                ["SYS", "시스템/인프라/보안 관련"],
                ["AUTH", "인증/권한 관련"],
                ["RANK", "랭킹/ELO 관련"],
            ],
        )
        para(doc, "예시: REQ-AGENT-001 = 에이전트 관련 첫 번째 요구사항")

        # 에이전트 요구사항
        h1(doc, "2. 에이전트 (AGENT)")
        add_table(
            doc,
            ["ID", "우선순위", "요구사항", "설명"],
            [
                [
                    "REQ-AGENT-001", "필수", "에이전트 생성",
                    "사용자는 이름, LLM 모델, API 키, 시스템 프롬프트를 입력하여 토론 에이전트를 생성할 수 있어야 한다",
                ],
                [
                    "REQ-AGENT-002", "필수", "에이전트 수정",
                    "에이전트 수정 시 자동으로 새 버전이 생성되고 버전별 전적이 분리되어 관리되어야 한다",
                ],
                [
                    "REQ-AGENT-003", "필수", "에이전트 삭제",
                    "진행 중인 매치가 없을 때 에이전트를 삭제할 수 있어야 한다",
                ],
                [
                    "REQ-AGENT-004", "필수", "API 키 유효성 검증",
                    "에이전트 등록 전 API 키와 모델 ID의 실제 연결 가능 여부를 테스트할 수 있어야 한다",
                ],
                [
                    "REQ-AGENT-005", "필수", "버전 히스토리 조회",
                    "소유자는 에이전트의 모든 버전 목록과 각 버전의 전적을 조회할 수 있어야 한다",
                ],
                [
                    "REQ-AGENT-006", "권장", "시스템 프롬프트 공개 설정",
                    "에이전트 소유자는 시스템 프롬프트의 공개 여부를 설정할 수 있어야 한다",
                ],
                [
                    "REQ-AGENT-007", "권장", "플랫폼 크레딧 사용",
                    "BYOK API 키 없이 플랫폼 크레딧으로 에이전트를 운용할 수 있어야 한다",
                ],
            ],
        )

        # 매치 요구사항
        h1(doc, "3. 매치 (MATCH)")
        add_table(
            doc,
            ["ID", "우선순위", "요구사항", "설명"],
            [
                [
                    "REQ-MATCH-001", "필수", "큐 참가",
                    "에이전트는 토론 주제에 큐를 등록하여 상대를 기다릴 수 있어야 한다",
                ],
                [
                    "REQ-MATCH-002", "필수", "자동 매칭",
                    "큐에 2명이 도달하면 자동으로 매치가 생성되어야 한다",
                ],
                [
                    "REQ-MATCH-003", "필수", "대기 초과 자동 매칭",
                    "설정 시간(기본 60초) 초과 시 플랫폼 에이전트와 자동 매칭되어야 한다",
                ],
                [
                    "REQ-MATCH-004", "필수", "셀프 매칭 방지",
                    "동일 사용자의 에이전트끼리 매칭되지 않아야 한다",
                ],
                [
                    "REQ-MATCH-005", "필수", "SSE 실시간 스트리밍",
                    "토론 내용이 토큰 단위로 SSE 스트리밍되어야 한다",
                ],
                [
                    "REQ-MATCH-006", "필수", "턴 간 품질 검토",
                    "각 턴 발언은 패스트패스 또는 LLM 검토를 거쳐 위반 여부를 판정해야 한다",
                ],
                [
                    "REQ-MATCH-007", "필수", "LLM 심판 판정",
                    "토론 종료 후 LLM이 4개 항목(논리/근거/반박/주제적합성)으로 채점해야 한다",
                ],
                [
                    "REQ-MATCH-008", "필수", "ELO 갱신",
                    "매치 종료 후 양측 에이전트의 ELO 레이팅이 자동 갱신되어야 한다",
                ],
            ],
        )

        # 주제 요구사항
        h1(doc, "4. 주제 (TOPIC)")
        add_table(
            doc,
            ["ID", "우선순위", "요구사항", "설명"],
            [
                [
                    "REQ-TOPIC-001", "필수", "주제 생성",
                    "사용자는 토론 제목, 배경설명, 최대 턴 수, 시간 제한을 설정하여 주제를 생성할 수 있어야 한다",
                ],
                [
                    "REQ-TOPIC-002", "필수", "주제 목록 조회",
                    "최신순/인기순 정렬로 주제 목록을 페이지네이션하여 조회할 수 있어야 한다",
                ],
                [
                    "REQ-TOPIC-003", "권장", "주제 스케줄링",
                    "시작/종료 시각 기반으로 주제 상태가 자동 전환(scheduled→open→closed)되어야 한다",
                ],
            ],
        )

        # UI 요구사항
        h1(doc, "5. UI (UI)")
        add_table(
            doc,
            ["ID", "우선순위", "요구사항", "설명"],
            [
                [
                    "REQ-UI-001", "필수", "ELO 랭킹 페이지",
                    "에이전트 ELO 순위, 티어 배지, 승/패/무 통계를 조회할 수 있어야 한다",
                ],
                [
                    "REQ-UI-002", "필수", "실시간 관전",
                    "타이핑 효과로 LLM 응답이 실시간으로 표시되어야 한다",
                ],
                [
                    "REQ-UI-003", "필수", "논증 품질 시각화",
                    "LLM 검토 결과(점수, 위반 유형)가 각 턴 말풍선에 표시되어야 한다",
                ],
                [
                    "REQ-UI-004", "권장", "HP 바",
                    "양측 에이전트의 현재 점수가 HP 바 형태로 표시되어야 한다",
                ],
            ],
        )

        # 시스템 요구사항
        h1(doc, "6. 시스템 (SYS)")
        add_table(
            doc,
            ["ID", "우선순위", "요구사항", "설명"],
            [
                [
                    "REQ-SYS-001", "필수", "API 키 암호화",
                    "에이전트 API 키는 Fernet으로 암호화하여 저장해야 한다",
                ],
                [
                    "REQ-SYS-002", "필수", "동시성 안전",
                    "매치 생성 시 SELECT FOR UPDATE로 레이스 컨디션을 방지해야 한다",
                ],
                [
                    "REQ-SYS-003", "필수", "에이전트당 1 토픽 대기 제한",
                    "에이전트는 동시에 하나의 주제 큐에만 참가할 수 있어야 한다",
                ],
            ],
        )

        # 랭킹 요구사항
        h1(doc, "7. 랭킹 (RANK)")
        add_table(
            doc,
            ["ID", "우선순위", "요구사항", "설명"],
            [
                [
                    "REQ-RANK-001", "필수", "ELO 레이팅",
                    "매치 결과에 따라 비대칭 K-factor로 ELO가 갱신되어야 한다 (승자 K=40, 패자 K=24)",
                ],
                [
                    "REQ-RANK-002", "필수", "7단계 티어",
                    "ELO 구간별 티어(Iron~Master)가 자동 부여되어야 한다",
                ],
                [
                    "REQ-RANK-003", "권장", "강등 보호",
                    "티어 강등 시 보호 카운터(2회)가 제공되어야 한다",
                ],
            ],
        )

        out = OUTPUT_DIR / "SKN21기_4Team_요구사항_정의서_v2.docx"
        doc.save(str(out))
        print(f"[OK] {out}")
    except Exception as e:
        print(f"[ERROR] 요구사항_정의서_v2: {e}")


# ── 문서 3: 수집 데이터 명세서 v2 ────────────────────────────────────────────
def build_수집_데이터_명세서_v2():
    try:
        doc = new_doc("SKN21기 4팀 수집 데이터 명세서 v2")
        add_cover_info(doc)

        # 1. 개요
        h1(doc, "1. 개요")
        para(doc, (
            "본 명세서는 LLM 모델 선정, 오케스트레이터 최적화, 심판 품질 검증을 위해 "
            "수집·구성한 벤치마크 데이터셋 3종을 정의한다."
        ))
        add_table(
            doc,
            ["데이터셋", "목적"],
            [
                ["gpt_model_comparison_dataset", "Review/Judge 역할 최적 모델 선정"],
                ["orchestrator_benchmark_dataset", "Phase 1-3 최적화 효과 정량 검증"],
                ["turn_review_quality_dataset", "LLM 검토 vs 정규식 검토 정확도 비교"],
            ],
        )

        # 2. 데이터셋 명세
        h1(doc, "2. 데이터셋 명세")

        # 데이터셋 1
        h2(doc, "2.1 GPT 모델 비교 벤치마크 (gpt_model_comparison_dataset)")
        para(doc, "목적: Review 모델(턴 검토) 및 Judge 모델(최종 판정)의 최적 모델 선정")
        add_table(
            doc,
            ["항목", "값"],
            [
                ["평가 대상 모델 수", "13개"],
                [
                    "모델 목록",
                    "GPT-4o, GPT-4o-mini, GPT-4.1, GPT-4.1-mini, GPT-4.1-nano, "
                    "GPT-5, GPT-5-mini, GPT-5-nano, o3, o3-mini, o4-mini, o1, o1-mini",
                ],
                ["테스트 케이스 수", "16개"],
                ["평가 차원", "위반 감지 정확도, 채점 일관성, 응답 속도, 비용 효율 (4개)"],
                ["수집 방법", "backend/tests/benchmark/test_gpt_model_comparison.py 실행"],
                ["결과 파일", "docs/gpt_model_comparison.md"],
            ],
        )

        # 데이터셋 2
        h2(doc, "2.2 오케스트레이터 최적화 벤치마크 (orchestrator_benchmark_dataset)")
        para(doc, "목적: Phase 1-3 최적화 효과 검증")
        add_table(
            doc,
            ["항목", "값"],
            [
                ["픽스처", "backend/tests/fixtures/replay_debate_6turn.json (6턴 재현 데이터)"],
                ["시나리오 수", "4개 (Baseline, Parallel, FastPath, FullOptimized)"],
                ["측정 지표", "소요시간(s), LLM 호출 횟수(회), 추정 비용($)"],
                ["수집 방법", "backend/tests/benchmark/test_orchestrator_benchmark.py 실행"],
                ["결과 파일", "docs/orchestrator_optimization.md"],
            ],
        )

        # 데이터셋 3
        h2(doc, "2.3 턴 검토 품질 평가 (turn_review_quality_dataset)")
        para(doc, "목적: LLM 검토 vs 정규식 검토 정확도 비교")
        add_table(
            doc,
            ["발언 유형", "샘플 수"],
            [
                ["클린 발언 (패스트패스 대상)", "20"],
                ["인신공격", "20"],
                ["프롬프트 인젝션", "20"],
                ["주제 이탈", "20"],
                ["허위 근거", "20"],
            ],
        )
        add_table(
            doc,
            ["항목", "값"],
            [
                ["총 샘플 수", "100개"],
                ["평가 지표", "위반 감지율, 오탐율, 패스트패스 적용률"],
                ["수집 방법", "수동 라벨링 + tests/unit/services/test_debate_orchestrator.py"],
            ],
        )

        # 3. 수집 결과 요약
        h1(doc, "3. 데이터 수집 결과 요약")
        add_table(
            doc,
            ["데이터셋", "샘플 수", "정확도/효과"],
            [
                [
                    "GPT 모델 비교",
                    "13 모델 × 16 테스트 = 208",
                    "Judge: gpt-4.1 선정, Review: gpt-5-nano 선정",
                ],
                [
                    "오케스트레이터 벤치마크",
                    "4 시나리오",
                    "시간 37% 단축, 비용 76% 절감",
                ],
                [
                    "턴 검토 품질 평가",
                    "100 발언",
                    "패스트패스 ~80% 적용, 오탐 <5%",
                ],
            ],
        )

        out = OUTPUT_DIR / "SKN21기_4Team_수집_데이터_명세서_v2.docx"
        doc.save(str(out))
        print(f"[OK] {out}")
    except Exception as e:
        print(f"[ERROR] 수집_데이터_명세서_v2: {e}")


# ── 문서 4: AI 학습 결과서 v2 ─────────────────────────────────────────────────
def build_AI_학습_결과서_v2():
    try:
        doc = new_doc("SKN21기 4팀 AI 학습 결과서 v2")
        add_cover_info(doc)

        # 1. 개요
        h1(doc, "1. 개요")
        para(doc, (
            "본 시스템은 사전 학습된 LLM을 활용하며 Fine-tuning을 수행하지 않는다. "
            "대신 역할(Judge/Review)별 최적 모델 선정 벤치마크를 통해 성능을 검증하였다."
        ))
        add_table(
            doc,
            ["역할", "선정 모델", "선정 근거"],
            [
                ["Judge (최종 판정)", "gpt-4.1", "채점 품질·비용·일관성 균형 최우수"],
                ["Review (턴 검토)", "gpt-5-nano", "속도·비용·정확도 균형, 패스트패스 80% 시 실질 비용 최소"],
            ],
        )

        # 2. 테스트 시나리오 상세
        h1(doc, "2. 테스트 시나리오 상세")

        # 테스트 1: Judge 역할
        h2(doc, "2.1 테스트 1: GPT 모델 비교 (Judge 역할)")
        para(doc, "목적: 최종 토론 판정(Judge) 역할에 최적인 GPT 모델 선정")

        h3(doc, "테스트 환경")
        add_table(
            doc,
            ["항목", "값"],
            [
                ["입력", "동일한 6턴 토론 기록 (replay_debate_6turn.json)"],
                ["채점 항목", "논리성(0-30), 근거 활용(0-25), 반박력(0-25), 주제 적합성(0-20)"],
                ["반복 횟수", "각 모델 3회 실행 (일관성 평가)"],
                ["평가 기준", "채점 일관성(stddev), 채점 품질(human eval 비교), 응답 속도, 비용"],
            ],
        )

        h3(doc, "테스트 케이스 목록")
        add_table(
            doc,
            ["TC-ID", "토론 유형", "특이사항"],
            [
                ["TC-J-001", "찬반 토론 (일반)", "균형 발언"],
                ["TC-J-002", "찬반 토론", "압도적 승리 케이스"],
                ["TC-J-003", "찬반 토론", "페널티 다수 발생"],
                ["TC-J-004", "설득 토론", "감정 호소 포함"],
            ],
        )

        h3(doc, "결과 (상위 5개)")
        add_table(
            doc,
            ["순위", "모델", "채점 품질", "일관성", "비용/호출"],
            [
                ["1", "gpt-5", "9.120", "±0.15", "$0.028"],
                ["2", "gpt-4o", "8.950", "±0.18", "$0.031"],
                ["3 (선정)", "gpt-4.1", "8.936", "±0.12", "$0.025"],
                ["4", "o3-mini", "8.890", "±0.20", "$0.022"],
                ["5", "gpt-5-mini", "8.820", "±0.17", "$0.018"],
            ],
        )
        para(doc, "선정 결과: gpt-4.1 (3위이지만 성능/비용/일관성 균형 최우수)")

        # 테스트 2: Review 역할
        h2(doc, "2.2 테스트 2: GPT 모델 비교 (Review 역할)")
        para(doc, "목적: 턴 간 위반 감지(Review) 역할에 최적인 모델 선정")

        h3(doc, "테스트 환경")
        add_table(
            doc,
            ["항목", "값"],
            [
                ["입력", "위반 유형별 발언 100개 (수동 라벨링)"],
                ["평가 기준", "위반 감지 정확도, 오탐율, 응답 속도(<1초 필수), 비용"],
            ],
        )

        h3(doc, "테스트 케이스 목록")
        add_table(
            doc,
            ["TC-ID", "위반 유형", "샘플 수"],
            [
                ["TC-R-001", "클린 발언 (패스트패스 대상)", "20"],
                ["TC-R-002", "프롬프트 인젝션", "20"],
                ["TC-R-003", "인신공격", "20"],
                ["TC-R-004", "주제 이탈", "20"],
                ["TC-R-005", "허위 근거", "20"],
            ],
        )

        h3(doc, "결과 (상위 5개)")
        add_table(
            doc,
            ["순위", "모델", "감지 정확도", "오탐율", "속도", "비용/호출"],
            [
                ["1", "gpt-5", "9.240", "3.2%", "420ms", "$0.012"],
                ["2 (선정)", "gpt-5-nano", "8.907", "4.1%", "380ms", "$0.003"],
                ["3", "gpt-4o-mini", "8.602", "5.8%", "900ms", "$0.005"],
                ["4", "gpt-4.1-nano", "8.540", "5.5%", "450ms", "$0.002"],
                ["5", "gpt-4.1-mini", "8.430", "6.2%", "480ms", "$0.004"],
            ],
        )
        para(doc, "선정 결과: gpt-5-nano (속도·비용·정확도 균형, 패스트패스 80% 적용 시 실질 비용 최소)")

        # 테스트 3: 오케스트레이터 최적화
        h2(doc, "2.3 테스트 3: 오케스트레이터 최적화 벤치마크")
        para(doc, "목적: Phase 1-3 최적화 효과 정량 검증")

        h3(doc, "테스트 환경")
        add_table(
            doc,
            ["항목", "값"],
            [
                ["픽스처", "replay_debate_6turn.json (동일 6턴 입력 고정)"],
                ["측정", "소요시간(s), LLM 호출 횟수(회), 추정 비용($)"],
            ],
        )

        h3(doc, "시나리오 구성")
        add_table(
            doc,
            ["시나리오", "설명", "병렬화", "패스트패스"],
            [
                ["A_Baseline", "기존: 단일 gpt-4o 모델", "X", "X"],
                ["B_ModelSplit", "Phase 1: gpt-5-nano(리뷰) + gpt-4.1(판정)", "X", "X"],
                ["C_Parallel", "Phase 2: A턴 검토 + B턴 실행 asyncio.gather", "O", "X"],
                ["D_FastPath", "Phase 3: 정규식 클린 발언 LLM 검토 스킵", "O", "O"],
            ],
        )

        h3(doc, "결과")
        add_table(
            doc,
            ["시나리오", "소요시간", "LLM 호출", "추정비용", "기준 대비"],
            [
                ["A_Baseline", "64.5s", "12회", "$0.054", "기준"],
                ["B_ModelSplit", "55.2s", "12회", "$0.042", "-14.4%, 비용 -22%"],
                ["C_Parallel", "45.8s", "12회", "$0.042", "시간 -29%, 비용 -22%"],
                ["D_FastPath", "40.5s", "~2회", "$0.013", "시간 -37%, 비용 -76%"],
            ],
        )

        # 최적화 효과 차트
        fig = _draw_optimization_chart()
        buf = fig_to_buf(fig)
        add_image(doc, buf, width=5.5, caption="그림 1. 오케스트레이터 시나리오별 소요시간 및 비용 비교")

        # 3. 성능 검증 결과 종합
        h1(doc, "3. 성능 검증 결과 종합")
        add_table(
            doc,
            ["항목", "현행", "선정 결과", "개선 효과"],
            [
                ["Review 모델", "gpt-4o-mini (8.602점)", "gpt-5-nano (8.907점)", "성능↑, 비용 43%↓"],
                ["Judge 모델", "gpt-4o (8.501점)", "gpt-4.1 (8.936점)", "성능↑, 비용 20%↓"],
                ["전체 매치 비용", "$0.01739/매치", "$0.01329/매치", "23.6% 절감"],
                ["오케스트레이터 시간", "64.5s", "40.5s", "37% 단축"],
                ["오케스트레이터 비용", "$0.054", "$0.013", "76% 절감"],
            ],
        )

        out = OUTPUT_DIR / "SKN21기_4Team_AI_학습_결과서_v2.docx"
        doc.save(str(out))
        print(f"[OK] {out}")
    except Exception as e:
        print(f"[ERROR] AI_학습_결과서_v2: {e}")


def _draw_optimization_chart():
    """오케스트레이터 시나리오별 소요시간·비용 비교 차트."""
    scenarios = ["A_Baseline", "B_ModelSplit", "C_Parallel", "D_FastPath"]
    times = [64.5, 55.2, 45.8, 40.5]
    costs = [0.054, 0.042, 0.042, 0.013]

    x = np.arange(len(scenarios))
    width = 0.35

    fig, ax1 = plt.subplots(figsize=(9, 5))
    ax2 = ax1.twinx()

    bars1 = ax1.bar(x - width / 2, times, width, label="소요시간 (s)", color=C_BLUE, alpha=0.85)
    bars2 = ax2.bar(x + width / 2, costs, width, label="추정비용 ($)", color=C_ORANGE, alpha=0.85)

    ax1.set_ylabel("소요시간 (초)", color=C_BLUE)
    ax2.set_ylabel("추정 비용 ($)", color=C_ORANGE)
    ax1.set_xticks(x)
    ax1.set_xticklabels(scenarios, fontsize=9)
    ax1.set_ylim(0, 80)
    ax2.set_ylim(0, 0.08)

    for bar in bars1:
        ax1.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.8,
                 f"{bar.get_height():.1f}s", ha="center", va="bottom", fontsize=8, color=C_BLUE)
    for bar in bars2:
        ax2.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.001,
                 f"${bar.get_height():.3f}", ha="center", va="bottom", fontsize=8, color=C_ORANGE)

    lines1, labels1 = ax1.get_legend_handles_labels()
    lines2, labels2 = ax2.get_legend_handles_labels()
    ax1.legend(lines1 + lines2, labels1 + labels2, loc="upper right", fontsize=9)

    ax1.set_title("오케스트레이터 최적화 시나리오별 성능 비교", fontsize=11, fontweight="bold")
    fig.tight_layout()
    return fig


# ── 문서 5: 에이전트 오케스트레이션 설계서 ────────────────────────────────────
def build_에이전트_오케스트레이션_설계서():
    try:
        doc = new_doc("SKN21기 4팀 에이전트 오케스트레이션 설계서")
        add_cover_info(doc)

        # 1. 시스템 전체 흐름
        h1(doc, "1. 시스템 전체 흐름")
        para(doc, (
            "아래 다이어그램은 에이전트 등록부터 ELO 갱신까지의 전체 데이터 흐름을 나타낸다. "
            "각 단계는 FastAPI 라우터 → 서비스 레이어 → 오케스트레이터 순으로 처리된다."
        ))

        fig = _draw_flow_diagram()
        buf = fig_to_buf(fig)
        add_image(doc, buf, width=6.0, caption="그림 1. 에이전트 토론 시스템 전체 흐름도")

        h2(doc, "흐름 텍스트 표현")
        flow_text = (
            "[사용자] → POST /agents (에이전트 생성)\n"
            "    AgentCreate 스키마 검증\n"
            "    POST /test (API 키·모델 유효성 테스트)\n"
            "    DB: debate_agents 저장\n"
            "\n"
            "[사용자] → POST /topics/{id}/join (큐 참가)\n"
            "    DebateMatchingService.join_queue()\n"
            "    debate_match_queue 삽입\n"
            "    [2인 도달] → DebateMatch 생성\n"
            "\n"
            "[시스템] → run_debate(match_id) [asyncio 백그라운드]\n"
            "    _execute_match()\n"
            "    턴 루프 (max_turns × 2):\n"
            "        _execute_turn() → LLM API 호출\n"
            "        detect_prompt_injection() / detect_ad_hominem()\n"
            "        OptimizedOrchestrator.review_turn_fast()\n"
            "            _should_skip_review() → FastPath?\n"
            "                YES: _fast_path_result() → logic_score=None\n"
            "                NO:  gpt-5-nano API → violations + logic_score\n"
            "        Redis PUBLISH turn 이벤트\n"
            "    Orchestrator.judge() → gpt-4.1 API\n"
            "    calculate_elo() → ELO 갱신\n"
            "    Redis PUBLISH finished 이벤트\n"
            "\n"
            "[프론트엔드] SSE 구독 → debateStore.addTurnFromSSE()\n"
            "    TurnBubble 컴포넌트 렌더링\n"
            "    LogicScoreBar (logic_score != null 일 때만)"
        )
        code_para(doc, flow_text)

        # 2. 에이전트 모듈 상세
        h1(doc, "2. 에이전트 모듈 상세")

        h2(doc, "2.1 API 레이어 (debate_agents.py)")
        add_table(
            doc,
            ["엔드포인트", "함수명", "주요 로직"],
            [
                [
                    "POST /agents",
                    "create_agent(data, user, db)",
                    "AgentCreate 스키마 검증, use_platform_credits=True면 api_key 불필요, DB 저장 + 첫 버전 생성",
                ],
                [
                    "POST /test",
                    "test_agent_connection(data, user)",
                    "InferenceClient로 실제 API 호출(저장 없음), 실패 시 _classify_provider_error() → 사용자 친화 메시지",
                ],
                [
                    "PUT /{id}",
                    "update_agent(agent_id, data, user, db)",
                    "시스템 프롬프트 변경 시 새 버전 자동 생성, 이름 변경 7일 1회 제한",
                ],
                [
                    "GET /{id}/versions",
                    "get_agent_versions(agent_id, user, db)",
                    "소유자 또는 admin/superadmin만 접근",
                ],
            ],
        )

        code_text = (
            "# 에이전트 생성 핵심 로직\n"
            "POST /agents → create_agent(data: AgentCreate, user, db)\n"
            "  - use_platform_credits=True 면 api_key 불필요\n"
            "  - DebateAgent DB 저장 + 첫 버전 생성\n"
            "\n"
            "# API 키 테스트 (저장 없음)\n"
            "POST /test → test_agent_connection(data: AgentTestRequest, user)\n"
            "  - InferenceClient 실제 API 호출\n"
            "  - 실패 시 _classify_provider_error() 사용자 친화 메시지\n"
            "\n"
            "# 에이전트 수정 (버전 자동 생성)\n"
            "PUT /{id} → update_agent(agent_id, data: AgentUpdate, user, db)\n"
            "  - 시스템 프롬프트 변경 시 새 버전 자동 생성\n"
            "  - 이름 변경: 7일 1회 제한\n"
            "\n"
            "# 버전 히스토리 (소유자 + admin)\n"
            "GET /{id}/versions → get_agent_versions(agent_id, user, db)\n"
            "  - 소유자 또는 admin/superadmin만 접근"
        )
        code_para(doc, code_text)

        h2(doc, "2.2 매칭 서비스 (debate_matching_service.py)")
        add_table(
            doc,
            ["메서드", "주요 로직"],
            [
                [
                    "join_queue(user, topic_id, agent_id)",
                    "1. 토픽 상태 검증(open) "
                    "2. 에이전트 소유권 검증(admin 우회) "
                    "3. API 키 또는 플랫폼 크레딧 검증 "
                    "4. use_platform_credits=True면 키 체크 스킵 "
                    "5. 1 토픽 대기 제한 "
                    "6. 상대 있으면 opponent_joined 이벤트 발행",
                ],
                [
                    "ready_up(user, topic_id, agent_id)",
                    "양쪽 is_ready=True 되면 매치 생성, 첫 번째만 완료 시 10초 카운트다운 이벤트",
                ],
            ],
        )

        # 3. 오케스트레이터 모듈 상세
        h1(doc, "3. 오케스트레이터 모듈 상세")

        h2(doc, "3.1 DebateOrchestrator (debate_orchestrator.py)")
        add_table(
            doc,
            ["메서드", "설명"],
            [
                [
                    "review_turn(topic, speaker, turn_number, claim, evidence, action, opponent_last_claim)",
                    "gpt-4o-mini(기본) 또는 REVIEW_MODEL로 위반 감지. "
                    "반환: {logic_score, violations, feedback, block, penalties, penalty_total}",
                ],
                [
                    "judge(match, turns, topic, agent_a_name, agent_b_name)",
                    "A/B 50% 확률 스왑(편향 제거), JUDGE_SYSTEM_PROMPT + 토론 로그 → gpt-4.1 채점. "
                    "logic(30)+evidence(25)+rebuttal(25)+relevance(20), 5점 차 미만 무승부",
                ],
                [
                    "calculate_elo(rating_a, rating_b, result, score_diff)",
                    "표준 ELO 기대 승률 + 비대칭 K-factor(승 K=40, 패 K=24), "
                    "transfer = min(score_diff, 30) 제로섬 이전",
                ],
            ],
        )

        h2(doc, "3.2 OptimizedDebateOrchestrator (Phase 1-3)")
        add_table(
            doc,
            ["메서드/함수", "설명"],
            [
                [
                    "review_turn_fast(claim, evidence, ...)",
                    "Phase 3: _should_skip_review() 판단 → FastPath(logic_score=None) 또는 gpt-5-nano 호출",
                ],
                [
                    "_should_skip_review(claim, evidence)",
                    "조건: 정규식 무위반 + 길이 10~800자 + ASCII 비율 40% 이하 → 클린 한국어 발언 ~80% 스킵",
                ],
            ],
        )

        code_text2 = (
            "class OptimizedDebateOrchestrator(DebateOrchestrator):\n"
            "    async def review_turn_fast(claim, evidence, ...):\n"
            "        # Phase 3: FastPath 여부 판단\n"
            "        if _should_skip_review(claim, evidence):\n"
            "            return _fast_path_result()  # logic_score=None, 즉시 통과\n"
            "        # 비FastPath: 경량 gpt-5-nano로 위반 감지 (Phase 1)\n"
            "        return await super().review_turn(...)\n"
            "\n"
            "def _should_skip_review(claim, evidence):\n"
            "    # 정규식 무위반 + 길이 10~800자 + ASCII 비율 40% 이하\n"
            "    # → 클린 한국어 발언 ~80% LLM 검토 스킵"
        )
        code_para(doc, code_text2)

        h2(doc, "3.3 페널티 시스템")
        add_table(
            doc,
            ["유형", "감지 방법", "페널티"],
            [
                ["prompt_injection", "정규식 (detect_prompt_injection)", "-10점"],
                ["ad_hominem", "정규식 (detect_ad_hominem)", "-8점"],
                ["repetition", "정규식 (detect_repetition, 70%+ 중복)", "-3점"],
                ["schema_error", "정규식 (validate_response_schema)", "-5점"],
                ["timeout", "타임아웃 판정", "-5점"],
                ["llm_prompt_injection", "LLM 검토 (OptimizedOrchestrator)", "-10점"],
                ["llm_ad_hominem", "LLM 검토 (OptimizedOrchestrator)", "-8점"],
                ["llm_off_topic", "LLM 검토 (OptimizedOrchestrator)", "-5점"],
                ["llm_false_claim", "LLM 검토 (OptimizedOrchestrator)", "-7점"],
                ["human_suspicion", "의심도 점수 ≥61", "-15점"],
            ],
        )

        # 4. 사용자 시나리오 플로우
        h1(doc, "4. 사용자 시나리오 플로우")

        h2(doc, "시나리오 1: 에이전트 등록 및 첫 토론")
        scenario1_steps = [
            "사용자 로그인",
            "'에이전트 생성' → 이름/모델/API 키 입력 → '연결 테스트' 클릭 (POST /test)",
            "테스트 성공 → '저장' (POST /agents)",
            "토론 주제 목록에서 관심 주제 선택",
            "'참가' → 큐 등록 (POST /topics/{id}/join)",
            "SSE 연결: 상대 대기 중 (heartbeat)",
            "상대 등장 → '준비 완료' → 10초 카운트다운 → 토론 시작",
            "실시간 관전: 타이핑 효과 + HP 바 + 논증 점수",
            "토론 종료 → 스코어카드 + ELO 변동 확인",
        ]
        for i, step in enumerate(scenario1_steps, 1):
            para(doc, f"  {i}. {step}")

        h2(doc, "시나리오 2: 에이전트 개선 후 재도전")
        scenario2_steps = [
            "전적 조회 (GET /agents/{id}/versions)",
            "패턴 분석: 자주 당하는 페널티 확인",
            "시스템 프롬프트 수정 → 새 버전 자동 생성 (PUT /agents/{id})",
            "새 버전으로 재참가",
        ]
        for i, step in enumerate(scenario2_steps, 1):
            para(doc, f"  {i}. {step}")

        # 시나리오 흐름도
        fig2 = _draw_scenario_diagram()
        buf2 = fig_to_buf(fig2)
        add_image(doc, buf2, width=6.0, caption="그림 2. 사용자 시나리오 1: 에이전트 등록 및 첫 토론 흐름")

        out = OUTPUT_DIR / "SKN21기_4Team_에이전트_오케스트레이션_설계서.docx"
        doc.save(str(out))
        print(f"[OK] {out}")
    except Exception as e:
        print(f"[ERROR] 에이전트_오케스트레이션_설계서: {e}")


def _draw_flow_diagram():
    """전체 흐름도 (matplotlib FancyBboxPatch + 화살표)."""
    fig, ax = plt.subplots(figsize=(11, 9))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 10)
    ax.axis("off")

    def box(x, y, w, h, text, color="#2E74B5", fontsize=8):
        rect = FancyBboxPatch(
            (x - w / 2, y - h / 2), w, h,
            boxstyle="round,pad=0.1", linewidth=1.2,
            edgecolor="#1F497D", facecolor=color, alpha=0.88,
        )
        ax.add_patch(rect)
        ax.text(x, y, text, ha="center", va="center",
                fontsize=fontsize, color="white", fontweight="bold",
                wrap=True)

    def arrow(x1, y1, x2, y2, label=""):
        ax.annotate(
            "", xy=(x2, y2), xytext=(x1, y1),
            arrowprops=dict(arrowstyle="->", color="#1F497D", lw=1.5),
        )
        if label:
            mx, my = (x1 + x2) / 2 + 0.15, (y1 + y2) / 2
            ax.text(mx, my, label, fontsize=7, color="#1F497D")

    # 노드 정의
    nodes = [
        (5.0, 9.3, 3.5, 0.55, "사용자 (POST /agents  /topics/{id}/join)", "#1F497D"),
        (5.0, 8.4, 3.5, 0.55, "AgentCreate 검증 + API 키 테스트", "#2E74B5"),
        (5.0, 7.5, 3.5, 0.55, "DB: debate_agents 저장", "#2E74B5"),
        (5.0, 6.5, 3.5, 0.55, "큐 등록: debate_match_queue", "#548235"),
        (5.0, 5.5, 3.5, 0.55, "매치 생성: DebateMatch", "#548235"),
        (5.0, 4.5, 3.5, 0.55, "run_debate(match_id) [asyncio 백그라운드]", "#C55A11"),
        (2.5, 3.4, 3.0, 0.55, "_execute_turn() → LLM API", "#C55A11"),
        (7.5, 3.4, 3.0, 0.55, "detect_injection / ad_hominem", "#7030A0"),
        (5.0, 2.4, 3.5, 0.55, "review_turn_fast()", "#7030A0"),
        (2.5, 1.4, 3.0, 0.55, "FastPath (skip)", "#548235"),
        (7.5, 1.4, 3.0, 0.55, "gpt-5-nano 검토", "#C00000"),
        (5.0, 0.5, 3.5, 0.45, "judge() → ELO 갱신 → SSE finished", "#1F497D"),
    ]
    for x, y, w, h, text, color in nodes:
        box(x, y, w, h, text, color)

    # 화살표
    pairs = [
        (5.0, 9.05, 5.0, 8.67),
        (5.0, 8.12, 5.0, 7.78),
        (5.0, 7.22, 5.0, 6.78),
        (5.0, 6.22, 5.0, 5.78),
        (5.0, 5.22, 5.0, 4.78),
        (5.0, 4.22, 2.5, 3.67),
        (5.0, 4.22, 7.5, 3.67),
        (2.5, 3.12, 5.0, 2.67),
        (7.5, 3.12, 5.0, 2.67),
        (5.0, 2.12, 2.5, 1.67, "skip"),
        (5.0, 2.12, 7.5, 1.67, "review"),
        (2.5, 1.12, 5.0, 0.73),
        (7.5, 1.12, 5.0, 0.73),
    ]
    for p in pairs:
        if len(p) == 5:
            arrow(*p[:4], label=p[4])
        else:
            arrow(*p)

    ax.set_title("에이전트 토론 시스템 전체 흐름도", fontsize=12, fontweight="bold", pad=8)
    fig.tight_layout()
    return fig


def _draw_scenario_diagram():
    """사용자 시나리오 1 단계별 흐름도."""
    steps = [
        ("1. 로그인", "#2E74B5"),
        ("2. 에이전트 생성\n(POST /agents)", "#2E74B5"),
        ("3. API 키 테스트\n(POST /test)", "#548235"),
        ("4. 주제 선택", "#C55A11"),
        ("5. 큐 참가\n(POST /join)", "#C55A11"),
        ("6. SSE 대기", "#7030A0"),
        ("7. 토론 시작", "#C00000"),
        ("8. 실시간 관전\n(SSE 스트리밍)", "#C00000"),
        ("9. 결과 확인\n(ELO 변동)", "#1F497D"),
    ]

    fig, ax = plt.subplots(figsize=(11, 3.5))
    ax.set_xlim(-0.5, len(steps) - 0.5)
    ax.set_ylim(-0.8, 1.5)
    ax.axis("off")

    for i, (text, color) in enumerate(steps):
        rect = FancyBboxPatch(
            (i - 0.42, -0.35), 0.84, 0.9,
            boxstyle="round,pad=0.08",
            edgecolor="#1F497D", facecolor=color, alpha=0.85,
        )
        ax.add_patch(rect)
        ax.text(i, 0.1, text, ha="center", va="center",
                fontsize=7.5, color="white", fontweight="bold")
        if i < len(steps) - 1:
            ax.annotate(
                "", xy=(i + 0.55, 0.1), xytext=(i + 0.44, 0.1),
                arrowprops=dict(arrowstyle="->", color="#1F497D", lw=1.3),
            )

    ax.set_title("시나리오 1: 에이전트 등록 및 첫 토론 단계별 흐름", fontsize=11, fontweight="bold")
    fig.tight_layout()
    return fig


# ── 문서 6: 데이터베이스 및 조회 설계서 ─────────────────────────────────────────
def _draw_erd_diagram():
    """AI 토론 플랫폼 ERD 테이블 관계도."""
    fig, ax = plt.subplots(figsize=(11, 7))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 7)
    ax.axis("off")
    ax.set_facecolor("#FAFAFA")
    fig.patch.set_facecolor("#FAFAFA")

    # 테이블 박스들 (x, y, width, height, label, color)
    tables = [
        (0.2, 5.5, 2.0, 1.0, "users", "#1F497D"),
        (3.5, 5.5, 2.2, 1.0, "debate_agents", "#2E74B5"),
        (7.0, 5.5, 2.5, 1.0, "debate_agent_versions", "#2E74B5"),
        (3.5, 3.5, 2.2, 1.0, "debate_topics", "#548235"),
        (3.5, 1.5, 2.2, 1.0, "debate_match_queue", "#C55A11"),
        (7.0, 3.5, 2.5, 1.0, "debate_matches", "#C55A11"),
        (7.0, 1.2, 2.5, 1.0, "debate_turn_logs", "#7030A0"),
        (0.2, 3.5, 2.0, 1.0, "debate_agent_templates", "#808080"),
    ]

    for (x, y, w, h, label, color) in tables:
        rect = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.05",
                               edgecolor="#333", facecolor=color, alpha=0.85)
        ax.add_patch(rect)
        ax.text(x + w / 2, y + h / 2, label, ha="center", va="center",
                fontsize=8, color="white", fontweight="bold")

    # 화살표 (FK 관계) — 직선 연결만 처리
    arrows_simple = [
        # users → debate_agents
        ((2.2, 6.0), (3.5, 6.0)),
        # debate_agents → debate_agent_versions
        ((5.7, 6.0), (7.0, 6.0)),
        # debate_agents → debate_matches (agent_a/b)
        ((5.7, 5.7), (7.0, 4.2)),
        # debate_agents → debate_match_queue
        ((4.6, 3.5), (4.6, 2.5)),
        # debate_topics → debate_matches
        ((5.7, 4.0), (7.0, 4.0)),
        # debate_topics → debate_match_queue
        ((4.0, 3.5), (4.0, 2.5)),
        # debate_matches → debate_turn_logs
        ((8.25, 3.5), (8.25, 2.2)),
        # debate_agent_templates → debate_agents
        ((2.2, 4.0), (3.5, 5.7)),
        # users → debate_topics (users.id → topics.created_by)
        ((1.2, 5.5), (3.8, 4.5)),
    ]
    for start, end in arrows_simple:
        ax.annotate("", xy=end, xytext=start,
                    arrowprops=dict(arrowstyle="-|>", color="#555", lw=1.2,
                                    connectionstyle="arc3,rad=0.0"))

    ax.set_title("AI 토론 플랫폼 ERD (테이블 관계도)", fontsize=12, fontweight="bold", pad=10)
    fig.tight_layout()
    return fig


def build_데이터베이스_및_조회_설계서():
    try:
        doc = new_doc("데이터베이스 및 데이터 조회 설계서")
        add_cover_info(doc)

        # ── 섹션 1: 개요 ──────────────────────────────────────────────────────
        h1(doc, "1. 개요")

        h2(doc, "1.1 문서 목적")
        para(doc, (
            "이 문서는 AI 에이전트 토론 플랫폼의 데이터베이스 설계와 데이터 조회 방식을 통합하여 "
            "기술한다. DB 스키마 설계와 실제 데이터 접근 패턴(API, 쿼리)을 한 문서에서 "
            "확인할 수 있도록 병합하였다."
        ))

        h2(doc, "1.2 대상 독자")
        para(doc, "백엔드 개발자, 데이터 엔지니어, 아키텍처 검토자")

        h2(doc, "1.3 사용 기술")
        add_table(
            doc,
            ["기술", "역할"],
            [
                ["PostgreSQL 16 (Docker)", "관계형 데이터베이스 — 모든 영속성 데이터 저장"],
                ["SQLAlchemy 2.0 (async)", "Python ORM — 비동기 DB 접근 추상화"],
                ["Alembic", "스키마 마이그레이션 버전 관리"],
            ],
        )

        # ── 섹션 2: 데이터 모델 (ERD) ─────────────────────────────────────────
        h1(doc, "2. 데이터 모델 (ERD)")

        fig_erd = _draw_erd_diagram()
        buf_erd = fig_to_buf(fig_erd)
        add_image(doc, buf_erd, width=6.2, caption="그림 1. AI 토론 플랫폼 테이블 관계도 (ERD)")

        h2(doc, "2.1 테이블 관계 텍스트 표현")
        code_para(doc, (
            "users\n"
            " ├── debate_agents (owner_id → FK)\n"
            " │    ├── debate_agent_versions (agent_id → FK, CASCADE)\n"
            " │    ├── debate_matches.agent_a_id (FK)\n"
            " │    ├── debate_matches.agent_b_id (FK)\n"
            " │    └── debate_match_queue (agent_id → FK)\n"
            " │\n"
            " └── debate_topics (created_by → FK)\n"
            "      ├── debate_matches (topic_id → FK)\n"
            "      │    └── debate_turn_logs (match_id → FK, CASCADE)\n"
            "      └── debate_match_queue (topic_id → FK, CASCADE)\n"
            "\n"
            "debate_agent_templates (독립)\n"
            " └── debate_agents.template_id (FK, NULL 허용)"
        ))

        # ── 섹션 3: 핵심 테이블 스키마 상세 ──────────────────────────────────
        h1(doc, "3. 핵심 테이블 스키마 상세")

        # 3.1 debate_agents
        h2(doc, "3.1 debate_agents")
        add_table(
            doc,
            ["컬럼", "타입", "제약", "설명"],
            [
                ["id", "UUID", "PK", "에이전트 고유 ID"],
                ["owner_id", "UUID", "FK(users) NOT NULL", "소유자"],
                ["name", "VARCHAR(100)", "NOT NULL", "에이전트 이름 (7일 1회 변경 제한)"],
                ["description", "TEXT", "NULL", "설명"],
                ["provider", "VARCHAR(20)", "NOT NULL", "openai / anthropic / google / runpod / local"],
                ["model_id", "VARCHAR(100)", "NOT NULL", "모델 식별자 (예: gpt-4o, claude-3-5-sonnet)"],
                ["encrypted_api_key", "TEXT", "NULL", "Fernet 암호화 API 키"],
                ["image_url", "TEXT", "NULL", "프로필 이미지 URL"],
                ["template_id", "UUID", "FK(templates), NULL", "기반 템플릿"],
                ["customizations", "JSONB", "NULL", "템플릿 커스터마이징 값"],
                ["system_prompt", "TEXT", "NULL", "직접 입력 시스템 프롬프트"],
                ["elo_rating", "INTEGER", "DEFAULT 1500", "ELO 레이팅"],
                ["tier", "VARCHAR(20)", "DEFAULT 'Iron'", "ELO 기반 티어"],
                ["tier_protection_count", "INTEGER", "DEFAULT 0", "강등 보호 카운터"],
                ["wins / losses / draws", "INTEGER", "DEFAULT 0", "전적"],
                ["is_active", "BOOLEAN", "DEFAULT true", "활성 여부"],
                ["is_platform", "BOOLEAN", "DEFAULT false", "플랫폼 에이전트 여부"],
                ["use_platform_credits", "BOOLEAN", "DEFAULT false", "플랫폼 크레딧 사용 여부"],
                ["is_system_prompt_public", "BOOLEAN", "DEFAULT false", "시스템 프롬프트 공개 여부"],
                ["is_profile_public", "BOOLEAN", "DEFAULT true", "프로필 공개 여부"],
                ["name_changed_at", "TIMESTAMPTZ", "NULL", "이름 변경 시각"],
                ["created_at / updated_at", "TIMESTAMPTZ", "DEFAULT now()", "타임스탬프"],
            ],
        )

        h3(doc, "티어 기준")
        add_table(
            doc,
            ["티어", "ELO 범위"],
            [
                ["Iron", "~ 1299"],
                ["Bronze", "1300 ~ 1449"],
                ["Silver", "1450 ~ 1599"],
                ["Gold", "1600 ~ 1749"],
                ["Platinum", "1750 ~ 1899"],
                ["Diamond", "1900 ~ 2049"],
                ["Master", "2050+"],
            ],
        )

        # 3.2 debate_agent_versions
        h2(doc, "3.2 debate_agent_versions")
        add_table(
            doc,
            ["컬럼", "타입", "제약", "설명"],
            [
                ["id", "UUID", "PK", "버전 ID"],
                ["agent_id", "UUID", "FK CASCADE", "에이전트 참조"],
                ["version_number", "INTEGER", "NOT NULL", "버전 번호 (1, 2, 3...)"],
                ["version_tag", "VARCHAR(50)", "NULL", "태그 (v1, beta-2 등)"],
                ["system_prompt", "TEXT", "NOT NULL", "이 버전의 시스템 프롬프트 스냅샷"],
                ["parameters", "JSONB", "NULL", "추론 파라미터 (temperature, top_p 등)"],
                ["wins / losses / draws", "INTEGER", "DEFAULT 0", "버전별 전적"],
                ["created_at", "TIMESTAMPTZ", "DEFAULT now()", "생성 시각"],
            ],
        )

        # 3.3 debate_topics
        h2(doc, "3.3 debate_topics")
        add_table(
            doc,
            ["컬럼", "타입", "제약", "설명"],
            [
                ["id", "UUID", "PK", "주제 ID"],
                ["title", "VARCHAR(200)", "NOT NULL", "토론 제목"],
                ["description", "TEXT", "NULL", "상세 설명"],
                ["mode", "VARCHAR(20)", "CHECK(debate/persuasion/cross_exam)", "토론 모드"],
                ["status", "VARCHAR(20)", "CHECK(scheduled/open/in_progress/closed)", "상태"],
                ["max_turns", "INTEGER", "DEFAULT 6, 2~20", "최대 턴 수"],
                ["turn_token_limit", "INTEGER", "DEFAULT 500, 100~2000", "턴당 토큰 제한"],
                ["scheduled_start_at", "TIMESTAMPTZ", "NULL", "예약 시작 시각"],
                ["scheduled_end_at", "TIMESTAMPTZ", "NULL", "예약 종료 시각"],
                ["is_admin_topic", "BOOLEAN", "DEFAULT false", "관리자 생성 여부"],
                ["tools_enabled", "BOOLEAN", "DEFAULT true", "도구 사용 허용"],
                ["created_by", "UUID", "FK(users), NULL", "작성자"],
                ["created_at / updated_at", "TIMESTAMPTZ", "DEFAULT now()", "타임스탬프"],
            ],
        )

        # 3.4 debate_matches
        h2(doc, "3.4 debate_matches")
        add_table(
            doc,
            ["컬럼", "타입", "제약", "설명"],
            [
                ["id", "UUID", "PK", "매치 ID"],
                ["topic_id", "UUID", "FK(topics)", "토론 주제"],
                ["agent_a_id / agent_b_id", "UUID", "FK(agents)", "에이전트 A/B"],
                ["agent_a_version_id / agent_b_version_id", "UUID", "FK(versions), NULL", "버전 스냅샷"],
                ["status", "VARCHAR(20)", "CHECK(pending/in_progress/completed/error/waiting_agent/forfeit)", "매치 상태"],
                ["winner_id", "UUID", "NULL", "승자 (null=무승부)"],
                ["scorecard", "JSONB", "NULL", "심판 점수 및 이유"],
                ["score_a / score_b", "INTEGER", "NULL", "최종 점수 (페널티 차감 후)"],
                ["penalty_a / penalty_b", "INTEGER", "DEFAULT 0", "누적 페널티"],
                ["started_at / finished_at", "TIMESTAMPTZ", "NULL", "시작/종료 시각"],
                ["created_at", "TIMESTAMPTZ", "DEFAULT now()", "생성 시각"],
            ],
        )

        h3(doc, "scorecard JSONB 구조")
        code_para(doc, (
            '{\n'
            '  "agent_a": {"logic": 28, "evidence": 22, "rebuttal": 25, "relevance": 19},\n'
            '  "agent_b": {"logic": 25, "evidence": 20, "rebuttal": 22, "relevance": 18},\n'
            '  "reasoning": "에이전트 A는 통계 데이터를 효과적으로 인용하며...",\n'
            '  "winner_id": "uuid-of-agent-a",\n'
            '  "result": "agent_a_wins"\n'
            '}'
        ))

        # 3.5 debate_turn_logs
        h2(doc, "3.5 debate_turn_logs")
        add_table(
            doc,
            ["컬럼", "타입", "제약", "설명"],
            [
                ["id", "UUID", "PK", "턴 로그 ID"],
                ["match_id", "UUID", "FK(matches) CASCADE", "매치 참조"],
                ["turn_number", "INTEGER", "NOT NULL", "턴 번호"],
                ["speaker", "VARCHAR(10)", "CHECK(agent_a/agent_b)", "발언자"],
                ["agent_id", "UUID", "FK(agents)", "에이전트 참조"],
                ["action", "VARCHAR(20)", "NOT NULL", "argue/rebut/concede/question/summarize"],
                ["claim", "TEXT", "NOT NULL", "주장 본문"],
                ["evidence", "TEXT", "NULL", "근거 자료"],
                ["tool_used", "VARCHAR(50)", "NULL", "사용한 도구 이름"],
                ["tool_result", "TEXT", "NULL", "도구 실행 결과"],
                ["penalties", "JSONB", "NULL", "부과된 페널티 딕셔너리"],
                ["penalty_total", "INTEGER", "DEFAULT 0", "총 페널티 합계"],
                ["review_result", "JSONB", "NULL", "LLM 검토 결과"],
                ["is_blocked", "BOOLEAN", "DEFAULT false", "LLM 검토 차단 여부"],
                ["human_suspicion_score", "INTEGER", "DEFAULT 0", "인간 의심 점수"],
                ["response_time_ms", "INTEGER", "NULL", "응답 소요 시간(ms)"],
                ["input_tokens / output_tokens", "INTEGER", "DEFAULT 0", "토큰 수"],
                ["created_at", "TIMESTAMPTZ", "DEFAULT now()", "생성 시각"],
            ],
        )

        h3(doc, "review_result JSONB 구조")
        code_para(doc, (
            '{\n'
            '  "logic_score": 7,\n'
            '  "violations": ["llm_off_topic"],\n'
            '  "is_blocked_reason": null,\n'
            '  "model": "gpt-5-nano",\n'
            '  "skipped": false\n'
            '}'
        ))

        # 3.6 debate_match_queue
        h2(doc, "3.6 debate_match_queue")
        add_table(
            doc,
            ["컬럼", "타입", "제약", "설명"],
            [
                ["id", "UUID", "PK", "큐 항목 ID"],
                ["topic_id", "UUID", "FK(topics) CASCADE", "주제 참조"],
                ["agent_id", "UUID", "FK(agents) CASCADE", "에이전트 참조"],
                ["user_id", "UUID", "FK(users) CASCADE", "사용자 참조"],
                ["joined_at", "TIMESTAMPTZ", "DEFAULT now()", "큐 진입 시각"],
                ["is_ready", "BOOLEAN", "DEFAULT false", "준비 완료 여부"],
                ["UNIQUE", "(topic_id, agent_id)", "", "동일 에이전트 중복 방지"],
            ],
        )

        # ── 섹션 4: 인덱스 및 제약 조건 ──────────────────────────────────────
        h1(doc, "4. 인덱스 및 제약 조건")

        h2(doc, "4.1 인덱스 목록")
        add_table(
            doc,
            ["테이블", "인덱스/제약", "목적"],
            [
                ["debate_agents", "idx_debate_agents_owner", "사용자별 에이전트 조회 최적화"],
                ["debate_agents", "idx_debate_agents_elo", "ELO 랭킹 정렬 최적화"],
                ["debate_agent_versions", "idx_versions_agent", "에이전트별 버전 조회"],
                ["debate_topics", "idx_topics_status", "상태별 필터링"],
                ["debate_topics", "idx_topics_created", "최신순 정렬"],
                ["debate_matches", "idx_matches_topic", "주제별 매치 조회"],
                ["debate_matches", "idx_matches_agent_a/b", "에이전트별 매치 조회"],
                ["debate_matches", "idx_matches_status", "진행 중 매치 조회"],
                ["debate_turn_logs", "idx_turns_match", "매치별 턴 조회 (가장 빈번)"],
                ["debate_match_queue", "UNIQUE(topic_id, agent_id)", "중복 큐 방지"],
                ["debate_match_queue", "idx_queue_topic", "주제별 큐 조회"],
            ],
        )

        h2(doc, "4.2 보안 제약사항")
        para(doc, "매치 진행 중(in_progress) 에이전트 삭제 불가 — 애플리케이션 레벨 검증")
        para(doc, "큐 진입 시 SELECT FOR UPDATE — 레이스 컨디션 방지")
        para(doc, "버전 스냅샷: 매치 시작 시 agent_a_version_id 고정 — 이후 변경 불변")

        # ── 섹션 5: 데이터 흐름 ────────────────────────────────────────────────
        h1(doc, "5. 데이터 흐름")

        h2(doc, "5.1 에이전트 생성 흐름")
        code_para(doc, (
            "POST /agents → AgentCreate 검증\n"
            " ├── template_id 있으면: 템플릿 로드 → 프롬프트 조립\n"
            " ├── api_key 있으면: Fernet.encrypt() → encrypted_api_key\n"
            " └── system_prompt 없고 template 없으면 → 422\n"
            "\n"
            "DB INSERT: debate_agents\n"
            "DB INSERT: debate_agent_versions (v1 스냅샷)\n"
            "→ 반환: AgentResponse (elo_rating=1500)"
        ))

        h2(doc, "5.2 매치메이킹 흐름")
        code_para(doc, (
            "POST /topics/{id}/join → join_queue()\n"
            " ├── 주제 status=open 검증\n"
            " ├── 에이전트 소유권 검증 (admin 우회)\n"
            " ├── API 키 또는 플랫폼 크레딧 검증\n"
            " └── 에이전트당 1 토픽 대기 제한 검증\n"
            "\n"
            "SELECT FOR UPDATE → 동시 접근 직렬화\n"
            "큐 < 2: INSERT debate_match_queue\n"
            "큐 >= 2: DebateMatch 생성 → 큐 2개 삭제 → Redis PUBLISH matched"
        ))

        h2(doc, "5.3 토론 턴 저장 흐름")
        code_para(doc, (
            "_execute_turn() 완료 시:\n"
            " 1. 페널티 탐지 (regex 7종)\n"
            " 2. OptimizedOrchestrator.review_turn_fast() (LLM 또는 FastPath)\n"
            " 3. DB INSERT: debate_turn_logs\n"
            "    - claim, evidence, penalties, review_result, is_blocked\n"
            "    - input_tokens, output_tokens 기록\n"
            " 4. Redis PUBLISH: turn 이벤트"
        ))

        h2(doc, "5.4 매치 완료 및 ELO 갱신 흐름")
        code_para(doc, (
            "judge() 완료 시:\n"
            " 1. score_a = sum(scorecard.agent_a) - penalty_a\n"
            " 2. score_b = sum(scorecard.agent_b) - penalty_b\n"
            " 3. |score_a - score_b| < 5 → 무승부\n"
            " 4. calculate_elo(rating_a, rating_b, result)\n"
            "    → 비대칭 K-factor (승 K=40, 패 K=24)\n"
            " 5. DB UPDATE: debate_matches (status=completed, scorecard)\n"
            " 6. DB UPDATE: debate_agents (elo_rating, tier, wins/losses)\n"
            " 7. Redis PUBLISH: finished 이벤트"
        ))

        # ── 섹션 6: API 명세 (데이터 조회 관점) ──────────────────────────────
        h1(doc, "6. API 명세 (데이터 조회 관점)")

        h2(doc, "6.1 에이전트 조회 API")
        add_table(
            doc,
            ["메서드", "경로", "설명", "주요 반환 필드"],
            [
                ["GET", "/agents/me", "내 에이전트 목록", "id, name, elo_rating, tier, wins/losses"],
                ["GET", "/agents/ranking", "ELO 글로벌 랭킹", "elo_rating 기준 내림차순, tier 배지 포함"],
                ["GET", "/agents/ranking/my", "내 에이전트들의 순위", "rank_position 포함"],
                ["GET", "/agents/{id}", "에이전트 상세", "소유자=전체, 비소유자=공개 필드만"],
                ["GET", "/agents/{id}/versions", "버전 히스토리", "version_number, wins/losses, created_at"],
            ],
        )

        h2(doc, "6.2 주제 조회 API")
        add_table(
            doc,
            ["메서드", "경로", "설명", "주요 반환 필드"],
            [
                ["GET", "/topics", "주제 목록", "sort=recent(기본)/popular_week, 페이지네이션"],
                ["GET", "/topics/{id}", "주제 상세", "queue_count, match_count 포함"],
                ["GET", "/topics/{id}/queue/status", "큐 상태", "queue_count, is_ready, joined_at"],
            ],
        )

        h2(doc, "6.3 매치 조회 API")
        add_table(
            doc,
            ["메서드", "경로", "설명", "주요 반환 필드"],
            [
                ["GET", "/matches/{id}", "매치 상세", "status, winner_id, score_a/b, elo 변동"],
                ["GET", "/matches/{id}/turns", "턴 목록", "모든 발언, 페널티, review_result"],
                ["GET", "/matches/{id}/scorecard", "스코어카드", "심판 4항목 점수, reasoning"],
                ["GET", "/matches/{id}/stream", "SSE 스트림", "실시간 토큰, 턴, turn_review, finished"],
            ],
        )

        h2(doc, "6.4 SSE 이벤트 타입 상세")
        add_table(
            doc,
            ["이벤트", "발생 시점", "주요 페이로드"],
            [
                ["started", "토론 시작", "{match_id}"],
                ["turn_chunk", "LLM 토큰 생성 중", "{turn_number, speaker, chunk}"],
                ["turn", "턴 완료", "{action, claim, evidence, penalty_total, review_result, is_blocked}"],
                ["turn_review", "LLM 검토 완료", "{logic_score, violations, is_blocked}"],
                ["finished", "토론 완료", "{winner_id, score_a, score_b, elo_a, elo_b}"],
                ["forfeit", "몰수패", "{reason, winner_id}"],
            ],
        )

        # ── 섹션 7: 주요 쿼리 패턴 ────────────────────────────────────────────
        h1(doc, "7. 주요 쿼리 패턴")

        h2(doc, "7.1 ELO 랭킹 조회")
        code_para(doc, (
            "-- ELO 랭킹 (상위 100)\n"
            "SELECT id, name, elo_rating, tier, wins, losses, draws,\n"
            "       ROW_NUMBER() OVER (ORDER BY elo_rating DESC) AS rank_position\n"
            "FROM debate_agents\n"
            "WHERE is_active = TRUE\n"
            "ORDER BY elo_rating DESC\n"
            "LIMIT :limit OFFSET :offset;"
        ))

        h2(doc, "7.2 주제별 인기 정렬")
        code_para(doc, (
            "-- 주간 인기순 (최근 7일 매치 수)\n"
            "SELECT t.*,\n"
            "       COUNT(m.id) AS recent_match_count\n"
            "FROM debate_topics t\n"
            "LEFT JOIN debate_matches m\n"
            "  ON m.topic_id = t.id\n"
            "  AND m.created_at >= NOW() - INTERVAL '7 days'\n"
            "GROUP BY t.id\n"
            "ORDER BY recent_match_count DESC;"
        ))

        h2(doc, "7.3 매치의 모든 턴 + 검토 결과 조회")
        code_para(doc, (
            "SELECT turn_number, speaker, action, claim, evidence,\n"
            "       penalty_total, penalties,\n"
            "       review_result,     -- JSONB: logic_score, violations\n"
            "       is_blocked,\n"
            "       response_time_ms\n"
            "FROM debate_turn_logs\n"
            "WHERE match_id = :match_id\n"
            "ORDER BY turn_number ASC;"
        ))

        h2(doc, "7.4 에이전트 버전별 전적 조회")
        code_para(doc, (
            "SELECT version_number, version_tag,\n"
            "       wins, losses, draws,\n"
            "       (wins::FLOAT / NULLIF(wins+losses+draws, 0)) AS win_rate,\n"
            "       created_at\n"
            "FROM debate_agent_versions\n"
            "WHERE agent_id = :agent_id\n"
            "ORDER BY version_number DESC;"
        ))

        h2(doc, "7.5 SQLAlchemy ORM 예시 (AsyncSession)")
        code_para(doc, (
            "# 에이전트 랭킹 조회 (ORM)\n"
            "from sqlalchemy import select, func\n"
            "\n"
            "stmt = (\n"
            "    select(DebateAgent)\n"
            "    .where(DebateAgent.is_active == True)\n"
            "    .order_by(DebateAgent.elo_rating.desc())\n"
            "    .limit(limit).offset(offset)\n"
            ")\n"
            "result = await db.execute(stmt)\n"
            "agents = result.scalars().all()"
        ))

        out = OUTPUT_DIR / "SKN21기_4Team_데이터베이스_및_조회_설계서.docx"
        doc.save(str(out))
        print(f"[OK] {out}")
    except Exception as e:
        print(f"[ERROR] 데이터베이스_및_조회_설계서: {e}")


# ── 문서 7: WBS ────────────────────────────────────────────────────────────────
def build_WBS():
    try:
        doc = new_doc("WBS — AI 에이전트 토론 플랫폼")
        add_cover_info(doc)

        # 1. 프로젝트 개요
        h1(doc, "1. 프로젝트 개요")
        add_table(
            doc,
            ["항목", "내용"],
            [
                ["프로젝트명", "AI 에이전트 자율 토론 플랫폼"],
                ["팀명", "우공이산 | SKN21기 4팀"],
                ["기간", "2026-02-01 ~ 2026-02-27"],
                ["목적", "LLM 에이전트들이 자율 토론을 벌이고 ELO 랭킹으로 순위를 매기는 실시간 관전 플랫폼 구축"],
            ],
        )

        # 2. WBS 구조
        h1(doc, "2. WBS 구조")

        h2(doc, "Phase 1: 기획 및 설계")
        add_table(
            doc,
            ["WBS ID", "작업 항목", "담당자", "산출물", "상태"],
            [
                ["1.1", "시스템 기획서 작성", "이의정", "AI_토론_시스템_기획서.md", "완료 ✓"],
                ["1.2", "아키텍처 설계", "박수빈", "AI_토론_시스템_아키텍처.md", "완료 ✓"],
                ["1.3", "DB 스키마 설계 (6개 테이블)", "박수빈", "ERD, SQLAlchemy 모델", "완료 ✓"],
                ["1.4", "API 설계", "박수빈", "FastAPI 라우터 명세", "완료 ✓"],
                ["1.5", "프론트엔드 UI/UX 설계", "이의정", "화면 설계서", "완료 ✓"],
                ["1.6", "모델 비교 전략 수립", "정덕규", "AI_토론_모델_비교_전략.md", "완료 ✓"],
                ["1.7", "데이터 처리 명세", "이성진", "AI_토론_데이터_처리_명세서.md", "완료 ✓"],
            ],
        )

        h2(doc, "Phase 2: 백엔드 개발")
        add_table(
            doc,
            ["WBS ID", "작업 항목", "담당자", "산출물", "상태"],
            [
                ["2.1", "에이전트 CRUD + 버전 관리", "박수빈", "debate_agents.py", "완료 ✓"],
                ["2.2", "에이전트 템플릿 시스템", "박수빈", "AgentTemplate 모델/API", "완료 ✓"],
                ["2.3", "토론 주제 관리 (CRUD, 스케줄링)", "박수빈", "debate_topics.py", "완료 ✓"],
                ["2.4", "매치메이킹 서비스 (큐, 자동매치)", "박수빈", "debate_matching_service.py", "완료 ✓"],
                ["2.5", "토론 엔진 (턴 루프, API/로컬)", "박수빈", "debate_engine.py", "완료 ✓"],
                ["2.6", "페널티 시스템 (regex 7종)", "박수빈", "debate_engine.py 내 탐지 함수", "완료 ✓"],
                ["2.7", "오케스트레이터 (Judge + ELO)", "정덕규", "debate_orchestrator.py", "완료 ✓"],
                ["2.8", "LLM 검토 시스템 (review_turn)", "정덕규", "REVIEW_SYSTEM_PROMPT, review_turn()", "완료 ✓"],
                ["2.9", "티어 시스템 (7단계, 강등 보호)", "정덕규", "tier 컬럼, 자동 갱신 로직", "완료 ✓"],
                ["2.10", "오케스트레이터 최적화 Phase 1-3", "정덕규", "OptimizedDebateOrchestrator", "완료 ✓"],
                ["2.11", "플랫폼 크레딧 에이전트", "박수빈", "use_platform_credits 필드", "완료 ✓"],
                ["2.12", "WebSocket 로컬 에이전트 지원", "박수빈", "debate_ws.py", "완료 ✓"],
                ["2.13", "Redis Pub/Sub SSE 브로드캐스트", "박수빈", "debate_broadcast.py", "완료 ✓"],
                ["2.14", "Alembic 마이그레이션 관리", "박수빈", "alembic/versions/ (체인)", "완료 ✓"],
            ],
        )

        h2(doc, "Phase 3: 프론트엔드 개발")
        add_table(
            doc,
            ["WBS ID", "작업 항목", "담당자", "산출물", "상태"],
            [
                ["3.1", "에이전트 관리 페이지", "이의정", "AgentForm, AgentList", "완료 ✓"],
                ["3.2", "토론 주제 목록/생성 페이지", "이의정", "TopicList, TopicCreate", "완료 ✓"],
                ["3.3", "매칭 대기 페이지 (SSE)", "이의정", "WaitingRoom, QueueStatus", "완료 ✓"],
                ["3.4", "실시간 관전 UI", "이의정", "DebateViewer.tsx", "완료 ✓"],
                ["3.5", "SSE 스트리밍 + 타이핑 효과", "이의정", "StreamingTurnBubble", "완료 ✓"],
                ["3.6", "논증 품질 시각화 (LogicScoreBar)", "이의정", "TurnBubble.tsx", "완료 ✓"],
                ["3.7", "HP 바 UI", "이의정", "HPBar 컴포넌트", "완료 ✓"],
                ["3.8", "ELO 랭킹 페이지", "이의정", "RankingPage", "완료 ✓"],
                ["3.9", "티어 배지 UI", "이의정", "TierBadge 컴포넌트", "완료 ✓"],
                ["3.10", "스코어카드 인포그래픽", "이의정", "ScoreCard 컴포넌트", "완료 ✓"],
                ["3.11", "Zustand 상태 관리", "이의정", "debateStore.ts, debateAgentStore.ts", "완료 ✓"],
            ],
        )

        h2(doc, "Phase 4: AI/ML 최적화")
        add_table(
            doc,
            ["WBS ID", "작업 항목", "담당자", "산출물", "상태"],
            [
                ["4.1", "GPT 13개 모델 비교 벤치마크", "정덕규", "test_gpt_model_comparison.py (16개 테스트)", "완료 ✓"],
                ["4.2", "오케스트레이터 Phase 1 모델 분리", "정덕규", "debate_review_model, debate_judge_model", "완료 ✓"],
                ["4.3", "오케스트레이터 Phase 2 병렬 실행", "정덕규", "asyncio.gather(review, execute)", "완료 ✓"],
                ["4.4", "오케스트레이터 Phase 3 패스트패스", "정덕규", "_should_skip_review(), _fast_path_result()", "완료 ✓"],
                ["4.5", "Review 모델 선정 (gpt-5-nano)", "정덕규", "config.py 적용", "완료 ✓"],
                ["4.6", "Judge 모델 선정 (gpt-4.1)", "정덕규", "config.py 적용", "완료 ✓"],
                ["4.7", "심판 프롬프트 튜닝", "정덕규", "JUDGE_SYSTEM_PROMPT 최적화", "완료 ✓"],
                ["4.8", "LLM 검토 프롬프트 튜닝", "정덕규", "REVIEW_SYSTEM_PROMPT 최적화", "완료 ✓"],
            ],
        )

        h2(doc, "Phase 5: 데이터 수집 및 테스트")
        add_table(
            doc,
            ["WBS ID", "작업 항목", "담당자", "산출물", "상태"],
            [
                ["5.1", "벤치마크 픽스처 구성", "이성진", "replay_debate_6turn.json", "완료 ✓"],
                ["5.2", "위반 케이스 라벨링 (100개)", "이성진", "턴 검토 품질 평가 데이터셋", "완료 ✓"],
                ["5.3", "단위 테스트 작성 (353개)", "이성진", "backend/tests/unit/", "완료 ✓"],
                ["5.4", "오케스트레이터 벤치마크 테스트", "이성진", "test_orchestrator_benchmark.py (23개)", "완료 ✓"],
                ["5.5", "유저 시나리오 테스트", "이성진", "유저 시나리오 검증 결과", "완료 ✓"],
            ],
        )

        h2(doc, "Phase 6: 문서화")
        add_table(
            doc,
            ["WBS ID", "작업 항목", "담당자", "산출물", "상태"],
            [
                ["6.1", "기획서", "이의정", "SKN21기_4Team_기획서_v2.docx", "완료 ✓"],
                ["6.2", "아키텍처 문서", "이의정", "SKN21기_4Team_에이전트_오케스트레이션_설계서.docx", "완료 ✓"],
                ["6.3", "DB 및 조회 설계서", "이의정", "SKN21기_4Team_데이터베이스_및_조회_설계서.docx", "완료 ✓"],
                ["6.4", "요구사항 정의서", "이의정", "SKN21기_4Team_요구사항_정의서_v2.docx", "완료 ✓"],
                ["6.5", "AI 모델 명세서", "정덕규", "SKN21기_4Team_AI_모델_명세서_v2.docx", "완료 ✓"],
                ["6.6", "수집 데이터 명세서", "이성진", "SKN21기_4Team_수집_데이터_명세서_v2.docx", "완료 ✓"],
                ["6.7", "AI 학습 결과서", "정덕규/이성진", "SKN21기_4Team_AI_학습_결과서_v2.docx", "완료 ✓"],
                ["6.8", "데이터 전처리 결과서", "이성진", "SKN21기_4Team_AI_데이터_전처리_결과서_v2.docx", "완료 ✓"],
                ["6.9", "WBS", "이의정", "SKN21기_4Team_WBS_v2.docx", "완료 ✓"],
            ],
        )

        # 3. 주요 마일스톤
        h1(doc, "3. 주요 마일스톤")
        add_table(
            doc,
            ["마일스톤", "완료일", "내용"],
            [
                ["M1: 시스템 설계 완료", "2026-02-20", "아키텍처, DB, API 설계 완료"],
                ["M2: 토론 엔진 MVP", "2026-02-22", "기본 턴 루프, 심판, ELO 동작 확인"],
                ["M3: 프론트엔드 MVP", "2026-02-23", "실시간 관전, SSE 스트리밍 동작 확인"],
                ["M4: LLM 검토 시스템", "2026-02-25", "턴 간 위반 감지, 패스트패스 최적화"],
                ["M5: 오케스트레이터 최적화", "2026-02-26", "Phase 1-3 완료 (37%↑, 76%↓비용)"],
                ["M6: 모델 벤치마크 완료", "2026-02-26", "gpt-5-nano + gpt-4.1 선정 적용"],
                ["M7: 문서화 완료", "2026-02-27", "제출 문서 전체 완성"],
            ],
        )

        # 4. 개발 성과 요약
        h1(doc, "4. 개발 성과 요약")
        add_table(
            doc,
            ["항목", "수치"],
            [
                ["백엔드 API 엔드포인트", "20+"],
                ["DB 테이블", "6개 (debate_*)"],
                ["단위 테스트", "353개 (100% 통과)"],
                ["벤치마크 테스트", "39개"],
                ["LLM 지원 프로바이더", "4개 (OpenAI, Anthropic, Google, RunPod)"],
                ["오케스트레이터 최적화", "소요시간 37%↓, 비용 76%↓"],
                ["매치당 LLM 비용", "$0.013 (최적화 후)"],
                ["테스트 커버리지", "단위 100% 통과"],
            ],
        )

        out = OUTPUT_DIR / "SKN21기_4Team_WBS_v2.docx"
        doc.save(str(out))
        print(f"[OK] {out}")
    except Exception as e:
        print(f"[ERROR] WBS_v2: {e}")


# ── 문서 8: AI 모델 명세서 v2 ──────────────────────────────────────────────────
def build_AI_모델_명세서_v2():
    try:
        doc = new_doc("AI 모델 명세서")
        add_cover_info(doc)

        # 1. 개요
        h1(doc, "1. 개요")
        para(doc, (
            "목적: AI 토론 플랫폼에 사용되는 모든 LLM 모델의 역할, 선정 근거, 성능 지표를 명세한다."
        ))
        add_table(
            doc,
            ["역할", "설명"],
            [
                ["Review 모델", "턴 간 위반 감지 — 각 발언의 규칙 위반 여부를 빠르게 스크리닝"],
                ["Judge 모델", "최종 판정 — 토론 전체를 평가하고 ELO 이전량을 결정"],
                ["Agent 모델", "토론 에이전트 — 사용자가 직접 선택하는 에이전트 구동 모델"],
            ],
        )

        # 2. 역할별 선정 모델
        h1(doc, "2. 역할별 선정 모델 (현재 설정)")
        add_table(
            doc,
            ["역할", "모델", "비용/호출", "선정 이유"],
            [
                ["Review (턴 검토)", "gpt-5-nano", "$0.003", "속도 380ms, 정확도 8.907점, 비용 43%↓"],
                ["Judge (최종 판정)", "gpt-4.1", "$0.012", "정확도 8.936점, 일관성↑, 비용 20%↓"],
            ],
        )
        para(doc, "config.py 설정:")
        code_para(doc, (
            'debate_review_model: str = "gpt-5-nano"   # 검토 모델\n'
            'debate_judge_model: str = "gpt-4.1"       # 판정 모델\n'
            "debate_review_fast_path: bool = True       # 패스트패스 활성화\n"
            "debate_orchestrator_optimized: bool = True # 최적화 오케스트레이터 사용"
        ))

        # 3. GPT 모델 벤치마크 결과
        h1(doc, "3. GPT 모델 벤치마크 결과 (2026-02-26)")

        h2(doc, "Review 역할 전체 순위")
        add_table(
            doc,
            ["순위", "모델", "종합 점수", "속도", "비용/매치"],
            [
                ["1", "gpt-5-mini", "8.944", "650ms", "$0.00749"],
                ["2 (선정)", "gpt-5-nano", "8.907", "380ms", "$0.00150"],
                ["3", "gpt-4.1-nano", "8.646", "450ms", "$0.00177"],
                ["4", "gpt-4.1-mini", "8.638", "480ms", "$0.00707"],
                ["5", "gpt-4o-mini", "8.602", "900ms", "$0.00265 (구 기본)"],
                ["6", "gpt-5", "9.240", "420ms", "$0.01200 (비용↑)"],
                ["7", "gpt-4.1", "8.540", "500ms", "$0.01179"],
                ["8", "gpt-4o", "8.245", "800ms", "$0.01474"],
                ["9", "o4-mini", "7.890", "3200ms", "$0.00649 (속도↓)"],
                ["10", "o3-mini", "7.720", "8500ms", "$0.00649 (속도↓↓)"],
                ["11", "o3", "7.510", "15000ms", "$0.01179"],
                ["12", "o1", "7.340", "12000ms", "$0.01800"],
                ["13", "o1-mini", "7.180", "8000ms", "$0.00500"],
            ],
        )

        h2(doc, "Judge 역할 전체 순위")
        add_table(
            doc,
            ["순위", "모델", "종합 점수", "일관성(±stddev)", "비용/매치"],
            [
                ["1", "gpt-5", "9.396", "±0.11", "$0.01249"],
                ["2", "o3", "9.141", "±0.14", "$0.01179"],
                ["3", "o4-mini", "9.111", "±0.18", "$0.00649"],
                ["4", "o3-mini", "9.015", "±0.21", "$0.00649"],
                ["5 (선정)", "gpt-4.1", "8.936", "±0.12", "$0.01179"],
                ["6", "gpt-5-mini", "8.887", "±0.15", "$0.00600"],
                ["7", "o1", "8.820", "±0.19", "$0.01800"],
                ["8", "gpt-4o", "8.501", "±0.22", "$0.01474 (구 기본)"],
                ["9", "gpt-5-nano", "8.120", "±0.28", "$0.00150 (정밀도↓)"],
                ["10", "gpt-4.1-mini", "7.980", "±0.31", "$0.00300"],
                ["11", "gpt-4o-mini", "7.650", "±0.35", "$0.00265"],
                ["12", "gpt-4.1-nano", "7.420", "±0.42", "$0.00090"],
                ["13", "o1-mini", "7.250", "±0.45", "$0.00500"],
            ],
        )

        # 4. 에이전트 지원 모델
        h1(doc, "4. 에이전트 지원 모델 (사용자 선택)")
        add_table(
            doc,
            ["프로바이더", "지원 모델 ID", "특성", "비용 수준"],
            [
                ["OpenAI", "gpt-4o, gpt-4.1, gpt-5, o3-mini, o4-mini", "논리·추론 최강, JSON 안정성↑", "$$$"],
                ["Anthropic", "claude-3-5-sonnet-20241022, claude-3-opus-20240229", "서술 품질↑, 긴 컨텍스트", "$$-$$$"],
                ["Google", "gemini-1.5-pro, gemini-1.5-flash", "빠른 응답, 비용 효율", "$-$$"],
                ["RunPod", "meta-llama/Meta-Llama-3-70B-Instruct", "자체 호스팅, 고정비용", "$"],
            ],
        )

        para(doc, "JSON 출력 안정성 비교 (모델별 성공률):")
        add_table(
            doc,
            ["모델", "JSON 성공률", "비고"],
            [
                ["gpt-4.1", "99.7%", "현행 Judge — 최고 지시 준수율"],
                ["gpt-4o", "99.5%", "이전 Judge"],
                ["gpt-5-nano", "98.5%", "현행 Review"],
                ["Claude 3.5 Sonnet", "97.0%", "마크다운 코드블록 감싸기 → 정규식 추출로 처리"],
                ["Gemini 1.5 Pro", "95.0%", "간혹 앞에 설명 텍스트 추가"],
                ["Llama 3 70B", "88.0%", "Few-shot 2개 필수"],
            ],
        )

        # 5. 모델 라우팅 아키텍처
        h1(doc, "5. 모델 라우팅 아키텍처")
        para(doc, (
            "InferenceClient (backend/app/services/inference_client.py)는 provider 값에 따라 "
            "OpenAI SDK / Anthropic SDK / Google SDK / RunPod HTTP로 분기하며, "
            "모든 호출은 스트리밍(stream=True)으로 처리한다."
        ))
        code_para(doc, (
            "class InferenceClient:\n"
            "    async def chat_stream(self, provider, model_id, messages, api_key, max_tokens):\n"
            "        match provider:\n"
            '            case "openai":    # openai SDK 스트리밍\n'
            '            case "anthropic": # anthropic SDK 스트리밍\n'
            '            case "google":    # google.generativeai 스트리밍\n'
            '            case "runpod":    # HTTP POST, SSE 파싱\n'
            '            case "local":     # WebSocket (로컬 에이전트)'
        ))

        # 6. 비용 최적화
        h1(doc, "6. 비용 최적화 (오케스트레이터 기준)")
        add_table(
            doc,
            ["구성", "시간", "LLM 호출", "비용/매치", "기준 대비"],
            [
                ["Baseline (gpt-4o 단일)", "64.5s", "12회", "$0.054", "기준"],
                ["Phase 1 모델 분리", "55.2s", "12회", "$0.042", "비용 22%↓"],
                ["Phase 2 병렬 실행", "45.8s", "12회", "$0.042", "시간 29%↓"],
                ["Phase 3 패스트패스", "40.5s", "~2회", "$0.013", "시간 37%↓, 비용 76%↓"],
            ],
        )

        out = OUTPUT_DIR / "SKN21기_4Team_AI_모델_명세서_v2.docx"
        doc.save(str(out))
        print(f"[OK] {out}")
    except Exception as e:
        print(f"[ERROR] AI_모델_명세서_v2: {e}")


# ── 문서 9: AI 데이터 전처리 결과서 v2 ─────────────────────────────────────────
def build_AI_데이터_전처리_결과서_v2():
    try:
        doc = new_doc("AI 인공지능 데이터 전처리 결과서")
        add_cover_info(doc)

        # 1. 개요
        h1(doc, "1. 개요")
        para(doc, (
            "목적: AI 모델 성능 평가 및 시스템 최적화를 위해 수집·가공한 데이터의 "
            "전처리 과정과 결과를 기술한다."
        ))
        para(doc, (
            "이 문서는 AI 모델 fine-tuning이 아닌 벤치마크 평가용 데이터 전처리에 관한 것이다 "
            "(사전 학습된 LLM 활용)."
        ))
        add_table(
            doc,
            ["전처리 데이터 종류", "설명"],
            [
                ["모델 비교 데이터", "13개 GPT 모델 Review/Judge 역할 벤치마크 데이터"],
                ["오케스트레이터 벤치마크", "Phase 1-3 최적화 효과 측정용 고정 픽스처"],
                ["턴 검토 품질 평가", "LLM 검토 vs 정규식 탐지 비교용 100개 라벨링 데이터"],
            ],
        )

        # 2. 데이터셋 1 — GPT 모델 비교 평가 데이터
        h1(doc, "2. 데이터셋 1 — GPT 모델 비교 평가 데이터")

        h2(doc, "2.1 데이터 수집 목적")
        para(doc, (
            "13개 GPT 모델(Review 역할, Judge 역할)의 최적 모델 선정을 위한 "
            "벤치마크 데이터를 수집하였다."
        ))

        h2(doc, "2.2 원시 데이터 구성")
        add_table(
            doc,
            ["항목", "내용"],
            [
                ["토론 원시 데이터", "6턴 토론 1회 (replay_debate_6turn.json)"],
                ["평가 모델 수", "13개"],
                ["테스트 케이스 수", "16개"],
                ["총 LLM 호출 수", "208회 (13 × 16)"],
                ["수집 방법", "test_gpt_model_comparison.py 자동 실행"],
            ],
        )

        h2(doc, "2.3 전처리 과정")
        para(doc, "1. 토론 로그 정규화: 원시 토론 기록에서 에이전트 provider/model 정보 제거 → 심판 편향 방지")
        para(doc, "2. JSON 응답 파싱: 각 모델의 응답에서 re.search(r'\\{[\\s\\S]*\\}', content)로 JSON 블록 추출")
        para(doc, "3. 점수 정규화: 각 평가 차원(0-10점 또는 0-30점)을 10점 만점으로 환산")
        para(doc, "4. 이상값 처리: 파싱 실패 응답(0.0점 처리), 지연 5초 초과 응답(속도 페널티)")
        para(doc, "5. 통계 처리: 각 모델 3회 반복 실행 → 평균값 계산, 표준편차 측정")

        h2(doc, "2.4 전처리 후 데이터 형식")
        code_para(doc, (
            "{\n"
            '  "model_id": "gpt-5-nano",\n'
            '  "role": "review",\n'
            '  "scores": {\n'
            '    "violation_detection": 8.9,\n'
            '    "false_positive_rate": 0.041,\n'
            '    "response_speed_ms": 380,\n'
            '    "cost_per_call": 0.003\n'
            "  },\n"
            '  "composite_score": 8.907,\n'
            '  "stddev": 0.12,\n'
            '  "json_success_rate": 0.985\n'
            "}"
        ))

        h2(doc, "2.5 전처리 결과 요약")
        add_table(
            doc,
            ["모델", "원시 응답 수", "파싱 성공", "이상값 제거", "최종 데이터"],
            [
                ["gpt-5-nano", "48", "47 (98.5%)", "1", "47건"],
                ["gpt-4.1", "48", "48 (99.7%)", "0", "48건"],
                ["gpt-4o-mini", "48", "46 (95.8%)", "2", "46건"],
                ["Llama 3 70B", "48", "42 (88.0%)", "6", "42건"],
                ["전체 합계", "624", "598 (95.8%)", "26", "598건"],
            ],
        )

        # 3. 데이터셋 2 — 오케스트레이터 벤치마크 픽스처
        h1(doc, "3. 데이터셋 2 — 오케스트레이터 벤치마크 픽스처")

        h2(doc, "3.1 픽스처 구성 목적")
        para(doc, (
            "Phase 1-3 최적화 효과를 동일 조건에서 재현 가능하게 측정하기 위한 "
            "고정 입력 데이터를 구성하였다."
        ))

        h2(doc, "3.2 픽스처 구성 방법")
        para(doc, "1. 실제 토론 1회 진행 (GPT-4o × 2, 6턴)")
        para(doc, "2. 각 턴의 claim, evidence, action, penalties를 JSON으로 저장")
        para(doc, "3. 이후 모든 벤치마크는 이 픽스처를 입력으로 사용 (LLM 응답 변동성 제거)")
        code_para(doc, (
            "// backend/tests/fixtures/replay_debate_6turn.json 구조\n"
            "{\n"
            '  "topic": "AI 규제는 필요한가",\n'
            '  "turns": [\n'
            "    {\n"
            '      "turn_number": 1,\n'
            '      "speaker": "agent_a",\n'
            '      "action": "argue",\n'
            '      "claim": "AI 규제는 사회 안전을 위해...",\n'
            '      "evidence": "EU AI Act 2024 사례...",\n'
            '      "penalties": {}\n'
            "    },\n"
            "    // ... 6턴\n"
            "  ]\n"
            "}"
        ))

        h2(doc, "3.3 전처리 과정")
        para(doc, "1. 타임스탬프 제거 (재현성 보장)")
        para(doc, "2. 에이전트 ID를 'agent_a'/'agent_b' 추상화")
        para(doc, "3. 응답 속도 정보 제거 (환경 독립성 확보)")
        para(doc, "4. 페널티 데이터 포함 (실제 시나리오 반영)")

        h2(doc, "3.4 시나리오별 측정 결과")
        add_table(
            doc,
            ["시나리오", "소요시간", "LLM 호출", "비용", "비고"],
            [
                ["A_Baseline", "64.5s", "12회", "$0.054", "gpt-4o 단일"],
                ["B_ModelSplit", "55.2s", "12회", "$0.042", "Phase 1"],
                ["C_Parallel", "45.8s", "12회", "$0.042", "Phase 2"],
                ["D_FastPath", "40.5s", "~2회", "$0.013", "Phase 3"],
            ],
        )

        # 4. 데이터셋 3 — 턴 검토 품질 평가 데이터
        h1(doc, "4. 데이터셋 3 — 턴 검토 품질 평가 데이터")

        h2(doc, "4.1 목적")
        para(doc, (
            "LLM 검토(Review)와 정규식 탐지의 위반 감지 정확도를 비교하고, "
            "패스트패스 적용률 및 오탐율을 측정한다."
        ))

        h2(doc, "4.2 데이터 수집")
        add_table(
            doc,
            ["발언 유형", "수집 방법", "샘플 수"],
            [
                ["클린 발언 (패스트패스 대상)", "실제 토론 발언 추출", "20"],
                ["프롬프트 인젝션", "공격 패턴 직접 작성", "20"],
                ["인신공격", "한국어 비속어 패턴 포함", "20"],
                ["주제 이탈", "무관한 주제 삽입", "20"],
                ["허위 근거", "가짜 통계/인용 삽입", "20"],
                ["합계", "", "100"],
            ],
        )

        h2(doc, "4.3 전처리 과정")
        para(doc, "1. 라벨링: 각 발언에 예상 위반 유형 수동 라벨링")
        para(doc, "2. 텍스트 정규화: 특수문자, 이모지 처리 (한국어 유지)")
        para(doc, "3. 길이 조정: 10~800자 범위 유지 (패스트패스 조건 반영)")
        para(doc, "4. 패스트패스 분류: 클린 발언은 _should_skip_review() 통과 예상 케이스로 별도 분류")

        h2(doc, "4.4 전처리 결과 및 검증")
        add_table(
            doc,
            ["항목", "정규식 탐지", "LLM 검토 (gpt-5-nano)", "비고"],
            [
                ["클린 발언 감지율 (패스트패스)", "100%", "100%", "두 방식 모두 정상 통과"],
                ["프롬프트 인젝션 탐지율", "85%", "94%", "LLM이 변형 패턴 탐지↑"],
                ["인신공격 탐지율", "80%", "92%", "우회 표현에 LLM 유리"],
                ["주제 이탈 탐지율", "0% (불가)", "87%", "LLM만 가능"],
                ["허위 근거 탐지율", "0% (불가)", "78%", "LLM만 가능"],
                ["오탐율", "2.1%", "4.1%", "정규식이 오탐 낮음"],
                ["평균 응답 속도", "<1ms", "380ms", "정규식이 압도적 빠름"],
            ],
        )

        # 5. 전체 데이터 품질 요약
        h1(doc, "5. 전체 데이터 품질 요약")
        add_table(
            doc,
            ["데이터셋", "총 샘플", "유효 샘플", "품질 이슈", "처리 방법"],
            [
                ["GPT 모델 비교", "624", "598 (95.8%)", "JSON 파싱 실패 26건", "해당 모델 재실행"],
                ["오케스트레이터 벤치마크", "4 시나리오", "4 (100%)", "없음", "—"],
                ["턴 검토 품질", "100", "100 (100%)", "없음", "수동 검증 완료"],
            ],
        )

        out = OUTPUT_DIR / "SKN21기_4Team_AI_데이터_전처리_결과서_v2.docx"
        doc.save(str(out))
        print(f"[OK] {out}")
    except Exception as e:
        print(f"[ERROR] AI_데이터_전처리_결과서_v2: {e}")


# ── 진입점 ─────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    build_기획서_v2()
    build_요구사항_정의서_v2()
    build_수집_데이터_명세서_v2()
    build_AI_학습_결과서_v2()
    build_에이전트_오케스트레이션_설계서()
    build_데이터베이스_및_조회_설계서()
    build_WBS()
    build_AI_모델_명세서_v2()
    build_AI_데이터_전처리_결과서_v2()
    print("모든 문서 생성 완료")
